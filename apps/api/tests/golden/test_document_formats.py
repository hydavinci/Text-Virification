from __future__ import annotations

import io
import shutil
from collections.abc import Callable
from pathlib import Path
from uuid import uuid4

import pymupdf
import pytest
from docx import Document

from text_verification.application.factory import build_default_verification_pipeline
from text_verification.application.verification_pipeline import VerificationCommand
from text_verification.compatibility.exporters import export_original
from text_verification.config import Settings
from text_verification.domain.documents import FileType
from text_verification.domain.verification import (
    VerificationExecutionMode,
    VerificationOptions,
)
from text_verification.infrastructure.storage import JobStorage

SAMPLE_TEXT = "Contact test@example.com for review."
MATRIX_TEXT = "cat test@example.com"


def _txt(target: Path) -> bytes:
    del target
    return SAMPLE_TEXT.encode("utf-8")


def _markdown(target: Path) -> bytes:
    del target
    return f"# Review\n\n{SAMPLE_TEXT}".encode()


def _csv(target: Path) -> bytes:
    del target
    return f"kind,value\ncontact,{SAMPLE_TEXT}\n".encode()


def _rtf(target: Path) -> bytes:
    del target
    return ("{\\rtf1\\ansi " + SAMPLE_TEXT + "}").encode("ascii")


def _docx(target: Path) -> bytes:
    stream = io.BytesIO()
    document = Document()
    document.add_heading("Review", level=1)
    document.add_paragraph(SAMPLE_TEXT)
    document.save(stream)
    return stream.getvalue()


def _pdf(target: Path) -> bytes:
    document = pymupdf.open()
    try:
        page = document.new_page(width=500, height=200)
        page.insert_text((40, 80), SAMPLE_TEXT, fontsize=14, fontname="helv")
        return document.tobytes()
    finally:
        document.close()


def _matrix_payload(file_type: FileType, target: Path) -> bytes:
    if file_type is FileType.TXT:
        return MATRIX_TEXT.encode()
    if file_type is FileType.MARKDOWN:
        return f"# Review\n\n{MATRIX_TEXT}".encode()
    if file_type is FileType.CSV:
        return f"kind,value\ncontact,{MATRIX_TEXT}\n".encode()
    if file_type is FileType.RTF:
        return ("{\\rtf1\\ansi " + MATRIX_TEXT + "}").encode("ascii")
    if file_type is FileType.DOCX:
        stream = io.BytesIO()
        document = Document()
        document.add_paragraph(MATRIX_TEXT)
        document.save(stream)
        return stream.getvalue()
    if file_type is FileType.PDF:
        document = pymupdf.open()
        try:
            page = document.new_page(width=500, height=200)
            page.insert_text((40, 80), MATRIX_TEXT, fontsize=14, fontname="helv")
            return document.tobytes()
        finally:
            document.close()
    if file_type is FileType.DOC:
        converter = (
            shutil.which("textutil")
            or shutil.which("soffice")
            or shutil.which("libreoffice")
        )
        if converter is None:
            pytest.skip(
                "Legacy DOC golden behavior requires textutil or LibreOffice; "
                "neither production converter is available."
            )
        from text_verification.compatibility.exporters import (
            _convert_docx_bytes_to_doc,
        )

        return _convert_docx_bytes_to_doc(
            _matrix_payload(FileType.DOCX, target),
            target.parent,
        )
    raise AssertionError(file_type)


FORMAT_BUILDERS: dict[FileType, Callable[[Path], bytes]] = {
    FileType.TXT: _txt,
    FileType.DOCX: _docx,
    FileType.PDF: _pdf,
    FileType.RTF: _rtf,
    FileType.MARKDOWN: _markdown,
    FileType.CSV: _csv,
}


@pytest.mark.parametrize("file_type", list(FileType))
def test_async_storage_accepts_all_seven_formats(
    tmp_path: Path,
    file_type: FileType,
) -> None:
    storage = JobStorage(tmp_path, max_upload_bytes=5 * 1024 * 1024)
    if file_type is FileType.DOC:
        payload = bytes.fromhex("D0CF11E0A1B11AE1") + b"\x00" * 16
    else:
        payload = FORMAT_BUILDERS[file_type](tmp_path / f"fixture.{file_type.value}")

    stored = storage.save_bytes(uuid4(), f"fixture.{file_type.value}", payload)

    assert stored.file_type is file_type


@pytest.mark.parametrize("file_type", list(FORMAT_BUILDERS))
def test_six_self_contained_golden_formats_produce_equivalent_issue_semantics(
    tmp_path: Path,
    file_type: FileType,
) -> None:
    source = tmp_path / f"fixture.{file_type.value}"
    source.write_bytes(FORMAT_BUILDERS[file_type](source))
    pipeline = build_default_verification_pipeline(Settings(llm_api_key=""))

    result = pipeline.run(
        VerificationCommand(
            document_id=uuid4(),
            source_path=source,
            direct_text=None,
            source_name=source.name,
            file_type=file_type,
            options=VerificationOptions(),
            execution_mode=VerificationExecutionMode.ASYNCHRONOUS,
        )
    )

    issue = next(issue for issue in result.issues if issue.type == "pii_email")
    assert SAMPLE_TEXT in result.text
    assert issue.original == "test@example.com"
    assert result.text[issue.start : issue.end] == issue.original
    assert issue.block_id is not None
    assert result.execution_mode is VerificationExecutionMode.ASYNCHRONOUS
    assert result.source_version.startswith("sha256:")


def test_legacy_doc_golden_is_explicitly_converter_limited(tmp_path: Path) -> None:
    source = tmp_path / "legacy.doc"
    source.write_bytes(_matrix_payload(FileType.DOC, source))
    pipeline = build_default_verification_pipeline(Settings(llm_api_key=""))

    result = pipeline.run(
        VerificationCommand(
            document_id=uuid4(),
            source_path=source,
            direct_text=None,
            source_name=source.name,
            file_type=FileType.DOC,
            options=VerificationOptions(),
            execution_mode=VerificationExecutionMode.ASYNCHRONOUS,
        )
    )

    assert MATRIX_TEXT in result.text
    assert any(issue.type == "pii_email" for issue in result.issues)


@pytest.mark.parametrize("file_type", list(FileType))
def test_seven_format_parse_verify_export_reparse_semantics(
    tmp_path: Path,
    file_type: FileType,
) -> None:
    source = tmp_path / f"matrix.{file_type.value}"
    source.write_bytes(_matrix_payload(file_type, source))
    pipeline = build_default_verification_pipeline(Settings(llm_api_key=""))
    options = VerificationOptions(
        custom_glossary=({"original": "cat", "standard": "dog"},),
    )

    result = pipeline.run(
        VerificationCommand(
            document_id=uuid4(),
            source_path=source,
            direct_text=None,
            source_name=source.name,
            file_type=file_type,
            options=options,
            execution_mode=VerificationExecutionMode.ASYNCHRONOUS,
        )
    )
    glossary_issue = next(issue for issue in result.issues if issue.original == "cat")
    assert glossary_issue.suggestion == "dog"
    assert any(issue.type == "pii_email" for issue in result.issues)
    exported = export_original(
        source,
        file_type.value,
        [
            (
                glossary_issue.original,
                glossary_issue.suggestion or "",
                glossary_issue.start,
                glossary_issue.end,
            )
        ],
        False,
        original_text=result.text,
    )
    exported_path = tmp_path / f"exported.{exported.extension}"
    exported_path.write_bytes(exported.content)
    exported_type = FileType(exported.extension)

    reparsed = pipeline.run(
        VerificationCommand(
            document_id=uuid4(),
            source_path=exported_path,
            direct_text=None,
            source_name=exported_path.name,
            file_type=exported_type,
            options=VerificationOptions(),
            execution_mode=VerificationExecutionMode.ASYNCHRONOUS,
        )
    )

    assert "dog" in reparsed.text
    assert "cat" not in reparsed.text
    assert any(issue.type == "pii_email" for issue in reparsed.issues)
