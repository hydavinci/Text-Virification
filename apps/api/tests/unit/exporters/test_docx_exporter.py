from __future__ import annotations

from pathlib import Path
from uuid import UUID, uuid4
from zipfile import ZipFile

import pytest
from docx import Document as WordDocument

import text_verification.exporters as exporters
from text_verification.exporters import (
    DocxExporter,
    ExportError,
    ExportWarning,
    Replacement,
    ReplacementPlan,
)
from text_verification.parsers.docx import DocxParser


@pytest.fixture
def fixture_path() -> Path:
    return Path(__file__).resolve().parents[2] / "fixtures" / "documents"


def test_docx_export_uses_current_sample_paragraph_locator_for_safe_replacement(
    fixture_path: Path,
    tmp_path: Path,
) -> None:
    source = fixture_path / "sample.docx"
    parsed = parse_docx(source)
    paragraph = next(block for block in parsed.blocks if block.kind == "paragraph")
    assert paragraph.block_id == "p-000001"

    result = DocxExporter().export(
        source,
        parsed,
        build_plan(
            Replacement(
                block_id=paragraph.block_id,
                start=2,
                end=6,
                original="示例文本",
                value="专业",
                issue_id=UUID("00000000-0000-0000-0000-000000000001"),
            )
        ),
        tmp_path / "modified.docx",
    )

    reparsed = parse_docx(result.path, source_name="modified.docx")
    exported = WordDocument(str(result.path))

    assert reparsed.blocks[1].text == "核验专业"
    assert exported.paragraphs[1].runs[0].text == "核验"
    assert exported.paragraphs[1].runs[1].text == "专业"
    assert result.warnings == []


def test_docx_export_preserves_run_formatting_for_single_run_replacement(
    tmp_path: Path,
) -> None:
    source = tmp_path / "styled.docx"
    document = WordDocument()
    paragraph = document.add_paragraph()
    paragraph.add_run("核验")
    styled_run = paragraph.add_run("示例文本")
    styled_run.bold = True
    document.save(source)

    parsed = parse_docx(source)
    block = parsed.blocks[0]

    result = DocxExporter().export(
        source,
        parsed,
        build_plan(
            Replacement(
                block_id=block.block_id,
                start=2,
                end=6,
                original="示例文本",
                value="专业",
                issue_id=UUID("00000000-0000-0000-0000-000000000002"),
            )
        ),
        tmp_path / "formatted.docx",
    )

    exported = WordDocument(str(result.path))

    assert exported.paragraphs[0].runs[1].text == "专业"
    assert exported.paragraphs[0].runs[1].bold is True
    assert not (tmp_path / "formatted.docx.tmp").exists()


def test_docx_export_warns_and_byte_copies_when_no_safe_sample_replacement_applies(
    fixture_path: Path,
    tmp_path: Path,
) -> None:
    source = fixture_path / "sample.docx"
    parsed = parse_docx(source)
    paragraph = next(block for block in parsed.blocks if block.kind == "paragraph")
    issue_id = UUID("00000000-0000-0000-0000-000000000003")

    result = DocxExporter().export(
        source,
        parsed,
        build_plan(
            Replacement(
                block_id=paragraph.block_id,
                start=1,
                end=4,
                original="验示例",
                value="替换",
                issue_id=issue_id,
            )
        ),
        tmp_path / "cross-run.docx",
    )

    exported = WordDocument(str(result.path))

    assert exported.paragraphs[1].text == "核验示例文本"
    assert result.path.read_bytes() == source.read_bytes()
    assert result.warnings == [
        ExportWarning(
            code="unsafe_docx_run_boundary",
            message=(
                "修改范围跨越多个 DOCX 文本运行，为保留格式已跳过；"
                "请在原文中手动修改后重新导出。"
            ),
            issue_id=issue_id,
            block_id=paragraph.block_id,
        )
    ]
    assert not (tmp_path / "cross-run.docx.tmp").exists()


def test_docx_export_skips_mixed_content_run_without_dropping_drawing_or_field_xml(
    fixture_path: Path,
    tmp_path: Path,
) -> None:
    source = fixture_path / "mixed-content-run.docx"
    parsed = parse_docx(source)
    block = parsed.blocks[0]
    issue_id = UUID("00000000-0000-0000-0000-000000000009")
    plan = build_plan(
        Replacement(
            block_id=block.block_id,
            start=2,
            end=4,
            original="正文",
            value="文稿",
            issue_id=issue_id,
        )
    )
    source_bytes = source.read_bytes()

    evaluated = exporters.DocxApplicabilityEvaluator().evaluate(source, parsed, plan)
    result = DocxExporter().export(
        source,
        parsed,
        plan,
        tmp_path / "mixed-content-run.docx",
    )

    assert source.read_bytes() == source_bytes
    assert evaluated.applicable == []
    assert evaluated.warnings == [
        ExportWarning(
            code="unsafe_docx_mixed_content",
            message=(
                "修改所在的 DOCX 文本运行还包含图片、域或其他非文本内容，"
                "为避免破坏文档已跳过；请在原文中手动修改后重新导出。"
            ),
            issue_id=issue_id,
            block_id=block.block_id,
        )
    ]
    assert result.warnings == evaluated.warnings
    assert result.path.read_bytes() == source_bytes
    with ZipFile(result.path) as archive:
        document_xml = archive.read("word/document.xml")
    assert b"drawing" in document_xml
    assert b"fldChar" in document_xml
    assert b"instrText" in document_xml
    assert b" DATE " in document_xml
    assert len(WordDocument(str(result.path)).inline_shapes) == 1


def test_docx_export_resolves_table_cell_locator_and_applies_descending_same_run_edits(
    tmp_path: Path,
) -> None:
    source = tmp_path / "table.docx"
    document = WordDocument()
    table = document.add_table(rows=1, cols=1)
    run = table.cell(0, 0).paragraphs[0].add_run("abcdefg")
    run.italic = True
    document.save(source)

    parsed = parse_docx(source)
    block = parsed.blocks[0]

    result = DocxExporter().export(
        source,
        parsed,
        build_plan(
            Replacement(
                block_id=block.block_id,
                start=1,
                end=3,
                original="bc",
                value="X",
                issue_id=UUID("00000000-0000-0000-0000-000000000004"),
            ),
            Replacement(
                block_id=block.block_id,
                start=4,
                end=6,
                original="ef",
                value="YZ",
                issue_id=UUID("00000000-0000-0000-0000-000000000005"),
            ),
        ),
        tmp_path / "table-modified.docx",
    )

    exported = WordDocument(str(result.path))
    cell_run = exported.tables[0].cell(0, 0).paragraphs[0].runs[0]

    assert exported.tables[0].cell(0, 0).text == "aXdYZg"
    assert cell_run.text == "aXdYZg"
    assert cell_run.italic is True
    assert result.warnings == []


def test_docx_export_resolves_second_paragraph_table_cell_locator(
    tmp_path: Path,
) -> None:
    source = tmp_path / "table-second-paragraph.docx"
    document = WordDocument()
    table = document.add_table(rows=1, cols=1)
    first_paragraph = table.cell(0, 0).paragraphs[0]
    first_paragraph.add_run("甲")
    first_paragraph.add_run("乙")
    second_paragraph = table.cell(0, 0).add_paragraph()
    second_paragraph.add_run("丙")
    styled_run = second_paragraph.add_run("丁戊")
    styled_run.bold = True
    document.save(source)

    parsed = parse_docx(source)
    block = parsed.blocks[0]
    assert block.text == "甲乙\n丙丁戊"

    result = DocxExporter().export(
        source,
        parsed,
        build_plan(
            Replacement(
                block_id=block.block_id,
                start=4,
                end=5,
                original="丁",
                value="庚",
                issue_id=UUID("00000000-0000-0000-0000-000000000007"),
            )
        ),
        tmp_path / "table-second-paragraph-modified.docx",
    )

    exported = WordDocument(str(result.path))
    exported_cell = exported.tables[0].cell(0, 0)

    assert exported_cell.paragraphs[0].text == "甲乙"
    assert exported_cell.paragraphs[1].runs[0].text == "丙"
    assert exported_cell.paragraphs[1].runs[1].text == "庚戊"
    assert exported_cell.paragraphs[1].runs[1].bold is True
    assert result.warnings == []


def test_docx_export_warns_for_cross_paragraph_table_cell_replacement(
    tmp_path: Path,
) -> None:
    source = tmp_path / "table-cross-paragraph.docx"
    document = WordDocument()
    table = document.add_table(rows=1, cols=1)
    first_paragraph = table.cell(0, 0).paragraphs[0]
    first_paragraph.add_run("甲")
    first_paragraph.add_run("乙")
    second_paragraph = table.cell(0, 0).add_paragraph()
    second_paragraph.add_run("丙")
    second_paragraph.add_run("丁")
    document.save(source)

    parsed = parse_docx(source)
    block = parsed.blocks[0]
    issue_id = UUID("00000000-0000-0000-0000-000000000008")
    assert block.text == "甲乙\n丙丁"

    result = DocxExporter().export(
        source,
        parsed,
        build_plan(
            Replacement(
                block_id=block.block_id,
                start=1,
                end=4,
                original="乙\n丙",
                value="替换",
                issue_id=issue_id,
            )
        ),
        tmp_path / "table-cross-paragraph-modified.docx",
    )

    exported = WordDocument(str(result.path))
    exported_cell = exported.tables[0].cell(0, 0)

    assert exported_cell.paragraphs[0].text == "甲乙"
    assert exported_cell.paragraphs[1].text == "丙丁"
    assert result.warnings == [
        ExportWarning(
            code="unsafe_docx_run_boundary",
            message=(
                "修改范围跨越多个 DOCX 文本运行，为保留格式已跳过；"
                "请在原文中手动修改后重新导出。"
            ),
            issue_id=issue_id,
            block_id=block.block_id,
        )
    ]


def test_docx_export_cleans_temp_and_raises_explicit_error_when_validation_reopen_fails(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    source = tmp_path / "source.docx"
    document = WordDocument()
    document.add_paragraph("核验示例文本")
    document.save(source)

    parsed = parse_docx(source)
    target = tmp_path / "broken.docx"
    temp_target = target.with_name(f"{target.name}.tmp")

    import text_verification.exporters.docx as docx_module

    original_loader = docx_module.WordDocument

    def failing_loader(path: str):
        if path.endswith(".tmp"):
            raise RuntimeError("cannot reopen temp")
        return original_loader(path)

    monkeypatch.setattr(docx_module, "WordDocument", failing_loader)

    with pytest.raises(ExportError) as raised:
        DocxExporter().export(
            source,
            parsed,
            build_plan(
                Replacement(
                    block_id=parsed.blocks[0].block_id,
                    start=0,
                    end=2,
                    original="核验",
                    value="检查",
                    issue_id=UUID("00000000-0000-0000-0000-000000000006"),
                )
            ),
            target,
        )

    assert raised.value.code == "docx_export_failed"
    assert raised.value.public_message == "无法导出 DOCX 文件。"
    assert not temp_target.exists()
    assert not target.exists()


def parse_docx(source: Path, *, source_name: str | None = None):
    return DocxParser().parse(
        source,
        document_id=uuid4(),
        source_name=source_name or source.name,
    )


def build_plan(*replacements: Replacement) -> ReplacementPlan:
    return ReplacementPlan(applicable=list(replacements), warnings=[])
