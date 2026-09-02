from __future__ import annotations

import io
import signal
import subprocess
import zipfile
from pathlib import Path

import pytest

from text_verification.compatibility import parser as parser_module


def _docx_bytes(text: str = "converted") -> bytes:
    from docx import Document

    output = io.BytesIO()
    document = Document()
    document.add_paragraph(text)
    document.save(output)
    return output.getvalue()


class _CompletedProcess:
    pid = 4242

    def __init__(self, on_wait=None) -> None:
        self.returncode: int | None = None
        self._on_wait = on_wait

    def wait(self, timeout: float | None = None) -> int:
        del timeout
        if self._on_wait is not None:
            self._on_wait()
            self._on_wait = None
        self.returncode = 0
        return 0

    def poll(self) -> int | None:
        return self.returncode


def test_libreoffice_conversion_uses_isolated_safe_headless_command(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    source = tmp_path / "source.doc"
    source.write_bytes(bytes.fromhex("D0CF11E0A1B11AE1") + b"\x00" * 16)
    captured: dict[str, object] = {}

    def which(name: str) -> str | None:
        return "/usr/bin/soffice" if name == "soffice" else None

    def popen(command, **kwargs):
        captured["command"] = command
        captured["kwargs"] = kwargs
        outdir = Path(command[command.index("--outdir") + 1])
        return _CompletedProcess(
            lambda: (outdir / "source.docx").write_bytes(_docx_bytes())
        )

    monkeypatch.setattr(parser_module.shutil, "which", which)
    monkeypatch.setattr(parser_module.subprocess, "Popen", popen)
    monkeypatch.setenv("HTTPS_PROXY", "http://proxy.invalid")

    converted = Path(parser_module._convert_doc_to_docx(str(source), str(tmp_path)))

    command = captured["command"]
    assert isinstance(command, list)
    assert command[0] == "/usr/bin/soffice"
    assert {
        "--headless",
        "--invisible",
        "--nologo",
        "--nodefault",
        "--nolockcheck",
        "--nofirststartwizard",
        "--norestore",
        "--safe-mode",
    } <= set(command)
    assert any(item.startswith("-env:UserInstallation=file://") for item in command)
    assert "--accept" not in command
    assert converted.name == "converted.docx"
    assert converted.read_bytes() == _docx_bytes()
    kwargs = captured["kwargs"]
    assert isinstance(kwargs, dict)
    assert kwargs["stdout"] is subprocess.DEVNULL
    assert kwargs["stderr"] is subprocess.DEVNULL
    assert kwargs["start_new_session"] is True
    environment = kwargs["env"]
    assert isinstance(environment, dict)
    assert "HTTPS_PROXY" not in environment
    assert Path(environment["HOME"]) == converted.parent
    assert environment["TMPDIR"] == environment["HOME"]


def test_doc_conversion_rejects_oversized_converted_output_before_parsing(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    source = tmp_path / "source.doc"
    source.write_bytes(bytes.fromhex("D0CF11E0A1B11AE1") + b"\x00" * 16)
    monkeypatch.setattr(parser_module, "MAX_DOC_CONVERTED_BYTES", 8)
    monkeypatch.setattr(
        parser_module.shutil,
        "which",
        lambda name: "/usr/bin/textutil" if name == "textutil" else None,
    )

    def popen(command, **kwargs):
        del kwargs
        output = Path(command[command.index("-output") + 1])
        return _CompletedProcess(lambda: output.write_bytes(b"x" * 9))

    monkeypatch.setattr(parser_module.subprocess, "Popen", popen)

    with pytest.raises(ValueError, match="size limit"):
        parser_module._convert_doc_to_docx(str(source), str(tmp_path))


def test_doc_conversion_rejects_oversized_source_before_process_start(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    source = tmp_path / "source.doc"
    source.write_bytes(b"x" * 9)
    monkeypatch.setattr(parser_module, "MAX_DOC_SOURCE_BYTES", 8)
    monkeypatch.setattr(
        parser_module.subprocess,
        "Popen",
        lambda *args, **kwargs: pytest.fail("converter must not start"),
    )

    with pytest.raises(ValueError, match="source exceeds"):
        parser_module._convert_doc_to_docx(str(source), str(tmp_path))


@pytest.mark.parametrize("payload", [None, b"not-a-docx"], ids=["missing", "malformed"])
def test_doc_conversion_rejects_missing_or_malformed_output(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
    payload: bytes | None,
) -> None:
    source = tmp_path / "source.doc"
    source.write_bytes(bytes.fromhex("D0CF11E0A1B11AE1") + b"\x00" * 16)
    monkeypatch.setattr(
        parser_module.shutil,
        "which",
        lambda name: "/usr/bin/textutil" if name == "textutil" else None,
    )

    def popen(command, **kwargs):
        del kwargs
        output = Path(command[command.index("-output") + 1])
        return _CompletedProcess(
            None if payload is None else lambda: output.write_bytes(payload)
        )

    monkeypatch.setattr(parser_module.subprocess, "Popen", popen)

    with pytest.raises(ValueError, match="valid DOCX"):
        parser_module._convert_doc_to_docx(str(source), str(tmp_path))


def test_doc_conversion_timeout_terminates_process_group(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    output = tmp_path / "converted.docx"
    process = _CompletedProcess()
    process.wait = lambda timeout=None: (_ for _ in ()).throw(  # type: ignore[method-assign]
        subprocess.TimeoutExpired(["converter"], timeout)
    )
    killed: list[tuple[int, signal.Signals]] = []

    def killpg(pid: int, sig: signal.Signals) -> None:
        killed.append((pid, sig))
        process.returncode = -int(sig)

    monkeypatch.setattr(parser_module.subprocess, "Popen", lambda *args, **kwargs: process)
    monkeypatch.setattr(parser_module.os, "killpg", killpg)

    with pytest.raises(ValueError, match="timed out"):
        parser_module._run_conversion_process(
            ["converter"],
            output_path=output,
            cwd=tmp_path,
            environment={},
            timeout_seconds=0.01,
        )

    assert killed == [(process.pid, signal.SIGTERM)]


def test_docx_parser_bounds_extracted_text_size(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    source = tmp_path / "large.docx"
    source.write_bytes(_docx_bytes("12345"))
    monkeypatch.setattr(parser_module, "MAX_DOC_PARSED_TEXT_CHARS", 4)

    with pytest.raises(ValueError, match="text size limit"):
        parser_module._parse_docx(str(source))


def test_docx_parser_bounds_extracted_element_count(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    from docx import Document

    source = tmp_path / "many-elements.docx"
    document = Document()
    document.add_paragraph("one")
    document.add_paragraph("two")
    document.save(source)
    monkeypatch.setattr(parser_module, "MAX_DOC_PARSED_ELEMENTS", 1)

    with pytest.raises(ValueError, match="element limit"):
        parser_module._parse_docx(str(source))


def test_converted_docx_rejects_unsafe_archive_entry(
    tmp_path: Path,
) -> None:
    source = tmp_path / "unsafe.docx"
    with zipfile.ZipFile(source, "w") as archive:
        archive.writestr("[Content_Types].xml", "<Types/>")
        archive.writestr("word/document.xml", "<w:document/>")
        archive.writestr("../escape", "bad")

    with pytest.raises(ValueError, match="unsafe"):
        parser_module._validate_converted_docx(source)


def test_doc_conversion_rejects_symlinked_source(
    tmp_path: Path,
) -> None:
    source = tmp_path / "source.doc"
    source.write_bytes(bytes.fromhex("D0CF11E0A1B11AE1") + b"\x00" * 16)
    linked = tmp_path / "linked.doc"
    try:
        linked.symlink_to(source)
    except OSError as error:
        pytest.skip(f"symlink creation unavailable: {error}")

    with pytest.raises(ValueError, match="regular file"):
        parser_module._convert_doc_to_docx(str(linked), str(tmp_path))
