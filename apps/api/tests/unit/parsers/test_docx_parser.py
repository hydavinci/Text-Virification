from pathlib import Path
from uuid import uuid4

import pytest
from docx import Document as WordDocument

from text_verification.domain.documents import ParseError
from text_verification.parsers.docx import DocxParser


@pytest.fixture
def fixture_path() -> Path:
    return Path(__file__).resolve().parents[2] / "fixtures" / "documents"


@pytest.fixture
def merged_cells_docx(fixture_path: Path) -> Path:
    return fixture_path / "merged-cells.docx"


def test_docx_parser_preserves_paragraph_table_and_run_mapping(fixture_path: Path) -> None:
    document = DocxParser().parse(
        fixture_path / "sample.docx",
        document_id=uuid4(),
        source_name="sample.docx",
    )

    assert [block.block_id for block in document.blocks] == [
        "h-000001",
        "p-000001",
        "t-000001-000001",
        "t-000001-000002",
    ]
    assert [block.kind for block in document.blocks] == [
        "heading",
        "paragraph",
        "table_cell",
        "table_cell",
    ]
    paragraph = document.blocks[1]
    assert paragraph.text == "核验示例文本"
    assert paragraph.source_locator["runs"] == [
        {"run_index": 0, "start": 0, "end": 2},
        {"run_index": 1, "start": 2, "end": 6},
    ]


def test_docx_parser_preserves_table_cell_paragraph_and_run_locators(
    tmp_path: Path,
) -> None:
    source = tmp_path / "table-runs.docx"
    document = WordDocument()
    table = document.add_table(rows=1, cols=1)
    first_paragraph = table.cell(0, 0).paragraphs[0]
    first_paragraph.add_run("甲")
    first_paragraph.add_run("乙")
    second_paragraph = table.cell(0, 0).add_paragraph()
    second_paragraph.add_run("丙")
    second_paragraph.add_run("丁")
    document.save(source)

    parsed = DocxParser().parse(
        source,
        document_id=uuid4(),
        source_name="table-runs.docx",
    )

    assert [block.block_id for block in parsed.blocks] == ["t-000001-000001"]
    assert parsed.blocks[0].text == "甲乙\n丙丁"
    assert parsed.blocks[0].source_locator == {
        "table_index": 0,
        "row_index": 0,
        "column_index": 0,
        "paragraphs": [
            {
                "paragraph_index": 0,
                "start": 0,
                "end": 2,
                "runs": [
                    {"run_index": 0, "start": 0, "end": 1},
                    {"run_index": 1, "start": 1, "end": 2},
                ],
            },
            {
                "paragraph_index": 1,
                "start": 3,
                "end": 5,
                "runs": [
                    {"run_index": 0, "start": 3, "end": 4},
                    {"run_index": 1, "start": 4, "end": 5},
                ],
            },
        ],
    }


def test_docx_parser_preserves_body_order_for_interleaved_blocks(tmp_path: Path) -> None:
    source = tmp_path / "interleaved.docx"
    document = WordDocument()
    document.add_paragraph("段落一")
    table = document.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "表格一"
    document.add_paragraph("段落二")
    document.save(source)

    parsed = DocxParser().parse(
        source,
        document_id=uuid4(),
        source_name="interleaved.docx",
    )

    assert [block.block_id for block in parsed.blocks] == [
        "p-000001",
        "t-000001-000001",
        "p-000002",
    ]
    assert [block.kind for block in parsed.blocks] == [
        "paragraph",
        "table_cell",
        "paragraph",
    ]
    assert [block.text for block in parsed.blocks] == ["段落一", "表格一", "段落二"]


def test_docx_parser_emits_each_merged_physical_cell_once(
    merged_cells_docx: Path,
) -> None:
    parsed = DocxParser().parse(
        merged_cells_docx,
        document_id=uuid4(),
        source_name="merged-cells.docx",
    )

    assert [block.block_id for block in parsed.blocks] == [
        "p-000001",
        "t-000001-000001",
        "t-000001-000003",
        "t-000001-000008",
        "p-000002",
    ]
    assert [block.text for block in parsed.blocks] == [
        "表格前",
        "横向合并",
        "纵向合并",
        "尾部单元格",
        "表格后",
    ]
    assert parsed.blocks[1].source_locator == {
        "table_index": 0,
        "row_index": 0,
        "column_index": 0,
        "paragraphs": [
            {
                "paragraph_index": 0,
                "start": 0,
                "end": 4,
                "runs": [
                    {"run_index": 0, "start": 0, "end": 2},
                    {"run_index": 1, "start": 2, "end": 4},
                ],
            }
        ],
    }
    assert parsed.blocks[2].source_locator["row_index"] == 0
    assert parsed.blocks[2].source_locator["column_index"] == 2


def test_docx_parser_rejects_missing_document_xml(tmp_path: Path) -> None:
    source = tmp_path / "broken.docx"
    source.write_bytes(b"not-a-docx")

    with pytest.raises(ParseError, match="无法解析 DOCX 文件"):
        DocxParser().parse(source, document_id=uuid4(), source_name="broken.docx")
