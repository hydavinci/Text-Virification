from __future__ import annotations

import base64
import struct
import time
import zlib
from collections.abc import Iterable
from pathlib import Path, PurePosixPath
from uuid import uuid4

import pytest
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn

from text_verification.compatibility.exporters import ExportError
from text_verification.document_processing.ocr_provider import OcrTextBox
from text_verification.document_processing.pdf_models import (
    PdfDocumentMetadata,
    PdfPageKind,
    PdfPageMetadata,
    PdfTable,
    PdfTableCell,
)
from text_verification.domain.documents import (
    DocumentMetadata,
    DocumentModel,
    FileType,
    TextBlock,
)
from text_verification.domain.ports import ResolvedSourcePath
from text_verification.exporters import docx_reconstruction as reconstruction_module
from text_verification.exporters.docx_reconstruction import (
    DOCX_RECONSTRUCTION,
    DocxReconstructionExporter,
    DocxReconstructionLimits,
)
from text_verification.exporters.registry import ExporterRegistry
from text_verification.parsers.pdf_parser import PdfParser


class _Block:
    def __init__(
        self,
        kind: str,
        text: str,
        *,
        page: int | None,
        y: float,
        x: float = 10.0,
        width: float = 100.0,
        style: dict[str, object] | None = None,
        table_index: int | None = None,
        row_index: int | None = None,
        cell_index: int | None = None,
        source_locator: dict[str, object] | None = None,
    ) -> None:
        self.kind = kind
        self.text = text
        self.page = page
        self.y = y
        self.x = x
        self.width = width
        self.style = style or {}
        self.table_index = table_index
        self.row_index = row_index
        self.cell_index = cell_index
        self.source_locator = source_locator or {}


class _StaticAnchoredSourcePathResolver:
    def __init__(self, path: Path) -> None:
        self.path = path

    def resolve_anchored(
        self,
        document: DocumentModel,
        *,
        source_path: Path | None = None,
    ) -> ResolvedSourcePath:
        del document
        return ResolvedSourcePath.from_path(source_path or self.path)


class _AnchoredSourcePathResolver:
    def __init__(self, root: Path, relative_path: PurePosixPath) -> None:
        self.root = root
        self.relative_path = relative_path

    def resolve_anchored(
        self,
        document: DocumentModel,
        *,
        source_path: Path | None = None,
    ) -> ResolvedSourcePath:
        del document, source_path
        return ResolvedSourcePath(
            root=self.root,
            relative_path=self.relative_path,
        )


class _LegacyPathResolver:
    def __init__(self, path: Path) -> None:
        self.path = path

    def resolve(
        self,
        document: DocumentModel,
        *,
        source_path: Path | None = None,
    ) -> Path:
        del document, source_path
        return self.path


class _FakeOcr:
    def __init__(self, output: list[OcrTextBox]) -> None:
        self._output = output

    def recognize(self, image: object, language: str) -> list[OcrTextBox]:
        del image, language
        return self._output


def _ocr_box(text: str, bbox: tuple[float, float, float, float]) -> OcrTextBox:
    x0, y0, x1, y1 = bbox
    return OcrTextBox(
        text=text,
        confidence=0.99,
        bbox=((x0, y0), (x1, y0), (x1, y1), (x0, y1)),
    )


def _png(width: int, height: int) -> bytes:
    def chunk(name: bytes, data: bytes) -> bytes:
        return (
            struct.pack(">I", len(data))
            + name
            + data
            + struct.pack(">I", zlib.crc32(name + data) & 0xFFFFFFFF)
        )

    rows = b"".join(b"\x00" + b"\x00\x00\x00" * width for _ in range(height))
    return (
        b"\x89PNG\r\n\x1a\n"
        + chunk(b"IHDR", struct.pack(">IIBBBBB", width, height, 8, 2, 0, 0, 0))
        + chunk(b"IDAT", zlib.compress(rows))
        + chunk(b"IEND", b"")
    )


def _png_with_corrupt_crc() -> bytes:
    content = bytearray(_png(2, 1))
    chunk_type = content.index(b"IDAT")
    chunk_length = int.from_bytes(content[chunk_type - 4 : chunk_type], "big")
    crc_offset = chunk_type + 4 + chunk_length
    content[crc_offset] ^= 0x01
    return bytes(content)


def _png_with_invalid_color_depth() -> bytes:
    content = bytearray(_png(2, 1))
    content[24] = 3
    content[29:33] = struct.pack(
        ">I",
        zlib.crc32(content[12:29]) & 0xFFFFFFFF,
    )
    return bytes(content)


def _png_with_decompression_overflow() -> bytes:
    def chunk(name: bytes, data: bytes) -> bytes:
        return (
            struct.pack(">I", len(data))
            + name
            + data
            + struct.pack(">I", zlib.crc32(name + data) & 0xFFFFFFFF)
        )

    return (
        b"\x89PNG\r\n\x1a\n"
        + chunk(b"IHDR", struct.pack(">IIBBBBB", 1, 1, 8, 2, 0, 0, 0))
        + chunk(b"IDAT", zlib.compress(b"\x00" + b"\x00" * 100))
        + chunk(b"IEND", b"")
    )


def _truncated_jpeg() -> bytes:
    import pymupdf

    pixmap = pymupdf.Pixmap(pymupdf.csRGB, pymupdf.IRect(0, 0, 8, 8), False)
    pixmap.clear_with(0x7F3F1F)
    content = pixmap.tobytes("jpeg")
    return content[:-20]


def test_reconstructs_structured_content_and_document_metadata(tmp_path: Path) -> None:
    document = _document(
        [
            _Block(
                "heading",
                "标题\n副题",
                page=1,
                y=10,
                style={
                    "level": 0,
                    "font": {
                        "name": "Noto Sans CJK SC",
                        "size": 16.0,
                        "bold": True,
                        "italic": True,
                    },
                },
            ),
            _Block("paragraph", "第一行\n第二行", page=1, y=40),
            _Block(
                "table_cell",
                "名称",
                page=1,
                y=80,
                table_index=0,
                row_index=0,
                cell_index=0,
                source_locator={"table_shape": {"rows": 2, "columns": 2}},
            ),
            _Block(
                "table_cell",
                "值",
                page=1,
                y=80,
                table_index=0,
                row_index=0,
                cell_index=1,
            ),
            _Block(
                "table_cell",
                "甲",
                page=1,
                y=100,
                table_index=0,
                row_index=1,
                cell_index=0,
            ),
            _Block(
                "table_cell",
                "",
                page=1,
                y=100,
                table_index=0,
                row_index=1,
                cell_index=1,
            ),
            _Block("paragraph", "表后正文", page=1, y=170),
        ],
        source_name="源标题.docx",
    )

    target = DocxReconstructionExporter().export(document, tmp_path / "rebuilt.docx")

    rebuilt = Document(target)
    assert rebuilt.paragraphs[0].style.name == "Heading 1"
    assert rebuilt.paragraphs[0].text == "标题\n副题"
    heading_run = rebuilt.paragraphs[0].runs[0]
    assert heading_run.bold is True
    assert heading_run.italic is True
    assert heading_run.font.size is not None
    assert heading_run.font.size.pt == 16.0
    run_fonts = heading_run._element.get_or_add_rPr().rFonts
    assert run_fonts is not None
    assert run_fonts.get(qn("w:eastAsia")) == "Noto Sans CJK SC"
    assert rebuilt.paragraphs[1].text == "第一行\n第二行"
    assert rebuilt.tables[0].cell(0, 0).text == "名称"
    assert rebuilt.tables[0].cell(0, 1).text == "值"
    assert rebuilt.tables[0].cell(1, 0).text == "甲"
    assert rebuilt.tables[0].cell(1, 1).text == ""
    assert rebuilt.paragraphs[-1].text == "表后正文"
    assert "名称" not in "\n".join(paragraph.text for paragraph in rebuilt.paragraphs)
    assert rebuilt.core_properties.title == "源标题"
    assert rebuilt.sections[0].top_margin.pt == 54.0
    assert rebuilt.sections[0].bottom_margin.pt == 54.0
    assert _body_kinds(rebuilt) == ["paragraph", "paragraph", "table", "paragraph"]


def test_orders_shuffled_blocks_by_page_and_vertical_position(tmp_path: Path) -> None:
    source_blocks = [
        _Block("paragraph", "最后", page=2, y=200),
        _Block(
            "table_cell",
            "右",
            page=1,
            y=100,
            table_index=4,
            row_index=0,
            cell_index=1,
            source_locator={"table_shape": {"rows": 1, "columns": 2}},
        ),
        _Block("paragraph", "最先", page=1, y=10),
        _Block("paragraph", "中间", page=1, y=150),
        _Block(
            "table_cell",
            "左",
            page=1,
            y=100,
            table_index=4,
            row_index=0,
            cell_index=0,
        ),
    ]
    document = _document(source_blocks)

    target = DocxReconstructionExporter().export(document, tmp_path / "ordered.docx")

    rebuilt = Document(target)
    assert _body_semantics(rebuilt) == [
        ("paragraph", "最先"),
        ("table", (("左", "右"),)),
        ("paragraph", "中间"),
        ("paragraph", "最后"),
    ]


def test_registry_resolves_explicit_docx_reconstruction_format() -> None:
    exporter = DocxReconstructionExporter()
    registry = ExporterRegistry([exporter])

    assert registry.get(DOCX_RECONSTRUCTION) is exporter


def test_reconstruction_has_deterministic_semantic_roundtrip(tmp_path: Path) -> None:
    document = _document(
        [
            _Block("heading", "标题", page=1, y=10, style={"level": 2}),
            _Block("paragraph", "正文", page=1, y=40),
            _Block(
                "table_cell",
                "单元格",
                page=1,
                y=80,
                table_index=0,
                row_index=0,
                cell_index=0,
            ),
        ]
    )
    exporter = DocxReconstructionExporter()

    first = Document(exporter.export(document, tmp_path / "first.docx"))
    second = Document(exporter.export(document, tmp_path / "second.docx"))

    assert _body_semantics(first) == _body_semantics(second)
    assert _body_semantics(first) == [
        ("paragraph", "标题"),
        ("paragraph", "正文"),
        ("table", (("单元格",),)),
    ]


def test_reconstruction_bytes_are_deterministic_across_zip_timestamps(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    document = _document([_Block("paragraph", "正文", page=1, y=10)])
    current = [time.struct_time((2024, 1, 1, 0, 0, 0, 0, 1, -1))]
    monkeypatch.setattr(time, "localtime", lambda *_args: current[0])
    exporter = DocxReconstructionExporter()

    first = exporter.export(document, tmp_path / "first-bytes.docx").read_bytes()
    current[0] = time.struct_time((2026, 9, 2, 12, 0, 0, 2, 245, -1))
    second = exporter.export(document, tmp_path / "second-bytes.docx").read_bytes()

    assert first == second


def test_reconstructs_real_parser_image_after_json_roundtrip(
    tmp_path: Path,
) -> None:
    source = _pdf_with_image(tmp_path / "source.pdf")
    parsed = PdfParser().parse(source)
    persisted = DocumentModel.model_validate_json(parsed.model_dump_json())

    target = DocxReconstructionExporter(
        anchored_source_resolver=_StaticAnchoredSourcePathResolver(source)
    ).export(persisted, tmp_path / "parser-image.docx")

    rebuilt = Document(target)
    assert len(rebuilt.inline_shapes) == 1
    assert any(paragraph.text == "After image" for paragraph in rebuilt.paragraphs)


def test_reconstructs_pdf_image_with_indirect_stream_length(tmp_path: Path) -> None:
    source = _pdf_with_indirect_image_length(
        tmp_path / "indirect-length.pdf",
        "valid",
    )
    parsed = PdfParser().parse(source)

    rebuilt = Document(
        DocxReconstructionExporter(
            anchored_source_resolver=_StaticAnchoredSourcePathResolver(source)
        ).export(parsed, tmp_path / "indirect-length.docx")
    )

    assert len(rebuilt.inline_shapes) == 1


@pytest.mark.parametrize(
    "length_kind",
    ["chain", "non_integer", "oversized"],
)
def test_ignores_untrusted_pdf_length_declarations_and_bounds_actual_stream(
    tmp_path: Path,
    length_kind: str,
) -> None:
    source = _pdf_with_indirect_image_length(
        tmp_path / f"{length_kind}.pdf",
        length_kind,
    )
    parsed = PdfParser().parse(source)

    rebuilt = Document(
        DocxReconstructionExporter(
            anchored_source_resolver=_StaticAnchoredSourcePathResolver(source)
        ).export(parsed, tmp_path / f"{length_kind}.docx")
    )

    assert len(rebuilt.inline_shapes) == 1


def test_rejects_actual_raw_pdf_image_stream_before_extraction(tmp_path: Path) -> None:
    source = _pdf_with_indirect_image_length(
        tmp_path / "raw-stream-limit.pdf",
        "valid",
    )
    parsed = PdfParser().parse(source)

    with pytest.raises(ExportError, match="size limit"):
        DocxReconstructionExporter(
            limits=DocxReconstructionLimits(max_image_bytes=10),
            anchored_source_resolver=_StaticAnchoredSourcePathResolver(source),
        ).export(parsed, tmp_path / "raw-stream-limit.docx")


def test_reconstructs_real_native_pdf_spans_as_distinct_runs(tmp_path: Path) -> None:
    source = _pdf_with_styled_spans(tmp_path / "styled.pdf")
    parsed = PdfParser().parse(source)
    persisted = DocumentModel.model_validate_json(parsed.model_dump_json())

    rebuilt = Document(
        DocxReconstructionExporter().export(
            persisted,
            tmp_path / "styled.docx",
        )
    )

    paragraph = next(item for item in rebuilt.paragraphs if item.text == "BoldItalic")
    assert [run.text for run in paragraph.runs] == ["Bold", "Italic"]
    assert paragraph.runs[0].bold is True
    assert paragraph.runs[0].font.size is not None
    assert paragraph.runs[0].font.size.pt == 18.0
    assert paragraph.runs[1].italic is True
    assert paragraph.runs[1].font.size is not None
    assert paragraph.runs[1].font.size.pt == 11.0


def test_reconstructs_real_fake_ocr_heading_table_and_image(tmp_path: Path) -> None:
    source = _scanned_pdf(tmp_path / "scanned.pdf")
    ocr = _FakeOcr(
        [
            _ocr_box("Visual heading", (40, 20, 300, 52)),
            _ocr_box("Body one", (40, 80, 200, 100)),
            _ocr_box("Body two", (40, 112, 200, 132)),
            _ocr_box("A1", (40, 200, 100, 220)),
            _ocr_box("B1", (240, 200, 300, 220)),
            _ocr_box("A2", (40, 270, 100, 290)),
            _ocr_box("B2", (240, 270, 300, 290)),
        ]
    )
    parsed = PdfParser(ocr=ocr).parse(source)
    persisted = DocumentModel.model_validate_json(parsed.model_dump_json())
    heading = next(block for block in persisted.blocks if block.kind == "heading")
    assert heading.style["level"] == 3
    assert heading.style["estimated_font_size"] == pytest.approx(16.0)
    table_cell = next(block for block in persisted.blocks if block.kind == "table_cell")
    assert table_cell.source_locator["table_shape"] == {"rows": 2, "columns": 2}
    assert table_cell.source_locator["table_bbox"] == pytest.approx(
        [20.0, 100.0, 150.0, 145.0]
    )
    row_bands = table_cell.source_locator["table_row_bands"]
    assert row_bands[0] == pytest.approx([20.0, 100.0, 150.0, 110.0])
    assert row_bands[1] == pytest.approx([20.0, 135.0, 150.0, 145.0])

    rebuilt = Document(
        DocxReconstructionExporter(
            anchored_source_resolver=_StaticAnchoredSourcePathResolver(source)
        ).export(persisted, tmp_path / "ocr.docx")
    )

    heading_paragraph = next(
        paragraph
        for paragraph in rebuilt.paragraphs
        if paragraph.text == "Visual heading"
    )
    assert heading_paragraph.style.name == "Heading 3"
    assert heading_paragraph.runs[0].font.size is not None
    assert heading_paragraph.runs[0].font.size.pt == 16.0
    assert rebuilt.tables[0].cell(0, 0).text == "A1"
    assert rebuilt.tables[0].cell(1, 1).text == "B2"
    assert len(rebuilt.inline_shapes) == 1


def test_real_parser_image_requires_injected_source_resolver(tmp_path: Path) -> None:
    source = _pdf_with_image(tmp_path / "source.pdf")
    parsed = PdfParser().parse(source)

    with pytest.raises(ExportError, match="source resolver"):
        DocxReconstructionExporter().export(parsed, tmp_path / "missing-source.docx")


def test_reconstruction_rejects_legacy_path_resolver_injection(
    tmp_path: Path,
) -> None:
    source = _pdf_with_image(tmp_path / "relative-source.pdf")

    with pytest.raises(TypeError, match="anchored_source_resolver"):
        DocxReconstructionExporter(
            anchored_source_resolver=_LegacyPathResolver(source)  # type: ignore[arg-type, call-arg]
        )


def test_reads_normal_nested_anchored_source_path(tmp_path: Path) -> None:
    root = tmp_path / "repository"
    nested = root / "nested"
    nested.mkdir(parents=True)
    source = _pdf_with_image(nested / "source.pdf")
    parsed = PdfParser().parse(source)

    rebuilt = Document(
        DocxReconstructionExporter(
            anchored_source_resolver=_AnchoredSourcePathResolver(
                root,
                PurePosixPath("nested/source.pdf"),
            )
        ).export(parsed, tmp_path / "nested-source.docx")
    )

    assert len(rebuilt.inline_shapes) == 1


@pytest.mark.parametrize(
    "relative_path",
    [
        PurePosixPath("../source.pdf"),
        PurePosixPath("/source.pdf"),
        PurePosixPath("."),
        PurePosixPath("bad\x00.pdf"),
        PurePosixPath("C:/source.pdf"),
    ],
)
def test_rejects_unsafe_anchored_source_components(
    tmp_path: Path,
    relative_path: PurePosixPath,
) -> None:
    with pytest.raises(ValueError, match="safe relative"):
        ResolvedSourcePath(root=tmp_path, relative_path=relative_path)


def test_rejects_intermediate_symlink_in_anchored_source_path(tmp_path: Path) -> None:
    root = tmp_path / "repository"
    root.mkdir()
    outside = tmp_path / "outside"
    outside.mkdir()
    source = _pdf_with_image(outside / "source.pdf")
    parsed = PdfParser().parse(source)
    (root / "linked").symlink_to(outside, target_is_directory=True)

    with pytest.raises(ExportError, match="unsafe"):
        DocxReconstructionExporter(
            anchored_source_resolver=_AnchoredSourcePathResolver(
                root,
                PurePosixPath("linked/source.pdf"),
            )
        ).export(parsed, tmp_path / "symlinked-parent.docx")


def test_source_resolution_fails_closed_without_secure_platform_flags(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    root = tmp_path / "repository"
    root.mkdir()
    source = _pdf_with_image(root / "source.pdf")
    parsed = PdfParser().parse(source)
    target = tmp_path / "unsupported-source.docx"
    monkeypatch.delattr(reconstruction_module.os, "O_NOFOLLOW")

    with pytest.raises(ExportError, match="Secure source access is unavailable"):
        DocxReconstructionExporter(
            anchored_source_resolver=_AnchoredSourcePathResolver(
                root,
                PurePosixPath("source.pdf"),
            )
        ).export(parsed, target)

    assert not target.exists()
    assert not list(tmp_path.glob(".*.uploading"))


@pytest.mark.parametrize(
    "missing_capability",
    ["dir_fd", "follow_symlinks"],
)
def test_atomic_publication_fails_closed_without_secure_platform_capabilities(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
    missing_capability: str,
) -> None:
    document = _document([_Block("paragraph", "正文", page=1, y=10)])
    target = tmp_path / f"unsupported-{missing_capability}.docx"
    if missing_capability == "dir_fd":
        monkeypatch.delattr(reconstruction_module.os, "supports_dir_fd")
    else:
        monkeypatch.delattr(reconstruction_module.os, "supports_follow_symlinks")

    with pytest.raises(ExportError, match="Secure DOCX publication is unavailable"):
        DocxReconstructionExporter().export(document, target)

    assert not target.exists()
    assert not list(tmp_path.glob(".*.uploading"))


def test_rejects_mismatched_and_symlinked_resolved_sources(tmp_path: Path) -> None:
    source = _pdf_with_image(tmp_path / "canonical.pdf")
    parsed = PdfParser().parse(source)
    different = _pdf_with_image(tmp_path / "different.pdf")
    with different.open("ab") as output:
        output.write(b"\n")

    with pytest.raises(ExportError, match="does not match"):
        DocxReconstructionExporter(
            anchored_source_resolver=_StaticAnchoredSourcePathResolver(different)
        ).export(parsed, tmp_path / "mismatched.docx")

    symlink = tmp_path / "linked.pdf"
    symlink.symlink_to(source)
    with pytest.raises(ExportError, match="unsafe"):
        DocxReconstructionExporter(
            anchored_source_resolver=_StaticAnchoredSourcePathResolver(symlink)
        ).export(parsed, tmp_path / "symlinked.docx")


@pytest.mark.parametrize(
    "payload",
    [
        {
            "kind": "bytes",
            "mime_type": "image/png",
            "data": _png(1, 1),
        },
        {
            "kind": "base64",
            "mime_type": "image/png",
            "data": base64.b64encode(_png(1, 1)).decode("ascii"),
        },
    ],
    ids=["raw-bytes", "base64"],
)
def test_rejects_embedded_image_payloads_in_canonical_metadata(
    tmp_path: Path,
    payload: dict[str, object],
) -> None:
    source = _pdf_with_image(tmp_path / "embedded-source.pdf")
    document = PdfParser().parse(source)
    image_block = next(block for block in document.blocks if block.kind == "image")
    image_block.source_locator["image_payload"] = payload

    with pytest.raises(ExportError, match="canonical image reference"):
        DocxReconstructionExporter(
            anchored_source_resolver=_StaticAnchoredSourcePathResolver(source)
        ).export(
            document,
            tmp_path / "embedded-image.docx",
        )


def test_exports_horizontally_disjoint_side_note_between_table_rows(
    tmp_path: Path,
) -> None:
    source = _pdf_with_interleaved_table(
        tmp_path / "side-note.pdf",
        "side",
    )
    parsed = PdfParser().parse(source)

    rebuilt = Document(
        DocxReconstructionExporter().export(
            parsed,
            tmp_path / "side-note.docx",
        )
    )

    assert rebuilt.tables[0].cell(1, 1).text == "B2"
    assert any(paragraph.text == "SIDE" for paragraph in rebuilt.paragraphs)


@pytest.mark.parametrize(
    ("placement", "text"),
    [("caption", "CAPTION"), ("below", "BELOW")],
)
def test_exports_blocks_near_but_outside_table_row_span(
    tmp_path: Path,
    placement: str,
    text: str,
) -> None:
    source = _pdf_with_interleaved_table(
        tmp_path / f"{placement}.pdf",
        placement,
    )
    parsed = PdfParser().parse(source)

    rebuilt = Document(
        DocxReconstructionExporter().export(
            parsed,
            tmp_path / f"{placement}.docx",
        )
    )

    assert rebuilt.tables[0].cell(0, 0).text == "A1"
    assert any(paragraph.text == text for paragraph in rebuilt.paragraphs)


def test_rejects_spatially_embedded_image_between_table_rows(
    tmp_path: Path,
) -> None:
    source = _pdf_with_interleaved_table(
        tmp_path / "embedded-image.pdf",
        "image",
    )
    parsed = PdfParser().parse(source)

    with pytest.raises(ExportError, match="interleaved"):
        DocxReconstructionExporter(
            anchored_source_resolver=_StaticAnchoredSourcePathResolver(source)
        ).export(parsed, tmp_path / "embedded-image.docx")


def test_rejects_embedded_image_in_blank_terminal_table_row(
    tmp_path: Path,
) -> None:
    source = _pdf_with_blank_terminal_row(
        tmp_path / "blank-row-image.pdf",
        "image",
    )
    parsed = PdfParser().parse(source)

    with pytest.raises(ExportError, match="interleaved"):
        DocxReconstructionExporter(
            anchored_source_resolver=_StaticAnchoredSourcePathResolver(source)
        ).export(parsed, tmp_path / "blank-row-image.docx")


def test_exports_sparse_blank_row_table_with_disjoint_side_note(
    tmp_path: Path,
) -> None:
    source = _pdf_with_blank_terminal_row(
        tmp_path / "blank-row-side.pdf",
        "side",
    )
    parsed = PdfParser().parse(source)

    rebuilt = Document(
        DocxReconstructionExporter().export(
            parsed,
            tmp_path / "blank-row-side.docx",
        )
    )

    assert len(rebuilt.tables[0].rows) == 2
    assert rebuilt.tables[0].cell(1, 0).text == ""
    assert any(paragraph.text == "SIDE" for paragraph in rebuilt.paragraphs)


@pytest.mark.parametrize(
    ("block_x", "should_reject"),
    [(80.0, True), (220.0, False)],
    ids=["embedded", "disjoint"],
)
def test_vertical_merge_incomplete_geometry_fails_closed_only_on_overlap(
    tmp_path: Path,
    block_x: float,
    should_reject: bool,
) -> None:
    table = PdfTable(
        table_index=0,
        bbox=(24.0, 40.0, 184.0, 160.0),
        row_count=2,
        column_count=1,
        rows=(
            (
                PdfTableCell(
                    text="merged",
                    bbox=(24.0, 40.0, 184.0, 160.0),
                    table_index=0,
                    row_index=0,
                    cell_index=0,
                ),
            ),
            (
                PdfTableCell(
                    text="",
                    bbox=None,
                    table_index=0,
                    row_index=1,
                    cell_index=0,
                ),
            ),
        ),
    )
    document = _document(
        [
            _Block(
                "table_cell",
                "merged",
                page=1,
                y=40,
                x=24,
                width=160,
                table_index=0,
                row_index=0,
                cell_index=0,
                source_locator={
                    "merge": {"row_span": 2, "column_span": 1},
                },
            ),
            _Block(
                "paragraph",
                "NOTE",
                page=1,
                y=100,
                x=block_x,
                width=40,
            ),
        ],
        metadata=_metadata_with_table(table),
    )
    target = tmp_path / f"vertical-merge-{block_x}.docx"

    if should_reject:
        with pytest.raises(ExportError, match="interleaved"):
            DocxReconstructionExporter().export(document, target)
    else:
        rebuilt = Document(DocxReconstructionExporter().export(document, target))
        assert rebuilt.tables[0].cell(0, 0).text == "merged"
        assert any(paragraph.text == "NOTE" for paragraph in rebuilt.paragraphs)


@pytest.mark.parametrize(
    "blocks",
    [
        [
            _Block(
                "table_cell",
                "A",
                page=1,
                y=10,
                table_index=0,
                row_index=0,
                cell_index=0,
            ),
            _Block(
                "table_cell",
                "B",
                page=1,
                y=20,
                table_index=0,
                row_index=0,
                cell_index=0,
            ),
        ],
        [
            _Block(
                "table_cell",
                "A",
                page=1,
                y=10,
                table_index=0,
                row_index=0,
                cell_index=1,
            )
        ],
        [
            _Block(
                "table_cell",
                "A",
                page=1,
                y=10,
                table_index=0,
                row_index=-1,
                cell_index=0,
            )
        ],
        [
            _Block(
                "table_cell",
                "A",
                page=1,
                y=10,
                table_index=0,
                row_index=0,
                cell_index=0,
                source_locator={"table_shape": {"rows": 1, "columns": 1}},
            ),
            _Block(
                "table_cell",
                "B",
                page=1,
                y=20,
                table_index=0,
                row_index=0,
                cell_index=1,
                source_locator={"table_shape": {"rows": 1, "columns": 2}},
            ),
        ],
        [
            _Block(
                "table_cell",
                "A",
                page=1,
                y=10,
                table_index=1_000_000,
                row_index=0,
                cell_index=0,
            )
        ],
    ],
    ids=[
        "duplicate",
        "sparse",
        "negative-index",
        "conflicting-dimensions",
        "huge-index",
    ],
)
def test_rejects_invalid_table_shapes(tmp_path: Path, blocks: list[_Block]) -> None:
    with pytest.raises(ExportError):
        DocxReconstructionExporter().export(
            _document(blocks),
            tmp_path / "invalid-table.docx",
        )


def test_preserves_explicit_table_merge(tmp_path: Path) -> None:
    document = _document(
        [
            _Block(
                "table_cell",
                "合并",
                page=1,
                y=10,
                table_index=0,
                row_index=0,
                cell_index=0,
                source_locator={
                    "table_shape": {"rows": 2, "columns": 2},
                    "merge": {"row_span": 2, "column_span": 2},
                },
            )
        ]
    )

    rebuilt = Document(
        DocxReconstructionExporter().export(document, tmp_path / "merged.docx")
    )

    top_left = rebuilt.tables[0].cell(0, 0)
    assert top_left.text == "合并"
    assert top_left._tc.tcPr.gridSpan is not None
    assert top_left._tc.tcPr.gridSpan.val == 2
    assert top_left._tc.tcPr.vMerge is not None


@pytest.mark.parametrize(
    "payload",
    [
        None,
        {
            "kind": "base64",
            "mime_type": "image/png",
            "data": "not base64!",
        },
        {
            "kind": "bytes",
            "mime_type": "image/jpeg",
            "data": _png(1, 1),
        },
        {
            "kind": "bytes",
            "mime_type": "image/png",
            "data": b"\x89PNG\r\n\x1a\nspoofed",
        },
        {
            "kind": "bytes",
            "mime_type": "image/png",
            "data": _png(1, 1)[:33],
        },
        {
            "kind": "url",
            "mime_type": "image/png",
            "url": "https://example.invalid/image.png",
        },
        {
            "kind": "bytes",
            "mime_type": "image/png",
            "data": _png(1, 1),
            "unexpected": True,
        },
    ],
    ids=[
        "missing",
        "invalid-base64",
        "mime-spoof",
        "invalid-signature",
        "truncated-image",
        "url",
        "extra-contract-field",
    ],
)
def test_rejects_invalid_image_payloads(
    tmp_path: Path,
    payload: dict[str, object] | None,
) -> None:
    source_locator = {} if payload is None else {"image_payload": payload}
    document = _document(
        [_Block("image", "", page=1, y=10, source_locator=source_locator)]
    )

    with pytest.raises(ExportError):
        DocxReconstructionExporter().export(
            document,
            tmp_path / "invalid-image.docx",
        )


def test_rejects_oversized_image_bytes_and_dimensions() -> None:
    image = _png(4, 4)

    with pytest.raises(ExportError, match="size limit"):
        reconstruction_module._validated_image(
            image,
            DocxReconstructionLimits(max_image_bytes=len(image) - 1),
        )
    with pytest.raises(ExportError, match="dimensions"):
        reconstruction_module._validated_image(
            image,
            DocxReconstructionLimits(max_image_pixels=15),
        )


@pytest.mark.parametrize(
    "content",
    [
        pytest.param(_png_with_corrupt_crc(), id="png-crc"),
        pytest.param(_png_with_invalid_color_depth(), id="png-color-depth"),
        pytest.param(_png_with_decompression_overflow(), id="png-decompression"),
        pytest.param(_truncated_jpeg(), id="jpeg-truncation"),
    ],
)
def test_fully_decodes_and_rejects_corrupt_supported_images(content: bytes) -> None:
    with pytest.raises(ExportError):
        reconstruction_module._validated_image(
            content,
            DocxReconstructionLimits(),
        )


def test_rejects_decoded_image_byte_limit_before_decode() -> None:
    with pytest.raises(ExportError, match="decoded image"):
        reconstruction_module._validated_image(
            _png(20, 20),
            DocxReconstructionLimits(max_decoded_image_bytes=1_000),
        )


def test_applies_pdf_page_size_orientation_and_safe_title(tmp_path: Path) -> None:
    document = _document(
        [_Block("paragraph", "正文", page=1, y=10)],
        source_name="标题\x01.pdf",
        metadata=DocumentMetadata(
            pdf=PdfDocumentMetadata(
                pages=(
                    PdfPageMetadata(
                        page=1,
                        kind=PdfPageKind.TEXT,
                        page_bbox=(0.0, 0.0, 792.0, 612.0),
                        text_length=2,
                        text_density=0.1,
                        image_coverage=0.0,
                        ocr_required=False,
                    ),
                )
            )
        ),
    )

    rebuilt = Document(
        DocxReconstructionExporter().export(document, tmp_path / "landscape.docx")
    )

    section = rebuilt.sections[0]
    assert section.orientation == WD_ORIENT.LANDSCAPE
    assert section.page_width.pt == 792.0
    assert section.page_height.pt == 612.0
    assert rebuilt.core_properties.title == "标题\ufffd"


def test_normalizes_unrepresentable_xml_controls_in_block_text(tmp_path: Path) -> None:
    document = _document([_Block("paragraph", "正文\x01内容", page=1, y=10)])

    rebuilt = Document(
        DocxReconstructionExporter().export(document, tmp_path / "safe-text.docx")
    )

    assert rebuilt.paragraphs[0].text == "正文\ufffd内容"


def test_uses_estimated_font_size_without_mutating_normal_style(tmp_path: Path) -> None:
    document = _document(
        [
            _Block(
                "heading",
                "估算字号",
                page=1,
                y=10,
                style={
                    "level": 99,
                    "estimated_font_size": 19.0,
                    "east_asia_font": "SimSun",
                },
            )
        ]
    )

    rebuilt = Document(
        DocxReconstructionExporter().export(document, tmp_path / "font.docx")
    )

    paragraph = rebuilt.paragraphs[0]
    assert paragraph.style.name == "Heading 9"
    assert paragraph.runs[0].font.size is not None
    assert paragraph.runs[0].font.size.pt == 19.0
    run_fonts = paragraph.runs[0]._element.get_or_add_rPr().rFonts
    assert run_fonts is not None
    assert run_fonts.get(qn("w:eastAsia")) == "SimSun"
    normal_fonts = rebuilt.styles["Normal"].element.get_or_add_rPr().rFonts
    assert normal_fonts is None or normal_fonts.get(qn("w:eastAsia")) != "SimSun"


def test_enforces_table_and_total_element_bounds(tmp_path: Path) -> None:
    table = _document(
        [
            _Block(
                "table_cell",
                "A",
                page=1,
                y=10,
                table_index=0,
                row_index=1,
                cell_index=0,
                source_locator={"table_shape": {"rows": 2, "columns": 1}},
            )
        ]
    )
    paragraphs = _document(
        [
            _Block("paragraph", "A", page=1, y=10),
            _Block("paragraph", "B", page=1, y=20),
        ]
    )

    with pytest.raises(ExportError, match="dimensions"):
        DocxReconstructionExporter(
            limits=DocxReconstructionLimits(max_table_rows=1)
        ).export(table, tmp_path / "table-bound.docx")
    with pytest.raises(ExportError, match="element limit"):
        DocxReconstructionExporter(
            limits=DocxReconstructionLimits(max_output_elements=1)
        ).export(paragraphs, tmp_path / "element-bound.docx")


def test_preflight_rejects_aggregate_text_characters(tmp_path: Path) -> None:
    document = _document(
        [
            _Block("paragraph", "AAAA", page=1, y=10),
            _Block("paragraph", "BBBB", page=1, y=20),
        ]
    )

    with pytest.raises(ExportError, match="text character"):
        DocxReconstructionExporter(
            limits=DocxReconstructionLimits(max_text_chars=7)
        ).export(document, tmp_path / "text-bound.docx")

    assert not (tmp_path / "text-bound.docx").exists()


def test_preflight_rejects_native_run_count(tmp_path: Path) -> None:
    source = _pdf_with_styled_spans(tmp_path / "styled-run-limit.pdf")
    document = PdfParser().parse(source)

    with pytest.raises(ExportError, match="run count"):
        DocxReconstructionExporter(
            limits=DocxReconstructionLimits(max_runs=1)
        ).export(document, tmp_path / "run-bound.docx")


def test_preflight_rejects_raw_span_count_before_model_validation(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    document = _document(
        [
            _Block(
                "paragraph",
                "A",
                page=1,
                y=10,
                style={"spans": [{}, {}]},
            )
        ]
    )

    def unexpected_validation(value: object) -> object:
        del value
        raise AssertionError("span validation must not run over budget")

    monkeypatch.setattr(
        reconstruction_module.PdfTextSpan,
        "model_validate",
        unexpected_validation,
    )
    with pytest.raises(ExportError, match="run count"):
        DocxReconstructionExporter(
            limits=DocxReconstructionLimits(max_runs=1)
        ).export(document, tmp_path / "raw-run-bound.docx")


def test_preflight_rejects_malformed_span_collection_without_length_access(
    tmp_path: Path,
) -> None:
    class PoisonedCollection:
        def __len__(self) -> int:
            raise AssertionError("malformed span collection length was accessed")

    document = _document(
        [
            _Block(
                "paragraph",
                "A",
                page=1,
                y=10,
                style={"spans": PoisonedCollection()},
            )
        ]
    )

    with pytest.raises(ExportError, match="span metadata"):
        DocxReconstructionExporter().export(
            document,
            tmp_path / "malformed-spans.docx",
        )


def test_preflight_counts_image_render_run(tmp_path: Path) -> None:
    source = _pdf_with_image(tmp_path / "image-run-limit.pdf")
    document = PdfParser().parse(source)

    with pytest.raises(ExportError, match="run count"):
        DocxReconstructionExporter(
            limits=DocxReconstructionLimits(max_runs=1),
            anchored_source_resolver=_StaticAnchoredSourcePathResolver(source),
        ).export(document, tmp_path / "image-run-bound.docx")


def test_preflight_rejects_image_count_and_aggregate_media(tmp_path: Path) -> None:
    source = _pdf_with_images(tmp_path / "two-images.pdf")
    document = PdfParser().parse(source)
    resolver = _StaticAnchoredSourcePathResolver(source)

    with pytest.raises(ExportError, match="image count"):
        DocxReconstructionExporter(
            limits=DocxReconstructionLimits(max_images=1),
            anchored_source_resolver=resolver,
        ).export(document, tmp_path / "image-count.docx")
    with pytest.raises(ExportError, match="aggregate image media"):
        DocxReconstructionExporter(
            limits=DocxReconstructionLimits(max_total_image_bytes=1),
            anchored_source_resolver=resolver,
        ).export(document, tmp_path / "image-media.docx")


def test_bounded_docx_output_leaves_no_partial_target(tmp_path: Path) -> None:
    document = _document([_Block("paragraph", "正文", page=1, y=10)])
    target = tmp_path / "bounded-output.docx"

    with pytest.raises(ExportError, match="output size"):
        DocxReconstructionExporter(
            limits=DocxReconstructionLimits(max_output_bytes=1_000)
        ).export(document, target)

    assert not target.exists()
    assert not list(tmp_path.glob(".*.uploading"))


def test_rejects_invalid_limit_configuration() -> None:
    with pytest.raises(ValueError):
        DocxReconstructionLimits(max_table_rows=0)


def test_does_not_overwrite_existing_or_symlink_targets(tmp_path: Path) -> None:
    document = _document([_Block("paragraph", "正文", page=1, y=10)])
    target = tmp_path / "existing.docx"
    target.write_bytes(b"existing")

    with pytest.raises(ExportError, match="already exists"):
        DocxReconstructionExporter().export(document, target)

    assert target.read_bytes() == b"existing"
    assert not list(tmp_path.glob(".*.uploading"))

    symlink_target = tmp_path / "linked.docx"
    symlink_target.symlink_to(target)
    with pytest.raises(ExportError, match="already exists"):
        DocxReconstructionExporter().export(document, symlink_target)
    assert target.read_bytes() == b"existing"
    assert not list(tmp_path.glob(".*.uploading"))


def test_rejects_symlinked_or_missing_target_parent(tmp_path: Path) -> None:
    document = _document([_Block("paragraph", "正文", page=1, y=10)])
    real_parent = tmp_path / "real"
    real_parent.mkdir()
    linked_parent = tmp_path / "linked"
    linked_parent.symlink_to(real_parent, target_is_directory=True)

    with pytest.raises(ExportError, match="parent"):
        DocxReconstructionExporter().export(document, linked_parent / "output.docx")
    with pytest.raises(ExportError, match="parent"):
        DocxReconstructionExporter().export(
            document,
            tmp_path / "missing" / "output.docx",
        )
    assert list(real_parent.iterdir()) == []


def _document(
    blocks: Iterable[_Block],
    *,
    source_name: str = "source.pdf",
    metadata: DocumentMetadata | None = None,
) -> DocumentModel:
    text_parts: list[str] = []
    canonical_blocks: list[TextBlock] = []
    cursor = 0
    for index, block in enumerate(blocks):
        if text_parts:
            text_parts.append("\n")
            cursor += 1
        start = cursor
        text_parts.append(block.text)
        cursor += len(block.text)
        canonical_blocks.append(
            TextBlock(
                block_id=f"block-{index}",
                kind=block.kind,
                text=block.text,
                global_start=start,
                global_end=cursor,
                block_start=0,
                block_end=len(block.text),
                page=block.page,
                paragraph_index=index if block.kind in {"paragraph", "heading"} else None,
                table_index=block.table_index,
                row_index=block.row_index,
                cell_index=block.cell_index,
                bbox=(block.x, block.y, block.x + block.width, block.y + 10.0),
                parent_id=None,
                style=block.style,
                source_locator={
                    "page": block.page,
                    **block.source_locator,
                },
            )
        )
    return DocumentModel(
        document_id=uuid4(),
        source_version="sha256:test",
        file_type=FileType.PDF,
        source_name=source_name,
        text="".join(text_parts),
        blocks=canonical_blocks,
        parser_name="test",
        parser_version="1",
        metadata=metadata or DocumentMetadata(),
    )


def _body_kinds(document: object) -> list[str]:
    kinds: list[str] = []
    for child in document.element.body.iterchildren():
        if child.tag == qn("w:p"):
            kinds.append("paragraph")
        elif child.tag == qn("w:tbl"):
            kinds.append("table")
    return kinds


def _body_semantics(document: object) -> list[tuple[str, object]]:
    paragraphs = iter(document.paragraphs)
    tables = iter(document.tables)
    semantics: list[tuple[str, object]] = []
    for kind in _body_kinds(document):
        if kind == "paragraph":
            semantics.append(("paragraph", next(paragraphs).text))
        else:
            table = next(tables)
            semantics.append(
                (
                    "table",
                    tuple(tuple(cell.text for cell in row.cells) for row in table.rows),
                )
            )
    return semantics


def _pdf_with_image(target: Path) -> Path:
    import pymupdf

    pdf = pymupdf.open()
    try:
        page = pdf.new_page(width=240, height=240)
        page.insert_image(pymupdf.Rect(24, 24, 124, 74), stream=_png(4, 2))
        page.insert_text((24, 120), "After image", fontsize=12, fontname="helv")
        pdf.save(target)
    finally:
        pdf.close()
    return target


def _pdf_with_styled_spans(target: Path) -> Path:
    import pymupdf

    pdf = pymupdf.open()
    try:
        page = pdf.new_page(width=400, height=200)
        page.insert_htmlbox(
            pymupdf.Rect(20, 20, 380, 80),
            (
                '<span style="font-size:18pt;font-weight:bold">Bold</span>'
                '<span style="font-size:11pt;font-style:italic">Italic</span>'
            ),
        )
        pdf.save(target)
    finally:
        pdf.close()
    return target


def _scanned_pdf(target: Path) -> Path:
    import pymupdf

    pdf = pymupdf.open()
    try:
        page = pdf.new_page(width=360, height=360)
        page.insert_image(page.rect, stream=_png(20, 20))
        pdf.save(target)
    finally:
        pdf.close()
    return target


def _pdf_with_interleaved_table(target: Path, interleaved_kind: str) -> Path:
    import pymupdf

    pdf = pymupdf.open()
    try:
        page = pdf.new_page(width=360, height=240)
        for x in (24, 104, 184):
            page.draw_line((x, 40), (x, 160))
        for y in (40, 100, 160):
            page.draw_line((24, y), (184, y))
        for text, x, y in (
            ("A1", 34, 70),
            ("B1", 114, 70),
            ("A2", 34, 140),
            ("B2", 114, 140),
        ):
            page.insert_text((x, y), text, fontsize=10, fontname="helv")
        if interleaved_kind == "side":
            page.insert_text((220, 90), "SIDE", fontsize=10, fontname="helv")
        elif interleaved_kind == "caption":
            page.insert_text((24, 30), "CAPTION", fontsize=10, fontname="helv")
        elif interleaved_kind == "below":
            page.insert_text((24, 190), "BELOW", fontsize=10, fontname="helv")
        else:
            page.insert_image(
                pymupdf.Rect(80, 80, 120, 95),
                stream=_png(2, 1),
            )
        pdf.save(target)
    finally:
        pdf.close()
    return target


def _pdf_with_blank_terminal_row(target: Path, block_kind: str) -> Path:
    import pymupdf

    pdf = pymupdf.open()
    try:
        page = pdf.new_page(width=360, height=240)
        for x in (24, 104, 184):
            page.draw_line((x, 40), (x, 160))
        for y in (40, 100, 160):
            page.draw_line((24, y), (184, y))
        page.insert_text((34, 70), "A1", fontsize=10, fontname="helv")
        page.insert_text((114, 70), "B1", fontsize=10, fontname="helv")
        if block_kind == "image":
            page.insert_image(
                pymupdf.Rect(80, 120, 120, 135),
                stream=_png(2, 1),
            )
        else:
            page.insert_text((220, 130), "SIDE", fontsize=10, fontname="helv")
        pdf.save(target)
    finally:
        pdf.close()
    return target


def _pdf_with_images(target: Path) -> Path:
    import pymupdf

    pdf = pymupdf.open()
    try:
        page = pdf.new_page(width=240, height=240)
        page.insert_image(pymupdf.Rect(20, 20, 60, 40), stream=_png(2, 1))
        page.insert_image(pymupdf.Rect(20, 80, 80, 110), stream=_png(3, 2))
        pdf.save(target)
    finally:
        pdf.close()
    return target


def _pdf_with_indirect_image_length(target: Path, length_kind: str) -> Path:
    image_stream = zlib.compress(b"\xff\x00\x00\x00\xff\x00")
    content_stream = b"q 100 0 0 50 20 100 cm /Im0 Do Q\n"
    actual_length = str(len(image_stream)).encode("ascii")
    if length_kind == "valid":
        length_object = actual_length
        extra_objects: list[bytes] = []
    elif length_kind == "chain":
        length_object = b"7 0 R"
        extra_objects = [actual_length]
    elif length_kind == "non_integer":
        length_object = b"(invalid)"
        extra_objects = []
    elif length_kind == "oversized":
        length_object = str(
            DocxReconstructionLimits().max_image_bytes + 1
        ).encode("ascii")
        extra_objects = []
    else:
        raise AssertionError(length_kind)
    objects = [
        b"<< /Type /Catalog /Pages 2 0 R >>",
        b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
        (
            b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 200 200] "
            b"/Resources << /XObject << /Im0 4 0 R >> >> /Contents 5 0 R >>"
        ),
        (
            b"<< /Type /XObject /Subtype /Image /Width 2 /Height 1 "
            b"/ColorSpace /DeviceRGB /BitsPerComponent 8 /Filter /FlateDecode "
            b"/Length 6 0 R >>\nstream\n"
            + image_stream
            + b"\nendstream"
        ),
        (
            b"<< /Length "
            + str(len(content_stream)).encode("ascii")
            + b" >>\nstream\n"
            + content_stream
            + b"endstream"
        ),
        length_object,
        *extra_objects,
    ]
    payload = bytearray(b"%PDF-1.4\n%\xe2\xe3\xcf\xd3\n")
    offsets = [0]
    for object_number, value in enumerate(objects, start=1):
        offsets.append(len(payload))
        payload.extend(
            f"{object_number} 0 obj\n".encode("ascii")
            + value
            + b"\nendobj\n"
        )
    xref_offset = len(payload)
    payload.extend(f"xref\n0 {len(objects) + 1}\n".encode("ascii"))
    payload.extend(b"0000000000 65535 f \n")
    for offset in offsets[1:]:
        payload.extend(f"{offset:010d} 00000 n \n".encode("ascii"))
    payload.extend(
        (
            f"trailer\n<< /Size {len(objects) + 1} /Root 1 0 R >>\n"
            f"startxref\n{xref_offset}\n%%EOF\n"
        ).encode("ascii")
    )
    target.write_bytes(payload)
    return target


def _metadata_with_table(table: PdfTable) -> DocumentMetadata:
    return DocumentMetadata(
        pdf=PdfDocumentMetadata(
            pages=(
                PdfPageMetadata(
                    page=1,
                    kind=PdfPageKind.TEXT,
                    page_bbox=(0.0, 0.0, 360.0, 240.0),
                    text_length=10,
                    text_density=0.1,
                    image_coverage=0.0,
                    ocr_required=False,
                    tables=(table,),
                ),
            )
        )
    )
