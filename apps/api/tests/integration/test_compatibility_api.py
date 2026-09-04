from __future__ import annotations

import hashlib
import json
import logging
from io import BytesIO
from pathlib import Path
from uuid import UUID, uuid4
from zipfile import BadZipFile, ZipFile

import pytest
from docx import Document
from fastapi import FastAPI
from fastapi.testclient import TestClient
from pydantic import BaseModel

from text_verification.api.dependencies import get_verification_pipeline
from text_verification.application.errors import VerificationError
from text_verification.application.verification_pipeline import VerificationCommand
from text_verification.compatibility.storage import CompatibilityStorage
from text_verification.config import Settings, get_settings
from text_verification.domain.documents import FileType, TextBlock
from text_verification.domain.text_edits import MAX_REVISION_TEXT_CODEPOINTS
from text_verification.domain.verification import (
    Scenario,
    VerificationAnalysisMode,
    VerificationDegradation,
    VerificationExecutionMode,
    VerificationOptions,
    VerificationResult,
    VerificationStatistics,
    VerificationSummary,
)
from text_verification.infrastructure.dictionary_loader import DictionaryLoadError
from text_verification.parsers import compatibility_parser as compatibility_parser_module

PDF_FIXTURE_DIRECTORY = Path(__file__).resolve().parents[1] / "fixtures" / "pdf"
MAX_COMPATIBILITY_EXPORT_REQUEST_BYTES = 32 * 1024 * 1024


def override_storage(app: FastAPI, storage_root: Path) -> None:
    app.dependency_overrides[get_settings] = lambda: Settings(
        storage_root=storage_root,
        max_upload_bytes=1024 * 1024,
    )


class RecordingPipeline:
    def __init__(self, result: VerificationResult) -> None:
        self._result = result
        self.commands: list[VerificationCommand] = []

    def run(self, command: VerificationCommand) -> VerificationResult:
        self.commands.append(command)
        return self._result


class FailingPipeline:
    def __init__(self, error: VerificationError) -> None:
        self._error = error
        self.commands: list[VerificationCommand] = []

    def run(self, command: VerificationCommand) -> VerificationResult:
        self.commands.append(command)
        cause = self._error.__cause__
        if cause is not None:
            raise self._error from cause
        raise self._error


def _verification_result(*, text: str = "检查文本") -> VerificationResult:
    document_id = uuid4()
    verification_run_id = uuid4()
    return VerificationResult(
        verification_run_id=verification_run_id,
        document_id=document_id,
        source_version="sha256:test-source-version",
        source_name="直接输入文本",
        file_type=FileType.TXT,
        scenario=Scenario.GENERAL,
        text=text,
        blocks=(
            TextBlock(
                block_id="p-0",
                kind="paragraph",
                text=text,
                global_start=0,
                global_end=len(text),
                block_start=0,
                block_end=len(text),
                page=None,
                paragraph_index=0,
                table_index=None,
                row_index=None,
                cell_index=None,
                bbox=None,
                parent_id=None,
                style={},
                source_locator={"paragraph_index": 0},
            ),
        ),
        parser_name="test-parser",
        parser_version="1",
        stats=VerificationStatistics(
            char_count=len(text),
            char_count_no_space=len(text),
            line_count=1,
            paragraph_count=1,
            language="zh",
            primary_count=len(text),
            primary_label="总字数",
        ),
        issues=(),
        summary=VerificationSummary(total=0),
        execution_mode=VerificationExecutionMode.SYNCHRONOUS,
        analysis_mode=VerificationAnalysisMode.LOCAL_ONLY,
        dictionary_versions={},
        degradation=VerificationDegradation(),
    )


def _dictionary_load_failure(detail: str) -> VerificationError:
    error = VerificationError(
        code="dictionary_load_failed",
        stage="checking",
        message="A verification dictionary could not be loaded.",
        retryable=False,
    )
    error.__cause__ = DictionaryLoadError(detail)
    return error


def test_analyze_direct_text_supports_legacy_options(
    app: FastAPI,
    client: TestClient,
    tmp_path: Path,
) -> None:
    override_storage(app, tmp_path)

    response = client.post(
        "/api/v1/analyze",
        data={
            "text": "帐号包含禁用词和AI，联系 test@example.com。台湾产品全球领先。",
            "scenario": "business",
            "enable_security": "true",
            "enable_sensitive": "true",
            "enable_ad_extreme": "true",
            "custom_glossary": json.dumps(
                [{"original": "AI", "standard": "人工智能"}],
                ensure_ascii=False,
            ),
            "banned_words": json.dumps(["禁用词"], ensure_ascii=False),
        },
    )

    assert response.status_code == 200
    payload = response.json()
    assert set(payload) == {
        "success",
        "filename",
        "source_name",
        "file_type",
        "text",
        "blocks",
        "parser_name",
        "parser_version",
        "stats",
        "issues",
        "summary",
        "file_id",
        "file_ext",
        "scenario",
        "document_id",
        "verification_run_id",
        "source_version",
        "execution_mode",
        "analysis_mode",
        "dictionary_versions",
        "degradation",
    }
    assert payload["success"] is True
    assert payload["scenario"] == "business"
    assert payload["file_id"] is None
    assert UUID(payload["document_id"])
    assert UUID(payload["verification_run_id"])
    assert payload["source_version"].startswith("sha256:")
    assert payload["execution_mode"] == "synchronous"
    assert payload["analysis_mode"] == "local_only"
    assert set(payload["dictionary_versions"]) == {"sensitive_rules", "ad_extreme_words"}
    assert payload["degradation"] == {"is_degraded": False, "reasons": []}
    assert payload["stats"]["char_count"] == len(payload["text"])
    assert payload["issues"]
    assert UUID(payload["issues"][0]["issue_id"])
    assert payload["issues"][0]["position"] < payload["issues"][0]["end_position"]
    issue_types = {issue["type"] for issue in payload["issues"]}
    assert {
        "typo",
        "custom_term",
        "banned_word",
        "pii_email",
        "sensitive_territory",
        "ad_extreme",
    } <= issue_types


def test_analyze_route_uses_injected_pipeline(
    app: FastAPI,
    client: TestClient,
) -> None:
    pipeline = RecordingPipeline(_verification_result())
    app.dependency_overrides[get_verification_pipeline] = lambda: pipeline

    response = client.post("/api/v1/analyze", data={"text": "检查文本"})

    assert response.status_code == 200
    assert len(pipeline.commands) == 1
    assert pipeline.commands[0].direct_text == "检查文本"
    assert pipeline.commands[0].source_path is None
    assert pipeline.commands[0].source_name == "直接输入文本"
    assert pipeline.commands[0].file_type is FileType.TXT
    assert pipeline.commands[0].options == VerificationOptions()
    assert pipeline.commands[0].execution_mode is VerificationExecutionMode.SYNCHRONOUS


def test_analyze_direct_text_masks_dictionary_load_failures(
    app: FastAPI,
    client: TestClient,
    tmp_path: Path,
) -> None:
    override_storage(app, tmp_path)
    app.dependency_overrides[get_verification_pipeline] = lambda: FailingPipeline(
        _dictionary_load_failure("dictionary missing at /secret/dictionaries/sensitive_rules.json")
    )

    response = client.post("/api/v1/analyze", data={"text": "待检测文本"})

    assert response.status_code == 500
    assert response.json() == {"detail": "Verification dictionaries are unavailable."}
    assert "secret" not in response.text.lower()
    assert "dictionary missing" not in response.text.lower()


def test_analyze_direct_text_normalizes_dictionary_encoding_failure(
    app: FastAPI,
    client: TestClient,
    tmp_path: Path,
) -> None:
    override_storage(app, tmp_path)
    app.dependency_overrides[get_verification_pipeline] = lambda: FailingPipeline(
        _dictionary_load_failure("dictionary decoding failed in invalid-dictionaries")
    )

    response = client.post("/api/v1/analyze", data={"text": "待检测文本"})

    assert response.status_code == 500
    assert response.json() == {"detail": "Verification dictionaries are unavailable."}
    assert "invalid-dictionaries" not in response.text


def test_upload_and_export_txt_from_uuid_scoped_storage(
    app: FastAPI,
    client: TestClient,
    tmp_path: Path,
) -> None:
    override_storage(app, tmp_path)

    analysis = client.post(
        "/api/v1/analyze",
        files={"file": ("../../unsafe.txt", "帐号测试".encode(), "text/plain")},
    )

    assert analysis.status_code == 200
    result = analysis.json()
    file_id = result["file_id"]
    assert result["filename"] == "unsafe.txt"
    assert (tmp_path / "compatibility" / file_id / "source.txt").read_text() == "帐号测试"

    exported = client.post(
        "/api/v1/export-original",
        json={
            "file_id": file_id,
            "filename": "../../unsafe.txt",
            "replacements": [
                {
                    "original": "帐号",
                    "suggestion": "账号",
                    "position": 0,
                    "end_position": 2,
                }
            ],
            "modified_text": "账号测试",
            "track_changes": False,
        },
    )

    assert exported.status_code == 200
    assert exported.content.decode() == "账号测试"
    assert "unsafe_%E4%BF%AE%E6%94%B9%E7%89%88_" in exported.headers["content-disposition"]
    assert exported.headers["content-disposition"].endswith(".txt")


def test_export_original_unexpected_failure_returns_stable_sanitized_500(
    app: FastAPI,
    client: TestClient,
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
    caplog: pytest.LogCaptureFixture,
) -> None:
    from text_verification.api.routes import compatibility as compatibility_routes

    override_storage(app, tmp_path)
    secret = "export-original-secret-never-reflect"

    def fail_export(*args: object, **kwargs: object) -> object:
        del args, kwargs
        raise RuntimeError(f"unexpected exporter leak {secret}")

    monkeypatch.setattr(compatibility_routes, "export_original", fail_export)

    with caplog.at_level(logging.WARNING, logger="text_verification.api.routes.compatibility"):
        analysis = client.post(
            "/api/v1/analyze",
            files={"file": ("source.txt", "帐号测试".encode(), "text/plain")},
        )
        assert analysis.status_code == 200

        response = client.post(
            "/api/v1/export-original",
            json={
                "file_id": analysis.json()["file_id"],
                "filename": "source.txt",
                "replacements": [
                    {
                        "original": "帐号",
                        "suggestion": "账号",
                        "position": 0,
                        "end_position": 2,
                    }
                ],
                "modified_text": "账号测试",
                "track_changes": False,
            },
        )

    assert response.status_code == 500
    assert response.json() == {
        "detail": "Export failed due to an internal server error."
    }
    assert secret not in response.text
    assert "RuntimeError" not in response.text
    export_records = [
        record
        for record in caplog.records
        if record.name == "text_verification.api.routes.compatibility"
    ]
    assert export_records == []
    for record in export_records:
        assert record.exc_info is None
        assert record.exc_text is None
        assert secret not in record.getMessage()
        assert "Traceback" not in record.getMessage()
        for value in record.__dict__.values():
            assert secret not in str(value)
            assert "Traceback" not in str(value)
    assert secret not in caplog.text
    assert "Traceback" not in caplog.text


def test_export_original_rejects_modified_text_above_configured_upload_limit(
    app: FastAPI,
    client: TestClient,
    tmp_path: Path,
) -> None:
    app.dependency_overrides[get_settings] = lambda: Settings(
        storage_root=tmp_path / "compatibility",
        max_upload_bytes=8,
    )
    analysis = client.post(
        "/api/v1/analyze",
        files={"file": ("source.txt", b"text", "text/plain")},
    )
    assert analysis.status_code == 200

    accepted = client.post(
        "/api/v1/export-original",
        json={
            "file_id": analysis.json()["file_id"],
            "filename": "source.txt",
            "modified_text": "12345678",
            "track_changes": False,
        },
    )
    rejected = client.post(
        "/api/v1/export-original",
        json={
            "file_id": analysis.json()["file_id"],
            "filename": "source.txt",
            "modified_text": "123456789",
            "track_changes": False,
        },
    )

    assert accepted.status_code == 200
    assert accepted.content == b"12345678"
    assert rejected.status_code == 413


def test_export_original_rejects_modified_text_above_the_model_codepoint_limit(
    app: FastAPI,
    client: TestClient,
    tmp_path: Path,
) -> None:
    settings = Settings(
        storage_root=tmp_path / "compatibility",
        max_upload_bytes=6 * 1024 * 1024,
    )
    app.dependency_overrides[get_settings] = lambda: settings
    file_id = uuid4()
    boundary_text = "x" * MAX_REVISION_TEXT_CODEPOINTS
    CompatibilityStorage(
        settings.storage_root,
        settings.max_upload_bytes,
    ).save_stream(
        file_id,
        "source.txt",
        BytesIO(boundary_text.encode()),
    )

    accepted = client.post(
        "/api/v1/export-original",
        json={
            "file_id": str(file_id),
            "filename": "source.txt",
            "modified_text": boundary_text,
            "track_changes": False,
        },
    )
    rejected = client.post(
        "/api/v1/export-original",
        json={
            "file_id": str(file_id),
            "filename": "source.txt",
            "modified_text": "x" * (MAX_REVISION_TEXT_CODEPOINTS + 1),
            "track_changes": False,
        },
    )

    assert accepted.status_code == 200
    assert len(accepted.content) == MAX_REVISION_TEXT_CODEPOINTS
    assert rejected.status_code == 422
    assert rejected.json()["detail"][0]["type"] == "string_too_long"


def test_export_original_rejects_an_oversized_declared_body_before_model_validation(
    client: TestClient,
) -> None:
    response = client.post(
        "/api/v1/export-original",
        content=b"{}",
        headers={
            "Content-Type": "application/json",
            "Content-Length": str(MAX_COMPATIBILITY_EXPORT_REQUEST_BYTES + 1),
        },
    )

    assert response.status_code == 413


def test_uploaded_source_version_hashes_source_bytes_not_extracted_text(
    app: FastAPI,
    client: TestClient,
    tmp_path: Path,
) -> None:
    override_storage(app, tmp_path)
    utf8_bytes = "帐号测试".encode()
    utf16_bytes = "帐号测试".encode("utf-16")

    utf8 = client.post(
        "/api/v1/analyze",
        files={"file": ("utf8.txt", utf8_bytes, "text/plain")},
    )
    utf16 = client.post(
        "/api/v1/analyze",
        files={"file": ("utf16.txt", utf16_bytes, "text/plain")},
    )

    assert utf8.status_code == 200
    assert utf16.status_code == 200
    assert utf8.json()["text"] == utf16.json()["text"] == "帐号测试"
    assert utf8.json()["source_version"] == (
        f"sha256:{hashlib.sha256(utf8_bytes).hexdigest()}"
    )
    assert utf16.json()["source_version"] == (
        f"sha256:{hashlib.sha256(utf16_bytes).hexdigest()}"
    )
    assert utf8.json()["source_version"] != utf16.json()["source_version"]


def test_uploaded_file_is_deleted_when_dictionary_load_fails(
    app: FastAPI,
    client: TestClient,
    tmp_path: Path,
) -> None:
    override_storage(app, tmp_path)
    app.dependency_overrides[get_verification_pipeline] = lambda: FailingPipeline(
        _dictionary_load_failure("invalid dictionary at C:\\secret\\rules.json")
    )

    response = client.post(
        "/api/v1/analyze",
        files={"file": ("source.txt", "帐号测试".encode(), "text/plain")},
    )

    compatibility_root = tmp_path / "compatibility"
    assert response.status_code == 500
    assert response.json() == {"detail": "Verification dictionaries are unavailable."}
    assert "secret" not in response.text.lower()
    assert not compatibility_root.exists() or list(compatibility_root.iterdir()) == []


def test_uploaded_file_is_deleted_when_dictionary_encoding_is_invalid(
    app: FastAPI,
    client: TestClient,
    tmp_path: Path,
) -> None:
    override_storage(app, tmp_path)
    app.dependency_overrides[get_verification_pipeline] = lambda: FailingPipeline(
        _dictionary_load_failure("dictionary decoding failed in invalid-dictionaries")
    )

    response = client.post(
        "/api/v1/analyze",
        files={"file": ("source.txt", "帐号测试".encode(), "text/plain")},
    )

    compatibility_root = tmp_path / "compatibility"
    assert response.status_code == 500
    assert response.json() == {"detail": "Verification dictionaries are unavailable."}
    assert not compatibility_root.exists() or list(compatibility_root.iterdir()) == []


def test_failed_upload_cleanup_warning_omits_secret_and_traceback(
    app: FastAPI,
    client: TestClient,
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
    caplog: pytest.LogCaptureFixture,
) -> None:
    override_storage(app, tmp_path)
    app.dependency_overrides[get_verification_pipeline] = lambda: FailingPipeline(
        _dictionary_load_failure("dictionary missing at /secret/dictionaries/sensitive_rules.json")
    )
    secret = "cleanup-secret-never-reflect"

    def fail_delete(self: CompatibilityStorage, file_id: UUID) -> None:
        del file_id
        raise PermissionError(f"locked cleanup path {secret}")

    monkeypatch.setattr(CompatibilityStorage, "delete", fail_delete)

    with caplog.at_level(logging.WARNING, logger="text_verification.api.routes.compatibility"):
        response = client.post(
            "/api/v1/analyze",
            files={"file": ("source.txt", "帐号测试".encode(), "text/plain")},
        )

    compatibility_root = tmp_path / "compatibility"
    [stored_upload] = list(compatibility_root.iterdir())
    record = caplog.records[0]
    standard_fields = set(logging.makeLogRecord({}).__dict__) | {"message"}

    assert response.status_code == 500
    assert response.json() == {"detail": "Verification dictionaries are unavailable."}
    assert set(record.__dict__) - standard_fields == {"file_id", "error_type"}
    assert [captured.getMessage() for captured in caplog.records] == [
        "compatibility_upload_cleanup_failed"
    ]
    assert record.file_id == stored_upload.name
    assert record.error_type == "PermissionError"
    assert record.exc_info is None
    assert record.exc_text is None
    for value in record.__dict__.values():
        assert secret not in str(value)
    assert secret not in response.text
    assert secret not in caplog.text
    assert "Traceback" not in caplog.text


def test_uploaded_value_error_preserves_legacy_400_detail_and_cleans_up(
    app: FastAPI,
    client: TestClient,
    tmp_path: Path,
    monkeypatch,
) -> None:
    override_storage(app, tmp_path)

    def fail_parse(*args: object) -> tuple[str, str, list[tuple[int, int, str]]]:
        del args
        raise ValueError("expected legacy parse failure")

    monkeypatch.setattr(compatibility_parser_module, "parse_file", fail_parse)

    response = client.post(
        "/api/v1/analyze",
        files={"file": ("source.txt", "帐号测试".encode(), "text/plain")},
    )

    compatibility_root = tmp_path / "compatibility"
    assert response.status_code == 400
    assert response.json() == {"detail": "expected legacy parse failure"}
    assert not compatibility_root.exists() or list(compatibility_root.iterdir()) == []


def test_uploaded_library_parse_failure_keeps_legacy_prefix_and_cleans_up(
    app: FastAPI,
    client: TestClient,
    tmp_path: Path,
    monkeypatch,
) -> None:
    override_storage(app, tmp_path)

    def fail_parse(*args: object) -> tuple[str, str, list[tuple[int, int, str]]]:
        del args
        raise BadZipFile("File is not a zip file")

    monkeypatch.setattr(compatibility_parser_module, "parse_file", fail_parse)
    source = BytesIO()
    document = Document()
    document.add_paragraph("帐号测试")
    document.save(source)

    response = client.post(
        "/api/v1/analyze",
        files={
            "file": (
                "source.docx",
                source.getvalue(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )

    compatibility_root = tmp_path / "compatibility"
    assert response.status_code == 400
    assert response.json() == {"detail": "File parsing failed: File is not a zip file"}
    assert not compatibility_root.exists() or list(compatibility_root.iterdir()) == []


def test_uploaded_validation_error_uses_prefixed_legacy_detail_and_cleans_up(
    app: FastAPI,
    client: TestClient,
    tmp_path: Path,
    monkeypatch,
) -> None:
    override_storage(app, tmp_path)

    class ParseProbe(BaseModel):
        count: int

    def fail_parse(*args: object) -> tuple[str, str, list[tuple[int, int, str]]]:
        del args
        ParseProbe.model_validate({"count": "bad"})
        raise AssertionError("unreachable")

    monkeypatch.setattr(compatibility_parser_module, "parse_file", fail_parse)

    response = client.post(
        "/api/v1/analyze",
        files={"file": ("source.txt", "帐号测试".encode(), "text/plain")},
    )

    compatibility_root = tmp_path / "compatibility"
    assert response.status_code == 400
    detail = response.json()["detail"]
    assert detail.startswith("File parsing failed: 1 validation error for ParseProbe")
    assert not detail.startswith("1 validation error for ParseProbe")
    assert not compatibility_root.exists() or list(compatibility_root.iterdir()) == []


def test_export_html_report_escapes_user_content(client: TestClient) -> None:
    response = client.post(
        "/api/v1/export",
        json={
            "filename": "<script>alert(1)</script>.txt",
            "stats": {"char_count": 1},
            "summary": {"total": 1},
            "issues": [
                {
                    "layer": "character",
                    "type": "typo",
                    "severity": "error",
                    "original": "<b>x</b>",
                    "suggestion": "y",
                    "description": "test",
                    "context": "context",
                }
            ],
        },
    )

    assert response.status_code == 200
    assert response.headers["content-type"].startswith("text/html")
    html = response.text
    assert "<script>" not in html
    assert "&lt;script&gt;" in html
    assert "&lt;b&gt;x&lt;/b&gt;" in html


def test_export_html_report_keeps_nullable_suggestion_display_safe(
    client: TestClient,
) -> None:
    response = client.post(
        "/api/v1/export",
        json={
            "filename": "report.txt",
            "stats": {},
            "summary": {"total": 1},
            "issues": [
                {
                    "layer": "security",
                    "type": "sensitive_politics",
                    "severity": "error",
                    "original": "敏感词",
                    "suggestion": None,
                    "description": "请人工处理",
                    "context": "敏感词",
                }
            ],
        },
    )

    assert response.status_code == 200
    assert '<td class="suggestion"></td>' in response.text
    assert ">None<" not in response.text


def test_export_report_preflights_raw_result_block_count_before_validation(
    client: TestClient,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    from text_verification.api.routes import compatibility as compatibility_routes

    monkeypatch.setattr(
        compatibility_routes,
        "MAX_CANONICAL_RESULT_BLOCKS",
        2,
        raising=False,
    )
    response = client.post(
        "/api/v1/export",
        json={
            "filename": "report.txt",
            "stats": {},
            "summary": {"total": 0},
            "issues": [],
            "text": "abc",
            "blocks": [
                {"block_id": "a", "text": "a"},
                {"block_id": "b", "text": "b"},
                {"block_id": "c", "text": "c"},
            ],
        },
    )

    assert response.status_code == 413


def test_export_report_preflights_raw_result_aggregate_text_size(
    client: TestClient,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    from text_verification.api.routes import compatibility as compatibility_routes

    monkeypatch.setattr(
        compatibility_routes,
        "MAX_CANONICAL_RESULT_TOTAL_CODEPOINTS",
        5,
        raising=False,
    )
    response = client.post(
        "/api/v1/export",
        json={
            "filename": "report.txt",
            "stats": {},
            "summary": {"total": 0},
            "issues": [],
            "text": "ab",
            "blocks": [
                {"block_id": "a", "text": "ab"},
                {"block_id": "b", "text": "ab"},
            ],
        },
    )

    assert response.status_code == 413


def test_export_report_preflights_raw_result_aggregate_utf8_size(
    client: TestClient,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    from text_verification.api.routes import compatibility as compatibility_routes

    monkeypatch.setattr(
        compatibility_routes,
        "MAX_CANONICAL_RESULT_TOTAL_UTF8_BYTES",
        7,
    )
    response = client.post(
        "/api/v1/export",
        json={
            "filename": "report.txt",
            "stats": {},
            "summary": {"total": 0},
            "issues": [],
            "text": "😀",
            "blocks": [{"block_id": "a", "text": "😀"}],
        },
    )

    assert response.status_code == 413


def test_export_report_rejects_malformed_raw_block_without_reflecting_it(
    client: TestClient,
) -> None:
    secret = "raw-block-secret-that-must-not-be-reflected"
    response = client.post(
        "/api/v1/export",
        json={
            "filename": "report.txt",
            "stats": {},
            "summary": {"total": 0},
            "issues": [],
            "text": "a",
            "blocks": [secret],
        },
    )

    assert response.status_code == 422
    assert secret not in response.text


def test_export_report_json_body_limit_is_inclusive_and_streaming(
    client: TestClient,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    from text_verification.api.routes import compatibility as compatibility_routes

    body = b'{"filename":"report.txt","stats":{},"summary":{},"issues":[]}'
    monkeypatch.setattr(
        compatibility_routes,
        "MAX_COMPATIBILITY_REPORT_REQUEST_BYTES",
        len(body),
    )
    accepted = client.post(
        "/api/v1/export",
        content=body,
        headers={"Content-Type": "application/json"},
    )
    monkeypatch.setattr(
        compatibility_routes,
        "MAX_COMPATIBILITY_REPORT_REQUEST_BYTES",
        len(body) - 1,
    )
    rejected = client.post(
        "/api/v1/export",
        content=iter((body[:7], body[7:])),
        headers={"Content-Type": "application/json"},
    )

    assert accepted.status_code == 200
    assert rejected.status_code == 413


def test_export_report_accepts_large_valid_legacy_full_result_payload(
    client: TestClient,
) -> None:
    text = "汉" * 3_000_000
    body = json.dumps(
        {
            "filename": "report.txt",
            "stats": {},
            "summary": {"total": 0},
            "issues": [],
            "text": text,
            "blocks": [{"block_id": "only", "text": text}],
        },
        separators=(",", ":"),
    ).encode()
    assert len(body) > 32 * 1024 * 1024

    response = client.post(
        "/api/v1/export",
        content=body,
        headers={"Content-Type": "application/json"},
    )

    assert response.status_code == 200


def test_export_report_reader_does_not_buffer_the_complete_raw_body(
    client: TestClient,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    from text_verification.api import body_readers

    calls = 0
    original = body_readers.read_bounded_body

    async def recording_reader(*args: object, **kwargs: object) -> bytes:
        nonlocal calls
        calls += 1
        return await original(*args, **kwargs)

    monkeypatch.setattr(body_readers, "read_bounded_body", recording_reader)

    response = client.post(
        "/api/v1/export",
        json={
            "filename": "report.txt",
            "stats": {},
            "summary": {"total": 0},
            "issues": [],
        },
    )

    assert response.status_code == 200
    assert calls == 0


def test_export_report_retained_budget_rejects_chunked_empty_array_entries(
    client: TestClient,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    from text_verification.api.routes import compatibility as compatibility_routes

    body = json.dumps(
        {
            "filename": "report.txt",
            "stats": {"values": [""] * 10_000},
            "summary": {},
            "issues": [],
        },
        separators=(",", ":"),
    ).encode()
    monkeypatch.setattr(
        compatibility_routes,
        "MAX_COMPATIBILITY_EXPORT_REQUEST_BYTES",
        100,
    )

    response = client.post(
        "/api/v1/export",
        content=iter((body[:17], body[17:103], body[103:])),
        headers={"Content-Type": "application/json"},
    )

    assert response.status_code == 413


@pytest.mark.parametrize(
    "stats",
    [
        {chr(0x100 + index): "" for index in range(24)},
        {"nested": [[[]]]},
    ],
)
def test_export_report_retained_budget_accounts_for_wide_empty_structures(
    client: TestClient,
    monkeypatch: pytest.MonkeyPatch,
    stats: dict[str, object],
) -> None:
    from text_verification.api.routes import compatibility as compatibility_routes

    monkeypatch.setattr(
        compatibility_routes,
        "MAX_COMPATIBILITY_EXPORT_REQUEST_BYTES",
        100,
    )

    response = client.post(
        "/api/v1/export",
        json={"stats": stats},
    )

    assert response.status_code == 413


def test_export_report_rejects_dotted_root_key_without_filename_spoof(
    client: TestClient,
) -> None:
    response = client.post(
        "/api/v1/export",
        content=(
            b'{"filename.x":"spoofed.txt","stats":{},'
            b'"summary":{},"issues":[]}'
        ),
        headers={"Content-Type": "application/json"},
    )

    assert response.status_code == 422
    assert "spoofed.txt" not in response.text


@pytest.mark.parametrize(
    "raw_block",
    [
        (
            b'{"block_id":"a","text":"safe","style":'
            b'{"marker":"nested-secret","marker":"replacement"}}'
        ),
        (
            b'{"block_id":"a","text":"safe","source_locator":'
            b'{"path":"nested-secret","path":"replacement"}}'
        ),
    ],
)
def test_export_report_rejects_duplicate_keys_in_skipped_block_subtrees(
    client: TestClient,
    caplog: pytest.LogCaptureFixture,
    raw_block: bytes,
) -> None:
    response = client.post(
        "/api/v1/export",
        content=(
            b'{"filename":"report.txt","stats":{},"summary":{},'
            b'"issues":[],"text":"safe","blocks":['
            + raw_block
            + b"]}"
        ),
        headers={"Content-Type": "application/json"},
    )

    assert response.status_code == 422
    assert "nested-secret" not in response.text
    assert "nested-secret" not in caplog.text


def test_export_report_charges_large_unique_keys_in_skipped_block_metadata(
    client: TestClient,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    from text_verification.api.routes import compatibility as compatibility_routes

    monkeypatch.setattr(
        compatibility_routes,
        "MAX_COMPATIBILITY_EXPORT_REQUEST_BYTES",
        400,
    )
    large_key = "k" * 500

    def export_with_style_key(style_key: str):
        return client.post(
            "/api/v1/export",
            json={
                "filename": "report.txt",
                "stats": {},
                "summary": {},
                "issues": [],
                "text": "safe",
                "blocks": [
                    {
                        "block_id": "a",
                        "text": "safe",
                        "style": {style_key: ""}
                    }
                ]
            },
        )

    accepted = export_with_style_key("k")
    rejected = export_with_style_key(large_key)

    assert accepted.status_code == 200
    assert rejected.status_code == 413


def test_export_report_accepts_dotted_string_values(
    client: TestClient,
) -> None:
    response = client.post(
        "/api/v1/export",
        json={
            "filename": "report.v1.txt",
            "stats": {"label": "section.one"},
            "summary": {},
            "issues": [],
        },
    )

    assert response.status_code == 200


def test_export_report_malformed_json_never_reflects_request_secrets(
    client: TestClient,
    caplog: pytest.LogCaptureFixture,
) -> None:
    secret = "report-body-secret-never-reflect"
    response = client.post(
        "/api/v1/export",
        content=b'{"filename":"' + secret.encode(),
        headers={"Content-Type": "application/json"},
    )

    assert response.status_code == 422
    assert secret not in response.text
    assert secret not in caplog.text


def test_docx_export_can_emit_word_track_changes(
    app: FastAPI,
    client: TestClient,
    tmp_path: Path,
) -> None:
    override_storage(app, tmp_path)
    source = BytesIO()
    document = Document()
    document.add_paragraph("帐号测试")
    document.save(source)

    analysis = client.post(
        "/api/v1/analyze",
        files={
            "file": (
                "source.docx",
                source.getvalue(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )
    assert analysis.status_code == 200

    exported = client.post(
        "/api/v1/export-original",
        json={
            "file_id": analysis.json()["file_id"],
            "filename": "source.docx",
            "replacements": [
                {
                    "original": "帐号",
                    "suggestion": "账号",
                    "position": 0,
                    "end_position": 2,
                }
            ],
            "modified_text": "账号测试",
            "track_changes": True,
        },
    )

    assert exported.status_code == 200
    with ZipFile(BytesIO(exported.content)) as archive:
        document_xml = archive.read("word/document.xml").decode()
    assert "<w:del " in document_xml
    assert "<w:ins " in document_xml
    assert "帐" in document_xml
    assert "账" in document_xml


def test_docx_analysis_maps_uploaded_issue_to_paragraph_block_offsets(
    app: FastAPI,
    client: TestClient,
    tmp_path: Path,
) -> None:
    override_storage(app, tmp_path)
    source = BytesIO()
    document = Document()
    document.add_paragraph("第一段正常")
    document.add_paragraph("第二段帐号待修正")
    document.save(source)

    analysis = client.post(
        "/api/v1/analyze",
        files={
            "file": (
                "source.docx",
                source.getvalue(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )

    assert analysis.status_code == 200
    issue = next(
        item for item in analysis.json()["issues"] if item["original"] == "帐号"
    )

    assert issue["page"] is None
    assert issue["block_id"] == "paragraph-1"
    assert issue["block_start"] == 3
    assert issue["block_end"] == 5


def test_export_changes_only_the_selected_duplicate(
    app: FastAPI,
    client: TestClient,
    tmp_path: Path,
) -> None:
    override_storage(app, tmp_path)
    analysis = client.post(
        "/api/v1/analyze",
        files={"file": ("source.txt", "帐号 A 帐号 B".encode(), "text/plain")},
    )
    result = analysis.json()

    exported = client.post(
        "/api/v1/export-original",
        json={
            "file_id": result["file_id"],
            "filename": "source.txt",
            "replacements": [
                {
                    "original": "帐号",
                    "suggestion": "账号",
                    "position": 0,
                    "end_position": 2,
                }
            ],
            "modified_text": "账号 A 帐号 B",
            "track_changes": False,
        },
    )

    assert exported.status_code == 200
    assert exported.content.decode() == "账号 A 帐号 B"

    manually_edited = client.post(
        "/api/v1/export-original",
        json={
            "file_id": result["file_id"],
            "filename": "source.txt",
            "modified_text": "帐号 A 帐号（已核对） B",
            "track_changes": False,
        },
    )
    assert manually_edited.status_code == 200
    assert manually_edited.content.decode() == "帐号 A 帐号（已核对） B"


def test_docx_export_preserves_unaffected_run_formatting(
    app: FastAPI,
    client: TestClient,
    tmp_path: Path,
) -> None:
    override_storage(app, tmp_path)
    source = BytesIO()
    document = Document()
    paragraph = document.add_paragraph()
    paragraph.add_run("帐").bold = True
    paragraph.add_run("号测试").italic = True
    document.save(source)
    analysis = client.post(
        "/api/v1/analyze",
        files={
            "file": (
                "source.docx",
                source.getvalue(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )

    exported = client.post(
        "/api/v1/export-original",
        json={
            "file_id": analysis.json()["file_id"],
            "filename": "source.docx",
            "modified_text": "账号测试",
            "track_changes": False,
        },
    )

    assert exported.status_code == 200
    modified = Document(BytesIO(exported.content))
    runs = modified.paragraphs[0].runs
    assert runs[0].text == "账"
    assert runs[0].bold is True
    assert runs[1].text == "号测试"
    assert runs[1].italic is True


def test_docx_export_preserves_unaffected_inline_content(
    app: FastAPI,
    client: TestClient,
    tmp_path: Path,
) -> None:
    override_storage(app, tmp_path)
    source = BytesIO()
    document = Document()
    paragraph = document.add_paragraph()
    paragraph.add_run("A")
    paragraph.add_run().add_tab()
    paragraph.add_run("B")
    document.save(source)
    analysis = client.post(
        "/api/v1/analyze",
        files={
            "file": (
                "source.docx",
                source.getvalue(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )

    exported = client.post(
        "/api/v1/export-original",
        json={
            "file_id": analysis.json()["file_id"],
            "filename": "source.docx",
            "modified_text": "C\tB",
            "track_changes": False,
        },
    )

    assert exported.status_code == 200
    with ZipFile(BytesIO(exported.content)) as archive:
        document_xml = archive.read("word/document.xml").decode()
    assert "<w:tab/>" in document_xml
    assert Document(BytesIO(exported.content)).paragraphs[0].text == "C\tB"


def test_rtf_analysis_decodes_unicode_and_ansi_hex(
    app: FastAPI,
    client: TestClient,
    tmp_path: Path,
) -> None:
    override_storage(app, tmp_path)
    rtf = rb"{\rtf1\ansi\ansicpg936 Unicode: \u36134?\u21495? ANSI: \'d5\'ca\'ba\'c5}"
    analysis = client.post(
        "/api/v1/analyze",
        files={"file": ("source.rtf", rtf, "application/rtf")},
    )

    assert analysis.status_code == 200
    assert analysis.json()["text"] == "Unicode: 账号 ANSI: 帐号"


def test_rtf_export_keeps_offsets_after_repeated_paragraph_breaks(
    app: FastAPI,
    client: TestClient,
    tmp_path: Path,
) -> None:
    override_storage(app, tmp_path)
    source = rb"{\rtf1\ansi A\par\par\par B}"
    analysis = client.post(
        "/api/v1/analyze",
        files={"file": ("source.rtf", source, "application/rtf")},
    )
    result = analysis.json()
    assert result["text"] == "A\n\n\nB"

    exported = client.post(
        "/api/v1/export-original",
        json={
            "file_id": result["file_id"],
            "filename": "source.rtf",
            "modified_text": "A\n\n\nC",
            "track_changes": False,
        },
    )
    assert exported.status_code == 200
    reparsed = client.post(
        "/api/v1/analyze",
        files={"file": ("modified.rtf", exported.content, "application/rtf")},
    )
    assert reparsed.status_code == 200
    assert reparsed.json()["text"] == "A\n\n\nC"

    unicode_export = client.post(
        "/api/v1/export-original",
        json={
            "file_id": result["file_id"],
            "filename": "source.rtf",
            "modified_text": "A\n\n\nB\n😀",
            "track_changes": False,
        },
    )
    assert unicode_export.status_code == 200
    unicode_reparsed = client.post(
        "/api/v1/analyze",
        files={"file": ("unicode.rtf", unicode_export.content, "application/rtf")},
    )
    assert unicode_reparsed.status_code == 200
    assert unicode_reparsed.json()["text"] == "A\n\n\nB\n😀"


def test_pdf_export_resolves_duplicates_and_rejects_unsafe_insertions(
    app: FastAPI,
    client: TestClient,
    tmp_path: Path,
) -> None:
    import fitz

    override_storage(app, tmp_path)
    document = fitz.open()
    page = document.new_page()
    page.insert_text((72, 72), "word word")
    source = document.tobytes()
    document.close()
    analysis = client.post(
        "/api/v1/analyze",
        files={"file": ("source.pdf", source, "application/pdf")},
    )
    result = analysis.json()
    assert result["text"] == "word word"

    replaced = client.post(
        "/api/v1/export-original",
        json={
            "file_id": result["file_id"],
            "filename": "source.pdf",
            "modified_text": "term term",
            "track_changes": False,
        },
    )
    inserted = client.post(
        "/api/v1/export-original",
        json={
            "file_id": result["file_id"],
            "filename": "source.pdf",
            "modified_text": "word extra word",
            "track_changes": False,
        },
    )

    assert replaced.status_code == 200
    assert inserted.status_code == 400
    assert "does not support insertion-only edits" in inserted.json()["detail"]


def test_pdf_modified_text_keeps_edits_across_newline_independent(
    app: FastAPI,
    client: TestClient,
    monkeypatch: pytest.MonkeyPatch,
    tmp_path: Path,
) -> None:
    import pymupdf

    override_storage(app, tmp_path)
    document = pymupdf.open()
    page = document.new_page()
    page.insert_text((72, 72), "A")
    page.insert_text((72, 100), "B")
    source = document.tobytes()
    document.close()
    analyzed = client.post(
        "/api/v1/analyze",
        files={"file": ("lines.pdf", source, "application/pdf")},
    )
    assert analyzed.status_code == 200
    assert analyzed.json()["text"] == "A\nB"

    inserted: list[str] = []
    original_insert_textbox = pymupdf.Page.insert_textbox

    def recording_insert_textbox(
        page: pymupdf.Page,
        rectangle: object,
        buffer: str | list[str],
        **kwargs: object,
    ) -> float:
        assert isinstance(buffer, str)
        inserted.append(buffer)
        return original_insert_textbox(page, rectangle, buffer, **kwargs)

    monkeypatch.setattr(pymupdf.Page, "insert_textbox", recording_insert_textbox)
    exported = client.post(
        "/api/v1/export-original",
        json={
            "file_id": analyzed.json()["file_id"],
            "filename": "lines.pdf",
            "modified_text": "X\nY",
            "track_changes": False,
        },
    )

    assert exported.status_code == 200
    assert inserted == ["X", "Y"]
    with pymupdf.open(stream=exported.content, filetype="pdf") as output:
        assert output[0].get_text("text").strip() == "X\nY"


def test_pdf_positioned_edits_apply_in_canonical_source_order(
    app: FastAPI,
    client: TestClient,
    monkeypatch: pytest.MonkeyPatch,
    tmp_path: Path,
) -> None:
    import pymupdf

    override_storage(app, tmp_path)
    document = pymupdf.open()
    page = document.new_page()
    page.insert_text((72, 72), "A B")
    source = document.tobytes()
    document.close()
    analyzed = client.post(
        "/api/v1/analyze",
        files={"file": ("ordered.pdf", source, "application/pdf")},
    )
    assert analyzed.status_code == 200

    inserted: list[str] = []
    original_insert_textbox = pymupdf.Page.insert_textbox

    def recording_insert_textbox(
        page: pymupdf.Page,
        rectangle: object,
        buffer: str | list[str],
        **kwargs: object,
    ) -> float:
        assert isinstance(buffer, str)
        inserted.append(buffer)
        return original_insert_textbox(page, rectangle, buffer, **kwargs)

    monkeypatch.setattr(pymupdf.Page, "insert_textbox", recording_insert_textbox)
    exported = client.post(
        "/api/v1/export-original",
        json={
            "file_id": analyzed.json()["file_id"],
            "filename": "ordered.pdf",
            "replacements": [
                {
                    "original": "B",
                    "suggestion": "Y",
                    "position": 2,
                    "end_position": 3,
                },
                {
                    "original": "A",
                    "suggestion": "X",
                    "position": 0,
                    "end_position": 1,
                },
            ],
            "track_changes": False,
        },
    )

    assert exported.status_code == 200
    assert inserted == ["X", "Y"]
    with pymupdf.open(stream=exported.content, filetype="pdf") as output:
        assert output[0].get_text("text").strip() == "X Y"


def test_pdf_analysis_maps_uploaded_issue_to_page_block_offsets(
    app: FastAPI,
    client: TestClient,
    tmp_path: Path,
) -> None:
    import fitz

    override_storage(app, tmp_path)
    document = fitz.open()
    first_page = document.new_page()
    first_page.insert_text((72, 72), "Clean page")
    second_page = document.new_page()
    second_page.insert_text((72, 72), "forbidden")
    source = document.tobytes()
    document.close()

    analysis = client.post(
        "/api/v1/analyze",
        data={"banned_words": json.dumps(["forbidden"])},
        files={"file": ("source.pdf", source, "application/pdf")},
    )

    assert analysis.status_code == 200
    issue = next(
        item for item in analysis.json()["issues"] if item["original"] == "forbidden"
    )

    assert issue["page"] == 2
    assert issue["block_id"] == "page-2"
    assert issue["block_start"] == 0
    assert issue["block_end"] == 9


def test_scan_only_pdf_returns_an_explicit_ocr_capability_error(
    app: FastAPI,
    client: TestClient,
    tmp_path: Path,
) -> None:
    override_storage(app, tmp_path)

    response = client.post(
        "/api/v1/analyze",
        files={
            "file": (
                "scan.pdf",
                PDF_FIXTURE_DIRECTORY.joinpath("scanned-page.pdf").read_bytes(),
                "application/pdf",
            )
        },
    )

    assert response.status_code == 422
    assert response.json() == {
        "detail": "OCR is required for scanned PDF pages: 1."
    }


def test_sync_and_async_reject_pdf_declared_as_text_plain_with_safe_shapes(
    app: FastAPI,
    client: TestClient,
    tmp_path: Path,
) -> None:
    override_storage(app, tmp_path)
    payload = PDF_FIXTURE_DIRECTORY.joinpath("text-page.pdf").read_bytes()

    synchronous = client.post(
        "/api/v1/analyze",
        files={"file": ("source.pdf", payload, "text/plain")},
    )
    asynchronous = client.post(
        "/api/v1/jobs",
        files={"file": ("source.pdf", payload, "text/plain")},
    )

    assert synchronous.status_code == 400
    assert isinstance(synchronous.json()["detail"], str)
    assert asynchronous.status_code == 415
    assert asynchronous.json()["detail"]["code"] == "unsupported_file_type"


@pytest.mark.parametrize(
    ("name", "payload", "mismatched_mime"),
    [
        ("source.docx", None, "application/pdf"),
        (
            "source.doc",
            bytes.fromhex("D0CF11E0A1B11AE1") + b"\x00" * 16,
            "application/pdf",
        ),
        ("source.pdf", None, "text/plain"),
        ("source.txt", b"text", "application/pdf"),
        ("source.rtf", br"{\rtf1 text}", "application/pdf"),
        ("source.md", b"# text", "application/pdf"),
        ("source.csv", b"name,value\ntext,1\n", "application/pdf"),
    ],
)
def test_all_seven_formats_have_sync_async_mime_mismatch_parity(
    app: FastAPI,
    client: TestClient,
    tmp_path: Path,
    name: str,
    payload: bytes | None,
    mismatched_mime: str,
) -> None:
    override_storage(app, tmp_path)
    if name.endswith(".docx"):
        stream = BytesIO()
        document = Document()
        document.add_paragraph("text")
        document.save(stream)
        payload = stream.getvalue()
    elif name.endswith(".pdf"):
        payload = PDF_FIXTURE_DIRECTORY.joinpath("text-page.pdf").read_bytes()
    assert payload is not None

    synchronous = client.post(
        "/api/v1/analyze",
        files={"file": (name, payload, mismatched_mime)},
    )
    asynchronous = client.post(
        "/api/v1/jobs",
        files={"file": (name, payload, mismatched_mime)},
    )

    assert synchronous.status_code == 400
    assert isinstance(synchronous.json()["detail"], str)
    assert asynchronous.status_code == 415
    assert asynchronous.json()["detail"]["code"] == "unsupported_file_type"


def test_mixed_pdf_reports_partial_ocr_requirement_in_compatibility_response(
    app: FastAPI,
    client: TestClient,
    tmp_path: Path,
) -> None:
    override_storage(app, tmp_path)

    response = client.post(
        "/api/v1/analyze",
        files={
            "file": (
                "mixed.pdf",
                PDF_FIXTURE_DIRECTORY.joinpath("mixed-page.pdf").read_bytes(),
                "application/pdf",
            )
        },
    )

    assert response.status_code == 200
    payload = response.json()
    assert payload["ocr_requirement"] == {"mode": "partial", "pages": [1]}
    assert payload["degradation"] == {
        "is_degraded": True,
        "reasons": ["ocr_required_pages"],
    }
    assert payload["pdf_metadata"]["pages"][0]["kind"] == "mixed"


def test_pdf_export_uses_the_same_canonical_visual_text_order_as_analysis(
    app: FastAPI,
    client: TestClient,
    tmp_path: Path,
) -> None:
    override_storage(app, tmp_path)
    source = PDF_FIXTURE_DIRECTORY.joinpath("layout-order.pdf").read_bytes()
    analyzed = client.post(
        "/api/v1/analyze",
        files={"file": ("styled.pdf", source, "application/pdf")},
    )

    assert analyzed.status_code == 200
    assert analyzed.json()["text"] == "Top\nAlphaBeta\nBottom"
    exported = client.post(
        "/api/v1/export-original",
        json={
            "file_id": analyzed.json()["file_id"],
            "filename": "styled.pdf",
            "modified_text": "Top\nAlphaBETA\nBottom",
            "track_changes": False,
        },
    )

    assert exported.status_code == 200
    reparsed = client.post(
        "/api/v1/analyze",
        files={"file": ("styled.pdf", exported.content, "application/pdf")},
    )
    assert reparsed.status_code == 200
    assert reparsed.json()["text"] == "Top\nAlphaBETA\nBottom"


def test_pdf_export_maps_repeated_text_by_canonical_visual_position(
    app: FastAPI,
    client: TestClient,
    monkeypatch: pytest.MonkeyPatch,
    tmp_path: Path,
) -> None:
    import pymupdf

    override_storage(app, tmp_path)
    source = PDF_FIXTURE_DIRECTORY.joinpath("duplicate-text.pdf").read_bytes()
    analyzed = client.post(
        "/api/v1/analyze",
        files={"file": ("duplicate.pdf", source, "application/pdf")},
    )
    assert analyzed.status_code == 200
    assert analyzed.json()["text"] == "Repeat\nRepeat"

    original_search = pymupdf.Page.search_for

    def reversed_search(page: pymupdf.Page, needle: str, *args: object, **kwargs: object):
        return list(reversed(original_search(page, needle, *args, **kwargs)))

    monkeypatch.setattr(pymupdf.Page, "search_for", reversed_search)
    exported = client.post(
        "/api/v1/export-original",
        json={
            "file_id": analyzed.json()["file_id"],
            "filename": "duplicate.pdf",
            "modified_text": "Repeat\nX",
            "track_changes": False,
        },
    )

    assert exported.status_code == 200
    reparsed = client.post(
        "/api/v1/analyze",
        files={"file": ("duplicate.pdf", exported.content, "application/pdf")},
    )
    assert reparsed.status_code == 200
    assert reparsed.json()["text"] == "Repeat\nX"


def test_pdf_export_changes_only_the_second_occurrence_within_one_span(
    app: FastAPI,
    client: TestClient,
    tmp_path: Path,
) -> None:
    override_storage(app, tmp_path)
    source = PDF_FIXTURE_DIRECTORY.joinpath("repeated-span.pdf").read_bytes()
    analyzed = client.post(
        "/api/v1/analyze",
        files={"file": ("repeat.pdf", source, "application/pdf")},
    )
    assert analyzed.status_code == 200
    assert analyzed.json()["text"] == "token token"

    exported = client.post(
        "/api/v1/export-original",
        json={
            "file_id": analyzed.json()["file_id"],
            "filename": "repeat.pdf",
            "replacements": [
                {
                    "original": "token",
                    "suggestion": "alter",
                    "position": 6,
                    "end_position": 11,
                }
            ],
            "track_changes": False,
        },
    )

    assert exported.status_code == 200
    reparsed = client.post(
        "/api/v1/analyze",
        files={"file": ("repeat.pdf", exported.content, "application/pdf")},
    )
    assert reparsed.status_code == 200
    assert reparsed.json()["text"] == "token alter"


def test_pdf_export_redacts_only_the_exact_glyph_in_a_proportional_span(
    app: FastAPI,
    client: TestClient,
    tmp_path: Path,
) -> None:
    import pymupdf

    override_storage(app, tmp_path)
    source = PDF_FIXTURE_DIRECTORY.joinpath("proportional-span.pdf").read_bytes()
    analyzed = client.post(
        "/api/v1/analyze",
        files={"file": ("proportional.pdf", source, "application/pdf")},
    )
    assert analyzed.status_code == 200
    assert analyzed.json()["text"] == "WWWWi"

    exported = client.post(
        "/api/v1/export-original",
        json={
            "file_id": analyzed.json()["file_id"],
            "filename": "proportional.pdf",
            "replacements": [
                {
                    "original": "i",
                    "suggestion": "X",
                    "position": 4,
                    "end_position": 5,
                }
            ],
            "track_changes": False,
        },
    )

    assert exported.status_code == 200
    with pymupdf.open(stream=exported.content, filetype="pdf") as output:
        assert output[0].get_text("text").strip() == "WWWWX"


def test_pdf_export_edits_table_cell_glyphs_without_duplicating_outside_text(
    app: FastAPI,
    client: TestClient,
    tmp_path: Path,
) -> None:
    override_storage(app, tmp_path)
    source = PDF_FIXTURE_DIRECTORY.joinpath("table-structure.pdf").read_bytes()
    analyzed = client.post(
        "/api/v1/analyze",
        files={"file": ("table.pdf", source, "application/pdf")},
    )
    assert analyzed.status_code == 200
    assert analyzed.json()["text"] == "A1\nA2\nB1\nB1"

    exported = client.post(
        "/api/v1/export-original",
        json={
            "file_id": analyzed.json()["file_id"],
            "filename": "table.pdf",
            "replacements": [
                {
                    "original": "A1",
                    "suggestion": "X1",
                    "position": 0,
                    "end_position": 2,
                }
            ],
            "track_changes": False,
        },
    )

    assert exported.status_code == 200
    reparsed = client.post(
        "/api/v1/analyze",
        files={"file": ("table.pdf", exported.content, "application/pdf")},
    )
    assert reparsed.status_code == 200
    assert reparsed.json()["text"] == "X1\nA2\nB1\nB1"
    assert reparsed.json()["text"].count("B1") == 2


def test_pdf_export_replaces_a_range_crossing_styled_source_spans(
    app: FastAPI,
    client: TestClient,
    tmp_path: Path,
) -> None:
    override_storage(app, tmp_path)
    source = PDF_FIXTURE_DIRECTORY.joinpath("layout-order.pdf").read_bytes()
    analyzed = client.post(
        "/api/v1/analyze",
        files={"file": ("styled.pdf", source, "application/pdf")},
    )
    assert analyzed.status_code == 200

    exported = client.post(
        "/api/v1/export-original",
        json={
            "file_id": analyzed.json()["file_id"],
            "filename": "styled.pdf",
            "modified_text": "Top\nMerged\nBottom",
            "track_changes": False,
        },
    )

    assert exported.status_code == 200
    reparsed = client.post(
        "/api/v1/analyze",
        files={"file": ("styled.pdf", exported.content, "application/pdf")},
    )
    assert reparsed.status_code == 200
    assert reparsed.json()["text"] == "Top\nMerged\nBottom"


def test_pdf_track_changes_annotates_every_mapped_styled_piece(
    app: FastAPI,
    client: TestClient,
    tmp_path: Path,
) -> None:
    import pymupdf

    override_storage(app, tmp_path)
    source = PDF_FIXTURE_DIRECTORY.joinpath("styled-space.pdf").read_bytes()
    analyzed = client.post(
        "/api/v1/analyze",
        files={"file": ("styled-space.pdf", source, "application/pdf")},
    )
    assert analyzed.status_code == 200
    assert analyzed.json()["text"] == "Alpha Beta"

    exported = client.post(
        "/api/v1/export-original",
        json={
            "file_id": analyzed.json()["file_id"],
            "filename": "styled-space.pdf",
            "replacements": [
                {
                    "original": "Alpha Beta",
                    "suggestion": "Merged",
                    "position": 0,
                    "end_position": 10,
                }
            ],
            "track_changes": True,
        },
    )

    assert exported.status_code == 200
    with pymupdf.open(stream=exported.content, filetype="pdf") as output:
        annotations = list(output[0].annots() or ())
        assert len(annotations) == 2
        assert all(
            annotation.info["content"] == "原文: Alpha Beta\n建议: Merged"
            for annotation in annotations
        )


def test_scenarios_and_formats_are_discoverable(client: TestClient) -> None:
    scenarios_response = client.get("/api/v1/scenarios")
    formats_response = client.get("/api/v1/formats")

    assert scenarios_response.status_code == 200
    assert {item["id"] for item in scenarios_response.json()["scenarios"]} == {
        "general",
        "academic",
        "business",
        "legal",
        "news",
        "technical",
    }
    assert [item["ext"] for item in formats_response.json()["formats"]] == [
        "txt",
        "docx",
        "doc",
        "pdf",
        "rtf",
        "md",
        "csv",
    ]
    assert formats_response.json()["formats"][0] == {
        "ext": "txt",
        "name": "纯文本文件",
        "accept": ".txt",
    }
