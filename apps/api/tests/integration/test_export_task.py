from __future__ import annotations

import logging
import os
from concurrent.futures import ThreadPoolExecutor
from datetime import UTC, datetime, timedelta
from io import BytesIO
from pathlib import Path
from threading import Event, Lock
from uuid import UUID, uuid4

import pytest
from docx import Document as WordDocument
from fastapi import FastAPI
from sqlalchemy import update
from sqlalchemy.orm import Session, sessionmaker

from text_verification.api.dependencies import get_db_session, get_job_storage
from text_verification.checkers.models import CheckCategory, CheckerFailure
from text_verification.domain.derived_content import derive_document
from text_verification.domain.documents import DocumentModel, FileType
from text_verification.domain.exports import (
    ExportCheckerFailureSnapshot,
    ExportIssueSummarySnapshot,
    ExportSnapshot,
    ExportStatus,
    ExportType,
)
from text_verification.domain.issues import (
    DecisionAction,
    DecisionCommand,
    Issue,
    IssueSeverity,
)
from text_verification.domain.jobs import JobStatus
from text_verification.exporters import ExportError, ReplacementPlanner, TxtExporter
from text_verification.infrastructure.analysis_repositories import AnalysisRepository
from text_verification.infrastructure.decision_repository import DecisionRepository
from text_verification.infrastructure.export_repository import ExportRepository
from text_verification.infrastructure.orm import ExportRow, IssueDecisionRow, JobRow
from text_verification.infrastructure.repositories import JobRepository
from text_verification.infrastructure.storage import JobStorage
from text_verification.parsers import DocxParser, TxtParser


@pytest.fixture
def celery_eager() -> None:
    from text_verification.workers.celery_app import celery_app

    original_values = {
        "task_always_eager": celery_app.conf.task_always_eager,
        "task_store_eager_result": celery_app.conf.task_store_eager_result,
        "task_eager_propagates": celery_app.conf.task_eager_propagates,
    }
    celery_app.conf.update(
        task_always_eager=True,
        task_store_eager_result=False,
        task_eager_propagates=False,
    )
    try:
        yield
    finally:
        celery_app.conf.update(**original_values)


@pytest.fixture
def export_storage(tmp_path: Path) -> JobStorage:
    root = tmp_path / "export-task-jobs"
    root.mkdir()
    return JobStorage(root, max_upload_bytes=25 * 1024 * 1024)


@pytest.fixture(autouse=True)
def export_task_dependencies(
    app: FastAPI,
    db_session_factory: sessionmaker[Session],
    export_storage: JobStorage,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    from text_verification.workers import export_tasks

    def fresh_session():
        session = db_session_factory()
        try:
            yield session
        finally:
            session.close()

    app.dependency_overrides[get_db_session] = fresh_session
    app.dependency_overrides[get_job_storage] = lambda: export_storage
    monkeypatch.setattr(
        export_tasks,
        "SESSION_FACTORY_PROVIDER",
        lambda: db_session_factory,
    )
    monkeypatch.setattr(export_tasks, "STORAGE_FACTORY", lambda: export_storage)


def test_txt_export_task_completes_round_trips_and_downloads(
    client,
    celery_eager,
    db_session: Session,
    db_session_factory: sessionmaker[Session],
    export_storage: JobStorage,
) -> None:
    job_id, issue = _seed_reviewed_txt_job(db_session, export_storage)
    other_job_id = _seed_empty_job(
        db_session,
        file_type=FileType.TXT,
        status=JobStatus.COMPLETED,
    )
    db_session.execute(
        update(IssueDecisionRow)
        .where(IssueDecisionRow.issue_id == issue.issue_id)
        .values(job_id=other_job_id)
    )
    db_session.commit()

    created = client.post(
        f"/api/v1/jobs/{job_id}/exports",
        json={"type": "modified_document"},
    )
    export_id = UUID(created.json()["export_id"])
    status_response = client.get(f"/api/v1/jobs/{job_id}/exports/{export_id}")
    download = client.get(f"/api/v1/jobs/{job_id}/exports/{export_id}/download")

    assert created.status_code == 202
    assert status_response.json()["status"] == "completed"
    assert download.status_code == 200
    assert download.content.decode("utf-8") == "修改后的正文\n"
    stored_path = export_storage.export_path(job_id, export_id, "txt")
    reparsed = TxtParser().parse(
        stored_path,
        document_id=uuid4(),
        source_name="modified_document.txt",
    )
    assert [block.text for block in reparsed.blocks] == ["修改后的正文"]
    stored = _load_export(db_session_factory, export_id)
    assert stored.status == ExportStatus.COMPLETED
    assert stored.warnings == []


def test_docx_export_task_round_trips_safe_changes_and_preserves_skipped_warning(
    client,
    celery_eager,
    db_session: Session,
    db_session_factory: sessionmaker[Session],
    export_storage: JobStorage,
) -> None:
    job_id, unsafe_issue = _seed_reviewed_docx_job(db_session, export_storage)

    created = client.post(
        f"/api/v1/jobs/{job_id}/exports",
        json={"type": "modified_document", "confirm_warnings": True},
    )
    export_id = UUID(created.json()["export_id"])
    download = client.get(f"/api/v1/jobs/{job_id}/exports/{export_id}/download")

    assert created.status_code == 202
    assert download.status_code == 200
    exported_path = export_storage.export_path(job_id, export_id, "docx")
    reparsed = DocxParser().parse(
        exported_path,
        document_id=uuid4(),
        source_name="modified_document.docx",
    )
    assert [block.text for block in reparsed.blocks] == ["核验示例正文"]
    stored = _load_export(db_session_factory, export_id)
    assert stored.status == ExportStatus.COMPLETED
    assert [warning.model_dump(mode="json") for warning in stored.warnings] == [
        {
            "code": "unsafe_docx_run_boundary",
            "message": (
                "修改范围跨越多个 DOCX 文本运行，为保留格式已跳过；"
                "请在原文中手动修改后重新导出。"
            ),
            "issue_id": str(unsafe_issue.issue_id),
            "block_id": unsafe_issue.block_id,
        }
    ]


def test_docx_html_report_uses_same_preflight_applicability_warnings(
    client,
    celery_eager,
    db_session: Session,
    db_session_factory: sessionmaker[Session],
    export_storage: JobStorage,
) -> None:
    job_id, unsafe_issue = _seed_reviewed_docx_job(db_session, export_storage)

    created = client.post(
        f"/api/v1/jobs/{job_id}/exports",
        json={"type": "html_report"},
    )
    export_id = UUID(created.json()["export_id"])
    download = client.get(f"/api/v1/jobs/{job_id}/exports/{export_id}/download")

    assert created.status_code == 202
    assert download.status_code == 200
    html = download.content.decode("utf-8")
    assert "unsafe_docx_run_boundary" in html
    assert str(unsafe_issue.issue_id) in html
    stored = _load_export(db_session_factory, export_id)
    assert [warning.code for warning in stored.warnings] == [
        "unsafe_docx_run_boundary"
    ]


def test_html_report_task_includes_checker_failures_for_reviewed_pdf_job(
    client,
    celery_eager,
    db_session: Session,
    db_session_factory: sessionmaker[Session],
    export_storage: JobStorage,
) -> None:
    issue = _build_issue(
        document_id=uuid4(),
        block_id="p-000001",
        original="原始",
        suggestion=None,
        start=0,
        end=2,
        page=1,
    )
    job_id = _seed_analyzed_job(
        db_session,
        file_type=FileType.PDF,
        document=_simple_document(
            document_id=issue.document_id,
            file_type=FileType.PDF,
            text="原始正文",
        ),
        issues=[issue],
        failures={
            CheckCategory.SECURITY: CheckerFailure(
                code="checker_failed",
                message="安全分类检查失败。",
            )
        },
    )
    _apply_decision(
        db_session,
        job_id,
        issue,
        DecisionAction.ACCEPTED,
        replacement="修订",
    )

    created = client.post(
        f"/api/v1/jobs/{job_id}/exports",
        json={"type": "html_report"},
    )
    export_id = UUID(created.json()["export_id"])
    download = client.get(f"/api/v1/jobs/{job_id}/exports/{export_id}/download")

    assert created.status_code == 202
    assert download.status_code == 200
    html = download.content.decode("utf-8")
    assert "checker_failed" in html
    assert "安全分类检查失败。" in html
    stored = _load_export(db_session_factory, export_id)
    assert stored.warnings == []


@pytest.mark.skipif(os.name == "nt", reason="WeasyPrint native runtime is provided by Docker")
def test_pdf_report_task_generates_pdf_for_pdf_job(
    client,
    celery_eager,
    db_session: Session,
    export_storage: JobStorage,
) -> None:
    job_id = _seed_analyzed_job(
        db_session,
        file_type=FileType.PDF,
        document=_simple_document(
            document_id=uuid4(),
            file_type=FileType.PDF,
            text="报告正文",
        ),
        issues=[],
        failures={},
    )

    created = client.post(
        f"/api/v1/jobs/{job_id}/exports",
        json={"type": "pdf_report"},
    )
    export_id = UUID(created.json()["export_id"])
    download = client.get(f"/api/v1/jobs/{job_id}/exports/{export_id}/download")

    assert created.status_code == 202
    assert download.status_code == 200
    assert download.content.startswith(b"%PDF")
    assert export_storage.export_path(job_id, export_id, "pdf").is_file()


def test_transient_task_failure_retries_and_completes_without_terminal_regression(
    celery_eager,
    db_session: Session,
    db_session_factory: sessionmaker[Session],
    export_storage: JobStorage,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    from text_verification.workers import export_tasks

    job_id, _issue = _seed_reviewed_txt_job(db_session, export_storage)
    export = ExportRepository(db_session).create(
        job_id,
        export_tasks.ExportType.MODIFIED_DOCUMENT,
        "txt",
        snapshot=_snapshot_for_job(db_session, job_id),
    )
    db_session.commit()
    real_export = TxtExporter.export
    attempts = 0

    def flaky_export(self, document, plan, target):
        nonlocal attempts
        attempts += 1
        if attempts == 1:
            raise RuntimeError("transient write failure")
        return real_export(self, document, plan, target)

    monkeypatch.setattr(TxtExporter, "export", flaky_export)

    result = export_tasks.process_export.delay(str(export.export_id))

    assert result.successful()
    assert attempts == 2
    stored = _load_export(db_session_factory, export.export_id)
    assert stored.status == ExportStatus.COMPLETED


def test_export_uses_decision_snapshot_taken_before_later_decision_change(
    client,
    celery_eager,
    db_session: Session,
    db_session_factory: sessionmaker[Session],
    export_storage: JobStorage,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    from text_verification.api.routes import exports as export_routes
    from text_verification.workers import export_tasks

    job_id, issue = _seed_reviewed_txt_job(db_session, export_storage)
    dispatched: list[str] = []
    monkeypatch.setattr(
        export_routes,
        "dispatch_process_export",
        lambda export_id: dispatched.append(export_id),
    )

    created = client.post(
        f"/api/v1/jobs/{job_id}/exports",
        json={"type": "modified_document"},
    )
    export_id = UUID(created.json()["export_id"])
    outcome = DecisionRepository(db_session).apply(
        job_id,
        DecisionCommand(
            issue_id=issue.issue_id,
            issue_version=1,
            expected_revision=1,
            action=DecisionAction.ACCEPTED,
            replacement="后改的",
        ),
    )
    assert outcome.decision is not None
    db_session.commit()

    result = export_tasks.process_export.delay(str(export_id))

    assert dispatched == [str(export_id)]
    assert result.successful()
    assert export_storage.export_path(job_id, export_id, "txt").read_text(
        encoding="utf-8"
    ) == "修改后的正文\n"
    assert _load_export(db_session_factory, export_id).status == ExportStatus.COMPLETED


def test_export_uses_analysis_snapshot_taken_before_reanalysis(
    client,
    celery_eager,
    db_session: Session,
    export_storage: JobStorage,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    from text_verification.api.routes import exports as export_routes
    from text_verification.workers import export_tasks

    old_issue = _build_issue(
        document_id=uuid4(),
        block_id="p-000001",
        original="旧问题",
        suggestion="旧建议",
        start=0,
        end=3,
    )
    job_id = _seed_analyzed_job(
        db_session,
        file_type=FileType.TXT,
        document=_simple_document(
            document_id=old_issue.document_id,
            file_type=FileType.TXT,
            text="旧问题正文",
        ),
        issues=[old_issue],
        failures={},
    )
    dispatched: list[str] = []
    monkeypatch.setattr(
        export_routes,
        "dispatch_process_export",
        lambda export_id: dispatched.append(export_id),
    )
    created = client.post(
        f"/api/v1/jobs/{job_id}/exports",
        json={"type": "html_report"},
    )
    export_id = UUID(created.json()["export_id"])
    new_issue = _build_issue(
        document_id=uuid4(),
        block_id="p-000001",
        original="新问题",
        suggestion="新建议",
        start=0,
        end=3,
    )
    new_document = _simple_document(
        document_id=new_issue.document_id,
        file_type=FileType.TXT,
        text="新问题正文",
    ).model_copy(update={"version": 2})
    AnalysisRepository(db_session).replace_analysis(
        job_id,
        new_document,
        [new_issue.model_copy(update={"document_version": 2})],
        {},
    )
    db_session.commit()

    result = export_tasks.process_export.delay(str(export_id))
    html = export_storage.export_path(job_id, export_id, "html").read_text(
        encoding="utf-8"
    )

    assert dispatched == [str(export_id)]
    assert result.successful()
    assert "旧问题" in html
    assert "新问题" not in html
    assert "发现问题：1" in html


def test_docx_export_rejects_source_bytes_changed_after_snapshot(
    client,
    celery_eager,
    db_session: Session,
    db_session_factory: sessionmaker[Session],
    export_storage: JobStorage,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    from text_verification.api.routes import exports as export_routes
    from text_verification.workers import export_tasks

    job_id, _issue = _seed_reviewed_docx_job(db_session, export_storage)
    dispatched: list[str] = []
    monkeypatch.setattr(
        export_routes,
        "dispatch_process_export",
        lambda export_id: dispatched.append(export_id),
    )
    created = client.post(
        f"/api/v1/jobs/{job_id}/exports",
        json={"type": "modified_document", "confirm_warnings": True},
    )
    export_id = UUID(created.json()["export_id"])
    export_storage.source_path(job_id, FileType.DOCX).write_bytes(b"changed source")

    result = export_tasks.process_export.delay(str(export_id))
    stored = _load_export(db_session_factory, export_id)

    assert dispatched == [str(export_id)]
    assert result.failed()
    assert stored.status == ExportStatus.FAILED
    assert stored.error_code == "export_source_changed"
    assert stored.error_message == "导出源文件校验失败，请重新创建。"
    assert [warning.code for warning in stored.warnings] == [
        "unsafe_docx_run_boundary"
    ]
    assert not export_storage.export_path(job_id, export_id, "docx").exists()


def test_publish_timeout_cannot_overwrite_worker_completed_terminal_state(
    client,
    celery_eager,
    db_session: Session,
    db_session_factory: sessionmaker[Session],
    export_storage: JobStorage,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    from text_verification.api.routes import exports as export_routes
    from text_verification.workers import export_tasks

    job_id = _seed_analyzed_job(
        db_session,
        file_type=FileType.TXT,
        document=_simple_document(
            document_id=uuid4(),
            file_type=FileType.TXT,
            text="报告正文",
        ),
        issues=[],
        failures={},
    )

    def complete_then_timeout(export_id: str) -> None:
        result = export_tasks.process_export.delay(export_id)
        assert result.successful()
        raise TimeoutError("publish acknowledgement timed out")

    monkeypatch.setattr(
        export_routes,
        "dispatch_process_export",
        complete_then_timeout,
    )

    created = client.post(
        f"/api/v1/jobs/{job_id}/exports",
        json={"type": "html_report"},
    )
    export_id = UUID(created.json()["export_id"])
    stored = _load_export(db_session_factory, export_id)

    assert created.status_code == 202
    assert created.json()["status"] == "queued"
    assert created.json()["dispatch_status"] == "deferred"
    assert stored.status == ExportStatus.COMPLETED
    assert stored.error_code is None
    assert export_storage.export_path(job_id, export_id, "html").is_file()


def test_repeat_delivery_keeps_completed_export_terminal(
    celery_eager,
    db_session: Session,
    db_session_factory: sessionmaker[Session],
    export_storage: JobStorage,
) -> None:
    from text_verification.workers import export_tasks

    job_id, _issue = _seed_reviewed_txt_job(db_session, export_storage)
    export = ExportRepository(db_session).create(
        job_id,
        export_tasks.ExportType.MODIFIED_DOCUMENT,
        "txt",
        snapshot=_snapshot_for_job(db_session, job_id),
    )
    db_session.commit()

    first_result = export_tasks.process_export.delay(str(export.export_id))
    first_stored = _load_export(db_session_factory, export.export_id)
    second_result = export_tasks.process_export.delay(str(export.export_id))
    second_stored = _load_export(db_session_factory, export.export_id)

    assert export_tasks.process_export.name == "text_verification.process_export"
    assert first_result.successful()
    assert second_result.successful()
    assert first_stored.status == ExportStatus.COMPLETED
    assert second_stored.status == ExportStatus.COMPLETED
    assert second_stored.updated_at == first_stored.updated_at
    assert export_storage.export_path(job_id, export.export_id, "txt").read_text(
        encoding="utf-8"
    ) == "修改后的正文\n"


def test_concurrent_deliveries_claim_export_once(
    celery_eager,
    db_session: Session,
    db_session_factory: sessionmaker[Session],
    export_storage: JobStorage,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    from text_verification.workers import export_tasks

    job_id, _issue = _seed_reviewed_txt_job(db_session, export_storage)
    export = ExportRepository(db_session).create(
        job_id,
        export_tasks.ExportType.MODIFIED_DOCUMENT,
        "txt",
        snapshot=_snapshot_for_job(db_session, job_id),
    )
    db_session.commit()
    render_started = Event()
    allow_render_to_finish = Event()
    counter_lock = Lock()
    render_count = 0
    finalizer_count = 0
    real_export = TxtExporter.export
    real_mark_completed = ExportRepository.mark_completed

    def blocking_export(self, document, plan, target):
        nonlocal render_count
        with counter_lock:
            render_count += 1
        render_started.set()
        if not allow_render_to_finish.wait(timeout=5):
            raise TimeoutError("timed out waiting to finish export rendering")
        return real_export(self, document, plan, target)

    def counting_mark_completed(self, export_id, *, warnings):
        nonlocal finalizer_count
        with counter_lock:
            finalizer_count += 1
        return real_mark_completed(self, export_id, warnings=warnings)

    monkeypatch.setattr(TxtExporter, "export", blocking_export)
    monkeypatch.setattr(ExportRepository, "mark_completed", counting_mark_completed)

    with ThreadPoolExecutor(max_workers=2) as executor:
        first_delivery = executor.submit(
            export_tasks.process_export.delay,
            str(export.export_id),
        )
        assert render_started.wait(timeout=3)
        assert _load_export(db_session_factory, export.export_id).status == ExportStatus.PROCESSING

        duplicate_delivery = executor.submit(
            export_tasks.process_export.delay,
            str(export.export_id),
        )
        try:
            duplicate_result = duplicate_delivery.result(timeout=2)
        finally:
            allow_render_to_finish.set()
        first_result = first_delivery.result(timeout=5)

    stored = _load_export(db_session_factory, export.export_id)
    assert first_result.successful()
    assert duplicate_result.successful()
    assert render_count == 1
    assert finalizer_count == 1
    assert stored.status == ExportStatus.COMPLETED
    assert stored.error_code is None
    assert stored.error_message is None


def test_exhausted_task_failure_is_safe_and_repeat_delivery_stays_failed(
    celery_eager,
    db_session: Session,
    db_session_factory: sessionmaker[Session],
    export_storage: JobStorage,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    from text_verification.workers import export_tasks

    job_id, _issue = _seed_reviewed_txt_job(db_session, export_storage)
    export = ExportRepository(db_session).create(
        job_id,
        export_tasks.ExportType.MODIFIED_DOCUMENT,
        "txt",
        snapshot=_snapshot_for_job(db_session, job_id),
    )
    db_session.commit()
    attempts = 0

    def failing_export(self, document, plan, target):
        nonlocal attempts
        attempts += 1
        raise RuntimeError(r"internal failure at C:\secret\jobs\artifact.txt")

    monkeypatch.setattr(TxtExporter, "export", failing_export)

    first_result = export_tasks.process_export.delay(str(export.export_id))
    first_stored = _load_export(db_session_factory, export.export_id)
    second_result = export_tasks.process_export.delay(str(export.export_id))
    second_stored = _load_export(db_session_factory, export.export_id)

    assert first_result.failed()
    assert second_result.failed()
    assert attempts == export_tasks.PROCESS_EXPORT_MAX_RETRIES + 1
    assert first_stored.status == ExportStatus.FAILED
    assert first_stored.error_code == "export_failed"
    assert first_stored.error_message == "导出失败，请稍后重试。"
    assert "secret" not in first_stored.error_message
    assert second_stored.status == ExportStatus.FAILED
    assert second_stored.updated_at == first_stored.updated_at
    assert not export_storage.export_path(job_id, export.export_id, "txt").exists()


def test_expected_export_error_persists_public_code_and_task_result_is_failed(
    celery_eager,
    db_session: Session,
    db_session_factory: sessionmaker[Session],
    export_storage: JobStorage,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    from text_verification.workers import export_tasks

    job_id, _issue = _seed_reviewed_txt_job(db_session, export_storage)
    export = ExportRepository(db_session).create(
        job_id,
        export_tasks.ExportType.MODIFIED_DOCUMENT,
        "txt",
        snapshot=_snapshot_for_job(db_session, job_id),
    )
    db_session.commit()

    def failing_export(self, document, plan, target):
        raise ExportError("txt_export_failed", "无法导出 TXT 文件。")

    monkeypatch.setattr(TxtExporter, "export", failing_export)

    result = export_tasks.process_export.delay(str(export.export_id))

    assert result.failed()
    stored = _load_export(db_session_factory, export.export_id)
    assert stored.status == ExportStatus.FAILED
    assert stored.error_code == "txt_export_failed"
    assert stored.error_message == "无法导出 TXT 文件。"


def test_recovery_redispatches_stale_queued_and_processing_but_not_fresh_or_terminal(
    celery_eager,
    db_session: Session,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    from text_verification.workers import export_tasks
    from text_verification.workers.celery_app import celery_app

    job_id = _seed_empty_job(
        db_session,
        file_type=FileType.TXT,
        status=JobStatus.COMPLETED,
    )
    repository = ExportRepository(db_session)
    snapshot = _snapshot_for_job(db_session, job_id)
    stale = repository.create(
        job_id, ExportType.HTML_REPORT, "html", snapshot=snapshot
    )
    fresh = repository.create(
        job_id, ExportType.HTML_REPORT, "html", snapshot=snapshot
    )
    stale_processing = repository.create(
        job_id, ExportType.HTML_REPORT, "html", snapshot=snapshot
    )
    fresh_processing = repository.create(
        job_id, ExportType.HTML_REPORT, "html", snapshot=snapshot
    )
    completed = repository.create(
        job_id, ExportType.HTML_REPORT, "html", snapshot=snapshot
    )
    failed = repository.create(
        job_id, ExportType.HTML_REPORT, "html", snapshot=snapshot
    )
    repository.mark_processing(stale_processing.export_id)
    repository.mark_processing(fresh_processing.export_id)
    repository.mark_completed(completed.export_id, warnings=[])
    repository.mark_failed(
        failed.export_id,
        error_code="export_failed",
        error_message="导出失败，请稍后重试。",
    )
    now = datetime.now(UTC)
    stale_at = now - timedelta(seconds=61)
    stale_processing_at = now - timedelta(
        seconds=int(celery_app.conf.task_time_limit) + 61
    )
    fresh_processing_at = now - timedelta(
        seconds=int(celery_app.conf.task_time_limit) - 1
    )
    db_session.execute(
        update(ExportRow)
        .where(
            ExportRow.export_id.in_(
                [stale.export_id, completed.export_id, failed.export_id]
            )
        )
        .values(created_at=stale_at, updated_at=stale_at)
    )
    db_session.execute(
        update(ExportRow)
        .where(ExportRow.export_id == stale_processing.export_id)
        .values(updated_at=stale_processing_at)
    )
    db_session.execute(
        update(ExportRow)
        .where(ExportRow.export_id == fresh_processing.export_id)
        .values(updated_at=fresh_processing_at)
    )
    db_session.commit()
    dispatched: list[str] = []
    monkeypatch.setattr(
        export_tasks,
        "dispatch_recovered_export",
        lambda export_id: dispatched.append(export_id),
        raising=False,
    )

    result = export_tasks.recover_stale_queued_exports.delay()

    assert export_tasks.recover_stale_queued_exports.name == (
        "text_verification.recover_stale_queued_exports"
    )
    assert result.successful()
    assert set(result.result) == {str(stale.export_id), str(stale_processing.export_id)}
    assert set(dispatched) == {str(stale.export_id), str(stale_processing.export_id)}
    assert str(fresh.export_id) not in dispatched
    assert str(fresh_processing.export_id) not in dispatched
    assert str(completed.export_id) not in dispatched
    assert str(failed.export_id) not in dispatched


def test_stale_queue_recovery_dispatch_failure_stays_queued_and_task_fails(
    celery_eager,
    db_session: Session,
    db_session_factory: sessionmaker[Session],
    monkeypatch: pytest.MonkeyPatch,
    caplog: pytest.LogCaptureFixture,
) -> None:
    from text_verification.workers import export_tasks

    job_id = _seed_empty_job(
        db_session,
        file_type=FileType.TXT,
        status=JobStatus.COMPLETED,
    )
    export = ExportRepository(db_session).create(
        job_id,
        ExportType.HTML_REPORT,
        "html",
        snapshot=_snapshot_for_job(db_session, job_id),
    )
    stale_at = datetime.now(UTC) - timedelta(seconds=61)
    db_session.execute(
        update(ExportRow)
        .where(ExportRow.export_id == export.export_id)
        .values(created_at=stale_at, updated_at=stale_at)
    )
    db_session.commit()

    def fail_dispatch(_export_id: str) -> None:
        raise ConnectionError("broker unavailable at internal endpoint")

    monkeypatch.setattr(
        export_tasks,
        "dispatch_recovered_export",
        fail_dispatch,
        raising=False,
    )
    caplog.set_level(logging.ERROR, logger=export_tasks.__name__)

    result = export_tasks.recover_stale_queued_exports.delay()

    stored = _load_export(db_session_factory, export.export_id)
    assert result.failed()
    assert type(result.result).__name__ == "QueuedExportRecoveryError"
    assert "internal endpoint" not in str(result.result)
    assert stored.status == ExportStatus.QUEUED
    assert "stale_queued_export_dispatch_failed" in caplog.messages


def test_stale_queue_recovery_is_scheduled_every_minute() -> None:
    from text_verification.workers.celery_app import celery_app

    assert celery_app.conf.task_reject_on_worker_lost is True
    assert celery_app.conf.beat_schedule["recover-stale-queued-exports-every-minute"] == {
        "task": "text_verification.recover_stale_queued_exports",
        "schedule": 60.0,
    }


def test_processing_recovery_threshold_is_strictly_beyond_task_hard_limit() -> None:
    from text_verification.workers import export_tasks
    from text_verification.workers.celery_app import celery_app

    assert (
        export_tasks.PROCESSING_EXPORT_RECOVERY_AGE_SECONDS
        > celery_app.conf.task_time_limit
    )


def _seed_reviewed_txt_job(
    session: Session,
    storage: JobStorage,
) -> tuple[UUID, Issue]:
    job_id = _seed_empty_job(
        session,
        file_type=FileType.TXT,
        status=JobStatus.COMPLETED,
    )
    stored = storage.save_bytes(job_id, "sample.txt", "原始正文".encode())
    document = TxtParser().parse(
        stored.path,
        document_id=uuid4(),
        source_name="sample.txt",
    )
    issue = _build_issue(
        document_id=document.document_id,
        block_id=document.blocks[0].block_id,
        original="原始",
        suggestion="修改后的",
        start=0,
        end=2,
    )
    AnalysisRepository(session).replace_analysis(job_id, document, [issue], {})
    session.commit()
    _apply_decision(session, job_id, issue, DecisionAction.ACCEPTED)
    return job_id, issue


def _seed_reviewed_docx_job(
    session: Session,
    storage: JobStorage,
) -> tuple[UUID, Issue]:
    job_id = _seed_empty_job(
        session,
        file_type=FileType.DOCX,
        status=JobStatus.COMPLETED,
    )
    source = WordDocument()
    paragraph = source.add_paragraph()
    paragraph.add_run("核验")
    paragraph.add_run("示例")
    paragraph.add_run("文本")
    payload = BytesIO()
    source.save(payload)
    stored = storage.save_bytes(job_id, "sample.docx", payload.getvalue())
    document = DocxParser().parse(
        stored.path,
        document_id=uuid4(),
        source_name="sample.docx",
    )
    unsafe_issue = _build_issue(
        document_id=document.document_id,
        block_id=document.blocks[0].block_id,
        original="验示",
        suggestion="审查",
        start=1,
        end=3,
    )
    safe_issue = _build_issue(
        document_id=document.document_id,
        block_id=document.blocks[0].block_id,
        original="文本",
        suggestion="正文",
        start=4,
        end=6,
    )
    AnalysisRepository(session).replace_analysis(
        job_id,
        document,
        [unsafe_issue, safe_issue],
        {},
    )
    session.commit()
    _apply_decision(session, job_id, unsafe_issue, DecisionAction.ACCEPTED)
    _apply_decision(session, job_id, safe_issue, DecisionAction.ACCEPTED)
    return job_id, unsafe_issue


def _seed_analyzed_job(
    session: Session,
    *,
    file_type: FileType,
    document: DocumentModel,
    issues: list[Issue],
    failures: dict[CheckCategory, CheckerFailure],
) -> UUID:
    job_id = _seed_empty_job(
        session,
        file_type=file_type,
        status=JobStatus.COMPLETED if not failures else JobStatus.PARTIAL,
    )
    AnalysisRepository(session).replace_analysis(job_id, document, issues, failures)
    session.commit()
    return job_id


def _seed_empty_job(
    session: Session,
    *,
    file_type: FileType,
    status: JobStatus,
) -> UUID:
    now = datetime.now(UTC)
    job_id = uuid4()
    repository = JobRepository(session)
    repository.create_job(
        job_id=job_id,
        source_name=f"sample.{file_type.value}",
        file_type=file_type.value,
        size_bytes=16,
        storage_key=str(job_id),
        created_at=now,
        expires_at=now + timedelta(hours=1),
    )
    if status != JobStatus.QUEUED:
        repository.transition(
            job_id,
            status,
            100,
            "处理完成",
        )
    repository.commit()
    return job_id


def _apply_decision(
    session: Session,
    job_id: UUID,
    issue: Issue,
    action: DecisionAction,
    *,
    replacement: str | None = None,
) -> None:
    resolved_replacement = replacement
    if action == DecisionAction.ACCEPTED and resolved_replacement is None:
        resolved_replacement = issue.suggestion
    outcome = DecisionRepository(session).apply(
        job_id,
        DecisionCommand(
            issue_id=issue.issue_id,
            issue_version=1,
            expected_revision=0,
            action=action,
            replacement=resolved_replacement,
        ),
    )
    assert outcome.decision is not None
    session.commit()


def _simple_document(
    *,
    document_id: UUID,
    file_type: FileType,
    text: str,
) -> DocumentModel:
    return DocumentModel(
        document_id=document_id,
        file_type=file_type,
        source_name=f"sample.{file_type.value}",
        version=1,
        blocks=[
            {
                "block_id": "p-000001",
                "kind": "paragraph",
                "text": text,
                "page": 1 if file_type == FileType.PDF else None,
                "paragraph_index": 0,
                "parent_id": None,
                "style": {},
                "source_locator": {"paragraph_index": 0},
            }
        ],
        metadata={},
    )


def _build_issue(
    *,
    document_id: UUID,
    block_id: str,
    original: str,
    suggestion: str | None,
    start: int,
    end: int,
    page: int | None = None,
) -> Issue:
    return Issue(
        issue_id=uuid4(),
        document_id=document_id,
        block_id=block_id,
        page=page,
        start=start,
        end=end,
        original=original,
        suggestion=suggestion,
        alternatives=[] if suggestion is None else [suggestion],
        type="literal",
        severity=IssueSeverity.WARNING,
        layer=CheckCategory.CHARACTER.value,
        message="命中规则。",
        rule_id="character-001",
        source="test",
        source_version="1",
        confidence=1.0,
        auto_fixable=suggestion is not None,
        context=original,
    )


def _snapshot_for_job(session: Session, job_id: UUID) -> ExportSnapshot:
    job = JobRepository(session).get_job(job_id)
    assert job is not None
    assert job.file_type != FileType.DOCX
    repository = AnalysisRepository(session)
    document = repository.get_document(job_id)
    if document is None:
        document = DocumentModel(
            document_id=uuid4(),
            file_type=job.file_type,
            source_name=job.source_name,
            version=1,
            blocks=[],
            metadata={},
        )
        issues: list[Issue] = []
        failures: dict[CheckCategory, CheckerFailure] = {}
        summary_total = 0
        summary_by_category: dict[CheckCategory, int] = {}
        summary_by_severity: dict[IssueSeverity, int] = {}
        summary_by_decision: dict[str, int] = {}
    else:
        issues = repository.list_all_issues(job_id)
        failures = repository.get_checker_failures(job_id)
        summary = repository.summarize_issues(job_id)
        summary_total = summary.total
        summary_by_category = summary.by_category
        summary_by_severity = summary.by_severity
        summary_by_decision = summary.by_decision
    version_id = session.get(JobRow, job_id).active_version_id
    snapshot_kwargs: dict[str, object]
    if version_id is None:
        warnings = ReplacementPlanner().build(document, issues).warnings
        snapshot_kwargs = {"schema_version": 1}
    else:
        derived = derive_document(version_id, document, issues)
        warnings = ReplacementPlanner().from_derived(derived).warnings
        snapshot_kwargs = {
            "schema_version": 2,
            "document_version_id": version_id,
            "decision_snapshot_sha256": derived.decision_snapshot_sha256,
        }
    return ExportSnapshot(
        **snapshot_kwargs,
        captured_at=datetime.now(UTC),
        source_name=job.source_name,
        source_type=job.file_type,
        source_size_bytes=job.size_bytes,
        source_sha256=None,
        scenario=job.scenario,
        enabled_categories=list(job.enabled_categories),
        completed_categories=[
            category for category in job.enabled_categories if category not in failures
        ],
        checker_failures=[
            ExportCheckerFailureSnapshot(
                category=category,
                code=failures[category].code,
                message=failures[category].message,
            )
            for category in job.enabled_categories
            if category in failures
        ],
        summary=ExportIssueSummarySnapshot(
            total=summary_total,
            by_category=summary_by_category,
            by_severity=summary_by_severity,
            by_decision=summary_by_decision,
        ),
        document=document,
        issues=issues,
        preflight_warnings=warnings,
    )


def _load_export(
    session_factory: sessionmaker[Session],
    export_id: UUID,
):
    session = session_factory()
    try:
        stored = ExportRepository(session).get(export_id)
        assert stored is not None
        return stored
    finally:
        session.close()
