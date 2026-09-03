from __future__ import annotations

from collections.abc import Iterator
from concurrent.futures import ThreadPoolExecutor
from contextlib import contextmanager
from dataclasses import asdict, dataclass, field, replace
from datetime import UTC, datetime, timedelta
from hashlib import sha256
from pathlib import Path
from threading import Event, RLock
from uuid import UUID, uuid4

import pytest
from docx import Document

from text_verification.application import (
    ArtifactLifecycleStatus,
    ArtifactReservation,
    ArtifactSnapshot,
    VerificationError,
)
from text_verification.application import reconstruction_export as reconstruction_export_module
from text_verification.application.artifact_service import ArtifactPersistenceService
from text_verification.application.factory import build_default_exporter_registry
from text_verification.application.recheck_provenance import (
    RecheckGrantBinding,
    RecheckProvenanceGrantService,
)
from text_verification.application.reconstruction_export import (
    ReconstructionExportService,
    _document_with_revision_text,
)
from text_verification.document_processing.ocr_provider import OcrTextBox
from text_verification.domain.artifacts import ArtifactFinalizationRejection
from text_verification.domain.documents import (
    DocumentMetadata,
    DocumentModel,
    ExportFormat,
    FileType,
    TextBlock,
)
from text_verification.domain.jobs import JobProgressStage, JobRead, JobStatus
from text_verification.domain.verification import (
    DocumentRevisionKind,
    PersistedDocumentRevision,
    RecheckProvenance,
    Scenario,
    StaleReviewRevisionError,
    VerificationAnalysisMode,
    VerificationDegradation,
    VerificationExecutionMode,
    VerificationResult,
    VerificationStatistics,
    VerificationSummary,
)
from text_verification.exporters.registry import ExporterRegistry
from text_verification.infrastructure.artifact_storage import (
    ArtifactRepairQuarantine,
)
from text_verification.infrastructure.storage import (
    JobOwnedSourcePathResolver,
    JobStorage,
)
from text_verification.infrastructure.verification_repository import (
    JobResultSnapshot,
    JobResultState,
)
from text_verification.parsers.pdf_parser import PdfParser

PDF_FIXTURES = Path(__file__).resolve().parents[1] / "fixtures" / "pdf"


class _FakeOcr:
    def recognize(self, image: object, language: str) -> list[OcrTextBox]:
        del image, language
        return [
            OcrTextBox(
                text="test@example.com",
                confidence=0.99,
                bbox=((40.0, 40.0), (300.0, 40.0), (300.0, 80.0), (40.0, 80.0)),
            )
        ]


@dataclass
class _RepositoryState:
    result_by_job: dict[UUID, VerificationResult]
    revisions: dict[UUID, PersistedDocumentRevision] = field(default_factory=dict)
    artifacts: dict[UUID, ArtifactSnapshot] = field(default_factory=dict)
    result_state_by_job: dict[UUID, JobResultState] = field(default_factory=dict)
    reserve_error: Exception | None = None
    expire_on_snapshot_read: int | None = None
    replacement_result_on_snapshot_read: tuple[int, VerificationResult] | None = None
    reject_finalize_once: bool = False
    stale_revision_on_finalize: PersistedDocumentRevision | None = None
    newer_revision_after_repair_commit: PersistedDocumentRevision | None = None
    newer_revision_on_artifact_read: PersistedDocumentRevision | None = None
    snapshot_reads: int = 0
    repair_transition_pending_commit: bool = False
    lock: RLock = field(default_factory=RLock)


class _InMemoryExportRepository:
    def __init__(self, state: _RepositoryState) -> None:
        self._state = state

    def read_result_snapshot(self, job_id: UUID) -> JobResultSnapshot:
        self._state.snapshot_reads += 1
        if self._state.expire_on_snapshot_read == self._state.snapshot_reads:
            return JobResultSnapshot(JobResultState.EXPIRED, None)
        replacement = self._state.replacement_result_on_snapshot_read
        if replacement is not None and replacement[0] == self._state.snapshot_reads:
            self._state.result_by_job[job_id] = replacement[1]
        result = self._state.result_by_job.get(job_id)
        state = self._state.result_state_by_job.get(
            job_id,
            JobResultState.MISSING if result is None else JobResultState.READY,
        )
        return JobResultSnapshot(
            state,
            result if state is JobResultState.READY else None,
        )

    def read_review_revision(
        self,
        review_revision_id: UUID,
    ) -> PersistedDocumentRevision | None:
        return self._state.revisions.get(review_revision_id)

    def read_export_revision(
        self,
        job_id: UUID,
        verification_run_id: UUID,
        review_revision_id: UUID | None,
    ) -> PersistedDocumentRevision | None:
        result = self._state.result_by_job[job_id]
        if result.verification_run_id != verification_run_id:
            raise LookupError("verification run was superseded")
        revisions = sorted(
            (
                revision
                for revision in self._state.revisions.values()
                if revision.verification_run_id == verification_run_id
            ),
            key=lambda revision: revision.revision_number,
        )
        latest = revisions[-1] if revisions else None
        if review_revision_id is None:
            if latest is not None:
                raise StaleReviewRevisionError("requested revision is stale")
            return None
        revision = self._state.revisions.get(review_revision_id)
        if revision is None:
            raise LookupError("revision missing")
        if latest is None or latest.revision_id != review_revision_id:
            raise StaleReviewRevisionError("requested revision is stale")
        return revision

    def begin_export_artifact_repair(
        self,
        expected: ArtifactReservation,
        *,
        consistency_check,
    ) -> ArtifactReservation | None:
        with self._state.lock:
            existing = self._state.artifacts.get(expected.export_artifact_id)
            if existing is None:
                return None
            if existing.content_sha256 != expected.content_sha256:
                raise ValueError("repair metadata conflict")
            consistency_check()
            pending = replace(
                existing,
                status=ArtifactLifecycleStatus.PENDING,
                reserved_at=expected.reserved_at,
                reservation_version=existing.reservation_version + 1,
                ready_at=None,
            )
            self._state.artifacts[expected.export_artifact_id] = pending
            self._state.repair_transition_pending_commit = True
            return ArtifactReservation(
                **{
                    key: value
                    for key, value in asdict(pending).items()
                    if key != "ready_at"
                }
            )

    def reserve_export_artifact(self, **values: object) -> ArtifactReservation:
        if self._state.reserve_error is not None:
            raise self._state.reserve_error
        artifact_id = values["export_artifact_id"]
        assert isinstance(artifact_id, UUID)
        verification_run_id = values["verification_run_id"]
        job_id = next(
            job_id
            for job_id, result in self._state.result_by_job.items()
            if result.verification_run_id == verification_run_id
        )
        with self._state.lock:
            revisions = sorted(
                (
                    revision
                    for revision in self._state.revisions.values()
                    if revision.verification_run_id == verification_run_id
                ),
                key=lambda revision: revision.revision_number,
            )
            latest = revisions[-1] if revisions else None
            requested_revision_id = values["review_revision_id"]
            if (requested_revision_id is None) != (latest is None) or (
                requested_revision_id is not None
                and latest is not None
                and requested_revision_id != latest.revision_id
            ):
                raise StaleReviewRevisionError("requested revision is stale")
            existing = self._state.artifacts.get(artifact_id)
            if existing is not None:
                assert existing.content_sha256 == values["content_sha256"]
                if existing.status is ArtifactLifecycleStatus.PENDING:
                    existing = replace(
                        existing,
                        reserved_at=values["reserved_at"],
                        reservation_version=existing.reservation_version + 1,
                    )
                    self._state.artifacts[artifact_id] = existing
                return ArtifactReservation(
                    **{
                        key: value
                        for key, value in asdict(existing).items()
                        if key != "ready_at"
                    }
                )
            reservation = ArtifactReservation(
                job_id=job_id,
                **values,
                status=ArtifactLifecycleStatus.PENDING,
                reservation_version=1,
            )
            snapshot_values = asdict(reservation)
            self._state.artifacts[artifact_id] = ArtifactSnapshot(
                **snapshot_values,
                ready_at=None,
            )
            return reservation

    def finalize_export_artifact(
        self,
        reservation: ArtifactReservation,
        *,
        ready_at: datetime,
        consistency_check,
        require_current_result: bool = False,
    ) -> ArtifactSnapshot | ArtifactFinalizationRejection | None:
        consistency_check()
        with self._state.lock:
            if require_current_result:
                revisions = sorted(
                    (
                        revision
                        for revision in self._state.revisions.values()
                        if revision.verification_run_id
                        == reservation.verification_run_id
                    ),
                    key=lambda revision: revision.revision_number,
                )
                latest = revisions[-1] if revisions else None
                if (
                    (reservation.review_revision_id is None) != (latest is None)
                    or (
                        reservation.review_revision_id is not None
                        and latest is not None
                        and reservation.review_revision_id != latest.revision_id
                    )
                ):
                    return ArtifactFinalizationRejection.STALE_REVISION
            if self._state.stale_revision_on_finalize is not None:
                stale = self._state.stale_revision_on_finalize
                self._state.revisions[stale.revision_id] = stale
                return ArtifactFinalizationRejection.STALE_REVISION
            if self._state.reject_finalize_once:
                self._state.reject_finalize_once = False
                return None
            snapshot_values = asdict(reservation)
            snapshot_values["status"] = ArtifactLifecycleStatus.READY
            snapshot = ArtifactSnapshot(
                **snapshot_values,
                ready_at=ready_at,
            )
            self._state.artifacts[reservation.export_artifact_id] = snapshot
            return snapshot

    def authorize_export_artifact_publication(
        self,
        reservation: ArtifactReservation,
        *,
        publish,
    ):
        with self._state.lock:
            existing = self._state.artifacts.get(
                reservation.export_artifact_id
            )
            if (
                existing is None
                or existing.status is not reservation.status
                or existing.reservation_version
                != reservation.reservation_version
            ):
                return None
            return publish()

    def compensate_pending_export_artifact(
        self,
        reservation: ArtifactReservation,
        *,
        delete_file,
    ) -> bool:
        with self._state.lock:
            existing = self._state.artifacts.get(
                reservation.export_artifact_id
            )
            if (
                existing is None
                or existing.status is not ArtifactLifecycleStatus.PENDING
                or existing.reservation_version
                != reservation.reservation_version
            ):
                return False
            if not delete_file():
                return False
            self._state.artifacts.pop(reservation.export_artifact_id, None)
            return True

    def read_export_artifact(self, export_artifact_id: UUID) -> ArtifactSnapshot | None:
        with self._state.lock:
            if self._state.newer_revision_on_artifact_read is not None:
                newer = self._state.newer_revision_on_artifact_read
                self._state.revisions[newer.revision_id] = newer
                self._state.newer_revision_on_artifact_read = None
            return self._state.artifacts.get(export_artifact_id)

    def commit(self) -> None:
        if self._state.repair_transition_pending_commit:
            self._state.repair_transition_pending_commit = False
            newer = self._state.newer_revision_after_repair_commit
            if newer is not None:
                self._state.revisions[newer.revision_id] = newer
                self._state.newer_revision_after_repair_commit = None
        return None

    def rollback(self) -> None:
        return None


def _repository_factory(state: _RepositoryState):
    @contextmanager
    def factory() -> Iterator[_InMemoryExportRepository]:
        yield _InMemoryExportRepository(state)

    return factory


def _job_and_result(storage: JobStorage) -> tuple[JobRead, VerificationResult]:
    job_id = uuid4()
    stored = storage.save_bytes(
        job_id,
        "scanned-page.pdf",
        (PDF_FIXTURES / "scanned-page.pdf").read_bytes(),
    )
    parsed = PdfParser(ocr=_FakeOcr()).parse(stored.path).model_copy(
        update={"document_id": job_id, "source_name": "scanned-page.pdf"}
    )
    now = datetime(2026, 9, 2, 4, 0, tzinfo=UTC)
    job = JobRead(
        job_id=job_id,
        source_name=parsed.source_name,
        file_type=FileType.PDF,
        size_bytes=stored.size_bytes,
        status=JobStatus.COMPLETED,
        progress=100,
        created_at=now,
        expires_at=now + timedelta(hours=24),
    )
    return job, _result(parsed)


def _result(document: DocumentModel) -> VerificationResult:
    return VerificationResult(
        verification_run_id=uuid4(),
        document_id=document.document_id,
        source_version=document.source_version,
        source_name=document.source_name,
        file_type=document.file_type,
        scenario=Scenario.GENERAL,
        text=document.text,
        blocks=tuple(document.blocks),
        parser_name=document.parser_name,
        parser_version=document.parser_version,
        metadata=document.metadata,
        ocr_requirement=document.metadata.pdf_ocr_requirement,
        stats=VerificationStatistics(
            char_count=len(document.text),
            char_count_no_space=len(document.text.replace(" ", "")),
            line_count=1,
            paragraph_count=1,
            language="en",
            primary_count=len(document.text.split()),
            primary_label="英文单词",
        ),
        issues=(),
        summary=VerificationSummary(total=0),
        execution_mode=VerificationExecutionMode.ASYNCHRONOUS,
        analysis_mode=VerificationAnalysisMode.LOCAL_ONLY,
        degradation=VerificationDegradation(),
    )


def _service(
    storage: JobStorage,
    state: _RepositoryState,
    *,
    registry_calls: list[object] | None = None,
    max_revision_bytes: int | None = None,
    max_revision_codepoints: int | None = None,
    recheck_grant_service: RecheckProvenanceGrantService | None = None,
) -> ReconstructionExportService:
    def registry_factory(resolver):
        if registry_calls is not None:
            registry_calls.append(resolver)
        return build_default_exporter_registry(
            anchored_source_resolver=resolver,
            max_output_bytes=storage.max_document_bytes,
        )

    kwargs = {}
    if max_revision_codepoints is not None:
        kwargs["max_revision_codepoints"] = max_revision_codepoints
    return ReconstructionExportService(
        storage,
        _repository_factory(state),
        exporter_registry_factory=registry_factory,
        max_revision_bytes=max_revision_bytes,
        recheck_grant_service=recheck_grant_service,
        **kwargs,
    )


def _projection_document(
    text: str,
    blocks: list[TextBlock],
) -> DocumentModel:
    return DocumentModel(
        document_id=uuid4(),
        source_version="sha256:projection",
        file_type=FileType.PDF,
        source_name="projection.pdf",
        text=text,
        blocks=blocks,
        parser_name="projection-test",
        parser_version="1",
        metadata=DocumentMetadata(),
    )


def _projection_block(
    block_id: str,
    kind: str,
    text: str,
    start: int,
    end: int,
    *,
    parent_id: str | None = None,
    table_index: int | None = None,
    row_index: int | None = None,
    cell_index: int | None = None,
    bbox: tuple[float, float, float, float] | None = (0, 0, 10, 10),
) -> TextBlock:
    return TextBlock(
        block_id=block_id,
        kind=kind,
        text=text,
        global_start=start,
        global_end=end,
        block_start=0,
        block_end=len(text),
        page=1,
        paragraph_index=0 if kind in {"paragraph", "heading"} else None,
        table_index=table_index,
        row_index=row_index,
        cell_index=cell_index,
        bbox=bbox,
        parent_id=parent_id,
        style={"spans": []},
        source_locator={},
    )


def _export_projected_revision(
    tmp_path: Path,
    *,
    source: str,
    revised: str,
    blocks: list[TextBlock],
) -> tuple[list[str], list[list[str]]]:
    storage = JobStorage(tmp_path / "jobs", max_upload_bytes=5 * 1024 * 1024)
    job, base_result = _job_and_result(storage)
    result = base_result.model_copy(
        update={
            "source_version": f"sha256:{sha256(source.encode()).hexdigest()}",
            "text": source,
            "blocks": tuple(
                block.model_copy(update={"style": {}})
                for block in blocks
            ),
        }
    )
    revision_id = uuid4()
    revision = PersistedDocumentRevision(
        revision_id=revision_id,
        document_id=result.document_id,
        verification_run_id=result.verification_run_id,
        source_version=result.source_version,
        revision_number=1,
        created_at=job.created_at + timedelta(minutes=1),
        parent_revision_id=None,
        persistence_state="persisted",
        kind=DocumentRevisionKind.MANUAL,
        text=revised,
    )
    state = _RepositoryState(
        {job.job_id: result},
        revisions={revision_id: revision},
    )

    artifact = _service(storage, state).export(
        job,
        ExportFormat.DOCX_RECONSTRUCTION,
        review_revision_id=revision_id,
    )
    download = _service(storage, state).download(
        job.job_id,
        artifact.export_artifact_id,
    )
    with download.handle:
        rebuilt = Document(download.path)
        paragraphs = [paragraph.text for paragraph in rebuilt.paragraphs]
        tables = [
            [cell.text for row in table.rows for cell in row.cells]
            for table in rebuilt.tables
        ]
        assert download.handle.read_bytes()
    return paragraphs, tables


def test_revision_projection_preserves_repeated_unicode_edits_within_blocks() -> None:
    source = "AAAA\n😀表😀"
    revised = "AAΩAA\n😀字表😀"
    document = _projection_document(
        source,
        [
            _projection_block("a", "paragraph", "AAAA", 0, 4),
            _projection_block(
                "b",
                "paragraph",
                "😀表😀",
                5,
                len(source),
            ),
        ],
    )

    projected = _document_with_revision_text(document, revised)

    assert projected.text == revised
    assert [
        (block.block_id, block.text, block.global_start, block.global_end)
        for block in projected.blocks
    ] == [
        ("a", "AAΩAA", 0, 5),
        ("b", "😀字表😀", 6, 10),
    ]
    original_by_id = {block.block_id: block for block in document.blocks}
    assert all(
        "spans" not in block.style
        for block in projected.blocks
        if block.text != original_by_id[block.block_id].text
    )


@pytest.mark.parametrize(
    ("source", "revised", "blocks", "expected_paragraphs"),
    [
        (
            "甲\n乙\n丙",
            "甲\n新\n乙\n丙",
            [
                _projection_block("a", "paragraph", "甲\n乙", 0, 3),
                _projection_block("b", "paragraph", "丙", 4, 5),
            ],
            ["甲\n新\n乙", "丙"],
        ),
        (
            "正文",
            "前😀正文",
            [_projection_block("only", "paragraph", "正文", 0, 2)],
            ["前😀正文"],
        ),
        (
            "正文",
            "正文后😀",
            [_projection_block("only", "paragraph", "正文", 0, 2)],
            ["正文后😀"],
        ),
        (
            "甲\n乙",
            "甲尾\n首乙",
            [
                _projection_block("a", "paragraph", "甲", 0, 1),
                _projection_block("b", "paragraph", "乙", 2, 3),
            ],
            ["甲尾", "首乙"],
        ),
        (
            "重复\n行\n重复\n😀重复",
            "重Ω复\n行\n重复\n😀字重复",
            [
                _projection_block("a", "paragraph", "重复\n行", 0, 4),
                _projection_block("b", "paragraph", "重复\n😀重复", 5, 11),
            ],
            ["重Ω复\n行", "重复\n😀字重复"],
        ),
    ],
)
def test_reconstruction_exports_source_anchored_paragraph_edits_exactly(
    tmp_path: Path,
    source: str,
    revised: str,
    blocks: list[TextBlock],
    expected_paragraphs: list[str],
) -> None:
    paragraphs, tables = _export_projected_revision(
        tmp_path,
        source=source,
        revised=revised,
        blocks=blocks,
    )

    assert paragraphs == expected_paragraphs
    assert "\n".join(paragraphs) == revised
    assert tables == []


@pytest.mark.parametrize(
    ("source", "revised", "expected_paragraphs"),
    [
        ("a\na", "\naa", ["", "aa"]),
        ("a\na", "aa\n", ["aa", ""]),
        ("same\nsame", "\nsamesame", ["", "samesame"]),
        ("same\nsame", "samesame\n", ["samesame", ""]),
        ("😀\n😀", "\n😀😀", ["", "😀😀"]),
        ("😀\n😀", "😀😀\n", ["😀😀", ""]),
    ],
)
def test_reconstruction_exports_repeated_adjacent_blocks_without_crossing_separator(
    tmp_path: Path,
    source: str,
    revised: str,
    expected_paragraphs: list[str],
) -> None:
    first_text, second_text = source.split("\n")
    paragraphs, tables = _export_projected_revision(
        tmp_path,
        source=source,
        revised=revised,
        blocks=[
            _projection_block(
                "first",
                "paragraph",
                first_text,
                0,
                len(first_text),
            ),
            _projection_block(
                "second",
                "paragraph",
                second_text,
                len(first_text) + 1,
                len(source),
            ),
        ],
    )

    assert paragraphs == expected_paragraphs
    assert "\n".join(paragraphs) == revised
    assert tables == []


def test_reconstruction_exports_boundary_edits_around_table_structure_exactly(
    tmp_path: Path,
) -> None:
    source = "段\n甲|乙"
    revised = "段尾\n首甲|乙后"
    paragraphs, tables = _export_projected_revision(
        tmp_path,
        source=source,
        revised=revised,
        blocks=[
            _projection_block(
                "p",
                "paragraph",
                "段",
                0,
                1,
                bbox=(0, -20, 10, -10),
            ),
            _projection_block(
                "a",
                "table_cell",
                "甲",
                2,
                3,
                table_index=0,
                row_index=0,
                cell_index=0,
            ),
            _projection_block(
                "b",
                "table_cell",
                "乙",
                4,
                5,
                table_index=0,
                row_index=0,
                cell_index=1,
            ),
        ],
    )

    assert paragraphs == ["段尾"]
    assert tables == [["首甲", "乙后"]]
    assert f"{paragraphs[0]}\n{tables[0][0]}|{tables[0][1]}" == revised


def test_revision_projection_rejects_many_small_block_diffs_by_cumulative_work() -> None:
    block_count = 101
    separators = [chr(0xE000 + index) for index in range(block_count - 1)]
    source_parts: list[str] = []
    revised_parts: list[str] = []
    blocks: list[TextBlock] = []
    cursor = 0
    for index in range(block_count):
        source_block = "a" * 100
        revised_block = "b" * 100
        source_parts.append(source_block)
        revised_parts.append(revised_block)
        blocks.append(
            _projection_block(
                f"p-{index}",
                "paragraph",
                source_block,
                cursor,
                cursor + len(source_block),
            )
        )
        cursor += len(source_block)
        if index < len(separators):
            separator = separators[index]
            source_parts.append(separator)
            revised_parts.append(separator)
            cursor += len(separator)
    document = _projection_document("".join(source_parts), blocks)

    with pytest.raises(VerificationError) as raised:
        _document_with_revision_text(document, "".join(revised_parts))

    assert raised.value.code == "revision_diff_too_complex"


def test_revision_projection_rejects_cumulative_edit_operations_before_next_matcher(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    monkeypatch.setattr(
        reconstruction_export_module,
        "MAX_TEXT_EDIT_OPERATIONS",
        3,
        raising=False,
    )
    document = _projection_document(
        "a|a~a^a",
        [
            _projection_block("p-0", "paragraph", "a", 0, 1),
            _projection_block("p-1", "paragraph", "a", 2, 3),
            _projection_block("p-2", "paragraph", "a", 4, 5),
            _projection_block("p-3", "paragraph", "a", 6, 7),
        ],
    )

    with pytest.raises(VerificationError) as raised:
        _document_with_revision_text(document, "b|b~b^b")

    assert raised.value.code == "revision_diff_too_complex"


def test_revision_projection_preflights_block_count_before_owner_projection(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    monkeypatch.setattr(
        reconstruction_export_module,
        "MAX_REVISION_PROJECTION_BLOCKS",
        2,
        raising=False,
    )
    monkeypatch.setattr(
        reconstruction_export_module,
        "_project_revision_owner_text",
        lambda *args, **kwargs: (_ for _ in ()).throw(
            AssertionError("owner projection must not start")
        ),
    )
    document = _projection_document(
        "a\nb\nc",
        [
            _projection_block("a", "paragraph", "a", 0, 1),
            _projection_block("b", "paragraph", "b", 2, 3),
            _projection_block("c", "paragraph", "c", 4, 5),
        ],
    )

    with pytest.raises(VerificationError) as raised:
        _document_with_revision_text(document, "a\nb\nC")

    assert raised.value.code == "revision_diff_too_complex"


def test_revision_projection_preflights_source_text_before_owner_projection(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    monkeypatch.setattr(
        reconstruction_export_module,
        "MAX_REVISION_TEXT_CODEPOINTS",
        2,
    )
    monkeypatch.setattr(
        reconstruction_export_module,
        "_project_revision_owner_text",
        lambda *args, **kwargs: (_ for _ in ()).throw(
            AssertionError("owner projection must not start")
        ),
    )
    document = _projection_document(
        "abc",
        [_projection_block("only", "paragraph", "abc", 0, 3)],
    )

    with pytest.raises(VerificationError) as raised:
        _document_with_revision_text(document, "abd")

    assert raised.value.code == "revision_text_too_large"


def test_revision_projection_charges_prefix_and_structural_projection_to_one_budget(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    monkeypatch.setattr(
        reconstruction_export_module,
        "MAX_TEXT_DIFF_WORK",
        5,
    )
    document = _projection_document(
        "aaaaaaaaaa",
        [_projection_block("only", "paragraph", "aaaaaaaaaa", 0, 10)],
    )

    with pytest.raises(VerificationError) as raised:
        _document_with_revision_text(document, "aaaaaaaaab")

    assert raised.value.code == "revision_diff_too_complex"


def test_revision_projection_owner_lookup_is_indexed_not_edit_times_owner(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    lookup_calls = 0
    real_bisect_right = reconstruction_export_module.bisect_right

    def counted_bisect_right(values, target):
        nonlocal lookup_calls
        lookup_calls += 1
        return real_bisect_right(values, target)

    monkeypatch.setattr(
        reconstruction_export_module,
        "bisect_right",
        counted_bisect_right,
    )
    block_count = 200
    source = "|".join("a" for _ in range(block_count))
    revised = source.replace("a", "b", 1)
    blocks = [
        _projection_block(
            f"p-{index}",
            "paragraph",
            "a",
            index * 2,
            index * 2 + 1,
        )
        for index in range(block_count)
    ]

    projected = _document_with_revision_text(
        _projection_document(source, blocks),
        revised,
    )

    assert projected.text == revised
    assert lookup_calls < block_count * 5


@pytest.mark.parametrize(
    ("source", "revised", "blocks"),
    [
        (
            "A\nB",
            "AB",
            [
                _projection_block("a", "paragraph", "A", 0, 1),
                _projection_block("b", "paragraph", "B", 2, 3),
            ],
        ),
        (
            "A\nB",
            "A|B",
            [
                _projection_block("a", "paragraph", "A", 0, 1),
                _projection_block("b", "paragraph", "B", 2, 3),
            ],
        ),
        (
            "A\n\nB",
            "A\nX\nB",
            [
                _projection_block("a", "paragraph", "A", 0, 1),
                _projection_block("b", "paragraph", "B", 3, 4),
            ],
        ),
        (
            "甲|乙",
            "甲乙",
            [
                _projection_block(
                    "a",
                    "table_cell",
                    "甲",
                    0,
                    1,
                    table_index=0,
                    row_index=0,
                    cell_index=0,
                ),
                _projection_block(
                    "b",
                    "table_cell",
                    "乙",
                    2,
                    3,
                    table_index=0,
                    row_index=0,
                    cell_index=1,
                ),
            ],
        ),
    ],
)
def test_revision_projection_rejects_cross_structure_gap_edits(
    source: str,
    revised: str,
    blocks: list[TextBlock],
) -> None:
    document = _projection_document(source, blocks)

    with pytest.raises(VerificationError) as raised:
        _document_with_revision_text(document, revised)

    assert raised.value.code == "revision_structure_conflict"


def test_reconstruction_exports_exact_repeated_unicode_revision_without_loss(
    tmp_path: Path,
) -> None:
    storage = JobStorage(tmp_path / "jobs", max_upload_bytes=5 * 1024 * 1024)
    job, base_result = _job_and_result(storage)
    source = "AAAA\n😀表😀"
    revised = "AAΩAA\n😀字表😀"
    result = base_result.model_copy(
        update={
            "source_version": "sha256:repeated-unicode",
            "text": source,
            "blocks": (
                _projection_block("a", "paragraph", "AAAA", 0, 4),
                _projection_block("b", "paragraph", "😀表😀", 5, 8),
            ),
        }
    )
    revision_id = uuid4()
    revision = PersistedDocumentRevision(
        revision_id=revision_id,
        document_id=result.document_id,
        verification_run_id=result.verification_run_id,
        source_version=result.source_version,
        revision_number=1,
        created_at=job.created_at + timedelta(minutes=1),
        parent_revision_id=None,
        persistence_state="persisted",
        kind=DocumentRevisionKind.MANUAL,
        text=revised,
    )
    state = _RepositoryState(
        {job.job_id: result},
        revisions={revision_id: revision},
    )

    artifact = _service(storage, state).export(
        job,
        ExportFormat.DOCX_RECONSTRUCTION,
        review_revision_id=revision_id,
    )
    download = _service(storage, state).download(
        job.job_id,
        artifact.export_artifact_id,
    )

    with download.handle:
        rebuilt = Document(download.path)
        assert "\n".join(paragraph.text for paragraph in rebuilt.paragraphs) == revised
        assert download.handle.read_bytes()


def test_reconstruction_rejects_structure_conflict_before_artifact_reservation(
    tmp_path: Path,
) -> None:
    storage = JobStorage(tmp_path / "jobs", max_upload_bytes=5 * 1024 * 1024)
    job, base_result = _job_and_result(storage)
    result = base_result.model_copy(
        update={
            "source_version": "sha256:paragraph-boundary",
            "text": "A\nB",
            "blocks": (
                _projection_block("a", "paragraph", "A", 0, 1),
                _projection_block("b", "paragraph", "B", 2, 3),
            ),
        }
    )
    revision_id = uuid4()
    revision = PersistedDocumentRevision(
        revision_id=revision_id,
        document_id=result.document_id,
        verification_run_id=result.verification_run_id,
        source_version=result.source_version,
        revision_number=1,
        created_at=job.created_at + timedelta(minutes=1),
        parent_revision_id=None,
        persistence_state="persisted",
        kind=DocumentRevisionKind.MANUAL,
        text="AB",
    )
    state = _RepositoryState(
        {job.job_id: result},
        revisions={revision_id: revision},
    )

    with pytest.raises(VerificationError) as raised:
        _service(storage, state).export(
            job,
            ExportFormat.DOCX_RECONSTRUCTION,
            review_revision_id=revision_id,
        )

    assert raised.value.code == "revision_structure_conflict"
    assert state.artifacts == {}
    assert not list((storage._root / "artifacts").rglob("*.docx"))


def test_reconstruction_rejects_true_table_separator_edit_before_reservation(
    tmp_path: Path,
) -> None:
    storage = JobStorage(tmp_path / "jobs", max_upload_bytes=5 * 1024 * 1024)
    job, base_result = _job_and_result(storage)
    result = base_result.model_copy(
        update={
            "source_version": "sha256:table-separator",
            "text": "甲|乙",
            "blocks": (
                _projection_block(
                    "a",
                    "table_cell",
                    "甲",
                    0,
                    1,
                    table_index=0,
                    row_index=0,
                    cell_index=0,
                ),
                _projection_block(
                    "b",
                    "table_cell",
                    "乙",
                    2,
                    3,
                    table_index=0,
                    row_index=0,
                    cell_index=1,
                ),
            ),
        }
    )
    revision_id = uuid4()
    state = _RepositoryState(
        {job.job_id: result},
        revisions={
            revision_id: PersistedDocumentRevision(
                revision_id=revision_id,
                document_id=result.document_id,
                verification_run_id=result.verification_run_id,
                source_version=result.source_version,
                revision_number=1,
                created_at=job.created_at + timedelta(minutes=1),
                parent_revision_id=None,
                persistence_state="persisted",
                kind=DocumentRevisionKind.MANUAL,
                text="甲乙",
            )
        },
    )

    with pytest.raises(VerificationError) as raised:
        _service(storage, state).export(
            job,
            ExportFormat.DOCX_RECONSTRUCTION,
            review_revision_id=revision_id,
        )

    assert raised.value.code == "revision_structure_conflict"
    assert state.artifacts == {}
    assert not list((storage._root / "artifacts").rglob("*.docx"))


def test_reconstruction_rejects_legacy_oversized_revision_before_diff_or_reservation(
    tmp_path: Path,
) -> None:
    storage = JobStorage(tmp_path / "jobs", max_upload_bytes=5 * 1024 * 1024)
    job, result = _job_and_result(storage)
    revision_id = uuid4()
    revision = PersistedDocumentRevision(
        revision_id=revision_id,
        document_id=result.document_id,
        verification_run_id=result.verification_run_id,
        source_version=result.source_version,
        revision_number=1,
        created_at=job.created_at + timedelta(minutes=1),
        parent_revision_id=None,
        persistence_state="persisted",
        kind=DocumentRevisionKind.MANUAL,
        text="😀a",
    )
    state = _RepositoryState(
        {job.job_id: result},
        revisions={revision_id: revision},
    )

    with pytest.raises(VerificationError) as raised:
        _service(storage, state, max_revision_bytes=4).export(
            job,
            ExportFormat.DOCX_RECONSTRUCTION,
            review_revision_id=revision_id,
        )

    assert raised.value.code == "revision_text_too_large"
    assert state.artifacts == {}


def test_revision_projection_expands_nested_parent_and_table_cell_boundaries() -> None:
    document = _projection_document(
        "AA|BB",
        [
            _projection_block("root", "header", "AA|BB", 0, 5),
            _projection_block("paragraph", "paragraph", "AA", 0, 2, parent_id="root"),
            _projection_block(
                "cell",
                "table_cell",
                "BB",
                3,
                5,
                parent_id="root",
                table_index=0,
                row_index=0,
                cell_index=0,
            ),
        ],
    )

    projected = _document_with_revision_text(document, "AΩA|BβB")

    by_id = {block.block_id: block for block in projected.blocks}
    assert (by_id["paragraph"].text, by_id["paragraph"].global_end) == ("AΩA", 3)
    assert (
        by_id["cell"].text,
        by_id["cell"].global_start,
        by_id["cell"].global_end,
    ) == ("BβB", 4, 7)
    assert (
        by_id["root"].text,
        by_id["root"].global_start,
        by_id["root"].global_end,
    ) == ("AΩA|BβB", 0, 7)


def test_revision_projection_fails_when_edited_text_has_no_block_owner() -> None:
    document = _projection_document("A", [])

    with pytest.raises(VerificationError) as raised:
        _document_with_revision_text(document, "AX")

    assert raised.value.code == "revision_text_unmappable"


def test_revision_projection_rejects_nested_renderable_blocks_that_would_duplicate_edits() -> None:
    document = _projection_document(
        "AB",
        [
            _projection_block("parent", "paragraph", "AB", 0, 2),
            _projection_block(
                "child",
                "paragraph",
                "A",
                0,
                1,
                parent_id="parent",
            ),
        ],
    )

    with pytest.raises(VerificationError) as raised:
        _document_with_revision_text(document, "AXB")

    assert raised.value.code == "revision_text_unmappable"


def test_reconstructs_persisted_ocr_document_and_downloads_verified_artifact(
    tmp_path: Path,
) -> None:
    storage = JobStorage(tmp_path / "jobs", max_upload_bytes=5 * 1024 * 1024)
    job, result = _job_and_result(storage)
    state = _RepositoryState({job.job_id: result})
    stages: list[JobProgressStage] = []

    artifact = _service(storage, state).export(
        job,
        ExportFormat.DOCX_RECONSTRUCTION,
        progress_observer=stages.append,
    )
    download = _service(storage, state).download(
        job.job_id,
        artifact.export_artifact_id,
    )

    assert stages == [JobProgressStage.EXPORTING, JobProgressStage.FINALIZING]
    assert artifact.job_id == job.job_id
    assert artifact.file_type is FileType.DOCX
    rebuilt = Document(download.path)
    assert any("test@example.com" in paragraph.text for paragraph in rebuilt.paragraphs)
    assert download.handle.content_sha256 == artifact.content_sha256
    download.handle.close()


def test_reconstructs_the_persisted_revision_text_and_keys_artifact_by_revision(
    tmp_path: Path,
) -> None:
    storage = JobStorage(tmp_path / "jobs", max_upload_bytes=5 * 1024 * 1024)
    job, result = _job_and_result(storage)
    revision_id = uuid4()
    revision = PersistedDocumentRevision(
        revision_id=revision_id,
        document_id=result.document_id,
        verification_run_id=result.verification_run_id,
        source_version=result.source_version,
        revision_number=1,
        created_at=job.created_at + timedelta(minutes=1),
        parent_revision_id=None,
        persistence_state="persisted",
        kind=DocumentRevisionKind.REVIEW,
        text=result.text.replace("test@example.com", "reviewed@example.com"),
    )
    state = _RepositoryState(
        {job.job_id: result},
        revisions={revision_id: revision},
    )

    artifact = _service(storage, state).export(
        job,
        ExportFormat.DOCX_RECONSTRUCTION,
        review_revision_id=revision_id,
    )
    download = _service(storage, state).download(
        job.job_id,
        artifact.export_artifact_id,
    )

    assert state.artifacts[artifact.export_artifact_id].review_revision_id == revision_id
    rebuilt = Document(download.path)
    text = "\n".join(paragraph.text for paragraph in rebuilt.paragraphs)
    assert "reviewed@example.com" in text
    assert "test@example.com" not in text
    download.handle.close()


def test_valid_recheck_grant_authorizes_revision_export(
    tmp_path: Path,
) -> None:
    storage = JobStorage(tmp_path / "jobs", max_upload_bytes=5 * 1024 * 1024)
    job, result = _job_and_result(storage)
    revision_id = uuid4()
    revision = PersistedDocumentRevision(
        revision_id=revision_id,
        document_id=result.document_id,
        verification_run_id=result.verification_run_id,
        source_version=result.source_version,
        revision_number=1,
        created_at=job.created_at + timedelta(minutes=1),
        parent_revision_id=None,
        persistence_state="persisted",
        kind=DocumentRevisionKind.MANUAL,
        text=result.text.replace("test@example.com", "rechecked@example.com"),
    )
    state = _RepositoryState(
        {job.job_id: result},
        revisions={revision_id: revision},
    )
    grants = RecheckProvenanceGrantService(
        "server-owned-recheck-grant-secret-32-bytes",
        now_factory=lambda: job.created_at + timedelta(minutes=2),
    )
    binding = RecheckGrantBinding(
        job_id=job.job_id,
        original_document_id=result.document_id,
        original_verification_run_id=result.verification_run_id,
        original_source_version=result.source_version,
        submitted_text="重新检查基线",
        result_document_id=uuid4(),
        result_verification_run_id=uuid4(),
        result_source_version="sha256:" + "b" * 64,
    )
    provenance = RecheckProvenance(
        grant=grants.issue(binding),
        result_document_id=binding.result_document_id,
        result_verification_run_id=binding.result_verification_run_id,
        result_source_version=binding.result_source_version,
        recheck_text=binding.submitted_text,
    )

    artifact = _service(
        storage,
        state,
        recheck_grant_service=grants,
    ).export(
        job,
        ExportFormat.DOCX_RECONSTRUCTION,
        review_revision_id=revision_id,
        recheck_provenance=provenance,
    )

    assert state.artifacts[artifact.export_artifact_id].review_revision_id == revision_id


def test_tampered_recheck_grant_result_is_rejected_before_export_artifact(
    tmp_path: Path,
) -> None:
    storage = JobStorage(tmp_path / "jobs", max_upload_bytes=5 * 1024 * 1024)
    job, result = _job_and_result(storage)
    grants = RecheckProvenanceGrantService(
        "server-owned-recheck-grant-secret-32-bytes",
        now_factory=lambda: job.created_at + timedelta(minutes=2),
    )
    binding = RecheckGrantBinding(
        job_id=job.job_id,
        original_document_id=result.document_id,
        original_verification_run_id=result.verification_run_id,
        original_source_version=result.source_version,
        submitted_text=result.text,
        result_document_id=uuid4(),
        result_verification_run_id=uuid4(),
        result_source_version="sha256:" + "b" * 64,
    )
    provenance = RecheckProvenance(
        grant=grants.issue(binding),
        result_document_id=uuid4(),
        result_verification_run_id=binding.result_verification_run_id,
        result_source_version=binding.result_source_version,
        recheck_text=binding.submitted_text,
    )
    state = _RepositoryState({job.job_id: result})

    with pytest.raises(VerificationError) as raised:
        _service(
            storage,
            state,
            recheck_grant_service=grants,
        ).export(
            job,
            ExportFormat.DOCX_RECONSTRUCTION,
            recheck_provenance=provenance,
        )

    assert raised.value.code == "recheck_provenance_invalid"
    assert state.artifacts == {}


def test_reconstruction_rejects_a_revision_from_another_result(
    tmp_path: Path,
) -> None:
    storage = JobStorage(tmp_path / "jobs", max_upload_bytes=5 * 1024 * 1024)
    job, result = _job_and_result(storage)
    revision_id = uuid4()
    state = _RepositoryState(
        {job.job_id: result},
        revisions={
            revision_id: PersistedDocumentRevision(
                revision_id=revision_id,
                document_id=uuid4(),
                verification_run_id=result.verification_run_id,
                source_version=result.source_version,
                revision_number=1,
                created_at=job.created_at,
                parent_revision_id=None,
                persistence_state="persisted",
                kind=DocumentRevisionKind.REVIEW,
                text="foreign",
            )
        },
    )

    with pytest.raises(VerificationError) as raised:
        _service(storage, state).export(
            job,
            ExportFormat.DOCX_RECONSTRUCTION,
            review_revision_id=revision_id,
        )

    assert raised.value.code == "revision_identity_mismatch"
    assert state.artifacts == {}


def test_reconstruction_rejects_a_persisted_revision_superseded_by_a_newer_one(
    tmp_path: Path,
) -> None:
    storage = JobStorage(tmp_path / "jobs", max_upload_bytes=5 * 1024 * 1024)
    job, result = _job_and_result(storage)
    first_id = uuid4()
    second_id = uuid4()
    state = _RepositoryState(
        {job.job_id: result},
        revisions={
            first_id: PersistedDocumentRevision(
                revision_id=first_id,
                document_id=result.document_id,
                verification_run_id=result.verification_run_id,
                source_version=result.source_version,
                revision_number=1,
                created_at=job.created_at + timedelta(minutes=1),
                parent_revision_id=None,
                persistence_state="persisted",
                kind=DocumentRevisionKind.REVIEW,
                text="first",
            ),
            second_id: PersistedDocumentRevision(
                revision_id=second_id,
                document_id=result.document_id,
                verification_run_id=result.verification_run_id,
                source_version=result.source_version,
                revision_number=2,
                created_at=job.created_at + timedelta(minutes=2),
                parent_revision_id=first_id,
                persistence_state="persisted",
                kind=DocumentRevisionKind.MANUAL,
                text="second",
            ),
        },
    )

    with pytest.raises(VerificationError) as raised:
        _service(storage, state).export(
            job,
            ExportFormat.DOCX_RECONSTRUCTION,
            review_revision_id=first_id,
        )

    assert raised.value.code == "revision_export_stale"
    assert raised.value.stage == "exporting"
    assert raised.value.retryable is False
    assert state.artifacts == {}


def test_job_owned_original_format_export_uses_the_persisted_revision(
    tmp_path: Path,
) -> None:
    storage = JobStorage(tmp_path / "jobs", max_upload_bytes=1024 * 1024)
    job_id = uuid4()
    stored = storage.save_bytes(job_id, "sample.txt", b"source text")
    now = datetime(2026, 9, 3, 4, 0, tzinfo=UTC)
    document = DocumentModel(
        document_id=job_id,
        source_version=f"sha256:{sha256(b'source text').hexdigest()}",
        file_type=FileType.TXT,
        source_name="sample.txt",
        text="source text",
        blocks=[
            _projection_block(
                "text",
                "paragraph",
                "source text",
                0,
                len("source text"),
            )
        ],
        parser_name="plain-text",
        parser_version="1",
        metadata=DocumentMetadata(),
    )
    result = _result(document)
    job = JobRead(
        job_id=job_id,
        source_name="sample.txt",
        file_type=FileType.TXT,
        size_bytes=stored.size_bytes,
        status=JobStatus.COMPLETED,
        progress=100,
        created_at=now,
        expires_at=now + timedelta(hours=24),
    )
    revision_id = uuid4()
    state = _RepositoryState(
        {job_id: result},
        revisions={
            revision_id: PersistedDocumentRevision(
                revision_id=revision_id,
                document_id=job_id,
                verification_run_id=result.verification_run_id,
                source_version=result.source_version,
                revision_number=1,
                created_at=now + timedelta(minutes=1),
                parent_revision_id=None,
                persistence_state="persisted",
                kind=DocumentRevisionKind.MANUAL,
                text="edited text",
            )
        },
    )

    artifact = _service(
        storage,
        state,
        max_revision_bytes=len(b"edited text"),
    ).export(
        job,
        ExportFormat.ORIGINAL_FORMAT,
        review_revision_id=revision_id,
        track_changes=False,
    )
    download = _service(storage, state).download(
        job_id,
        artifact.export_artifact_id,
    )

    with download.handle:
        assert download.handle.read_bytes() == b"edited text"
    assert artifact.file_type is FileType.TXT
    assert artifact.file_name.endswith(".txt")


@pytest.mark.parametrize(
    ("revision_text", "max_codepoints", "max_bytes", "expected_code"),
    [
        ("😀a", 2, 5, None),
        ("😀ab", 2, 6, "revision_text_too_large"),
        ("😀ab", 3, 5, "revision_text_too_large"),
    ],
)
def test_job_owned_txt_export_enforces_exact_revision_limits_before_artifact_creation(
    tmp_path: Path,
    revision_text: str,
    max_codepoints: int,
    max_bytes: int,
    expected_code: str | None,
) -> None:
    storage = JobStorage(tmp_path / "jobs", max_upload_bytes=1024 * 1024)
    job_id = uuid4()
    stored = storage.save_bytes(job_id, "sample.txt", b"source")
    now = datetime(2026, 9, 3, 4, 0, tzinfo=UTC)
    document = DocumentModel(
        document_id=job_id,
        source_version=f"sha256:{sha256(b'source').hexdigest()}",
        file_type=FileType.TXT,
        source_name="sample.txt",
        text="source",
        blocks=[
            _projection_block(
                "text",
                "paragraph",
                "source",
                0,
                len("source"),
            )
        ],
        parser_name="plain-text",
        parser_version="1",
        metadata=DocumentMetadata(),
    )
    result = _result(document)
    job = JobRead(
        job_id=job_id,
        source_name="sample.txt",
        file_type=FileType.TXT,
        size_bytes=stored.size_bytes,
        status=JobStatus.COMPLETED,
        progress=100,
        created_at=now,
        expires_at=now + timedelta(hours=24),
    )
    revision_id = uuid4()
    state = _RepositoryState(
        {job_id: result},
        revisions={
            revision_id: PersistedDocumentRevision(
                revision_id=revision_id,
                document_id=job_id,
                verification_run_id=result.verification_run_id,
                source_version=result.source_version,
                revision_number=1,
                created_at=now + timedelta(minutes=1),
                parent_revision_id=None,
                persistence_state="persisted",
                kind=DocumentRevisionKind.MANUAL,
                text=revision_text,
            )
        },
    )
    service = _service(
        storage,
        state,
        max_revision_bytes=max_bytes,
        max_revision_codepoints=max_codepoints,
    )

    if expected_code is None:
        artifact = service.export(
            job,
            ExportFormat.ORIGINAL_FORMAT,
            review_revision_id=revision_id,
        )
        download = service.download(job_id, artifact.export_artifact_id)
        with download.handle:
            assert download.handle.read_bytes() == revision_text.encode()
        return

    with pytest.raises(VerificationError) as raised:
        service.export(
            job,
            ExportFormat.ORIGINAL_FORMAT,
            review_revision_id=revision_id,
        )

    assert raised.value.code == expected_code
    assert state.artifacts == {}
    assert not list((storage._root / "artifacts").rglob("*.txt"))


def test_reconstruction_eligibility_uses_canonical_structure_not_filename(
    tmp_path: Path,
) -> None:
    storage = JobStorage(tmp_path / "jobs", max_upload_bytes=1024)
    job_id = uuid4()
    now = datetime(2026, 9, 2, 4, 0, tzinfo=UTC)
    job = JobRead(
        job_id=job_id,
        source_name="looks-scanned.pdf",
        file_type=FileType.PDF,
        size_bytes=4,
        status=JobStatus.COMPLETED,
        progress=100,
        created_at=now,
        expires_at=now + timedelta(hours=1),
    )
    flat = DocumentModel(
        document_id=job_id,
        source_version="sha256:flat",
        file_type=FileType.PDF,
        source_name=job.source_name,
        text="flat",
        blocks=[
            TextBlock(
                block_id="p-0",
                kind="paragraph",
                text="flat",
                global_start=0,
                global_end=4,
                block_start=0,
                block_end=4,
                page=None,
                paragraph_index=0,
                table_index=None,
                row_index=None,
                cell_index=None,
                bbox=None,
                parent_id=None,
                style={},
                source_locator={},
            )
        ],
        parser_name="compatibility-flat",
        parser_version="1",
        metadata=DocumentMetadata(),
    )
    state = _RepositoryState({job_id: _result(flat)})

    with pytest.raises(VerificationError) as raised:
        _service(storage, state).export(job, ExportFormat.DOCX_RECONSTRUCTION)

    assert raised.value.code == "document_not_reconstructable"
    assert raised.value.stage == "exporting"
    assert raised.value.retryable is False
    assert state.artifacts == {}


def test_export_expiry_before_render_returns_gone_without_artifact(
    tmp_path: Path,
) -> None:
    storage = JobStorage(tmp_path / "jobs", max_upload_bytes=5 * 1024 * 1024)
    job, result = _job_and_result(storage)
    state = _RepositoryState(
        {job.job_id: result},
        result_state_by_job={job.job_id: JobResultState.EXPIRED},
    )
    registry_calls: list[object] = []

    with pytest.raises(VerificationError) as raised:
        _service(
            storage,
            state,
            registry_calls=registry_calls,
        ).export(job, ExportFormat.DOCX_RECONSTRUCTION)

    assert raised.value.code == "job_result_expired"
    assert registry_calls == []
    assert state.artifacts == {}


def test_export_expiry_after_render_before_publish_leaves_no_artifact(
    tmp_path: Path,
) -> None:
    storage = JobStorage(tmp_path / "jobs", max_upload_bytes=5 * 1024 * 1024)
    job, result = _job_and_result(storage)
    state = _RepositoryState(
        {job.job_id: result},
        expire_on_snapshot_read=2,
    )

    with pytest.raises(VerificationError) as raised:
        _service(storage, state).export(job, ExportFormat.DOCX_RECONSTRUCTION)

    assert raised.value.code == "job_result_expired"
    assert state.artifacts == {}
    assert not list((storage._root / "artifacts").rglob("*.docx"))


def test_concurrent_cleanup_wins_after_render_barrier_before_publish(
    tmp_path: Path,
) -> None:
    storage = JobStorage(tmp_path / "jobs", max_upload_bytes=5 * 1024 * 1024)
    job, result = _job_and_result(storage)
    state = _RepositoryState({job.job_id: result})
    rendered = Event()
    release = Event()

    def registry_factory(resolver):
        delegate = build_default_exporter_registry(
            anchored_source_resolver=resolver,
            max_output_bytes=storage.max_document_bytes,
        ).get(ExportFormat.DOCX_RECONSTRUCTION)

        class BlockingExporter:
            file_type = ExportFormat.DOCX_RECONSTRUCTION

            def export(self, document: DocumentModel, target: Path) -> Path:
                exported = delegate.export(document, target)
                rendered.set()
                if not release.wait(timeout=2):
                    raise TimeoutError("cleanup barrier was not released")
                return exported

        return ExporterRegistry([BlockingExporter()])

    service = ReconstructionExportService(
        storage,
        _repository_factory(state),
        exporter_registry_factory=registry_factory,
    )
    with ThreadPoolExecutor(max_workers=1) as executor:
        future = executor.submit(
            service.export,
            job,
            ExportFormat.DOCX_RECONSTRUCTION,
        )
        assert rendered.wait(timeout=2)
        state.result_state_by_job[job.job_id] = JobResultState.EXPIRED
        release.set()
        with pytest.raises(VerificationError) as raised:
            future.result(timeout=5)

    assert raised.value.code == "job_result_expired"
    assert state.artifacts == {}
    assert not list((storage._root / "artifacts").rglob("*.docx"))


def test_export_expiry_after_publish_before_finalize_aborts_file_and_metadata(
    tmp_path: Path,
) -> None:
    storage = JobStorage(tmp_path / "jobs", max_upload_bytes=5 * 1024 * 1024)
    job, result = _job_and_result(storage)
    state = _RepositoryState(
        {job.job_id: result},
        reject_finalize_once=True,
    )

    with pytest.raises(VerificationError) as raised:
        _service(storage, state).export(job, ExportFormat.DOCX_RECONSTRUCTION)

    assert raised.value.code == "job_result_expired"
    assert state.artifacts == {}
    assert not list((storage._root / "artifacts").rglob("*.docx"))


def test_newer_revision_after_reservation_rejects_finalization_and_cleans_artifact(
    tmp_path: Path,
) -> None:
    storage = JobStorage(tmp_path / "jobs", max_upload_bytes=5 * 1024 * 1024)
    job, result = _job_and_result(storage)
    first_id = uuid4()
    second_id = uuid4()
    first = PersistedDocumentRevision(
        revision_id=first_id,
        document_id=result.document_id,
        verification_run_id=result.verification_run_id,
        source_version=result.source_version,
        revision_number=1,
        created_at=job.created_at + timedelta(minutes=1),
        parent_revision_id=None,
        persistence_state="persisted",
        kind=DocumentRevisionKind.REVIEW,
        text=result.text.replace("test@example.com", "first@example.com"),
    )
    state = _RepositoryState(
        {job.job_id: result},
        revisions={first_id: first},
        stale_revision_on_finalize=PersistedDocumentRevision(
            revision_id=second_id,
            document_id=result.document_id,
            verification_run_id=result.verification_run_id,
            source_version=result.source_version,
            revision_number=2,
            created_at=job.created_at + timedelta(minutes=2),
            parent_revision_id=first_id,
            persistence_state="persisted",
            kind=DocumentRevisionKind.MANUAL,
            text="newer",
        ),
    )

    with pytest.raises(VerificationError) as raised:
        _service(storage, state).export(
            job,
            ExportFormat.DOCX_RECONSTRUCTION,
            review_revision_id=first_id,
        )

    assert raised.value.code == "revision_export_stale"
    assert raised.value.stage == "finalizing"
    assert state.artifacts == {}
    assert not list((storage._root / "artifacts").rglob("*.docx"))


def test_export_result_supersession_after_render_prevents_publish(
    tmp_path: Path,
) -> None:
    storage = JobStorage(tmp_path / "jobs", max_upload_bytes=5 * 1024 * 1024)
    job, result = _job_and_result(storage)
    replacement = result.model_copy(
        update={
            "verification_run_id": uuid4(),
            "source_version": "sha256:replacement",
        }
    )
    state = _RepositoryState(
        {job.job_id: result},
        replacement_result_on_snapshot_read=(2, replacement),
    )

    with pytest.raises(VerificationError) as raised:
        _service(storage, state).export(job, ExportFormat.DOCX_RECONSTRUCTION)

    assert raised.value.code == "export_source_superseded"
    assert state.artifacts == {}
    assert not list((storage._root / "artifacts").rglob("*.docx"))


def test_repeated_and_concurrent_exports_share_one_artifact_reference(
    tmp_path: Path,
) -> None:
    storage = JobStorage(tmp_path / "jobs", max_upload_bytes=5 * 1024 * 1024)
    job, result = _job_and_result(storage)
    state = _RepositoryState({job.job_id: result})
    registry_calls: list[object] = []
    service = _service(storage, state, registry_calls=registry_calls)

    with ThreadPoolExecutor(max_workers=2) as executor:
        references = list(
            executor.map(
                lambda _: service.export(job, ExportFormat.DOCX_RECONSTRUCTION),
                range(2),
            )
        )
    repeated = service.export(job, ExportFormat.DOCX_RECONSTRUCTION)

    assert references[0] == references[1] == repeated
    assert len(state.artifacts) == 1
    assert len(list((storage._root / "artifacts" / str(job.job_id)).glob("*.docx"))) == 1
    assert not list(storage.job_directory(job.job_id).glob("*.reconstructing.docx"))
    assert len(registry_calls) <= 2


def test_ready_artifact_retry_rechecks_latest_revision_before_download_reference(
    tmp_path: Path,
) -> None:
    storage = JobStorage(tmp_path / "jobs", max_upload_bytes=5 * 1024 * 1024)
    job, result = _job_and_result(storage)
    first_id = uuid4()
    second_id = uuid4()
    first = PersistedDocumentRevision(
        revision_id=first_id,
        document_id=result.document_id,
        verification_run_id=result.verification_run_id,
        source_version=result.source_version,
        revision_number=1,
        created_at=job.created_at + timedelta(minutes=1),
        parent_revision_id=None,
        persistence_state="persisted",
        kind=DocumentRevisionKind.REVIEW,
        text=result.text.replace("test@example.com", "first@example.com"),
    )
    state = _RepositoryState(
        {job.job_id: result},
        revisions={first_id: first},
    )
    service = _service(storage, state)
    service.export(
        job,
        ExportFormat.DOCX_RECONSTRUCTION,
        review_revision_id=first_id,
    )
    state.newer_revision_on_artifact_read = PersistedDocumentRevision(
        revision_id=second_id,
        document_id=result.document_id,
        verification_run_id=result.verification_run_id,
        source_version=result.source_version,
        revision_number=2,
        created_at=job.created_at + timedelta(minutes=2),
        parent_revision_id=first_id,
        persistence_state="persisted",
        kind=DocumentRevisionKind.MANUAL,
        text="newer",
    )

    with pytest.raises(VerificationError) as raised:
        service.export(
            job,
            ExportFormat.DOCX_RECONSTRUCTION,
            review_revision_id=first_id,
        )

    assert raised.value.code == "revision_export_stale"


def test_job_owned_source_resolver_rejects_cross_job_document(
    tmp_path: Path,
) -> None:
    storage = JobStorage(tmp_path / "jobs", max_upload_bytes=5 * 1024 * 1024)
    job, result = _job_and_result(storage)
    document = DocumentModel(
        document_id=result.document_id,
        source_version=result.source_version,
        file_type=result.file_type,
        source_name=result.source_name,
        text=result.text,
        blocks=list(result.blocks),
        parser_name=result.parser_name,
        parser_version=result.parser_version,
        metadata=result.metadata,
    ).model_copy(update={"document_id": uuid4()})

    with pytest.raises(ValueError, match="does not belong"):
        JobOwnedSourcePathResolver(
            storage,
            job.job_id,
            FileType.PDF,
        ).resolve_anchored(document)


def test_export_persistence_failure_is_typed_and_leaves_no_orphan(
    tmp_path: Path,
) -> None:
    storage = JobStorage(tmp_path / "jobs", max_upload_bytes=5 * 1024 * 1024)
    job, result = _job_and_result(storage)
    state = _RepositoryState(
        {job.job_id: result},
        reserve_error=RuntimeError("database unavailable"),
    )

    with pytest.raises(VerificationError) as raised:
        _service(storage, state).export(job, ExportFormat.DOCX_RECONSTRUCTION)

    assert raised.value.code == "export_persistence_failed"
    assert raised.value.stage == "finalizing"
    assert raised.value.retryable is True
    assert state.artifacts == {}
    assert not list(storage.job_directory(job.job_id).glob("*.reconstructing.docx"))


def test_repeated_export_repairs_missing_ready_artifact_without_new_reference(
    tmp_path: Path,
) -> None:
    storage = JobStorage(tmp_path / "jobs", max_upload_bytes=5 * 1024 * 1024)
    job, result = _job_and_result(storage)
    state = _RepositoryState({job.job_id: result})
    service = _service(storage, state)
    first = service.export(job, ExportFormat.DOCX_RECONSTRUCTION)
    snapshot = state.artifacts[first.export_artifact_id]
    (storage._root / snapshot.storage_key).unlink()

    repaired = service.export(job, ExportFormat.DOCX_RECONSTRUCTION)
    download = service.download(job.job_id, repaired.export_artifact_id)

    assert repaired == first
    assert download.handle.read_bytes()
    download.handle.close()


def test_repeated_export_repairs_corrupt_ready_artifact_without_new_reference(
    tmp_path: Path,
) -> None:
    storage = JobStorage(tmp_path / "jobs", max_upload_bytes=5 * 1024 * 1024)
    job, result = _job_and_result(storage)
    state = _RepositoryState({job.job_id: result})
    service = _service(storage, state)
    first = service.export(job, ExportFormat.DOCX_RECONSTRUCTION)
    snapshot = state.artifacts[first.export_artifact_id]
    (storage._root / snapshot.storage_key).write_bytes(b"corrupt")

    repaired = service.export(job, ExportFormat.DOCX_RECONSTRUCTION)
    download = service.download(job.job_id, repaired.export_artifact_id)

    assert repaired == first
    assert download.handle.read_bytes()
    download.handle.close()
    assert not storage.artifact_repair_quarantine_path(
        job.job_id,
        first.export_artifact_id,
        FileType.DOCX,
    ).exists()


def test_concurrent_corrupt_ready_repairs_converge(
    tmp_path: Path,
) -> None:
    storage = JobStorage(tmp_path / "jobs", max_upload_bytes=5 * 1024 * 1024)
    job, result = _job_and_result(storage)
    state = _RepositoryState({job.job_id: result})
    service = _service(storage, state)
    first = service.export(job, ExportFormat.DOCX_RECONSTRUCTION)
    snapshot = state.artifacts[first.export_artifact_id]
    (storage._root / snapshot.storage_key).write_bytes(b"corrupt")

    def repair():
        try:
            return service.export(job, ExportFormat.DOCX_RECONSTRUCTION)
        except VerificationError as error:
            return error

    with ThreadPoolExecutor(max_workers=2) as executor:
        repaired = list(executor.map(lambda _: repair(), range(2)))

    references = [
        result
        for result in repaired
        if not isinstance(result, VerificationError)
    ]
    conflicts = [
        result
        for result in repaired
        if isinstance(result, VerificationError)
    ]
    assert references
    assert all(reference == first for reference in references)
    assert len(references) + len(conflicts) == 2
    assert all(
        conflict.code == "export_artifact_repair_pending"
        and conflict.retryable is True
        for conflict in conflicts
    )
    concurrent_download = service.download(job.job_id, first.export_artifact_id)
    with concurrent_download.handle:
        assert concurrent_download.handle.read_bytes()
    assert not storage.artifact_repair_quarantine_path(
        job.job_id,
        first.export_artifact_id,
        FileType.DOCX,
    ).exists()


def test_crash_after_corrupt_quarantine_is_recoverable_on_retry(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    storage = JobStorage(tmp_path / "jobs", max_upload_bytes=5 * 1024 * 1024)
    job, result = _job_and_result(storage)
    state = _RepositoryState({job.job_id: result})
    service = _service(storage, state)
    first = service.export(job, ExportFormat.DOCX_RECONSTRUCTION)
    snapshot = state.artifacts[first.export_artifact_id]
    (storage._root / snapshot.storage_key).write_bytes(b"corrupt")
    real_publish = storage.publish_verified_artifact
    attempts = 0

    def fail_once(*args, **kwargs):
        nonlocal attempts
        attempts += 1
        if attempts == 1:
            raise OSError("worker stopped")
        return real_publish(*args, **kwargs)

    monkeypatch.setattr(storage, "publish_verified_artifact", fail_once)

    with pytest.raises(VerificationError) as raised:
        service.export(job, ExportFormat.DOCX_RECONSTRUCTION)

    assert raised.value.code == "export_artifact_repair_pending"
    assert state.artifacts[first.export_artifact_id].status is ArtifactLifecycleStatus.PENDING
    assert storage.artifact_repair_quarantine_path(
        job.job_id,
        first.export_artifact_id,
        FileType.DOCX,
    ).exists()

    repaired = service.export(job, ExportFormat.DOCX_RECONSTRUCTION)

    assert repaired == first
    retry_download = service.download(job.job_id, first.export_artifact_id)
    with retry_download.handle:
        assert retry_download.handle.read_bytes()
    assert not storage.artifact_repair_quarantine_path(
        job.job_id,
        first.export_artifact_id,
        FileType.DOCX,
    ).exists()


def test_repair_stale_interleaving_leaves_no_pending_metadata_or_quarantine(
    tmp_path: Path,
) -> None:
    storage = JobStorage(tmp_path / "jobs", max_upload_bytes=5 * 1024 * 1024)
    job, result = _job_and_result(storage)
    first_revision_id = uuid4()
    first_revision = PersistedDocumentRevision(
        revision_id=first_revision_id,
        document_id=result.document_id,
        verification_run_id=result.verification_run_id,
        source_version=result.source_version,
        revision_number=1,
        created_at=job.created_at + timedelta(minutes=1),
        parent_revision_id=None,
        persistence_state="persisted",
        kind=DocumentRevisionKind.REVIEW,
        text=result.text.replace("test@example.com", "first@example.com"),
    )
    state = _RepositoryState(
        {job.job_id: result},
        revisions={first_revision_id: first_revision},
    )
    service = _service(storage, state)
    first = service.export(
        job,
        ExportFormat.DOCX_RECONSTRUCTION,
        review_revision_id=first_revision_id,
    )
    snapshot = state.artifacts[first.export_artifact_id]
    (storage._root / snapshot.storage_key).write_bytes(b"corrupt")
    second_revision_id = uuid4()
    state.newer_revision_after_repair_commit = PersistedDocumentRevision(
        revision_id=second_revision_id,
        document_id=result.document_id,
        verification_run_id=result.verification_run_id,
        source_version=result.source_version,
        revision_number=2,
        created_at=job.created_at + timedelta(minutes=2),
        parent_revision_id=first_revision_id,
        persistence_state="persisted",
        kind=DocumentRevisionKind.MANUAL,
        text="newer revision",
    )

    with pytest.raises(VerificationError) as raised:
        service.export(
            job,
            ExportFormat.DOCX_RECONSTRUCTION,
            review_revision_id=first_revision_id,
        )

    assert raised.value.code == "revision_export_stale"
    assert first.export_artifact_id not in state.artifacts
    assert not storage.artifact_repair_quarantine_path(
        job.job_id,
        first.export_artifact_id,
        FileType.DOCX,
    ).exists()
    assert not (storage._root / snapshot.storage_key).exists()


def test_two_repairs_only_clean_their_exact_quarantine_after_stale_revision(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    storage = JobStorage(tmp_path / "jobs", max_upload_bytes=5 * 1024 * 1024)
    job, result = _job_and_result(storage)
    first_revision_id = uuid4()
    first_revision = PersistedDocumentRevision(
        revision_id=first_revision_id,
        document_id=result.document_id,
        verification_run_id=result.verification_run_id,
        source_version=result.source_version,
        revision_number=1,
        created_at=job.created_at + timedelta(minutes=1),
        parent_revision_id=None,
        persistence_state="persisted",
        kind=DocumentRevisionKind.REVIEW,
        text=result.text.replace("test@example.com", "first@example.com"),
    )
    state = _RepositoryState(
        {job.job_id: result},
        revisions={first_revision_id: first_revision},
    )
    service = _service(storage, state)
    first = service.export(
        job,
        ExportFormat.DOCX_RECONSTRUCTION,
        review_revision_id=first_revision_id,
    )
    snapshot = state.artifacts[first.export_artifact_id]
    (storage._root / snapshot.storage_key).write_bytes(b"corrupt")

    owner_at_persist = Event()
    reuser_at_persist = Event()
    allow_owner_persist = Event()
    allow_reuser_persist = Event()
    persist_lock = RLock()
    persist_calls = 0
    deleted_quarantines: list[ArtifactRepairQuarantine] = []
    real_persist = ArtifactPersistenceService.persist
    real_delete = storage.delete_artifact_repair_quarantine

    def interleaved_persist(self, request, *, reservation=None):
        nonlocal persist_calls
        with persist_lock:
            persist_calls += 1
            call_number = persist_calls
        if call_number == 1:
            owner_at_persist.set()
            if not allow_owner_persist.wait(timeout=3):
                raise TimeoutError("timed out waiting to release quarantine owner")
        else:
            reuser_at_persist.set()
            if not allow_reuser_persist.wait(timeout=3):
                raise TimeoutError("timed out waiting to release quarantine reuser")
        return real_persist(self, request, reservation=reservation)

    def record_delete(*args, **kwargs):
        deleted_quarantines.append(args[0])
        return real_delete(*args, **kwargs)

    monkeypatch.setattr(ArtifactPersistenceService, "persist", interleaved_persist)
    monkeypatch.setattr(
        storage,
        "delete_artifact_repair_quarantine",
        record_delete,
    )

    def repair() -> VerificationError:
        try:
            service.export(
                job,
                ExportFormat.DOCX_RECONSTRUCTION,
                review_revision_id=first_revision_id,
            )
        except VerificationError as error:
            return error
        raise AssertionError("stale repair unexpectedly succeeded")

    with ThreadPoolExecutor(max_workers=2) as executor:
        owner_future = executor.submit(repair)
        assert owner_at_persist.wait(timeout=2)
        reuser_future = executor.submit(repair)
        assert reuser_at_persist.wait(timeout=2)
        second_revision_id = uuid4()
        state.revisions[second_revision_id] = PersistedDocumentRevision(
            revision_id=second_revision_id,
            document_id=result.document_id,
            verification_run_id=result.verification_run_id,
            source_version=result.source_version,
            revision_number=2,
            created_at=job.created_at + timedelta(minutes=2),
            parent_revision_id=first_revision_id,
            persistence_state="persisted",
            kind=DocumentRevisionKind.MANUAL,
            text="newer revision",
        )
        try:
            allow_reuser_persist.set()
            assert reuser_future.result(timeout=5).code == "revision_export_stale"
            assert len(deleted_quarantines) == 1
        finally:
            allow_reuser_persist.set()
            allow_owner_persist.set()
        assert (
            owner_future.result(timeout=5).code
            == "export_artifact_repair_pending"
        )

    assert len(deleted_quarantines) == 1
    assert first.export_artifact_id not in state.artifacts
    assert not storage.artifact_repair_quarantine_path(
        job.job_id,
        first.export_artifact_id,
        FileType.DOCX,
    ).exists()
    assert not (storage._root / snapshot.storage_key).exists()


def test_corrupt_ready_repair_rejects_changed_canonical_metadata(
    tmp_path: Path,
) -> None:
    storage = JobStorage(tmp_path / "jobs", max_upload_bytes=5 * 1024 * 1024)
    job, result = _job_and_result(storage)
    state = _RepositoryState({job.job_id: result})
    service = _service(storage, state)
    artifact = service.export(job, ExportFormat.DOCX_RECONSTRUCTION)
    snapshot = state.artifacts[artifact.export_artifact_id]
    (storage._root / snapshot.storage_key).write_bytes(b"corrupt")
    state.artifacts[artifact.export_artifact_id] = replace(
        snapshot,
        content_sha256="0" * 64,
    )

    with pytest.raises(VerificationError) as raised:
        service.export(job, ExportFormat.DOCX_RECONSTRUCTION)

    assert raised.value.code == "export_artifact_conflict"
    assert raised.value.retryable is False


def test_download_maps_corrupt_artifact_to_typed_unavailable_error(
    tmp_path: Path,
) -> None:
    storage = JobStorage(tmp_path / "jobs", max_upload_bytes=5 * 1024 * 1024)
    job, result = _job_and_result(storage)
    state = _RepositoryState({job.job_id: result})
    service = _service(storage, state)
    artifact = service.export(job, ExportFormat.DOCX_RECONSTRUCTION)
    snapshot = state.artifacts[artifact.export_artifact_id]
    (storage._root / snapshot.storage_key).write_bytes(b"corrupt")

    with pytest.raises(VerificationError) as raised:
        service.download(job.job_id, artifact.export_artifact_id)

    assert raised.value.code == "export_artifact_unavailable"
    assert raised.value.stage == "finalizing"
    assert raised.value.retryable is True


@pytest.mark.parametrize(
    ("result_state", "expected_code"),
    [
        (JobResultState.MISSING, "job_result_expired"),
        (JobResultState.PENDING, "job_result_unavailable"),
        (JobResultState.UNAVAILABLE, "job_result_expired"),
        (JobResultState.EXPIRED, "job_result_expired"),
    ],
)
def test_download_denies_lingering_artifact_without_current_ready_result(
    tmp_path: Path,
    result_state: JobResultState,
    expected_code: str,
) -> None:
    storage = JobStorage(tmp_path / "jobs", max_upload_bytes=5 * 1024 * 1024)
    job, result = _job_and_result(storage)
    state = _RepositoryState({job.job_id: result})
    service = _service(storage, state)
    artifact = service.export(job, ExportFormat.DOCX_RECONSTRUCTION)
    state.result_state_by_job[job.job_id] = result_state

    with pytest.raises(VerificationError) as raised:
        service.download(job.job_id, artifact.export_artifact_id)

    assert raised.value.code == expected_code


def test_download_denies_artifact_from_superseded_result_version(
    tmp_path: Path,
) -> None:
    storage = JobStorage(tmp_path / "jobs", max_upload_bytes=5 * 1024 * 1024)
    job, result = _job_and_result(storage)
    state = _RepositoryState({job.job_id: result})
    service = _service(storage, state)
    artifact = service.export(job, ExportFormat.DOCX_RECONSTRUCTION)
    state.result_by_job[job.job_id] = result.model_copy(
        update={
            "verification_run_id": uuid4(),
            "source_version": "sha256:superseded",
        }
    )

    with pytest.raises(VerificationError) as raised:
        service.download(job.job_id, artifact.export_artifact_id)

    assert raised.value.code == "export_artifact_not_found"


def test_download_denies_artifact_for_a_superseded_review_revision(
    tmp_path: Path,
) -> None:
    storage = JobStorage(tmp_path / "jobs", max_upload_bytes=5 * 1024 * 1024)
    job, result = _job_and_result(storage)
    first_id = uuid4()
    second_id = uuid4()
    first = PersistedDocumentRevision(
        revision_id=first_id,
        document_id=result.document_id,
        verification_run_id=result.verification_run_id,
        source_version=result.source_version,
        revision_number=1,
        created_at=job.created_at + timedelta(minutes=1),
        parent_revision_id=None,
        persistence_state="persisted",
        kind=DocumentRevisionKind.REVIEW,
        text=result.text.replace("test@example.com", "first@example.com"),
    )
    state = _RepositoryState(
        {job.job_id: result},
        revisions={first_id: first},
    )
    service = _service(storage, state)
    artifact = service.export(
        job,
        ExportFormat.DOCX_RECONSTRUCTION,
        review_revision_id=first_id,
    )
    state.revisions[second_id] = PersistedDocumentRevision(
        revision_id=second_id,
        document_id=result.document_id,
        verification_run_id=result.verification_run_id,
        source_version=result.source_version,
        revision_number=2,
        created_at=job.created_at + timedelta(minutes=2),
        parent_revision_id=first_id,
        persistence_state="persisted",
        kind=DocumentRevisionKind.MANUAL,
        text="newer",
    )

    with pytest.raises(VerificationError) as raised:
        service.download(job.job_id, artifact.export_artifact_id)

    assert raised.value.code == "revision_export_stale"


def test_authorized_download_descriptor_survives_retention_unlink(
    tmp_path: Path,
) -> None:
    storage = JobStorage(tmp_path / "jobs", max_upload_bytes=5 * 1024 * 1024)
    job, result = _job_and_result(storage)
    state = _RepositoryState({job.job_id: result})
    service = _service(storage, state)
    artifact = service.export(job, ExportFormat.DOCX_RECONSTRUCTION)
    download = service.download(job.job_id, artifact.export_artifact_id)
    artifact_path = download.handle.path
    artifact_path.unlink()

    content = download.handle.read_bytes(require_current_entry=False)

    assert content
    download.handle.close()


def test_export_workspace_cleanup_failure_is_typed(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    storage = JobStorage(tmp_path / "jobs", max_upload_bytes=5 * 1024 * 1024)
    job, result = _job_and_result(storage)
    state = _RepositoryState({job.job_id: result})
    real_unlink = Path.unlink

    def fail_reconstruction_cleanup(
        path: Path,
        *args: object,
        **kwargs: object,
    ) -> None:
        if path.name.endswith(".reconstructing.docx"):
            raise PermissionError("locked")
        real_unlink(path, *args, **kwargs)

    monkeypatch.setattr(Path, "unlink", fail_reconstruction_cleanup)

    with pytest.raises(VerificationError) as raised:
        _service(storage, state).export(job, ExportFormat.DOCX_RECONSTRUCTION)

    assert raised.value.code == "export_workspace_cleanup_failed"
    assert raised.value.stage == "finalizing"
    assert raised.value.retryable is True
    assert state.artifacts == {}
