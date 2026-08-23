from __future__ import annotations

from concurrent.futures import ThreadPoolExecutor
from datetime import UTC, datetime, timedelta
from io import BytesIO
from pathlib import Path
from threading import Event
from uuid import UUID, uuid4

import pytest
from docx import Document as WordDocument
from fastapi import FastAPI
from fastapi.testclient import TestClient
from sqlalchemy import func, select, update
from sqlalchemy.orm import Session, sessionmaker

from text_verification.api.dependencies import get_db_session, get_job_storage
from text_verification.checkers.models import CheckCategory
from text_verification.domain.derived_content import derive_document
from text_verification.domain.documents import DocumentModel, FileType, TextBlock
from text_verification.domain.exports import (
    ExportCheckerFailureSnapshot,
    ExportIssueSummarySnapshot,
    ExportSnapshot,
    ExportType,
)
from text_verification.domain.issues import (
    DecisionAction,
    DecisionCommand,
    Issue,
    IssueSeverity,
)
from text_verification.domain.jobs import JobStatus
from text_verification.exporters import ReplacementPlanner
from text_verification.infrastructure.analysis_repositories import AnalysisRepository
from text_verification.infrastructure.decision_repository import DecisionRepository
from text_verification.infrastructure.export_repository import ExportRepository
from text_verification.infrastructure.orm import ExportRow, IssueDecisionRow, JobRow
from text_verification.infrastructure.repositories import JobRepository
from text_verification.infrastructure.storage import JobStorage
from text_verification.parsers import DocxParser


@pytest.fixture
def export_storage(tmp_path: Path) -> JobStorage:
    root = tmp_path / "export-api-jobs"
    root.mkdir()
    return JobStorage(root, max_upload_bytes=25 * 1024 * 1024)


@pytest.fixture(autouse=True)
def export_api_dependencies(
    app: FastAPI,
    db_session_factory: sessionmaker[Session],
    export_storage: JobStorage,
    monkeypatch: pytest.MonkeyPatch,
) -> list[str]:
    from text_verification.api.routes import exports as export_routes

    dispatched: list[str] = []

    def fresh_session():
        session = db_session_factory()
        try:
            yield session
        finally:
            session.close()

    app.dependency_overrides[get_db_session] = fresh_session
    app.dependency_overrides[get_job_storage] = lambda: export_storage
    monkeypatch.setattr(
        export_routes,
        "dispatch_process_export",
        lambda export_id: dispatched.append(export_id),
    )
    return dispatched


def test_pdf_job_rejects_modified_document_export(client, db_session: Session) -> None:
    job_id = _seed_job_with_analysis(db_session, file_type=FileType.PDF)

    response = client.post(
        f"/api/v1/jobs/{job_id}/exports",
        json={"type": "modified_document"},
    )

    assert response.status_code == 422
    assert response.json()["detail"] == {
        "code": "unsupported_export_type",
        "message": "该文件类型不支持所选导出格式。",
    }


def test_modified_document_requires_an_available_decision(
    client,
    db_session: Session,
) -> None:
    issue = _build_issue(file_type=FileType.TXT)
    job_id = _seed_job_with_analysis(
        db_session,
        file_type=FileType.TXT,
        issues=[issue],
    )

    response = client.post(
        f"/api/v1/jobs/{job_id}/exports",
        json={"type": "modified_document"},
    )

    assert response.status_code == 409
    assert response.json()["detail"] == {
        "code": "export_decisions_required",
        "message": "请先处理至少一个问题，再导出修改版文件。",
    }


def test_modified_document_export_returns_conflict_for_overlapping_replacements(
    client,
    db_session: Session,
    export_api_dependencies: list[str],
) -> None:
    document_id = uuid4()
    first = _build_issue(
        file_type=FileType.TXT,
        document_id=document_id,
        start=0,
        end=2,
        suggestion="首次",
    )
    second = _build_issue(
        file_type=FileType.TXT,
        document_id=document_id,
        start=1,
        end=3,
        suggestion="重叠",
    )
    job_id = _seed_job_with_analysis(
        db_session,
        file_type=FileType.TXT,
        issues=[first, second],
    )
    version_id = db_session.get(JobRow, job_id).active_version_id
    assert version_id is not None
    _force_accepted_decision(db_session, job_id, version_id, first)
    _force_accepted_decision(db_session, job_id, version_id, second)
    db_session.commit()

    response = client.post(
        f"/api/v1/jobs/{job_id}/exports",
        json={"type": "modified_document"},
    )

    assert response.status_code == 409
    assert response.json()["detail"] == {
        "code": "overlapping_replacements",
        "message": "已接受的修改范围存在冲突，请先保留其中一个后重试。",
        "issue_ids": sorted([str(first.issue_id), str(second.issue_id)]),
    }
    assert export_api_dependencies == []


def test_reports_allow_unreviewed_issues_and_use_server_derived_extension(
    client,
    db_session: Session,
    export_api_dependencies: list[str],
) -> None:
    job_id = _seed_job_with_analysis(
        db_session,
        file_type=FileType.PDF,
        issues=[_build_issue(file_type=FileType.PDF)],
    )

    response = client.post(
        f"/api/v1/jobs/{job_id}/exports",
        json={"type": "html_report"},
    )

    assert response.status_code == 202
    payload = response.json()
    assert payload["job_id"] == str(job_id)
    assert payload["export_type"] == "html_report"
    assert payload["status"] == "queued"
    assert payload["dispatch_status"] == "dispatched"
    assert payload["file_name"] == "report.html"
    assert "storage_key" not in payload
    assert "snapshot" not in payload
    assert export_api_dependencies == [payload["export_id"]]


def test_create_export_snapshot_v2_records_requested_version_and_decision_hash(
    client,
    db_session: Session,
    export_api_dependencies: list[str],
) -> None:
    issue = _build_issue(file_type=FileType.TXT)
    job_id = _seed_job_with_analysis(
        db_session,
        file_type=FileType.TXT,
        issues=[issue],
    )
    first_version_id = db_session.get(JobRow, job_id).active_version_id
    assert first_version_id is not None
    _apply_decision(
        db_session,
        job_id,
        issue,
        DecisionAction.ACCEPTED,
        replacement="人工最终",
    )
    AnalysisRepository(db_session).replace_analysis(
        job_id,
        _build_document(file_type=FileType.TXT, version=2),
        [],
        {},
    )
    db_session.commit()

    response = client.post(
        f"/api/v1/jobs/{job_id}/exports",
        json={
            "type": "modified_document",
            "version_id": str(first_version_id),
        },
    )

    assert response.status_code == 202
    assert export_api_dependencies == [response.json()["export_id"]]
    stored = ExportRepository(db_session).get(UUID(response.json()["export_id"]))
    assert stored is not None
    stored_row = db_session.get(ExportRow, stored.export_id)
    assert stored_row is not None
    assert stored_row.version_id == first_version_id
    assert stored.snapshot is not None
    assert stored.snapshot.schema_version == 2
    assert stored.snapshot.document_version_id == first_version_id
    assert stored.snapshot.issues[0].decision is not None
    assert stored.snapshot.issues[0].decision.replacement == "人工最终"
    assert stored.snapshot.issues[0].suggestion == "修改后的"
    assert stored.snapshot.decision_snapshot_sha256 == derive_document(
        first_version_id,
        stored.snapshot.document,
        stored.snapshot.issues,
    ).decision_snapshot_sha256


def test_modified_docx_warnings_require_confirmation_without_creating_export(
    client,
    db_session: Session,
    export_storage: JobStorage,
    export_api_dependencies: list[str],
) -> None:
    job_id, issue = _seed_reviewed_docx_job(db_session, export_storage)

    response = client.post(
        f"/api/v1/jobs/{job_id}/exports",
        json={"type": "modified_document"},
    )
    export_count = db_session.scalar(
        select(func.count()).select_from(ExportRow).where(ExportRow.job_id == job_id)
    )

    assert response.status_code == 409
    assert response.json()["detail"] == {
        "code": "export_confirmation_required",
        "message": "检测到无法自动应用的 DOCX 修改，请确认警告后重试。",
        "warnings": [
            {
                "code": "unsafe_docx_run_boundary",
                "message": (
                    "修改范围跨越多个 DOCX 文本运行，为保留格式已跳过；"
                    "请在原文中手动修改后重新导出。"
                ),
                "issue_id": str(issue.issue_id),
                "block_id": issue.block_id,
            }
        ],
    }
    assert export_count == 0
    assert export_api_dependencies == []


def test_confirmed_docx_export_persists_structured_warnings_for_status(
    client,
    db_session: Session,
    export_storage: JobStorage,
    export_api_dependencies: list[str],
) -> None:
    job_id, issue = _seed_reviewed_docx_job(db_session, export_storage)

    created = client.post(
        f"/api/v1/jobs/{job_id}/exports",
        json={"type": "modified_document", "confirm_warnings": True},
    )
    export_id = created.json()["export_id"]
    status_response = client.get(f"/api/v1/jobs/{job_id}/exports/{export_id}")
    expected_warning = {
        "code": "unsafe_docx_run_boundary",
        "message": (
            "修改范围跨越多个 DOCX 文本运行，为保留格式已跳过；"
            "请在原文中手动修改后重新导出。"
        ),
        "issue_id": str(issue.issue_id),
        "block_id": issue.block_id,
    }

    assert created.status_code == 202
    assert created.json()["warnings"] == [expected_warning]
    assert created.json()["dispatch_status"] == "dispatched"
    assert status_response.status_code == 200
    assert status_response.json()["warnings"] == [expected_warning]
    assert "storage_key" not in status_response.json()
    assert "snapshot" not in status_response.json()
    assert export_api_dependencies == [export_id]


def test_docx_report_preflight_exposes_warnings_without_confirmation(
    client,
    db_session: Session,
    export_storage: JobStorage,
) -> None:
    job_id, issue = _seed_reviewed_docx_job(db_session, export_storage)

    response = client.post(
        f"/api/v1/jobs/{job_id}/exports",
        json={"type": "html_report"},
    )

    assert response.status_code == 202
    assert response.json()["warnings"] == [
        {
            "code": "unsafe_docx_run_boundary",
            "message": (
                "修改范围跨越多个 DOCX 文本运行，为保留格式已跳过；"
                "请在原文中手动修改后重新导出。"
            ),
            "issue_id": str(issue.issue_id),
            "block_id": issue.block_id,
        }
    ]


def test_export_dispatch_occurs_only_after_queued_row_is_committed(
    client,
    db_session: Session,
    db_session_factory: sessionmaker[Session],
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    from text_verification.api.routes import exports as export_routes

    job_id = _seed_job_with_analysis(db_session, file_type=FileType.TXT)
    observed_statuses: list[str] = []

    def observe_committed_export(export_id: str) -> None:
        verification_session = db_session_factory()
        try:
            stored = ExportRepository(verification_session).get(UUID(export_id))
            assert stored is not None
            observed_statuses.append(stored.status.value)
        finally:
            verification_session.close()

    monkeypatch.setattr(export_routes, "dispatch_process_export", observe_committed_export)

    response = client.post(
        f"/api/v1/jobs/{job_id}/exports",
        json={"type": "html_report"},
    )

    assert response.status_code == 202
    assert observed_statuses == ["queued"]


def test_dispatch_timeout_returns_deferred_queued_export_for_recovery(
    client,
    db_session: Session,
    monkeypatch: pytest.MonkeyPatch,
    caplog: pytest.LogCaptureFixture,
) -> None:
    from text_verification.api.routes import exports as export_routes

    job_id = _seed_job_with_analysis(db_session, file_type=FileType.TXT)

    def fail_dispatch(_export_id: str) -> None:
        raise ConnectionError("broker unavailable")

    monkeypatch.setattr(export_routes, "dispatch_process_export", fail_dispatch)
    caplog.set_level("ERROR", logger=export_routes.__name__)

    response = client.post(
        f"/api/v1/jobs/{job_id}/exports",
        json={"type": "html_report"},
    )
    db_session.expire_all()
    stored = db_session.scalar(
        select(ExportRow)
        .where(ExportRow.job_id == job_id)
        .order_by(ExportRow.created_at.desc())
    )

    assert response.status_code == 202
    assert response.json()["status"] == "queued"
    assert response.json()["dispatch_status"] == "deferred"
    assert stored is not None
    assert stored.status == "queued"
    assert stored.error_code is None
    assert stored.error_message is None
    assert "export_dispatch_deferred" in caplog.messages
    stale_at = datetime.now(UTC) - timedelta(seconds=61)
    db_session.execute(
        update(ExportRow)
        .where(ExportRow.export_id == stored.export_id)
        .values(updated_at=stale_at)
    )
    db_session.commit()
    recoverable = ExportRepository(db_session).list_stale_recoverable(
        queued_cutoff=datetime.now(UTC) - timedelta(seconds=60),
        processing_cutoff=datetime.now(UTC) - timedelta(minutes=16),
        limit=100,
    )
    assert stored.export_id in recoverable


def test_oversized_export_snapshot_is_rejected_without_persistence_or_dispatch(
    client,
    db_session: Session,
    export_api_dependencies: list[str],
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    from text_verification.api.routes import exports as export_routes

    job_id = _seed_job_with_analysis(db_session, file_type=FileType.TXT)
    monkeypatch.setattr(export_routes, "MAX_EXPORT_SNAPSHOT_BYTES", 1, raising=False)

    response = client.post(
        f"/api/v1/jobs/{job_id}/exports",
        json={"type": "html_report"},
    )
    export_count = db_session.scalar(
        select(func.count()).select_from(ExportRow).where(ExportRow.job_id == job_id)
    )

    assert response.status_code == 413
    assert response.json()["detail"] == {
        "code": "export_snapshot_too_large",
        "message": "导出快照过大，无法创建导出；请缩小文档或问题数量后重试。",
    }
    assert export_count == 0
    assert export_api_dependencies == []


@pytest.mark.parametrize(
    ("status_value", "expected_status", "expected_code"),
    [
        (JobStatus.QUEUED, 409, "analysis_not_ready"),
        (JobStatus.FAILED, 409, "analysis_failed"),
        (JobStatus.EXPIRED, 410, "job_expired"),
    ],
)
def test_create_export_requires_known_terminal_analysis(
    client,
    db_session: Session,
    status_value: JobStatus,
    expected_status: int,
    expected_code: str,
) -> None:
    job_id = _seed_job(db_session, file_type=FileType.TXT, status=status_value)

    response = client.post(
        f"/api/v1/jobs/{job_id}/exports",
        json={"type": "html_report"},
    )

    assert response.status_code == expected_status
    assert response.json()["detail"]["code"] == expected_code


def test_create_export_rejects_terminal_job_without_persisted_analysis(
    client,
    db_session: Session,
) -> None:
    job_id = _seed_job(
        db_session,
        file_type=FileType.TXT,
        status=JobStatus.COMPLETED,
    )

    response = client.post(
        f"/api/v1/jobs/{job_id}/exports",
        json={"type": "html_report"},
    )

    assert response.status_code == 409
    assert response.json()["detail"] == {
        "code": "analysis_not_ready",
        "message": "分析结果尚未就绪，请稍后重试。",
    }


@pytest.mark.parametrize(
    ("locked_change", "expected_status", "expected_detail"),
    [
        (
            "expires",
            410,
            {
                "code": "job_expired",
                "message": "作业已过期，请重新上传文件。",
            },
        ),
        (
            "transitions",
            409,
            {
                "code": "analysis_not_ready",
                "message": "分析结果尚未就绪，请稍后重试。",
            },
        ),
    ],
)
def test_create_export_validates_job_state_acquired_after_waiting_for_lock(
    client,
    db_session: Session,
    db_session_factory: sessionmaker[Session],
    export_api_dependencies: list[str],
    monkeypatch: pytest.MonkeyPatch,
    locked_change: str,
    expected_status: int,
    expected_detail: dict[str, str],
) -> None:
    job_id = _seed_job_with_analysis(db_session, file_type=FileType.TXT)
    lock_attempted = Event()
    real_lock_job = JobRepository.lock_job

    def observed_lock_job(repository: JobRepository, locked_job_id: UUID) -> JobRow:
        lock_attempted.set()
        return real_lock_job(repository, locked_job_id)

    monkeypatch.setattr(JobRepository, "lock_job", observed_lock_job)
    blocking_session = db_session_factory()
    try:
        locked_job = blocking_session.execute(
            select(JobRow).where(JobRow.job_id == job_id).with_for_update()
        ).scalar_one()

        with ThreadPoolExecutor(max_workers=1) as executor:
            request = executor.submit(
                client.post,
                f"/api/v1/jobs/{job_id}/exports",
                json={"type": "html_report"},
            )
            assert lock_attempted.wait(timeout=3)
            if locked_change == "expires":
                locked_job.expires_at = datetime.now(UTC) - timedelta(seconds=1)
            else:
                locked_job.status = JobStatus.PARSING.value
                locked_job.progress = 50
            blocking_session.commit()
            response = request.result(timeout=5)
    finally:
        blocking_session.rollback()
        blocking_session.close()

    verification_session = db_session_factory()
    try:
        export_count = verification_session.scalar(
            select(func.count())
            .select_from(ExportRow)
            .where(ExportRow.job_id == job_id)
        )
    finally:
        verification_session.close()

    assert response.status_code == expected_status
    assert response.json()["detail"] == expected_detail
    assert export_count == 0
    assert export_api_dependencies == []


def test_unknown_export_type_has_structured_error(client, db_session: Session) -> None:
    job_id = _seed_job_with_analysis(db_session, file_type=FileType.TXT)

    response = client.post(
        f"/api/v1/jobs/{job_id}/exports",
        json={"type": "server_archive"},
    )

    assert response.status_code == 422
    assert response.json()["detail"] == {
        "code": "unsupported_export_type",
        "message": "不支持所选导出格式。",
    }


def test_export_status_is_scoped_to_job_and_omits_internal_storage_key(
    client,
    db_session: Session,
) -> None:
    first_job = _seed_job_with_analysis(db_session, file_type=FileType.TXT)
    second_job = _seed_job_with_analysis(db_session, file_type=FileType.TXT)
    export = ExportRepository(db_session).create(
        first_job,
        ExportType.HTML_REPORT,
        "html",
        snapshot=_snapshot_for_job(db_session, first_job),
    )
    db_session.commit()

    found = client.get(f"/api/v1/jobs/{first_job}/exports/{export.export_id}")
    wrong_job = client.get(f"/api/v1/jobs/{second_job}/exports/{export.export_id}")

    assert found.status_code == 200
    assert found.json()["status"] == "queued"
    assert "storage_key" not in found.json()
    assert wrong_job.status_code == 404
    assert wrong_job.json()["detail"]["code"] == "export_not_found"


def test_status_and_download_ignore_large_invalid_private_snapshot(
    app: FastAPI,
    db_session: Session,
    export_storage: JobStorage,
) -> None:
    job_id = _seed_job_with_analysis(db_session, file_type=FileType.TXT)
    repository = ExportRepository(db_session)
    export = repository.create(
        job_id,
        ExportType.MODIFIED_DOCUMENT,
        "txt",
        snapshot=_snapshot_for_job(db_session, job_id),
    )
    repository.mark_processing(export.export_id)
    repository.mark_completed(export.export_id, warnings=[])
    db_session.commit()
    export_path = export_storage.export_path(job_id, export.export_id, "txt")
    export_path.parent.mkdir(parents=True, exist_ok=True)
    export_path.write_bytes("修改后的正文\n".encode())
    db_session.execute(
        update(ExportRow)
        .where(ExportRow.export_id == export.export_id)
        .values(snapshot_json=_large_invalid_snapshot_sentinel())
    )
    db_session.commit()

    with TestClient(app, raise_server_exceptions=False) as resilient_client:
        status_response = resilient_client.get(
            f"/api/v1/jobs/{job_id}/exports/{export.export_id}"
        )
        download_response = resilient_client.get(
            f"/api/v1/jobs/{job_id}/exports/{export.export_id}/download"
        )

    assert status_response.status_code == 200
    assert status_response.json()["status"] == "completed"
    assert status_response.json()["warnings"] == []
    assert "storage_key" not in status_response.json()
    assert "snapshot" not in status_response.json()
    assert download_response.status_code == 200
    assert download_response.content.decode("utf-8") == "修改后的正文\n"


def test_download_returns_conflict_until_completed(client, db_session: Session) -> None:
    job_id = _seed_job_with_analysis(db_session, file_type=FileType.TXT)
    export = ExportRepository(db_session).create(
        job_id,
        ExportType.MODIFIED_DOCUMENT,
        "txt",
        snapshot=_snapshot_for_job(db_session, job_id),
    )
    db_session.commit()

    response = client.get(
        f"/api/v1/jobs/{job_id}/exports/{export.export_id}/download"
    )

    assert response.status_code == 409
    assert response.json()["detail"] == {
        "code": "export_not_ready",
        "message": "导出文件尚未生成，请稍后重试。",
    }


def test_download_returns_gone_after_expiry(
    client,
    db_session: Session,
) -> None:
    job_id = _seed_job_with_analysis(
        db_session,
        file_type=FileType.TXT,
        expires_at=datetime.now(UTC) - timedelta(minutes=1),
    )
    repository = ExportRepository(db_session)
    export = repository.create(
        job_id,
        ExportType.MODIFIED_DOCUMENT,
        "txt",
        snapshot=_snapshot_for_job(db_session, job_id),
    )
    repository.mark_processing(export.export_id)
    repository.mark_completed(export.export_id, warnings=[])
    db_session.commit()

    response = client.get(
        f"/api/v1/jobs/{job_id}/exports/{export.export_id}/download"
    )

    assert response.status_code == 410
    assert response.json()["detail"] == {
        "code": "export_expired",
        "message": "导出文件已过期，请重新创建。",
    }


def test_download_serves_generated_path_with_rfc5987_filename(
    client,
    db_session: Session,
    export_storage: JobStorage,
) -> None:
    job_id = _seed_job_with_analysis(db_session, file_type=FileType.TXT)
    repository = ExportRepository(db_session)
    export = repository.create(
        job_id,
        ExportType.MODIFIED_DOCUMENT,
        "txt",
        snapshot=_snapshot_for_job(db_session, job_id),
    )
    repository.mark_processing(export.export_id)
    repository.mark_completed(export.export_id, warnings=[])
    db_session.commit()
    export_path = export_storage.export_path(job_id, export.export_id, "txt")
    export_path.parent.mkdir(parents=True, exist_ok=True)
    export_path.write_bytes("修改后的正文\n".encode())

    response = client.get(
        f"/api/v1/jobs/{job_id}/exports/{export.export_id}/download"
    )

    assert response.status_code == 200
    assert response.content.decode("utf-8") == "修改后的正文\n"
    assert response.headers["content-type"].startswith("text/plain")
    assert response.headers["content-disposition"] == (
        "attachment; filename=\"modified_document.txt\"; "
        "filename*=UTF-8''modified_document.txt"
    )
    assert str(export_storage.job_directory(job_id)) not in str(response.headers)


def test_download_fails_closed_when_job_directory_escapes_storage_root(
    client,
    db_session: Session,
    export_storage: JobStorage,
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    job_id = _seed_job_with_analysis(db_session, file_type=FileType.TXT)
    repository = ExportRepository(db_session)
    export = repository.create(
        job_id,
        ExportType.MODIFIED_DOCUMENT,
        "txt",
        snapshot=_snapshot_for_job(db_session, job_id),
    )
    repository.mark_processing(export.export_id)
    repository.mark_completed(export.export_id, warnings=[])
    db_session.commit()
    outside = tmp_path / "outside"
    outside.mkdir()
    (outside / f"{export.export_id}.txt").write_text("secret", encoding="utf-8")
    monkeypatch.setattr(export_storage, "job_directory", lambda _job_id: outside)

    response = client.get(
        f"/api/v1/jobs/{job_id}/exports/{export.export_id}/download"
    )

    assert response.status_code == 410
    assert response.json()["detail"] == {
        "code": "export_file_unavailable",
        "message": "导出文件不可用，请重新创建。",
    }


def _seed_job_with_analysis(
    session: Session,
    *,
    file_type: FileType,
    issues: list[Issue] | None = None,
    expires_at: datetime | None = None,
) -> UUID:
    job_id = _seed_job(
        session,
        file_type=file_type,
        status=JobStatus.COMPLETED,
        expires_at=expires_at,
    )
    document = _build_document(
        file_type=file_type,
        document_id=issues[0].document_id if issues else None,
    )
    AnalysisRepository(session).replace_analysis(job_id, document, issues or [], {})
    session.commit()
    return job_id


def _seed_job(
    session: Session,
    *,
    file_type: FileType,
    status: JobStatus,
    expires_at: datetime | None = None,
) -> UUID:
    now = datetime.now(UTC)
    job_id = uuid4()
    repository = JobRepository(session)
    repository.create_job(
        job_id=job_id,
        source_name=f"sample.{file_type.value}",
        file_type=file_type.value,
        size_bytes=16,
        storage_key=str(job_id),
        created_at=now,
        expires_at=expires_at or now + timedelta(hours=1),
    )
    if status != JobStatus.QUEUED:
        repository.transition(
            job_id,
            status,
            100 if status in {JobStatus.COMPLETED, JobStatus.PARTIAL} else 0,
            "处理完成",
            error_code="analysis_failed" if status == JobStatus.FAILED else None,
            error_message="分析失败。" if status == JobStatus.FAILED else None,
        )
    repository.commit()
    return job_id


def _build_document(
    *,
    file_type: FileType,
    document_id: UUID | None = None,
    version: int = 1,
) -> DocumentModel:
    return DocumentModel(
        document_id=document_id or uuid4(),
        file_type=file_type,
        source_name=f"sample.{file_type.value}",
        version=version,
        blocks=[
            TextBlock(
                block_id="p-000001",
                kind="paragraph",
                text="原始正文",
                page=1 if file_type == FileType.PDF else None,
                paragraph_index=0,
                parent_id=None,
                style={},
                source_locator={"paragraph_index": 0},
            )
        ],
        metadata={},
    )


def _build_issue(
    *,
    file_type: FileType,
    document_id: UUID | None = None,
    start: int = 0,
    end: int = 2,
    suggestion: str = "修改后的",
) -> Issue:
    return Issue(
        issue_id=uuid4(),
        document_id=document_id or uuid4(),
        block_id="p-000001",
        page=1 if file_type == FileType.PDF else None,
        start=start,
        end=end,
        original="原始正文"[start:end],
        suggestion=suggestion,
        alternatives=[suggestion],
        type="literal",
        severity=IssueSeverity.WARNING,
        layer=CheckCategory.CHARACTER.value,
        message="命中规则。",
        rule_id="character-001",
        source="test",
        source_version="1",
        confidence=1.0,
        auto_fixable=True,
        context="原始正文",
    )


def _seed_reviewed_docx_job(
    session: Session,
    storage: JobStorage,
) -> tuple[UUID, Issue]:
    job_id = _seed_job(
        session,
        file_type=FileType.DOCX,
        status=JobStatus.COMPLETED,
    )
    source = WordDocument()
    paragraph = source.add_paragraph()
    paragraph.add_run("核验")
    paragraph.add_run("示例")
    payload = BytesIO()
    source.save(payload)
    stored = storage.save_bytes(job_id, "sample.docx", payload.getvalue())
    document = DocxParser().parse(
        stored.path,
        document_id=uuid4(),
        source_name="sample.docx",
    )
    issue = Issue(
        issue_id=uuid4(),
        document_id=document.document_id,
        document_version=document.version,
        block_id=document.blocks[0].block_id,
        page=None,
        start=1,
        end=3,
        original="验示",
        suggestion="审查",
        alternatives=["审查"],
        type="literal",
        severity=IssueSeverity.WARNING,
        layer=CheckCategory.CHARACTER.value,
        message="命中规则。",
        rule_id="character-001",
        source="test",
        source_version="1",
        confidence=1.0,
        auto_fixable=True,
        context=document.blocks[0].text,
    )
    AnalysisRepository(session).replace_analysis(job_id, document, [issue], {})
    session.commit()
    outcome = DecisionRepository(session).apply(
        job_id,
        DecisionCommand(
            issue_id=issue.issue_id,
            issue_version=document.version,
            expected_revision=0,
            action=DecisionAction.ACCEPTED,
            replacement=issue.suggestion,
        ),
    )
    assert outcome.decision is not None
    session.commit()
    return job_id, issue


def _apply_decision(
    session: Session,
    job_id: UUID,
    issue: Issue,
    action: DecisionAction,
    *,
    replacement: str | None = None,
) -> None:
    resolved_replacement = replacement
    if action == DecisionAction.ACCEPTED and resolved_replacement is None:
        resolved_replacement = issue.suggestion
    outcome = DecisionRepository(session).apply(
        job_id,
        DecisionCommand(
            issue_id=issue.issue_id,
            issue_version=1,
            expected_revision=0,
            action=action,
            replacement=resolved_replacement,
        ),
    )
    assert outcome.decision is not None
    session.commit()


def _force_accepted_decision(
    session: Session,
    job_id: UUID,
    version_id: UUID,
    issue: Issue,
) -> None:
    session.add(
        IssueDecisionRow(
            issue_id=issue.issue_id,
            version_id=version_id,
            job_id=job_id,
            issue_version=issue.document_version or 1,
            revision=1,
            action=DecisionAction.ACCEPTED.value,
            replacement=issue.suggestion,
            final_replacement=issue.suggestion,
            suggestion_id=None,
            operation_batch_id=None,
            updated_at=datetime.now(UTC),
        )
    )


def _snapshot_for_job(session: Session, job_id: UUID) -> ExportSnapshot:
    job = JobRepository(session).get_job(job_id)
    document = AnalysisRepository(session).get_document(job_id)
    assert job is not None
    assert document is not None
    repository = AnalysisRepository(session)
    issues = repository.list_all_issues(job_id)
    summary = repository.summarize_issues(job_id)
    failures = repository.get_checker_failures(job_id)
    version_id = session.get(JobRow, job_id).active_version_id
    snapshot_kwargs: dict[str, object]
    if version_id is None:
        warnings = ReplacementPlanner().build(document, issues).warnings
        snapshot_kwargs = {"schema_version": 1}
    else:
        derived = derive_document(version_id, document, issues)
        warnings = ReplacementPlanner().from_derived(derived).warnings
        snapshot_kwargs = {
            "schema_version": 2,
            "document_version_id": version_id,
            "decision_snapshot_sha256": derived.decision_snapshot_sha256,
        }
    return ExportSnapshot(
        **snapshot_kwargs,
        captured_at=datetime.now(UTC),
        source_name=job.source_name,
        source_type=job.file_type,
        source_size_bytes=job.size_bytes,
        source_sha256=None,
        scenario=job.scenario,
        enabled_categories=list(job.enabled_categories),
        completed_categories=[
            category for category in job.enabled_categories if category not in failures
        ],
        checker_failures=[
            ExportCheckerFailureSnapshot(
                category=category,
                code=failures[category].code,
                message=failures[category].message,
            )
            for category in job.enabled_categories
            if category in failures
        ],
        summary=ExportIssueSummarySnapshot(
            total=summary.total,
            by_category=summary.by_category,
            by_severity=summary.by_severity,
            by_decision=summary.by_decision,
        ),
        document=document,
        issues=issues,
        preflight_warnings=warnings,
    )


def _large_invalid_snapshot_sentinel() -> dict[str, object]:
    return {
        "schema_version": 1,
        "unexpected_payload": "x" * (1024 * 1024),
    }
