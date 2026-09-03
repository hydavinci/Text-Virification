# mypy: ignore-errors
from __future__ import annotations

import copy
import io
import shutil
import subprocess
from dataclasses import dataclass
from datetime import UTC, datetime
from math import isfinite
from pathlib import Path
from uuid import uuid4

from text_verification.compatibility.parser import (
    _convert_doc_to_docx,
    decode_rtf_with_spans,
    strip_html,
)
from text_verification.domain.text_edits import (
    MAX_REVISION_TEXT_UTF8_BYTES,
    TextDiffLimitError,
    build_bounded_text_edits,
    validate_revision_text,
)
from text_verification.parsers.pdf_parser import PdfParser


class ExportError(ValueError):
    pass


class ExportTextLimitError(ExportError):
    pass


@dataclass(frozen=True)
class ExportedDocument:
    content: bytes
    extension: str
    media_type: str


@dataclass(frozen=True)
class TextEdit:
    start: int
    end: int
    replacement: str


@dataclass(frozen=True)
class _PdfGlyphRectangle:
    rectangle: object
    source_start: int
    source_end: int
    line_index: int
    line_direction: tuple[float, float]
    writing_mode: int


_PDF_MAPPING_VALIDATION_ERROR = (
    "PDF compatibility mapping validation failed: "
    "character metadata is incomplete or invalid."
)
_PDF_DIRECTIONAL_COMPATIBILITY_ERROR = (
    "PDF compatibility mapping validation failed: directional source text "
    "cannot be safely replaced."
)
_PDF_EDIT_CONFLICT_ERROR = "PDF edits contain overlapping or ambiguous source ranges."
_PDF_INSERTION_SPACE_ERROR = "PDF replacement text does not fit the mapped source area."


MEDIA_TYPES = {
    "csv": "text/csv",
    "doc": "application/msword",
    "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "md": "text/markdown",
    "pdf": "application/pdf",
    "rtf": "application/rtf",
    "txt": "text/plain",
}


def export_original(
    source_path: Path,
    extension: str,
    replacements: list[tuple[str, str, int | None, int | None]],
    track_changes: bool,
    *,
    original_text: str | None = None,
    modified_text: str | None = None,
    max_text_bytes: int | None = None,
) -> ExportedDocument:
    cleaned_replacements = [
        (strip_html(original), strip_html(suggestion))
        for original, suggestion, _, _ in replacements
        if strip_html(original)
    ]
    edits = _build_edits(
        replacements,
        original_text,
        modified_text,
        max_text_bytes=max_text_bytes,
    )
    if extension == "docx":
        content = (
            _export_docx_edits(source_path, edits, track_changes)
            if edits is not None
            else _export_docx(source_path, cleaned_replacements, track_changes)
        )
        return ExportedDocument(content, "docx", MEDIA_TYPES["docx"])
    if extension == "doc":
        converted_path = Path(_convert_doc_to_docx(str(source_path), str(source_path.parent)))
        try:
            content = (
                _export_docx_edits(converted_path, edits, track_changes)
                if edits is not None
                else _export_docx(converted_path, cleaned_replacements, track_changes)
            )
            doc_content = _convert_docx_bytes_to_doc(content, source_path.parent)
            return ExportedDocument(doc_content, "doc", MEDIA_TYPES["doc"])
        finally:
            converted_path.unlink(missing_ok=True)
    if extension in {"txt", "md", "csv"}:
        content = (
            _export_text_edits(source_path, original_text or "", edits, track_changes)
            if edits is not None
            else _export_text(source_path, cleaned_replacements, track_changes)
        )
        return ExportedDocument(content, extension, MEDIA_TYPES[extension])
    if extension == "pdf":
        content = (
            _export_pdf_edits(
                source_path,
                original_text or "",
                edits,
                track_changes,
                coalesce_modified_text=modified_text is not None and not replacements,
            )
            if edits is not None
            else _export_pdf(source_path, cleaned_replacements, track_changes)
        )
        return ExportedDocument(content, "pdf", MEDIA_TYPES["pdf"])
    if extension == "rtf":
        content = (
            _export_rtf_edits(source_path, original_text or "", edits, track_changes)
            if edits is not None
            else _export_rtf(source_path, cleaned_replacements, track_changes)
        )
        return ExportedDocument(content, "rtf", MEDIA_TYPES["rtf"])
    raise ExportError(f"Unsupported export format: .{extension}")


def _build_edits(
    replacements: list[tuple[str, str, int | None, int | None]],
    original_text: str | None,
    modified_text: str | None,
    *,
    max_text_bytes: int | None = None,
) -> list[TextEdit] | None:
    positioned: list[TextEdit] = []
    for original, suggestion, start, end in replacements:
        if start is None or end is None:
            raise ExportError("Each replacement must include its analyzed text position.")
        if original_text is None or end < start or original_text[start:end] != original:
            raise ExportError("A replacement no longer matches the analyzed source text.")
        positioned.append(TextEdit(start, end, strip_html(suggestion)))
    if modified_text is None:
        return positioned
    if original_text is None:
        raise ExportError("Original text is required for modified-text export.")
    try:
        validate_revision_text(original_text)
        validate_revision_text(
            modified_text,
            max_utf8_bytes=(
                min(max_text_bytes, MAX_REVISION_TEXT_UTF8_BYTES)
                if max_text_bytes is not None
                else MAX_REVISION_TEXT_UTF8_BYTES
            ),
        )
    except TextDiffLimitError as error:
        raise ExportTextLimitError(str(error)) from error
    if positioned:
        accepted_text = _apply_text_edits(original_text, positioned, False)
        if accepted_text != modified_text:
            raise ExportError(
                "Please recheck manual edits before combining them with accepted suggestions."
            )
        return positioned
    try:
        edits = build_bounded_text_edits(original_text, modified_text)
    except TextDiffLimitError as error:
        raise ExportTextLimitError(str(error)) from error
    return [
        TextEdit(edit.start, edit.end, edit.replacement)
        for edit in edits
    ]


def _coalesce_pdf_modified_text_edits(
    edits: list[TextEdit],
    original_text: str,
    document: object,
) -> list[TextEdit]:
    if not edits:
        return []
    groups: list[list[TextEdit]] = [[edits[0]]]
    for edit in edits[1:]:
        if _pdf_edits_can_coalesce(
            groups[-1][-1],
            edit,
            original_text,
            document,
        ):
            groups[-1].append(edit)
        else:
            groups.append([edit])
    coalesced: list[TextEdit] = []
    for group in groups:
        start = group[0].start
        end = group[-1].end
        local_edits = [
            TextEdit(
                edit.start - start,
                edit.end - start,
                edit.replacement,
            )
            for edit in group
        ]
        coalesced.append(
            TextEdit(
                start,
                end,
                _apply_text_edits(original_text[start:end], local_edits, False),
            )
        )
    return coalesced


def _pdf_edits_can_coalesce(
    first: TextEdit,
    second: TextEdit,
    original_text: str,
    document: object,
) -> bool:
    separator = original_text[first.end:second.start]
    if second.start < first.end or len(separator) > 1 or any(
        character.isspace() for character in separator
    ):
        return False
    identity = _pdf_source_line_identity(document, first.start, second.end)
    return identity is not None


def _pdf_source_line_identity(
    document: object,
    start: int,
    end: int,
) -> tuple[str, int, int] | None:
    for block in document.blocks:  # type: ignore[attr-defined]
        if not (block.global_start <= start and end <= block.global_end):
            continue
        page_number = block.source_locator.get("page")
        characters = block.source_locator.get("characters")
        if not isinstance(page_number, int) or not isinstance(characters, list):
            return None
        local_start = start - block.global_start
        local_end = end - block.global_start
        covered: set[int] = set()
        line_indices: set[int] = set()
        for value in characters:
            if not isinstance(value, dict):
                continue
            group_start = value.get("source_start")
            group_end = value.get("source_end")
            line_index = value.get("line_index")
            if not (
                isinstance(group_start, int)
                and isinstance(group_end, int)
                and isinstance(line_index, int)
                and group_start < local_end
                and group_end > local_start
                and value.get("mapping_state") == "glyph"
            ):
                continue
            covered.update(
                range(max(group_start, local_start), min(group_end, local_end))
            )
            line_indices.add(line_index)
        if covered != set(range(local_start, local_end)) or len(line_indices) != 1:
            return None
        return block.block_id, page_number, next(iter(line_indices))
    return None


def _document_text_paragraphs(document: object) -> list[object]:
    paragraphs = list(document.paragraphs)  # type: ignore[attr-defined]
    seen = {id(paragraph._element) for paragraph in paragraphs}
    for table in document.tables:  # type: ignore[attr-defined]
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    element_id = id(paragraph._element)
                    if element_id not in seen:
                        paragraphs.append(paragraph)
                        seen.add(element_id)
    return paragraphs


def _paragraph_edit_groups(
    document: object,
    edits: list[TextEdit],
) -> list[tuple[object, list[TextEdit]]]:
    blocks: list[tuple[object, int, int, int]] = []
    offset = 0
    for paragraph in _document_text_paragraphs(document):
        raw_text = "".join(run.text for run in paragraph.runs)  # type: ignore[attr-defined]
        text = strip_html(raw_text).strip()
        if not text:
            continue
        left_trim = len(raw_text) - len(raw_text.lstrip())
        blocks.append((paragraph, offset, offset + len(text), left_trim))
        offset += len(text) + 1

    groups: list[tuple[object, list[TextEdit]]] = []
    for paragraph, block_start, block_end, left_trim in blocks:
        local_edits: list[TextEdit] = []
        for edit in edits:
            is_insertion = edit.start == edit.end
            contained = (
                block_start <= edit.start <= block_end
                if is_insertion
                else block_start <= edit.start and edit.end <= block_end
            )
            if contained:
                local_edits.append(
                    TextEdit(
                        edit.start - block_start + left_trim,
                        edit.end - block_start + left_trim,
                        edit.replacement,
                    )
                )
            elif edit.start < block_end and edit.end > block_start:
                raise ExportError(
                    "Edits spanning multiple Word paragraphs cannot preserve the original layout."
                )
        if local_edits:
            groups.append((paragraph, local_edits))

    covered = sum(len(local_edits) for _, local_edits in groups)
    if covered != len(edits):
        raise ExportError("An edit could not be mapped back to the original Word document.")
    return groups


def _run_spans(paragraph: object) -> list[tuple[object, int, int, str]]:
    spans: list[tuple[object, int, int, str]] = []
    offset = 0
    for run in paragraph.runs:  # type: ignore[attr-defined]
        text = run.text
        spans.append((run, offset, offset + len(text), text))
        offset += len(text)
    return spans


def _run_has_non_text_content(run: object) -> bool:
    from docx.oxml.ns import qn

    allowed = {qn("w:rPr"), qn("w:t")}
    return any(child.tag not in allowed for child in run._element)  # type: ignore[attr-defined]


def _export_docx_edits(source_path: Path, edits: list[TextEdit], track_changes: bool) -> bytes:
    from docx import Document

    document = Document(source_path)
    groups = _paragraph_edit_groups(document, edits)
    if track_changes:
        for paragraph, local_edits in groups:
            _apply_docx_tracked_edits(paragraph, local_edits)
    else:
        for paragraph, local_edits in groups:
            _apply_docx_edits(paragraph, local_edits)
    output = io.BytesIO()
    document.save(output)
    return output.getvalue()


def _apply_docx_edits(paragraph: object, edits: list[TextEdit]) -> None:
    spans = _run_spans(paragraph)
    current = {id(run): text for run, _, _, text in spans}
    for edit in sorted(edits, key=lambda item: (item.start, item.end), reverse=True):
        insertion_run: object | None = None
        insertion_offset = 0
        for run, run_start, run_end, _original_run_text in spans:
            if insertion_run is None and run_start <= edit.start <= run_end:
                insertion_run = run
                insertion_offset = edit.start - run_start
            overlap_start = max(edit.start, run_start)
            overlap_end = min(edit.end, run_end)
            if overlap_start < overlap_end:
                text = current[id(run)]
                local_start = overlap_start - run_start
                local_end = overlap_end - run_start
                current[id(run)] = text[:local_start] + text[local_end:]
        if insertion_run is None and spans:
            insertion_run = spans[-1][0]
            insertion_offset = len(current[id(insertion_run)])
        if insertion_run is None:
            raise ExportError("Cannot map an edit into an empty Word paragraph.")
        text = current[id(insertion_run)]
        current[id(insertion_run)] = (
            text[:insertion_offset] + edit.replacement + text[insertion_offset:]
        )
    for run, _, _, _ in spans:
        replacement = strip_html(current[id(run)])
        if replacement == run.text:
            continue
        if _run_has_non_text_content(run):
            raise ExportError(
                "An edited Word run contains inline content that cannot be changed safely."
            )
        run.text = replacement


def _apply_docx_tracked_edits(paragraph: object, edits: list[TextEdit]) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    spans = _run_spans(paragraph)
    author = "啄木鸟·中英文字智能检查"
    date_string = datetime.now(UTC).strftime("%Y-%m-%dT%H:%M:%SZ")
    xml_space = "{http://www.w3.org/XML/1998/namespace}space"
    revision_id = 1000
    insertion_runs: dict[int, int] = {}
    for edit_index, edit in enumerate(edits):
        for run, run_start, run_end, _ in spans:
            if run_start <= edit.start < run_end or (
                edit.start == run_end and run is spans[-1][0]
            ):
                insertion_runs[edit_index] = id(run)
                break

    def make_run(text: str, properties: object | None) -> object:
        run = OxmlElement("w:r")
        if properties is not None:
            run.append(copy.deepcopy(properties))
        node = OxmlElement("w:t")
        node.set(xml_space, "preserve")
        node.text = strip_html(text)
        run.append(node)
        return run

    def make_revision(tag: str, text: str, properties: object | None) -> object:
        nonlocal revision_id
        revision_id += 1
        revision = OxmlElement(tag)
        revision.set(qn("w:id"), str(revision_id))
        revision.set(qn("w:author"), author)
        revision.set(qn("w:date"), date_string)
        run = OxmlElement("w:r")
        if properties is not None:
            run.append(copy.deepcopy(properties))
        node = OxmlElement("w:delText" if tag == "w:del" else "w:t")
        node.set(xml_space, "preserve")
        node.text = strip_html(text)
        run.append(node)
        revision.append(run)
        return revision

    for run, run_start, run_end, text in spans:
        run_edits = [
            (edit_index, edit)
            for edit_index, edit in enumerate(edits)
            if insertion_runs.get(edit_index) == id(run)
            or (edit.start < run_end and edit.end > run_start)
        ]
        if not run_edits:
            continue
        if _run_has_non_text_content(run):
            raise ExportError(
                "An edited Word run contains inline content that cannot be tracked safely."
            )

        properties = run._element.find(qn("w:rPr"))
        elements: list[object] = []
        cursor = run_start
        for edit_index, edit in sorted(
            run_edits,
            key=lambda item: (item[1].start, item[1].end),
        ):
            overlap_start = max(edit.start, run_start)
            overlap_end = min(edit.end, run_end)
            if overlap_start > cursor:
                elements.append(
                    make_run(text[cursor - run_start : overlap_start - run_start], properties)
                )
            if overlap_start < overlap_end:
                elements.append(
                    make_revision(
                        "w:del",
                        text[overlap_start - run_start : overlap_end - run_start],
                        properties,
                    )
                )
            if insertion_runs.get(edit_index) == id(run) and edit.replacement:
                elements.append(make_revision("w:ins", edit.replacement, properties))
            cursor = max(cursor, overlap_end)
        if cursor < run_end:
            elements.append(make_run(text[cursor - run_start :], properties))

        insert_after = run._element
        for element in elements:
            insert_after.addnext(element)
            insert_after = element
        run._element.getparent().remove(run._element)


def _export_docx(
    source_path: Path,
    replacements: list[tuple[str, str]],
    track_changes: bool,
) -> bytes:
    if track_changes:
        return _export_docx_track_changes(source_path, replacements)
    return _export_docx_inplace(source_path, replacements)


def _paragraph_collections(document: object) -> list[object]:
    collections: list[object] = [document.paragraphs]  # type: ignore[attr-defined]
    for table in document.tables:  # type: ignore[attr-defined]
        for row in table.rows:
            for cell in row.cells:
                collections.append(cell.paragraphs)
    for section in document.sections:  # type: ignore[attr-defined]
        for header_footer in (
            section.header,
            section.first_page_header,
            section.even_page_header,
            section.footer,
            section.first_page_footer,
            section.even_page_footer,
        ):
            collections.append(header_footer.paragraphs)
    return collections


def _export_docx_inplace(
    source_path: Path,
    replacements: list[tuple[str, str]],
) -> bytes:
    from docx import Document

    document = Document(source_path)
    for paragraphs in _paragraph_collections(document):
        for paragraph in paragraphs:
            for run in paragraph.runs:
                run.text = strip_html(run.text)
                for original, suggestion in replacements:
                    run.text = run.text.replace(original, suggestion)

            full_text = strip_html("".join(run.text for run in paragraph.runs))
            modified = _replace_all(full_text, replacements)
            if modified != full_text and paragraph.runs:
                paragraph.runs[0].text = modified
                for run in paragraph.runs[1:]:
                    run.text = ""

    output = io.BytesIO()
    document.save(output)
    return output.getvalue()


def _export_docx_track_changes(
    source_path: Path,
    replacements: list[tuple[str, str]],
) -> bytes:
    from docx import Document
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    document = Document(source_path)
    author = "啄木鸟·中英文字智能检查"
    date_string = datetime.now(UTC).strftime("%Y-%m-%dT%H:%M:%SZ")
    revision_id = 1000
    xml_space = "{http://www.w3.org/XML/1998/namespace}space"

    def next_revision_id() -> str:
        nonlocal revision_id
        revision_id += 1
        return str(revision_id)

    def make_run(text: str, run_properties: object | None) -> object:
        run = OxmlElement("w:r")
        if run_properties is not None:
            run.append(copy.deepcopy(run_properties))
        text_element = OxmlElement("w:t")
        text_element.set(xml_space, "preserve")
        text_element.text = strip_html(text)
        run.append(text_element)
        return run

    def make_revision(tag: str, text: str, run_properties: object | None) -> object:
        revision = OxmlElement(tag)
        revision.set(qn("w:id"), next_revision_id())
        revision.set(qn("w:author"), author)
        revision.set(qn("w:date"), date_string)
        run = OxmlElement("w:r")
        if run_properties is not None:
            run.append(copy.deepcopy(run_properties))
        text_element = OxmlElement("w:delText" if tag == "w:del" else "w:t")
        text_element.set(xml_space, "preserve")
        text_element.text = strip_html(text)
        run.append(text_element)
        revision.append(run)
        return revision

    def process_paragraph(paragraph: object) -> None:
        runs = list(paragraph.runs)  # type: ignore[attr-defined]
        for run in runs:
            text = strip_html(run.text)
            matches = _non_overlapping_matches(text, replacements)
            if not matches:
                run.text = text
                continue
            run_properties = run._element.find(qn("w:rPr"))
            insert_after = run._element
            position = 0
            for start, end, original, suggestion in matches:
                if start > position:
                    element = make_run(text[position:start], run_properties)
                    insert_after.addnext(element)
                    insert_after = element
                deleted = make_revision("w:del", original, run_properties)
                insert_after.addnext(deleted)
                insert_after = deleted
                if suggestion:
                    inserted = make_revision("w:ins", suggestion, run_properties)
                    insert_after.addnext(inserted)
                    insert_after = inserted
                position = end
            if position < len(text):
                element = make_run(text[position:], run_properties)
                insert_after.addnext(element)
            run._element.getparent().remove(run._element)

        remaining_text = strip_html("".join(run.text for run in paragraph.runs))  # type: ignore[attr-defined]
        matching = [(old, new) for old, new in replacements if old in remaining_text and old != new]
        if not matching or not paragraph.runs:  # type: ignore[attr-defined]
            return

        run_properties = paragraph.runs[0]._element.find(qn("w:rPr"))  # type: ignore[attr-defined]
        matches = _non_overlapping_matches(remaining_text, matching)
        if not matches:
            return
        elements: list[object] = []
        position = 0
        for start, end, original, suggestion in matches:
            if start > position:
                elements.append(make_run(remaining_text[position:start], run_properties))
            elements.append(make_revision("w:del", original, run_properties))
            if suggestion:
                elements.append(make_revision("w:ins", suggestion, run_properties))
            position = end
        if position < len(remaining_text):
            elements.append(make_run(remaining_text[position:], run_properties))
        paragraph_element = paragraph._element  # type: ignore[attr-defined]
        for run in list(paragraph.runs):  # type: ignore[attr-defined]
            paragraph_element.remove(run._element)
        for element in elements:
            paragraph_element.append(element)

    for paragraphs in _paragraph_collections(document):
        for paragraph in paragraphs:
            process_paragraph(paragraph)

    output = io.BytesIO()
    document.save(output)
    return output.getvalue()


def _non_overlapping_matches(
    text: str,
    replacements: list[tuple[str, str]],
) -> list[tuple[int, int, str, str]]:
    matches: list[tuple[int, int, str, str]] = []
    for original, suggestion in replacements:
        if not original or original == suggestion:
            continue
        start = 0
        while (index := text.find(original, start)) >= 0:
            matches.append((index, index + len(original), original, suggestion))
            start = index + len(original)
    matches.sort(key=lambda item: (item[0], -(item[1] - item[0])))
    filtered: list[tuple[int, int, str, str]] = []
    last_end = 0
    for match in matches:
        if match[0] >= last_end:
            filtered.append(match)
            last_end = match[1]
    return filtered


def _replace_all(text: str, replacements: list[tuple[str, str]]) -> str:
    for original, suggestion in replacements:
        text = text.replace(original, suggestion)
    return text


def _export_text(
    source_path: Path,
    replacements: list[tuple[str, str]],
    track_changes: bool,
) -> bytes:
    content, encoding = _read_text(source_path)
    for original, suggestion in replacements:
        if track_changes and original != suggestion:
            marker = (
                f"【删除:{original}→修改为:{suggestion}】"
                if suggestion
                else f"【删除:{original}】"
            )
            content = content.replace(original, marker)
        else:
            content = content.replace(original, suggestion)
    return content.encode(encoding)


def _read_text(source_path: Path) -> tuple[str, str]:
    for encoding in ("utf-8", "gbk", "gb2312", "big5", "utf-16", "latin-1"):
        try:
            return source_path.read_text(encoding=encoding), encoding
        except UnicodeError:
            continue
    return source_path.read_text(encoding="utf-8", errors="ignore"), "utf-8"


def _apply_text_edits(content: str, edits: list[TextEdit], track_changes: bool) -> str:
    for edit in sorted(edits, key=lambda item: (item.start, item.end), reverse=True):
        if edit.start < 0 or edit.end < edit.start or edit.end > len(content):
            raise ExportError("An edit position is outside the original document.")
        original = content[edit.start:edit.end]
        replacement = edit.replacement
        if track_changes:
            replacement = (
                f"【删除:{original}→修改为:{replacement}】"
                if replacement
                else f"【删除:{original}】"
            )
        content = content[:edit.start] + replacement + content[edit.end:]
    return content


def _export_text_edits(
    source_path: Path,
    original_text: str,
    edits: list[TextEdit],
    track_changes: bool,
) -> bytes:
    content, encoding = _read_text(source_path)
    if strip_html(content) != original_text:
        raise ExportError("The stored text no longer matches the analyzed document.")
    return _apply_text_edits(content, edits, track_changes).encode(encoding)


def _validate_and_order_pdf_edits(
    edits: list[TextEdit],
    text_length: int,
) -> list[TextEdit]:
    ordered = sorted(edits, key=lambda edit: (edit.start, edit.end, edit.replacement))
    active_range_start: int | None = None
    active_range_end: int | None = None
    previous_insertion_position: int | None = None
    for edit in ordered:
        if edit.start < 0 or edit.end < edit.start or edit.end > text_length:
            raise ExportError("An edit position is outside the original document.")
        is_insertion = edit.start == edit.end
        if is_insertion:
            if previous_insertion_position == edit.start:
                raise ExportError(_PDF_EDIT_CONFLICT_ERROR)
            if (
                active_range_start is not None
                and active_range_end is not None
                and active_range_start < edit.start < active_range_end
            ):
                raise ExportError(_PDF_EDIT_CONFLICT_ERROR)
            previous_insertion_position = edit.start
            continue
        if active_range_end is not None and edit.start < active_range_end:
            raise ExportError(_PDF_EDIT_CONFLICT_ERROR)
        active_range_start = edit.start
        active_range_end = edit.end
    return ordered


def _export_pdf(
    source_path: Path,
    replacements: list[tuple[str, str]],
    track_changes: bool,
) -> bytes:
    try:
        import fitz
    except ImportError as error:
        raise ExportError("PDF export requires PyMuPDF.") from error

    document = fitz.open(source_path)
    try:
        for page in document:
            if track_changes:
                _annotate_pdf_page(page, replacements)
            else:
                _replace_pdf_page(page, replacements)
        return document.tobytes(garbage=4, deflate=True)
    finally:
        document.close()


def _export_pdf_edits(
    source_path: Path,
    original_text: str,
    edits: list[TextEdit],
    track_changes: bool,
    *,
    coalesce_modified_text: bool = False,
) -> bytes:
    try:
        import fitz
    except ImportError as error:
        raise ExportError("PDF export requires PyMuPDF.") from error

    document = fitz.open(source_path)
    try:
        canonical_document = PdfParser().parse(source_path)
        if canonical_document.text != original_text:
            raise ExportError("The stored PDF no longer matches the analyzed document.")
        ordered_edits = _validate_and_order_pdf_edits(edits, len(original_text))
        if coalesce_modified_text:
            ordered_edits = _coalesce_pdf_modified_text_edits(
                ordered_edits,
                original_text,
                canonical_document,
            )
            ordered_edits = _validate_and_order_pdf_edits(
                ordered_edits,
                len(original_text),
            )
        resolved: list[tuple[object, list[object], object, str, str]] = []
        for edit in ordered_edits:
            original = original_text[edit.start:edit.end]
            if not original:
                raise ExportError(
                    "PDF original-format export does not support insertion-only edits; "
                    "replace existing text or export the edited text as TXT."
                )
            page, rectangles, anchor, replacement_compatible = _pdf_edit_location(
                document,
                canonical_document,
                edit,
            )
            if not track_changes and not replacement_compatible:
                raise ExportError(_PDF_DIRECTIONAL_COMPATIBILITY_ERROR)
            resolved.append((page, rectangles, anchor, original, edit.replacement))

        redactions: list[tuple[object, object, str, float]] = []
        for page, rectangles, anchor, original, suggestion in resolved:
            if track_changes:
                for rectangle in rectangles:
                    annotation = page.add_highlight_annot(rectangle)
                    annotation.set_info(
                        title="啄木鸟·中英文字智能检查",
                        content=f"原文: {original}\n建议: {suggestion or '删除'}",
                    )
                    annotation.update()
                continue
            blocks = page.get_text("dict").get("blocks", [])
            font_size = _font_size_at(blocks, anchor)
            for rectangle in rectangles:
                page.add_redact_annot(rectangle)
            redactions.append((page, anchor, suggestion, font_size))

        if not track_changes:
            for page in document:
                page.apply_redactions()
            for page, rectangle, suggestion, font_size in redactions:
                if not suggestion:
                    continue
                insertion_box = fitz.Rect(
                    rectangle.x0,
                    rectangle.y0,
                    page.rect.x1 - 2,
                    rectangle.y1 + font_size,
                )
                _insert_pdf_textbox(
                    page,
                    insertion_box,
                    suggestion,
                    font_size,
                )
        return document.tobytes(garbage=4, deflate=True)
    finally:
        document.close()


def _insert_pdf_textbox(
    page: object,
    insertion_box: object,
    suggestion: str,
    font_size: float,
) -> None:
    try:
        result = page.insert_textbox(  # type: ignore[attr-defined]
            insertion_box,
            suggestion,
            fontsize=font_size,
            fontname="china-s" if _has_cjk(suggestion) else "helv",
        )
    except Exception:
        result = page.insert_textbox(  # type: ignore[attr-defined]
            insertion_box,
            suggestion,
            fontsize=font_size,
            fontname="helv",
        )
    if (
        isinstance(result, bool)
        or not isinstance(result, int | float)
        or not isfinite(float(result))
        or result < 0
    ):
        raise ExportError(_PDF_INSERTION_SPACE_ERROR)


def _pdf_edit_location(
    pdf: object,
    document: object,
    edit: TextEdit,
) -> tuple[object, list[object], object, bool]:
    for block in document.blocks:  # type: ignore[attr-defined]
        if not (block.global_start <= edit.start and edit.end <= block.global_end):
            continue
        locator = block.source_locator
        page_number = locator.get("page")
        characters = locator.get("characters")
        if (
            isinstance(page_number, bool)
            or not isinstance(page_number, int)
            or not isinstance(characters, list)
        ):
            raise ExportError(_PDF_MAPPING_VALIDATION_ERROR)
        local_start = edit.start - block.global_start
        local_end = edit.end - block.global_start
        try:
            page = pdf[page_number - 1]  # type: ignore[index]
        except (IndexError, TypeError):
            raise ExportError(_PDF_MAPPING_VALIDATION_ERROR) from None
        glyphs = _validated_pdf_glyphs(
            page,
            block.text,
            characters,
            local_start,
            local_end,
        )
        rectangles = [
            glyph.rectangle
            for glyph in _coalesce_pdf_glyph_rectangles(glyphs)
        ]
        if not rectangles:
            raise ExportError(_PDF_MAPPING_VALIDATION_ERROR)
        return (
            page,
            rectangles,
            rectangles[0],
            all(_pdf_glyph_is_horizontal_ltr(glyph) for glyph in glyphs),
        )
    raise ExportError(_PDF_MAPPING_VALIDATION_ERROR)


def _validated_pdf_glyphs(
    page: object,
    block_text: str,
    characters: list[object],
    local_start: int,
    local_end: int,
) -> list[_PdfGlyphRectangle]:
    if local_start < 0 or local_end > len(block_text) or local_end <= local_start:
        raise ExportError(_PDF_MAPPING_VALIDATION_ERROR)

    selected: list[dict[str, object]] = []
    covered_offsets: set[int] = set()
    group_ids: set[str] = set()
    for value in characters:
        if not isinstance(value, dict):
            continue
        source_start = value.get("source_start")
        source_end = value.get("source_end")
        if (
            isinstance(source_start, bool)
            or not isinstance(source_start, int)
            or isinstance(source_end, bool)
            or not isinstance(source_end, int)
            or source_start >= local_end
            or source_end <= local_start
        ):
            continue
        text = value.get("text")
        line_index = value.get("line_index")
        group_id = value.get("group_id")
        if (
            not isinstance(text, str)
            or not text
            or source_end != source_start + len(text)
            or source_start < 0
            or source_end > len(block_text)
            or block_text[source_start:source_end] != text
            or isinstance(line_index, bool)
            or not isinstance(line_index, int)
            or line_index < 0
            or not isinstance(group_id, str)
            or not group_id
            or group_id in group_ids
            or source_start < local_start
            or source_end > local_end
        ):
            raise ExportError(_PDF_MAPPING_VALIDATION_ERROR)
        group_ids.add(group_id)
        for offset in range(source_start, source_end):
            if offset in covered_offsets:
                raise ExportError(_PDF_MAPPING_VALIDATION_ERROR)
            covered_offsets.add(offset)
        selected.append(value)

    if covered_offsets != set(range(local_start, local_end)):
        raise ExportError(_PDF_MAPPING_VALIDATION_ERROR)

    glyphs: list[_PdfGlyphRectangle] = []
    for value in sorted(selected, key=lambda group: int(group["source_start"])):
        source_start = int(value["source_start"])
        source_end = int(value["source_end"])
        text = value.get("text")
        mapping_state = value.get("mapping_state")
        bbox_value = value.get("bbox")
        line_index = int(value["line_index"])
        assert isinstance(text, str)
        if text.isspace():
            if mapping_state != "synthetic_space" or bbox_value is not None:
                raise ExportError(_PDF_MAPPING_VALIDATION_ERROR)
            continue
        if mapping_state != "glyph":
            raise ExportError(_PDF_MAPPING_VALIDATION_ERROR)
        bbox = _validated_pdf_bbox(bbox_value)
        line_direction, writing_mode = _validated_pdf_source_direction(value)
        glyphs.append(
            _PdfGlyphRectangle(
                rectangle=_pdf_page_rectangle(page, bbox),
                source_start=source_start,
                source_end=source_end,
                line_index=line_index,
                line_direction=line_direction,
                writing_mode=writing_mode,
            )
        )
    return glyphs


def _validated_pdf_source_direction(
    value: dict[str, object],
) -> tuple[tuple[float, float], int]:
    direction_value = value.get("line_direction", (1.0, 0.0))
    writing_mode = value.get("writing_mode", 0)
    if (
        not isinstance(direction_value, tuple | list)
        or len(direction_value) != 2
        or isinstance(writing_mode, bool)
        or not isinstance(writing_mode, int)
        or writing_mode not in {0, 1}
    ):
        raise ExportError(_PDF_MAPPING_VALIDATION_ERROR)
    direction: list[float] = []
    for component in direction_value:
        if (
            isinstance(component, bool)
            or not isinstance(component, int | float)
            or not isfinite(float(component))
        ):
            raise ExportError(_PDF_MAPPING_VALIDATION_ERROR)
        direction.append(float(component))
    if direction == [0.0, 0.0]:
        raise ExportError(_PDF_MAPPING_VALIDATION_ERROR)
    return (direction[0], direction[1]), writing_mode


def _pdf_glyph_is_horizontal_ltr(glyph: _PdfGlyphRectangle) -> bool:
    return glyph.writing_mode == 0 and glyph.line_direction[0] > 0 and abs(
        glyph.line_direction[1]
    ) <= 1e-6


def _validated_pdf_bbox(value: object) -> tuple[float, float, float, float]:
    if not isinstance(value, list | tuple) or len(value) != 4:
        raise ExportError(_PDF_MAPPING_VALIDATION_ERROR)
    coordinates: list[float] = []
    for coordinate in value:
        if (
            isinstance(coordinate, bool)
            or not isinstance(coordinate, int | float)
            or not isfinite(float(coordinate))
        ):
            raise ExportError(_PDF_MAPPING_VALIDATION_ERROR)
        coordinates.append(float(coordinate))
    x0, y0, x1, y1 = coordinates
    if x1 <= x0 or y1 <= y0:
        raise ExportError(_PDF_MAPPING_VALIDATION_ERROR)
    return x0, y0, x1, y1


def _pdf_page_rectangle(
    page: object,
    bbox: tuple[float, float, float, float],
) -> object:
    import fitz

    rectangle = fitz.Rect(bbox) * ~page.rotation_matrix  # type: ignore[operator, union-attr]
    if rectangle.is_empty or rectangle.is_infinite:
        raise ExportError(_PDF_MAPPING_VALIDATION_ERROR)
    return rectangle


def _coalesce_pdf_glyph_rectangles(
    glyphs: list[_PdfGlyphRectangle],
) -> list[_PdfGlyphRectangle]:
    if not glyphs:
        return []
    coalesced = [glyphs[0]]
    for glyph in glyphs[1:]:
        previous = coalesced[-1]
        if not _pdf_glyph_rectangles_are_adjacent(previous, glyph):
            coalesced.append(glyph)
            continue
        import fitz

        coalesced[-1] = _PdfGlyphRectangle(
            rectangle=fitz.Rect(
                min(previous.rectangle.x0, glyph.rectangle.x0),
                min(previous.rectangle.y0, glyph.rectangle.y0),
                max(previous.rectangle.x1, glyph.rectangle.x1),
                max(previous.rectangle.y1, glyph.rectangle.y1),
            ),
            source_start=previous.source_start,
            source_end=glyph.source_end,
            line_index=previous.line_index,
            line_direction=previous.line_direction,
            writing_mode=previous.writing_mode,
        )
    return coalesced


def _pdf_glyph_rectangles_are_adjacent(
    previous: _PdfGlyphRectangle,
    following: _PdfGlyphRectangle,
) -> bool:
    if (
        previous.source_end != following.source_start
        or previous.line_index != following.line_index
        or previous.line_direction != following.line_direction
        or previous.writing_mode != following.writing_mode
    ):
        return False
    previous_rectangle = previous.rectangle
    following_rectangle = following.rectangle
    height = min(
        previous_rectangle.y1 - previous_rectangle.y0,
        following_rectangle.y1 - following_rectangle.y0,
    )
    overlap = min(previous_rectangle.y1, following_rectangle.y1) - max(
        previous_rectangle.y0,
        following_rectangle.y0,
    )
    tolerance = max(0.25, height * 0.05)
    horizontal_gap = following_rectangle.x0 - previous_rectangle.x1
    return (
        height > 0
        and overlap >= height * 0.5
        and -tolerance <= horizontal_gap <= tolerance
        and following_rectangle.x0 >= previous_rectangle.x0 - tolerance
    )


def _replace_pdf_page(page: object, replacements: list[tuple[str, str]]) -> None:
    redactions: list[tuple[object, str, float]] = []
    blocks = page.get_text("dict").get("blocks", [])  # type: ignore[attr-defined]
    for original, suggestion in replacements:
        for rectangle in page.search_for(original):  # type: ignore[attr-defined]
            font_size = _font_size_at(blocks, rectangle)
            redactions.append((rectangle, suggestion, font_size))
            page.add_redact_annot(rectangle)  # type: ignore[attr-defined]
    page.apply_redactions()  # type: ignore[attr-defined]
    for rectangle, suggestion, font_size in redactions:
        if not suggestion:
            continue
        point = (rectangle.x0, rectangle.y1 - 2)
        try:
            page.insert_text(  # type: ignore[attr-defined]
                point,
                suggestion,
                fontsize=font_size,
                fontname="china-s" if _has_cjk(suggestion) else "helv",
            )
        except Exception:
            page.insert_text(point, suggestion, fontsize=font_size, fontname="helv")  # type: ignore[attr-defined]


def _annotate_pdf_page(page: object, replacements: list[tuple[str, str]]) -> None:
    for original, suggestion in replacements:
        if original == suggestion:
            continue
        for rectangle in page.search_for(original):  # type: ignore[attr-defined]
            annotation = page.add_highlight_annot(rectangle)  # type: ignore[attr-defined]
            recommendation = f"原文: {original}\n建议: {suggestion or '删除'}"
            annotation.set_info(title="啄木鸟·中英文字智能检查", content=recommendation)
            annotation.update()


def _font_size_at(blocks: list[dict[str, object]], rectangle: object) -> float:
    import fitz

    for block in blocks:
        for line in block.get("lines", []):  # type: ignore[union-attr]
            for span in line.get("spans", []):
                if fitz.Rect(span["bbox"]).intersects(rectangle):
                    return float(span["size"])
    return 11.0


def _has_cjk(text: str) -> bool:
    return any("\u3400" <= character <= "\u9fff" for character in text)


def _export_rtf(
    source_path: Path,
    replacements: list[tuple[str, str]],
    track_changes: bool,
) -> bytes:
    content, _ = _read_text(source_path)
    for original, suggestion in replacements:
        encoded_original = _rtf_unicode(original)
        encoded_suggestion = _rtf_unicode(suggestion)
        if track_changes and original != suggestion:
            replacement = (
                "{\\strike " + original + "}{\\uld " + suggestion + "}"
                if suggestion
                else "{\\strike " + original + "}"
            )
            encoded_replacement = (
                "{\\strike " + encoded_original + "}{\\uld " + encoded_suggestion + "}"
                if suggestion
                else "{\\strike " + encoded_original + "}"
            )
        else:
            replacement = suggestion
            encoded_replacement = encoded_suggestion
        if original in content:
            content = content.replace(original, replacement)
        else:
            content = content.replace(encoded_original, encoded_replacement)
    return content.encode("utf-8")


def _export_rtf_edits(
    source_path: Path,
    original_text: str,
    edits: list[TextEdit],
    track_changes: bool,
) -> bytes:
    raw = source_path.read_bytes().decode("latin-1")
    parsed, spans = decode_rtf_with_spans(raw)
    normalized = parsed.strip()
    if normalized != original_text or len(parsed) != len(spans):
        raise ExportError("The stored RTF no longer matches the analyzed document.")

    leading = len(parsed) - len(parsed.lstrip())
    for edit in sorted(edits, key=lambda item: (item.start, item.end), reverse=True):
        start = edit.start + leading
        end = edit.end + leading
        if start < 0 or end < start or end > len(spans):
            raise ExportError("An RTF edit position is outside the original document.")
        raw_start = spans[start][0] if start < len(spans) else raw.rfind("}")
        raw_end = spans[end - 1][1] if end > start else raw_start
        original = parsed[start:end]
        encoded_replacement = _rtf_unicode(edit.replacement)
        if track_changes:
            encoded_original = _rtf_unicode(original)
            encoded_replacement = (
                "{\\strike "
                + encoded_original
                + "}{\\uld "
                + encoded_replacement
                + "}"
                if edit.replacement
                else "{\\strike " + encoded_original + "}"
            )
        raw = raw[:raw_start] + encoded_replacement + raw[raw_end:]
    return raw.encode("latin-1")


def _rtf_unicode(text: str) -> str:
    result: list[str] = []
    index = 0
    while index < len(text):
        character = text[index]
        if character == "\r":
            if index + 1 < len(text) and text[index + 1] == "\n":
                index += 1
            result.append("\\par ")
            index += 1
            continue
        if character == "\n":
            result.append("\\par ")
            index += 1
            continue
        codepoint = ord(character)
        if codepoint > 0xFFFF:
            value = codepoint - 0x10000
            high = 0xD800 + (value >> 10)
            low = 0xDC00 + (value & 0x3FF)
            result.append(f"\\u{high - 65536}?\\u{low - 65536}?")
        elif codepoint > 127:
            signed_codepoint = codepoint if codepoint < 32768 else codepoint - 65536
            result.append(f"\\u{signed_codepoint}?")
        elif character in {"\\", "{", "}"}:
            result.append("\\" + character)
        else:
            result.append(character)
        index += 1
    return "".join(result)


def _convert_docx_bytes_to_doc(content: bytes, work_directory: Path) -> bytes:
    source_path = work_directory / f"export-{uuid4()}.docx"
    output_path = source_path.with_suffix(".doc")
    source_path.write_bytes(content)
    try:
        if shutil.which("textutil"):
            subprocess.run(
                ["textutil", "-convert", "doc", str(source_path), "-output", str(output_path)],
                check=True,
                capture_output=True,
                timeout=120,
            )
        else:
            office = shutil.which("soffice") or shutil.which("libreoffice")
            if not office:
                raise ExportError("DOC export requires textutil or LibreOffice.")
            subprocess.run(
                [
                    office,
                    "--headless",
                    "--convert-to",
                    "doc:MS Word 97",
                    "--outdir",
                    str(work_directory),
                    str(source_path),
                ],
                check=True,
                capture_output=True,
                timeout=180,
            )
        if not output_path.is_file() or output_path.stat().st_size == 0:
            raise ExportError("DOC conversion did not produce an output file.")
        return output_path.read_bytes()
    except (OSError, subprocess.CalledProcessError, subprocess.TimeoutExpired) as error:
        raise ExportError("Failed to convert the modified document back to DOC.") from error
    finally:
        source_path.unlink(missing_ok=True)
        output_path.unlink(missing_ok=True)
