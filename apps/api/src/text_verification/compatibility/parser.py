# mypy: ignore-errors
# ruff: noqa
"""
文件解析模块
支持 txt、docx、pdf、rtf 等格式的文件解析。
解析时追踪文本在原文中的位置信息（页码/段落号），方便用户定位。
"""

import os
import re
import signal
import shutil
import subprocess
import uuid
import codecs
from html.parser import HTMLParser
from pathlib import Path, PurePosixPath
from time import monotonic
from typing import Tuple, Optional, List
import zipfile


MAX_DOC_SOURCE_BYTES = 25 * 1024 * 1024
MAX_DOC_CONVERTED_BYTES = 100 * 1024 * 1024
MAX_DOC_EXPANDED_BYTES = 200 * 1024 * 1024
MAX_DOC_SINGLE_ENTRY_BYTES = 100 * 1024 * 1024
MAX_DOC_ARCHIVE_ENTRIES = 10_000
MAX_DOC_PARSED_ELEMENTS = 100_000
MAX_DOC_PARSED_TEXT_CHARS = 5_000_000
DOC_CONVERSION_TIMEOUT_SECONDS = 180.0
DOC_CONVERSION_POLL_SECONDS = 0.1


class _HTMLStripper(HTMLParser):
    """用于剥离 HTML/XML 标签，只保留可见文本。"""

    def __init__(self):
        super().__init__()
        self.text = []

    def handle_data(self, data):
        self.text.append(data)

    def get_data(self):
        return ''.join(self.text)


def _strip_html(text: str) -> str:
    """
    安全地移除文本中的 HTML/XML 标签，只保留可见文本。

    如果标准库解析失败，则退回到正则表达式兜底。
    同时修复标签被错误截断后留下的零散属性文本（如 class="..."）。
    """
    if not text or '<' not in text or '>' not in text:
        return text
    try:
        stripper = _HTMLStripper()
        stripper.feed(text)
        result = stripper.get_data()
    except Exception:
        result = re.sub(r'<[^>]+>', '', text)

    # 防御性清理：去除标签被错误截断后残留的零散 HTML 属性片段
    # 例如 "class=\"issue-highlight\""、"data-issue-id=..." 等
    result = re.sub(r'\b[a-zA-Z][\w\-]*\s*=\s*"[^"]*"', '', result)
    result = re.sub(r'\b[a-zA-Z][\w\-]*\s*=\s*\'[^\']*\'', '', result)
    result = re.sub(r'[<>]', '', result)
    return result


# page_map 中每项为 (start_offset, end_offset, label)
PageMap = List[Tuple[int, int, str]]


def parse_file(
    file_path: str,
    file_extension: Optional[str] = None,
    work_directory: Optional[str] = None,
) -> Tuple[str, str, PageMap]:
    """
    解析文件，提取纯文本内容，并记录各段文本在原文中的位置信息。

    Args:
        file_path: 文件路径
        file_extension: 文件扩展名（如未提供则从路径推断）

    Returns:
        (提取的文本内容, 文件格式名称, page_map)
        page_map 为 (起始偏移, 结束偏移, 位置标签) 列表，如 [(0, 120, '第1页'), ...]
        对于不支持位置追踪的格式，page_map 为空列表。
    """
    if file_extension is None:
        file_extension = os.path.splitext(file_path)[1].lower().lstrip('.')

    if file_extension == 'txt':
        return _parse_txt(file_path), 'txt', []
    elif file_extension == 'docx':
        text, page_map = _parse_docx(file_path)
        return text, 'docx', page_map
    elif file_extension == 'doc':
        # 旧版 Word 二进制格式（.doc）：先转为 .docx 再复用 docx 解析逻辑
        text, page_map = _parse_doc(file_path, work_directory)
        return text, 'doc', page_map
    elif file_extension == 'pdf':
        text, page_map = _parse_pdf(file_path)
        return text, 'pdf', page_map
    elif file_extension == 'rtf':
        return _parse_rtf(file_path), 'rtf', []
    elif file_extension == 'md':
        return _parse_txt(file_path), 'markdown', []
    elif file_extension == 'csv':
        return _parse_txt(file_path), 'csv', []
    else:
        # 默认尝试以文本方式读取
        try:
            return _parse_txt(file_path), file_extension, []
        except Exception:
            raise ValueError(f'不支持的文件格式: .{file_extension}，目前支持 txt、docx、doc、pdf、rtf、md、csv')


def _parse_txt(file_path: str) -> str:
    """解析纯文本文件，自动检测编码，并清理可能混入的 HTML 代码"""
    encodings = ['utf-8', 'gbk', 'gb2312', 'big5', 'utf-16', 'latin-1']
    content = None
    for encoding in encodings:
        try:
            with open(file_path, 'r', encoding=encoding) as f:
                content = f.read()
            break
        except (UnicodeDecodeError, UnicodeError):
            continue
    if content is None:
        # 如果所有编码都失败，用 utf-8 忽略错误
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            content = f.read()
    return _strip_html(content)


def _parse_docx(file_path: str) -> Tuple[str, PageMap]:
    """解析 Word (.docx) 文件，追踪段落号，并清理混入的 HTML 代码"""
    from docx import Document
    _validate_converted_docx(Path(file_path))
    doc = Document(file_path)
    paragraphs = []
    page_map: PageMap = []
    offset = 0
    para_num = 0
    parsed_elements = 0
    parsed_text_chars = 0

    document_paragraphs = []
    seen_paragraphs = set()

    def append_paragraph(para) -> None:
        nonlocal parsed_elements, parsed_text_chars
        element_id = id(para._element)
        if element_id in seen_paragraphs:
            return
        parsed_elements += 1
        if parsed_elements > MAX_DOC_PARSED_ELEMENTS:
            raise ValueError('DOCX parsed element limit exceeded')
        parsed_text_chars += len(para.text)
        if parsed_text_chars > MAX_DOC_PARSED_TEXT_CHARS:
            raise ValueError('DOCX parsed text size limit exceeded')
        document_paragraphs.append(para)
        seen_paragraphs.add(element_id)

    for para in doc.paragraphs:
        append_paragraph(para)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    append_paragraph(para)

    for para in document_paragraphs:
        text = _strip_html(para.text).strip()
        if text:
            para_num += 1
            page_map.append((offset, offset + len(text), f'第{para_num}段'))
            paragraphs.append(text)
            offset += len(text) + 1  # +1 for \n

    return '\n'.join(paragraphs), page_map


def _convert_doc_to_docx(file_path: str, work_directory: Optional[str] = None) -> str:
    """将旧版 Word (.doc) 二进制格式转换为 .docx，便于解析与导出。

    优先使用 macOS 内置的 textutil；若不可用则尝试 LibreOffice (soffice)。
    返回转换后的 .docx 临时文件路径（调用方负责清理）。
    """
    raw_source = Path(file_path)
    if raw_source.is_symlink():
        raise ValueError('DOC source must be a regular file')
    source = raw_source.resolve(strict=True)
    if not source.is_file():
        raise ValueError('DOC source must be a regular file')
    if source.stat().st_size <= 0 or source.stat().st_size > MAX_DOC_SOURCE_BYTES:
        raise ValueError('DOC source exceeds the configured size limit')

    work_dir = Path(work_directory or source.parent).resolve(strict=False)
    work_dir.mkdir(parents=True, exist_ok=True)
    workspace = work_dir / f'doc-conversion-{uuid.uuid4().hex}'
    workspace.mkdir(mode=0o700)
    staged_source = workspace / 'source.doc'
    _copy_bounded(source, staged_source, MAX_DOC_SOURCE_BYTES)
    output = workspace / 'converted.docx'
    environment = _conversion_environment(workspace)
    attempted = False
    last_error = None

    textutil = shutil.which('textutil')
    if textutil:
        attempted = True
        try:
            _run_conversion_process(
                [
                    textutil,
                    '-convert',
                    'docx',
                    str(staged_source),
                    '-output',
                    str(output),
                ],
                output_path=output,
                cwd=workspace,
                environment=environment,
                timeout_seconds=DOC_CONVERSION_TIMEOUT_SECONDS,
            )
            _validate_converted_docx(output)
            return str(output)
        except (OSError, ValueError) as error:
            last_error = error
            output.unlink(missing_ok=True)

    soffice = shutil.which('soffice') or shutil.which('libreoffice')
    if soffice:
        attempted = True
        candidate = workspace / 'source.docx'
        profile = workspace / 'libreoffice-profile'
        profile.mkdir(mode=0o700)
        try:
            _run_conversion_process(
                [
                    soffice,
                    f'-env:UserInstallation={profile.as_uri()}',
                    '--headless',
                    '--invisible',
                    '--nologo',
                    '--nodefault',
                    '--nolockcheck',
                    '--nofirststartwizard',
                    '--norestore',
                    '--safe-mode',
                    '--convert-to',
                    'docx',
                    '--outdir',
                    str(workspace),
                    str(staged_source),
                ],
                output_path=candidate,
                cwd=workspace,
                environment=environment,
                timeout_seconds=DOC_CONVERSION_TIMEOUT_SECONDS,
            )
            _validate_converted_docx(candidate)
            candidate.replace(output)
            return str(output)
        except (OSError, ValueError) as error:
            last_error = error
            candidate.unlink(missing_ok=True)

    shutil.rmtree(workspace, ignore_errors=True)
    if attempted:
        if isinstance(last_error, ValueError):
            raise last_error
        raise ValueError('DOC conversion did not produce a valid DOCX file') from last_error
    raise ValueError('无法处理 .doc 文件：本机未安装 textutil 或 LibreOffice，请先安装其中之一')


def _run_conversion_process(
    command,
    *,
    output_path: Path,
    cwd: Path,
    environment: dict[str, str],
    timeout_seconds: float,
) -> None:
    process = subprocess.Popen(
        command,
        cwd=cwd,
        env=environment,
        stdin=subprocess.DEVNULL,
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        start_new_session=True,
    )
    deadline = monotonic() + timeout_seconds
    while True:
        remaining = deadline - monotonic()
        if remaining <= 0:
            _terminate_conversion_process(process)
            raise ValueError('DOC conversion timed out')
        try:
            return_code = process.wait(
                timeout=min(DOC_CONVERSION_POLL_SECONDS, remaining)
            )
            break
        except subprocess.TimeoutExpired:
            if output_path.exists() and output_path.stat().st_size > MAX_DOC_CONVERTED_BYTES:
                _terminate_conversion_process(process)
                raise ValueError('Converted DOCX exceeds the configured size limit')
    if return_code != 0:
        raise ValueError('DOC conversion process failed')
    if not output_path.is_file():
        raise ValueError('DOC conversion did not produce a valid DOCX file')
    if output_path.stat().st_size <= 0 or output_path.stat().st_size > MAX_DOC_CONVERTED_BYTES:
        raise ValueError('Converted DOCX exceeds the configured size limit')


def _terminate_conversion_process(process) -> None:
    if process.poll() is not None:
        return
    if os.name == 'posix':
        os.killpg(process.pid, signal.SIGTERM)
    else:
        process.terminate()
    if process.poll() is not None:
        return
    try:
        process.wait(timeout=5)
    except subprocess.TimeoutExpired:
        if os.name == 'posix':
            os.killpg(process.pid, signal.SIGKILL)
        else:
            process.kill()
        process.wait(timeout=5)


def _conversion_environment(workspace: Path) -> dict[str, str]:
    blocked = {
        'http_proxy',
        'https_proxy',
        'all_proxy',
        'ftp_proxy',
        'no_proxy',
    }
    environment = {
        key: value
        for key, value in os.environ.items()
        if key.lower() not in blocked
    }
    environment.update(
        {
            'HOME': str(workspace),
            'TMPDIR': str(workspace),
            'SAL_USE_VCLPLUGIN': 'svp',
        }
    )
    return environment


def _copy_bounded(source: Path, target: Path, maximum: int) -> None:
    copied = 0
    with source.open('rb') as input_file, target.open('xb') as output_file:
        while chunk := input_file.read(1024 * 1024):
            copied += len(chunk)
            if copied > maximum:
                raise ValueError('DOC source exceeds the configured size limit')
            output_file.write(chunk)


def _validate_converted_docx(path: Path) -> None:
    if not path.is_file():
        raise ValueError('DOC conversion did not produce a valid DOCX file')
    size = path.stat().st_size
    if size <= 0 or size > MAX_DOC_CONVERTED_BYTES:
        raise ValueError('Converted DOCX exceeds the configured size limit')
    try:
        with zipfile.ZipFile(path) as archive:
            infos = archive.infolist()
            if len(infos) > MAX_DOC_ARCHIVE_ENTRIES:
                raise ValueError('Converted DOCX has too many archive entries')
            total = 0
            names = set()
            for info in infos:
                if info.flag_bits & 0x1:
                    raise ValueError('Converted DOCX contains an encrypted entry')
                if info.file_size > MAX_DOC_SINGLE_ENTRY_BYTES:
                    raise ValueError('Converted DOCX entry exceeds the size limit')
                total += info.file_size
                if total > MAX_DOC_EXPANDED_BYTES:
                    raise ValueError('Converted DOCX expanded size limit exceeded')
                if _unsafe_archive_name(info.filename):
                    raise ValueError('Converted DOCX contains an unsafe archive path')
                names.add(info.filename)
            if '[Content_Types].xml' not in names or 'word/document.xml' not in names:
                raise ValueError('DOC conversion did not produce a valid DOCX file')
            if archive.testzip() is not None:
                raise ValueError('DOC conversion did not produce a valid DOCX file')
    except zipfile.BadZipFile as error:
        raise ValueError('DOC conversion did not produce a valid DOCX file') from error


def _unsafe_archive_name(name: str) -> bool:
    if (
        not name
        or name.startswith(('/', '\\'))
        or '\\' in name
        or (len(name) >= 2 and name[1] == ':')
    ):
        return True
    return any(part in {'', '.', '..'} for part in PurePosixPath(name).parts)


def _parse_doc(file_path: str, work_directory: Optional[str] = None) -> Tuple[str, PageMap]:
    """解析旧版 Word (.doc) 文件：借助 textutil/LibreOffice 转为 docx 后复用 docx 解析。"""
    docx_path = _convert_doc_to_docx(file_path, work_directory)
    try:
        return _parse_docx(docx_path)
    finally:
        try:
            if docx_path:
                shutil.rmtree(Path(docx_path).parent)
        except OSError:
            pass


def _parse_pdf(file_path: str) -> Tuple[str, PageMap]:
    """解析 PDF 文件，追踪页码，并清理混入的 HTML 代码"""
    import pdfplumber
    texts = []
    page_map: PageMap = []
    offset = 0
    with pdfplumber.open(file_path) as pdf:
        for i, page in enumerate(pdf.pages, 1):
            text = _strip_html(page.extract_text() or '').strip()
            if text:
                page_map.append((offset, offset + len(text), f'第{i}页'))
                texts.append(text)
                offset += len(text) + 1  # +1 for \n join
    return '\n'.join(texts), page_map


def _parse_rtf(file_path: str) -> str:
    """解析常见 ANSI/Unicode RTF 文本。"""
    content = open(file_path, 'rb').read().decode('latin-1')
    text, _ = decode_rtf_with_spans(content)
    return text.strip()


def decode_rtf_with_spans(content: str) -> Tuple[str, List[Tuple[int, int]]]:
    """返回 RTF 可见文本及每个字符在原始 RTF 字符串中的范围。"""
    destinations = {
        'colortbl', 'datastore', 'filetbl', 'fonttbl', 'footer', 'footerf',
        'footerl', 'footerr', 'header', 'headerf', 'headerl', 'headerr',
        'info', 'listoverridetable', 'listtable', 'object', 'pict',
        'stylesheet', 'themedata', 'xmlnstbl',
    }
    states = [{'skip': False, 'uc': 1, 'codec': 'cp1252'}]
    output: List[str] = []
    spans: List[Tuple[int, int]] = []
    index = 0
    fallback = 0

    def emit(text: str, start: int, end: int) -> None:
        if states[-1]['skip']:
            return
        output.extend(text)
        spans.extend([(start, end)] * len(text))

    while index < len(content):
        character = content[index]
        if character == '{':
            states.append(dict(states[-1]))
            index += 1
            continue
        if character == '}':
            if len(states) > 1:
                states.pop()
            index += 1
            continue
        if character != '\\':
            if fallback:
                fallback -= 1
            elif character not in '\r\n' and not states[-1]['skip']:
                raw = character.encode('latin-1')
                try:
                    decoded = raw.decode(str(states[-1]['codec']))
                except UnicodeError:
                    decoded = character
                emit(decoded, index, index + 1)
            index += 1
            continue

        token_start = index
        index += 1
        if index >= len(content):
            break
        symbol = content[index]
        if symbol in '\\{}':
            if fallback:
                fallback -= 1
            else:
                emit(symbol, token_start, index + 1)
            index += 1
            continue
        if symbol == "'":
            byte_tokens: List[Tuple[int, int, int]] = []
            while index + 2 < len(content) and content[index] == "'":
                try:
                    value = int(content[index + 1:index + 3], 16)
                except ValueError:
                    break
                byte_tokens.append((value, index - 1, index + 3))
                index += 3
                if index >= len(content) or content[index] != '\\':
                    break
                index += 1
            if fallback:
                fallback = max(0, fallback - len(byte_tokens))
                continue
            decoder = codecs.getincrementaldecoder(str(states[-1]['codec']))(errors='replace')
            pending_start = byte_tokens[0][1] if byte_tokens else token_start
            for value, start, end in byte_tokens:
                decoded = decoder.decode(bytes([value]), final=False)
                if decoded:
                    emit(decoded, pending_start, end)
                    pending_start = end
            decoded = decoder.decode(b'', final=True)
            if decoded and byte_tokens:
                emit(decoded, pending_start, byte_tokens[-1][2])
            continue
        if symbol == '*':
            states[-1]['skip'] = True
            index += 1
            continue
        if not symbol.isalpha():
            controls = {'~': '\u00a0', '-': '\u00ad', '_': '\u2011'}
            if not fallback and symbol in controls:
                emit(controls[symbol], token_start, index + 1)
            elif fallback:
                fallback -= 1
            index += 1
            continue

        match = re.match(r'([a-zA-Z]+)(-?\d+)? ?', content[index:])
        if not match:
            index += 1
            continue
        word = match.group(1)
        parameter = int(match.group(2)) if match.group(2) is not None else None
        index += len(match.group(0))
        if word in destinations:
            states[-1]['skip'] = True
        elif word == 'ansicpg' and parameter:
            states[-1]['codec'] = f'cp{parameter}'
        elif word == 'uc' and parameter is not None:
            states[-1]['uc'] = max(parameter, 0)
        elif word == 'u' and parameter is not None:
            codepoint = parameter if parameter >= 0 else parameter + 65536
            emit(chr(codepoint), token_start, index)
            fallback = int(states[-1]['uc'])
        elif word in {'par', 'line'}:
            emit('\n', token_start, index)
        elif word == 'tab':
            emit('\t', token_start, index)
        elif word == 'emdash':
            emit('\u2014', token_start, index)
        elif word == 'endash':
            emit('\u2013', token_start, index)

    combined_output: List[str] = []
    combined_spans: List[Tuple[int, int]] = []
    index = 0
    while index < len(output):
        codepoint = ord(output[index])
        if (
            0xD800 <= codepoint <= 0xDBFF
            and index + 1 < len(output)
            and 0xDC00 <= ord(output[index + 1]) <= 0xDFFF
        ):
            low = ord(output[index + 1])
            combined_output.append(chr(0x10000 + ((codepoint - 0xD800) << 10) + low - 0xDC00))
            combined_spans.append((spans[index][0], spans[index + 1][1]))
            index += 2
            continue
        combined_output.append(output[index])
        combined_spans.append(spans[index])
        index += 1
    return ''.join(combined_output), combined_spans


# 对外导出 HTML 清理函数，供导出模块复用
strip_html = _strip_html


def get_supported_formats() -> list:
    """返回支持的文件格式列表"""
    from text_verification.domain.capabilities import (
        CapabilityProfile,
        default_capability_manifest,
    )

    return default_capability_manifest().api_formats(
        CapabilityProfile.SYNCHRONOUS_COMPATIBILITY
    )
