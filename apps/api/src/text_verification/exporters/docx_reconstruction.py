from __future__ import annotations

import hashlib
import io
import math
import os
import stat
import struct
import zlib
from collections.abc import Mapping
from dataclasses import dataclass, fields
from datetime import UTC, datetime
from pathlib import Path
from typing import Any, Literal
from uuid import uuid4

import pymupdf
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.image.exceptions import (
    InvalidImageStreamError,
    UnexpectedEndOfFileError,
    UnrecognizedImageError,
)
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from text_verification.compatibility.exporters import ExportError
from text_verification.document_processing.pdf_models import PdfTextSpan
from text_verification.domain.documents import (
    DocumentModel,
    ExportFormat,
    FileType,
    TextBlock,
)
from text_verification.domain.ports import SourcePathResolver

DOCX_RECONSTRUCTION = ExportFormat.DOCX_RECONSTRUCTION
_DEFAULT_PAGE_WIDTH_POINTS = 612.0
_DEFAULT_PAGE_HEIGHT_POINTS = 792.0
_DEFAULT_MARGIN_POINTS = 54.0
_FIXED_CORE_PROPERTY_TIME = datetime(2000, 1, 1, tzinfo=UTC)
_XML_REPLACEMENT_CHARACTER = "\ufffd"
_PYMUPDF: Any = pymupdf

type SupportedImageMime = Literal["image/png", "image/jpeg"]


@dataclass(frozen=True)
class DocxReconstructionLimits:
    max_text_chars: int = 5_000_000
    max_output_elements: int = 20_000
    max_runs: int = 50_000
    max_images: int = 500
    max_tables: int = 100
    max_table_index: int = 10_000
    max_table_rows: int = 1_000
    max_table_columns: int = 100
    max_table_cells: int = 10_000
    max_image_bytes: int = 20 * 1024 * 1024
    max_total_image_bytes: int = 50 * 1024 * 1024
    max_pdf_source_bytes: int = 200 * 1024 * 1024
    max_image_width: int = 20_000
    max_image_height: int = 20_000
    max_image_pixels: int = 40_000_000
    max_decoded_image_bytes: int = 64 * 1024 * 1024
    max_output_bytes: int = 100 * 1024 * 1024

    def __post_init__(self) -> None:
        for field in fields(self):
            value = getattr(self, field.name)
            if isinstance(value, bool) or not isinstance(value, int) or value <= 0:
                raise ValueError(f"{field.name} must be a positive integer")


@dataclass(frozen=True)
class _Image:
    content: bytes
    mime_type: SupportedImageMime
    width: int
    height: int


@dataclass(frozen=True)
class _PdfImageRequest:
    block_id: str
    xref: int


@dataclass(frozen=True)
class _Table:
    page: int | None
    table_index: int
    rows: int
    columns: int
    cells: dict[tuple[int, int], TextBlock]
    merges: dict[tuple[int, int], tuple[int, int]]
    merged_coordinates: frozenset[tuple[int, int]]
    first_source_index: int
    last_source_index: int
    y: float
    x: float
    bottom: float


type _RenderUnit = TextBlock | _Table


class _BoundedBytesIO(io.BytesIO):
    def __init__(self, max_bytes: int) -> None:
        super().__init__()
        self._max_bytes = max_bytes

    def write(self, data: Any) -> int:
        size = len(data)
        if max(len(self.getbuffer()), self.tell() + size) > self._max_bytes:
            raise ExportError("Reconstructed DOCX exceeds the configured output size limit.")
        return super().write(data)


@dataclass(frozen=True)
class DocxReconstructionExporter:
    limits: DocxReconstructionLimits = DocxReconstructionLimits()
    source_path_resolver: SourcePathResolver | None = None
    file_type: ExportFormat = DOCX_RECONSTRUCTION

    def export(self, document: DocumentModel, target: Path) -> Path:
        _preflight_document(document, self.limits)
        units = _render_units(document, self.limits)
        images = _resolve_images(
            document,
            units,
            self.limits,
            self.source_path_resolver,
        )

        output = Document()
        _set_core_properties(output, document)
        _set_page_setup(output, document)

        for unit in units:
            if isinstance(unit, _Table):
                _add_table(output, unit)
            elif unit.kind == "image":
                _add_image(output, images[unit.block_id])
            elif unit.kind in {"heading", "paragraph"}:
                _add_text_block(output, unit)

        stream = _BoundedBytesIO(self.limits.max_output_bytes)
        output.save(stream)
        content = stream.getvalue()
        _publish_atomic(target, content)
        return target


def _preflight_document(
    document: DocumentModel,
    limits: DocxReconstructionLimits,
) -> None:
    if len(document.blocks) > limits.max_output_elements:
        raise ExportError("Document exceeds the configured output element limit.")
    text_characters = 0
    image_count = 0
    run_count = 0
    for block in document.blocks:
        text_characters += len(block.text)
        if text_characters > limits.max_text_chars:
            raise ExportError("Document exceeds the configured text character limit.")
        if block.kind == "image":
            image_count += 1
            if image_count > limits.max_images:
                raise ExportError("Document exceeds the configured image count limit.")
        if block.kind in {"heading", "paragraph"}:
            raw_runs = block.style.get("runs", block.style.get("spans"))
            run_count += len(raw_runs) if isinstance(raw_runs, list) else 1
            if block.style.get("spans") is not None:
                _validated_pdf_spans(block)
        elif block.kind == "table_cell":
            run_count += 1
        if run_count > limits.max_runs:
            raise ExportError("Document exceeds the configured run count limit.")


def _resolve_images(
    document: DocumentModel,
    units: list[_RenderUnit],
    limits: DocxReconstructionLimits,
    source_path_resolver: SourcePathResolver | None,
) -> dict[str, _Image]:
    image_blocks = [
        unit
        for unit in units
        if isinstance(unit, TextBlock) and unit.kind == "image"
    ]
    if not image_blocks:
        return {}
    requests = [
        _canonical_pdf_image_request(document, block)
        for block in image_blocks
    ]
    if source_path_resolver is None:
        raise ExportError("Canonical PDF images require an injected source resolver.")
    try:
        source_path = source_path_resolver.resolve(document)
    except (OSError, ValueError) as error:
        raise ExportError("Canonical PDF source could not be resolved.") from error
    source = _read_resolved_source(
        document,
        source_path,
        limits.max_pdf_source_bytes,
    )
    try:
        pdf: Any = _PYMUPDF.open(stream=source, filetype="pdf")
    except (RuntimeError, ValueError) as error:
        raise ExportError("Resolved PDF image source is invalid.") from error
    try:
        _preflight_pdf_image_streams(pdf, requests, limits)
        extracted: list[tuple[bytes, SupportedImageMime]] = []
        total_bytes = 0
        for request in requests:
            content, mime_type = _extract_pdf_image_content(pdf, request.xref)
            total_bytes += len(content)
            if total_bytes > limits.max_total_image_bytes:
                raise ExportError(
                    "Document exceeds the configured aggregate image media limit."
                )
            extracted.append((content, mime_type))
    finally:
        pdf.close()

    images: dict[str, _Image] = {}
    for request, (content, mime_type) in zip(requests, extracted, strict=True):
        image = _validated_image(content, limits)
        if image.mime_type != mime_type:
            raise ExportError("PDF image MIME type does not match its signature.")
        images[request.block_id] = image
    return images


def _preflight_pdf_image_streams(
    pdf: Any,
    requests: list[_PdfImageRequest],
    limits: DocxReconstructionLimits,
) -> None:
    total_bytes = 0
    for request in requests:
        try:
            if not pdf.xref_is_image(request.xref):
                raise ExportError("Canonical PDF image reference is not an image object.")
            value_type, raw_value = pdf.xref_get_key(request.xref, "Length")
            width_type, raw_width = pdf.xref_get_key(request.xref, "Width")
            height_type, raw_height = pdf.xref_get_key(request.xref, "Height")
        except (RuntimeError, ValueError) as error:
            raise ExportError("Canonical PDF image reference is invalid.") from error
        if value_type != "int" or width_type != "int" or height_type != "int":
            raise ExportError("Canonical PDF image stream metadata is unavailable.")
        try:
            image_bytes = int(raw_value)
            width = int(raw_width)
            height = int(raw_height)
        except (TypeError, ValueError) as error:
            raise ExportError("Canonical PDF image stream metadata is invalid.") from error
        if image_bytes <= 0 or image_bytes > limits.max_image_bytes:
            raise ExportError("PDF image stream exceeds the configured size limit.")
        if (
            width <= 0
            or height <= 0
            or width > limits.max_image_width
            or height > limits.max_image_height
            or width * height > limits.max_image_pixels
        ):
            raise ExportError("PDF image dimensions exceed configured limits.")
        if width * height * 8 > limits.max_decoded_image_bytes:
            raise ExportError("Image exceeds the configured decoded image byte limit.")
        total_bytes += image_bytes
        if total_bytes > limits.max_total_image_bytes:
            raise ExportError("Document exceeds the configured aggregate image media limit.")


def _extract_pdf_image_content(
    pdf: Any,
    xref: int,
) -> tuple[bytes, SupportedImageMime]:
    try:
        extracted = pdf.extract_image(xref)
    except (RuntimeError, ValueError) as error:
        raise ExportError("PDF image reference could not be extracted.") from error
    content = extracted.get("image")
    extension = extracted.get("ext")
    if not isinstance(content, bytes):
        raise ExportError("PDF image reference did not contain image data.")
    if extension in {"jpg", "jpeg"}:
        return content, "image/jpeg"
    if extension == "png":
        return content, "image/png"
    raise ExportError("PDF image reference uses an unsupported format.")


def _render_units(
    document: DocumentModel,
    limits: DocxReconstructionLimits,
) -> list[_RenderUnit]:
    table_groups: dict[tuple[int | None, int], list[tuple[int, TextBlock]]] = {}
    non_table: list[tuple[int, TextBlock]] = []
    for source_index, block in enumerate(document.blocks):
        if block.kind == "table_cell":
            if block.table_index is None:
                raise ExportError("Table cells require a table index.")
            if (
                isinstance(block.table_index, bool)
                or block.table_index < 0
                or block.table_index >= limits.max_table_index
            ):
                raise ExportError("Table index is outside configured limits.")
            table_groups.setdefault((block.page, block.table_index), []).append(
                (source_index, block)
            )
        elif block.kind in {"heading", "paragraph", "image"}:
            non_table.append((source_index, block))

    if len(table_groups) > limits.max_tables:
        raise ExportError("Document exceeds the configured table limit.")

    tables = [
        _table_from_blocks(document, key, blocks, limits)
        for key, blocks in table_groups.items()
    ]
    _reject_interleaved_table_blocks(tables, non_table)
    if len(non_table) + sum(table.rows * table.columns for table in tables) > (
        limits.max_output_elements
    ):
        raise ExportError("Document exceeds the configured output element limit.")

    indexed_units: list[tuple[tuple[float, ...], _RenderUnit]] = []
    indexed_units.extend(
        (_block_order(block, source_index), block)
        for source_index, block in non_table
    )
    indexed_units.extend((_table_order(table), table) for table in tables)
    indexed_units.sort(key=lambda item: item[0])
    return [unit for _, unit in indexed_units]


def _table_from_blocks(
    document: DocumentModel,
    key: tuple[int | None, int],
    indexed_blocks: list[tuple[int, TextBlock]],
    limits: DocxReconstructionLimits,
) -> _Table:
    page, table_index = key
    cells: dict[tuple[int, int], TextBlock] = {}
    declared_shapes: set[tuple[int, int]] = set()
    for _, block in indexed_blocks:
        if (
            block.row_index is None
            or block.cell_index is None
            or isinstance(block.row_index, bool)
            or isinstance(block.cell_index, bool)
        ):
            raise ExportError("Table cells require row and cell indices.")
        if (
            block.row_index < 0
            or block.cell_index < 0
            or block.row_index >= limits.max_table_rows
            or block.cell_index >= limits.max_table_columns
        ):
            raise ExportError("Table dimensions or cell coordinates exceed configured limits.")
        coordinate = (block.row_index, block.cell_index)
        if coordinate in cells:
            raise ExportError("Table contains duplicate cell coordinates.")
        cells[coordinate] = block
        shape = _declared_table_shape(block)
        if shape is not None:
            declared_shapes.add(shape)

    pdf_table = _pdf_table_metadata(document, page, table_index)
    if pdf_table is not None:
        declared_shapes.add((pdf_table.row_count, pdf_table.column_count))
    if len(declared_shapes) > 1:
        raise ExportError("Table metadata contains conflicting dimensions.")

    derived_rows = max(row for row, _ in cells) + 1
    derived_columns = max(column for _, column in cells) + 1
    rows, columns = next(iter(declared_shapes), (derived_rows, derived_columns))
    if rows < derived_rows or columns < derived_columns:
        raise ExportError("Table cell coordinates exceed declared dimensions.")
    if rows > limits.max_table_rows or columns > limits.max_table_columns:
        raise ExportError("Table dimensions exceed configured limits.")
    if rows * columns > limits.max_table_cells:
        raise ExportError("Table cell count exceeds the configured limit.")
    if not declared_shapes and len(cells) != rows * columns:
        raise ExportError("Table cells do not form a rectangular shape.")

    merges, merged_coordinates = _table_merges(cells, rows, columns)
    first_source_index = min(index for index, _ in indexed_blocks)
    last_source_index = max(index for index, _ in indexed_blocks)
    bboxes = [block.bbox for _, block in indexed_blocks if block.bbox is not None]
    table_bbox = pdf_table.bbox if pdf_table is not None else _declared_table_bbox(cells)
    if table_bbox is None and bboxes:
        table_bbox = (
            min(bbox[0] for bbox in bboxes),
            min(bbox[1] for bbox in bboxes),
            max(bbox[2] for bbox in bboxes),
            max(bbox[3] for bbox in bboxes),
        )
    return _Table(
        page=page,
        table_index=table_index,
        rows=rows,
        columns=columns,
        cells=cells,
        merges=merges,
        merged_coordinates=frozenset(merged_coordinates),
        first_source_index=first_source_index,
        last_source_index=last_source_index,
        y=table_bbox[1] if table_bbox is not None else math.inf,
        x=table_bbox[0] if table_bbox is not None else math.inf,
        bottom=table_bbox[3] if table_bbox is not None else math.inf,
    )


def _declared_table_bbox(
    cells: Mapping[tuple[int, int], TextBlock],
) -> tuple[float, float, float, float] | None:
    declared: set[tuple[float, float, float, float]] = set()
    for block in cells.values():
        raw_bbox = block.source_locator.get("table_bbox")
        if raw_bbox is None:
            continue
        if (
            not isinstance(raw_bbox, list | tuple)
            or len(raw_bbox) != 4
            or any(
                isinstance(coordinate, bool)
                or not isinstance(coordinate, int | float)
                for coordinate in raw_bbox
            )
        ):
            raise ExportError("Table placement metadata is invalid.")
        declared.add(
            (
                float(raw_bbox[0]),
                float(raw_bbox[1]),
                float(raw_bbox[2]),
                float(raw_bbox[3]),
            )
        )
    if len(declared) > 1:
        raise ExportError("Table metadata contains conflicting placement bounds.")
    if not declared:
        return None
    x0, y0, x1, y1 = next(iter(declared))
    if not all(math.isfinite(value) for value in (x0, y0, x1, y1)):
        raise ExportError("Table placement metadata is invalid.")
    if x1 <= x0 or y1 <= y0:
        raise ExportError("Table placement metadata is invalid.")
    return x0, y0, x1, y1


def _reject_interleaved_table_blocks(
    tables: list[_Table],
    non_table: list[tuple[int, TextBlock]],
) -> None:
    for table in tables:
        if not math.isfinite(table.y) or not math.isfinite(table.bottom):
            continue
        for source_index, block in non_table:
            if block.page != table.page:
                continue
            within_source_span = (
                block.bbox is None
                and table.first_source_index < source_index < table.last_source_index
            )
            within_vertical_span = (
                block.bbox is not None
                and table.y <= block.bbox[1] < table.bottom
            )
            if within_source_span or within_vertical_span:
                raise ExportError(
                    "Canonical blocks interleaved within a table cannot be reconstructed safely."
                )


def _declared_table_shape(block: TextBlock) -> tuple[int, int] | None:
    declared: list[tuple[int, int]] = []
    raw_shape = block.source_locator.get("table_shape")
    if raw_shape is not None:
        if not isinstance(raw_shape, Mapping) or set(raw_shape) != {"rows", "columns"}:
            raise ExportError("Table shape metadata is invalid.")
        declared.append(
            _validated_table_shape(raw_shape.get("rows"), raw_shape.get("columns"))
        )
    row_count = block.source_locator.get("row_count")
    column_count = block.source_locator.get("column_count")
    if row_count is not None or column_count is not None:
        declared.append(_validated_table_shape(row_count, column_count))
    if len(set(declared)) > 1:
        raise ExportError("Table metadata contains conflicting dimensions.")
    return declared[0] if declared else None


def _validated_table_shape(rows: object, columns: object) -> tuple[int, int]:
    if (
        isinstance(rows, bool)
        or not isinstance(rows, int)
        or isinstance(columns, bool)
        or not isinstance(columns, int)
        or rows <= 0
        or columns <= 0
    ):
        raise ExportError("Table shape metadata is invalid.")
    return rows, columns


def _table_merges(
    cells: Mapping[tuple[int, int], TextBlock],
    rows: int,
    columns: int,
) -> tuple[dict[tuple[int, int], tuple[int, int]], set[tuple[int, int]]]:
    merges: dict[tuple[int, int], tuple[int, int]] = {}
    occupied: set[tuple[int, int]] = set()
    merged_coordinates: set[tuple[int, int]] = set()
    for coordinate, block in sorted(cells.items()):
        raw_merge = block.source_locator.get("merge")
        if raw_merge is None:
            continue
        if (
            not isinstance(raw_merge, Mapping)
            or set(raw_merge) != {"row_span", "column_span"}
        ):
            raise ExportError("Table merge metadata is invalid.")
        row_span = raw_merge.get("row_span")
        column_span = raw_merge.get("column_span")
        if (
            isinstance(row_span, bool)
            or not isinstance(row_span, int)
            or isinstance(column_span, bool)
            or not isinstance(column_span, int)
            or row_span <= 0
            or column_span <= 0
        ):
            raise ExportError("Table merge metadata is invalid.")
        row_index, column_index = coordinate
        if row_index + row_span > rows or column_index + column_span > columns:
            raise ExportError("Table merge exceeds declared dimensions.")
        covered = {
            (row, column)
            for row in range(row_index, row_index + row_span)
            for column in range(column_index, column_index + column_span)
        }
        if occupied.intersection(covered):
            raise ExportError("Table merge regions overlap.")
        for covered_coordinate in covered - {coordinate}:
            covered_block = cells.get(covered_coordinate)
            if covered_block is not None and (
                covered_block.text
                or covered_block.source_locator.get("merge") is not None
            ):
                raise ExportError("Merged table cells contain conflicting content.")
        occupied.update(covered)
        merged_coordinates.update(covered - {coordinate})
        if row_span > 1 or column_span > 1:
            merges[coordinate] = (row_span, column_span)
    return merges, merged_coordinates


def _pdf_table_metadata(
    document: DocumentModel,
    page: int | None,
    table_index: int,
) -> Any | None:
    if document.metadata.pdf is None or page is None:
        return None
    if not 1 <= page <= len(document.metadata.pdf.pages):
        return None
    page_metadata = document.metadata.pdf.pages[page - 1]
    for table in page_metadata.tables:
        if table.table_index == table_index:
            return table
    return None


def _block_order(block: TextBlock, source_index: int) -> tuple[float, ...]:
    page = float(block.page) if block.page is not None else math.inf
    if block.bbox is None:
        return page, math.inf, math.inf, float(source_index), 0.0
    return page, block.bbox[1], block.bbox[0], float(source_index), 0.0


def _table_order(table: _Table) -> tuple[float, ...]:
    page = float(table.page) if table.page is not None else math.inf
    return page, table.y, table.x, float(table.first_source_index), 1.0


def _add_text_block(document: Any, block: TextBlock) -> None:
    if block.kind == "heading":
        level = _heading_level(block.style.get("level"))
        paragraph = document.add_paragraph(style=f"Heading {level}")
    else:
        paragraph = document.add_paragraph()
    spans = _validated_pdf_spans(block)
    if not spans:
        run = paragraph.add_run()
        _set_run_text(run, block.text)
        _apply_run_style(run, block.style)
        return
    previous_line: int | None = None
    for span in spans:
        run = paragraph.add_run()
        if previous_line is not None and span.line_index != previous_line:
            run.add_break()
        _set_run_text(run, span.text)
        _apply_run_style(
            run,
            {
                "font": {
                    "name": span.font_name,
                    "size": span.font_size,
                    "flags": span.font_flags,
                    "color": span.color,
                }
            },
        )
        previous_line = span.line_index


def _validated_pdf_spans(block: TextBlock) -> tuple[PdfTextSpan, ...]:
    raw_spans = block.style.get("spans")
    if raw_spans is None:
        return ()
    if not isinstance(raw_spans, list):
        raise ExportError("Canonical PDF span metadata is invalid.")
    try:
        spans = tuple(PdfTextSpan.model_validate(value) for value in raw_spans)
    except (TypeError, ValueError) as error:
        raise ExportError("Canonical PDF span metadata is invalid.") from error
    reconstructed: list[str] = []
    previous_line: int | None = None
    for span in spans:
        if previous_line is not None and span.line_index != previous_line:
            reconstructed.append("\n")
        reconstructed.append(span.text)
        previous_line = span.line_index
    if "".join(reconstructed) != block.text:
        raise ExportError("Canonical PDF spans do not reconstruct their block text.")
    return spans


def _heading_level(value: object) -> int:
    if isinstance(value, bool) or not isinstance(value, int):
        return 1
    return min(9, max(1, value))


def _set_run_text(run: Any, text: str) -> None:
    lines = _xml_safe_text(text).split("\n")
    run.add_text(lines[0])
    for line in lines[1:]:
        run.add_break()
        run.add_text(line)


def _apply_run_style(run: Any, style: Mapping[str, object]) -> None:
    raw_font = style.get("font")
    font = raw_font if isinstance(raw_font, Mapping) else {}
    name = style.get(
        "east_asia_font",
        font.get("east_asia", font.get("eastAsia", font.get("name", font.get("family")))),
    )
    if isinstance(name, str) and name:
        run.font.name = name
        run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), name)

    size = style.get("estimated_font_size", style.get("font_size", font.get("size")))
    if (
        isinstance(size, int | float)
        and not isinstance(size, bool)
        and math.isfinite(float(size))
        and 1.0 <= float(size) <= 200.0
    ):
        run.font.size = Pt(round(float(size) * 2.0) / 2.0)

    bold = font.get("bold")
    italic = font.get("italic")
    flags = font.get("flags")
    color = font.get("color")
    if isinstance(bold, bool):
        run.bold = bold
    elif isinstance(flags, int) and not isinstance(flags, bool):
        run.bold = bool(flags & 16)
    if isinstance(italic, bool):
        run.italic = italic
    elif isinstance(flags, int) and not isinstance(flags, bool):
        run.italic = bool(flags & 2)
    if (
        isinstance(color, int)
        and not isinstance(color, bool)
        and 0 <= color <= 0xFFFFFF
    ):
        run.font.color.rgb = RGBColor(
            (color >> 16) & 0xFF,
            (color >> 8) & 0xFF,
            color & 0xFF,
        )


def _add_table(document: Any, source: _Table) -> None:
    table = document.add_table(rows=source.rows, cols=source.columns)
    for (row_index, column_index), (row_span, column_span) in sorted(
        source.merges.items()
    ):
        table.cell(row_index, column_index).merge(
            table.cell(
                row_index + row_span - 1,
                column_index + column_span - 1,
            )
        )
    for (row_index, column_index), block in sorted(source.cells.items()):
        if (row_index, column_index) in source.merged_coordinates:
            continue
        cell = table.cell(row_index, column_index)
        if (row_index, column_index) in source.merges:
            cell.text = ""
        paragraph = cell.paragraphs[0]
        run = paragraph.add_run()
        _set_run_text(run, block.text)
        _apply_run_style(run, block.style)


def _add_image(document: Any, image: _Image) -> None:
    section = document.sections[-1]
    content_width = section.page_width - section.left_margin - section.right_margin
    natural_width = Inches(image.width / 96.0)
    width = min(natural_width, content_width)
    try:
        document.add_picture(io.BytesIO(image.content), width=width)
    except (
        InvalidImageStreamError,
        UnexpectedEndOfFileError,
        UnrecognizedImageError,
    ) as error:
        raise ExportError("Image payload is malformed or unsupported.") from error


def _canonical_pdf_image_request(
    document: DocumentModel,
    block: TextBlock,
) -> _PdfImageRequest:
    if "image_payload" in block.source_locator:
        raise ExportError("Image block contains a non-canonical image reference.")
    if (
        document.file_type is not FileType.PDF
        or document.metadata.pdf is None
        or block.page is None
        or not 1 <= block.page <= len(document.metadata.pdf.pages)
    ):
        raise ExportError("Image block does not contain a valid canonical image reference.")
    image_index = block.source_locator.get("image_index")
    xref = block.source_locator.get("xref")
    if (
        isinstance(image_index, bool)
        or not isinstance(image_index, int)
        or image_index < 0
        or isinstance(xref, bool)
        or not isinstance(xref, int)
        or xref <= 0
    ):
        raise ExportError("Image block does not contain a valid canonical image reference.")
    page_metadata = document.metadata.pdf.pages[block.page - 1]
    matches = [
        image
        for image in page_metadata.images
        if image.image_index == image_index and image.xref == xref
    ]
    if len(matches) != 1 or matches[0].bbox != block.bbox:
        raise ExportError("Canonical PDF image reference conflicts with document metadata.")
    return _PdfImageRequest(block_id=block.block_id, xref=xref)


def _read_resolved_source(
    document: DocumentModel,
    source_path: Path,
    max_bytes: int,
) -> bytes:
    if not source_path.is_absolute() or ".." in source_path.parts:
        raise ExportError("Resolved source path must be absolute and normalized.")
    file_fd = -1
    try:
        file_fd = os.open(source_path, os.O_RDONLY | os.O_NOFOLLOW)
        file_stat = os.fstat(file_fd)
        if not stat.S_ISREG(file_stat.st_mode):
            raise ExportError("Resolved source must be a regular file.")
        if file_stat.st_size <= 0 or file_stat.st_size > max_bytes:
            raise ExportError("Resolved source exceeds the configured size limit.")
        content = _read_bounded(file_fd, max_bytes)
    except ExportError:
        raise
    except OSError as error:
        raise ExportError("Resolved source is unsafe or unavailable.") from error
    finally:
        if file_fd >= 0:
            os.close(file_fd)
    if f"sha256:{hashlib.sha256(content).hexdigest()}" != document.source_version:
        raise ExportError("Resolved source does not match the canonical document.")
    return content


def _read_bounded(file_fd: int, max_bytes: int) -> bytes:
    chunks: list[bytes] = []
    remaining = max_bytes + 1
    while remaining > 0:
        chunk = os.read(file_fd, min(1024 * 1024, remaining))
        if not chunk:
            break
        chunks.append(chunk)
        remaining -= len(chunk)
    content = b"".join(chunks)
    if len(content) > max_bytes:
        raise ExportError("Resolved source exceeds the configured size limit.")
    return content


def _validated_image(
    content: bytes,
    limits: DocxReconstructionLimits,
) -> _Image:
    if not content or len(content) > limits.max_image_bytes:
        raise ExportError("Image payload exceeds the configured size limit.")
    if content.startswith(b"\x89PNG\r\n\x1a\n"):
        width, height, decoded_size = _validate_png(content, limits)
        mime_type: SupportedImageMime = "image/png"
    if content.startswith(b"\xff\xd8") and content.endswith(b"\xff\xd9"):
        width, height = _jpeg_dimensions(content)
        decoded_size = width * height * 4
        mime_type = "image/jpeg"
    elif not content.startswith(b"\x89PNG\r\n\x1a\n"):
        raise ExportError("Image payload signature is unsupported.")
    if (
        width > limits.max_image_width
        or height > limits.max_image_height
        or width * height > limits.max_image_pixels
    ):
        raise ExportError("Image dimensions exceed configured limits.")
    if decoded_size > limits.max_decoded_image_bytes:
        raise ExportError("Image exceeds the configured decoded image byte limit.")
    try:
        pixmap: Any = _PYMUPDF.Pixmap(content)
    except (RuntimeError, ValueError) as error:
        raise ExportError("Image payload could not be fully decoded.") from error
    try:
        if pixmap.width != width or pixmap.height != height:
            raise ExportError("Decoded image dimensions do not match its header.")
        if pixmap.stride * pixmap.height > limits.max_decoded_image_bytes:
            raise ExportError("Image exceeds the configured decoded image byte limit.")
    finally:
        del pixmap
    return _Image(content, mime_type, width, height)


def _validate_png(
    content: bytes,
    limits: DocxReconstructionLimits,
) -> tuple[int, int, int]:
    offset = 8
    width = 0
    height = 0
    row_bytes = 0
    expected_decoded = 0
    color_type = -1
    seen_header = False
    seen_palette = False
    seen_image_data = False
    ended_image_data = False
    seen_end = False
    decoded = bytearray()
    decompressor: Any | None = None
    while offset < len(content):
        if len(content) - offset < 12:
            raise ExportError("PNG image payload is truncated.")
        chunk_length = int.from_bytes(content[offset : offset + 4], "big")
        chunk_type = content[offset + 4 : offset + 8]
        data_start = offset + 8
        data_end = data_start + chunk_length
        crc_end = data_end + 4
        if data_end < data_start or crc_end > len(content):
            raise ExportError("PNG image payload is truncated.")
        chunk_data = content[data_start:data_end]
        expected_crc = int.from_bytes(content[data_end:crc_end], "big")
        actual_crc = zlib.crc32(chunk_type)
        actual_crc = zlib.crc32(chunk_data, actual_crc) & 0xFFFFFFFF
        if expected_crc != actual_crc:
            raise ExportError("PNG image chunk CRC is invalid.")
        if len(chunk_type) != 4 or not all(
            65 <= value <= 90 or 97 <= value <= 122 for value in chunk_type
        ):
            raise ExportError("PNG image chunk type is invalid.")
        if not seen_header and chunk_type != b"IHDR":
            raise ExportError("PNG image header must be the first chunk.")
        if seen_image_data and chunk_type != b"IDAT":
            ended_image_data = True

        if chunk_type == b"IHDR":
            if seen_header or chunk_length != 13:
                raise ExportError("PNG image header is invalid.")
            width, height, bit_depth, color_type, compression, filtering, interlace = (
                struct.unpack(">IIBBBBB", chunk_data)
            )
            valid_depths = {
                0: {1, 2, 4, 8, 16},
                2: {8, 16},
                3: {1, 2, 4, 8},
                4: {8, 16},
                6: {8, 16},
            }
            if (
                width <= 0
                or height <= 0
                or bit_depth not in valid_depths.get(color_type, set())
                or compression != 0
                or filtering != 0
                or interlace != 0
            ):
                raise ExportError("PNG image header uses unsupported values.")
            channels = {0: 1, 2: 3, 3: 1, 4: 2, 6: 4}[color_type]
            row_bytes = 1 + ((width * channels * bit_depth + 7) // 8)
            expected_decoded = row_bytes * height
            if expected_decoded > limits.max_decoded_image_bytes:
                raise ExportError("Image exceeds the configured decoded image byte limit.")
            decompressor = zlib.decompressobj()
            seen_header = True
        elif chunk_type == b"PLTE":
            if (
                seen_palette
                or seen_image_data
                or color_type in {0, 4}
                or chunk_length == 0
                or chunk_length > 768
                or chunk_length % 3
            ):
                raise ExportError("PNG image palette is invalid.")
            seen_palette = True
        elif chunk_type == b"IDAT":
            if ended_image_data or decompressor is None:
                raise ExportError("PNG image data chunks are invalid.")
            if color_type == 3 and not seen_palette:
                raise ExportError("Indexed PNG image requires a palette.")
            _decompress_png_chunk(
                decompressor,
                chunk_data,
                decoded,
                expected_decoded,
            )
            seen_image_data = True
        elif chunk_type == b"IEND":
            if chunk_length != 0 or not seen_image_data or decompressor is None:
                raise ExportError("PNG image end chunk is invalid.")
            _finish_png_decompression(decompressor, decoded, expected_decoded)
            seen_end = True
            offset = crc_end
            break
        elif chunk_type[0] & 0x20 == 0:
            raise ExportError("PNG image contains an unsupported critical chunk.")
        offset = crc_end
    if not seen_end or offset != len(content) or len(decoded) != expected_decoded:
        raise ExportError("PNG image payload is incomplete.")
    if any(decoded[row * row_bytes] > 4 for row in range(height)):
        raise ExportError("PNG image uses an invalid scanline filter.")
    return width, height, max(expected_decoded, width * height * 4)


def _decompress_png_chunk(
    decompressor: Any,
    chunk_data: bytes,
    decoded: bytearray,
    maximum: int,
) -> None:
    pending = chunk_data
    while pending:
        before = len(pending)
        remaining = maximum - len(decoded)
        try:
            output = decompressor.decompress(pending, max(1, remaining + 1))
        except zlib.error as error:
            raise ExportError("PNG image compressed data is invalid.") from error
        decoded.extend(output)
        if len(decoded) > maximum:
            raise ExportError("PNG image decompression exceeds declared dimensions.")
        pending = decompressor.unconsumed_tail
        if pending and len(pending) >= before:
            raise ExportError("PNG image compressed data is invalid.")


def _finish_png_decompression(
    decompressor: Any,
    decoded: bytearray,
    maximum: int,
) -> None:
    try:
        decoded.extend(decompressor.flush(max(1, maximum - len(decoded) + 1)))
    except zlib.error as error:
        raise ExportError("PNG image compressed data is invalid.") from error
    if (
        len(decoded) > maximum
        or not decompressor.eof
        or decompressor.unused_data
        or decompressor.unconsumed_tail
    ):
        raise ExportError("PNG image compressed data is invalid.")


def _jpeg_dimensions(content: bytes) -> tuple[int, int]:
    offset = 2
    while offset + 4 <= len(content):
        if content[offset] != 0xFF:
            raise ExportError("JPEG image payload is malformed.")
        while offset < len(content) and content[offset] == 0xFF:
            offset += 1
        if offset >= len(content):
            break
        marker = content[offset]
        offset += 1
        if marker in {0xD8, 0xD9} or 0xD0 <= marker <= 0xD7:
            continue
        if offset + 2 > len(content):
            break
        segment_length = int.from_bytes(content[offset : offset + 2], "big")
        if segment_length < 2 or offset + segment_length > len(content):
            raise ExportError("JPEG image payload is malformed.")
        if marker in {
            0xC0,
            0xC1,
            0xC2,
            0xC3,
            0xC5,
            0xC6,
            0xC7,
            0xC9,
            0xCA,
            0xCB,
            0xCD,
            0xCE,
            0xCF,
        }:
            if segment_length < 7:
                raise ExportError("JPEG image payload is malformed.")
            height = int.from_bytes(content[offset + 3 : offset + 5], "big")
            width = int.from_bytes(content[offset + 5 : offset + 7], "big")
            if width <= 0 or height <= 0:
                raise ExportError("JPEG image dimensions are invalid.")
            return width, height
        offset += segment_length
    raise ExportError("JPEG image dimensions are missing.")


def _set_core_properties(output: Any, document: DocumentModel) -> None:
    output.core_properties.title = _xml_safe_text(Path(document.source_name).stem)
    output.core_properties.created = _FIXED_CORE_PROPERTY_TIME
    output.core_properties.modified = _FIXED_CORE_PROPERTY_TIME


def _xml_safe_text(value: str) -> str:
    normalized: list[str] = []
    for character in value:
        codepoint = ord(character)
        if (
            character in {"\t", "\n", "\r"}
            or 0x20 <= codepoint <= 0x7E
            or 0xA0 <= codepoint <= 0xD7FF
            or 0xE000 <= codepoint <= 0xFFFD
            or 0x10000 <= codepoint <= 0x10FFFF
        ):
            normalized.append(character)
        else:
            normalized.append(_XML_REPLACEMENT_CHARACTER)
    return "".join(normalized)


def _set_page_setup(output: Any, document: DocumentModel) -> None:
    width = _DEFAULT_PAGE_WIDTH_POINTS
    height = _DEFAULT_PAGE_HEIGHT_POINTS
    if document.metadata.pdf is not None and document.metadata.pdf.pages:
        page_bbox = document.metadata.pdf.pages[0].page_bbox
        candidate_width = page_bbox[2] - page_bbox[0]
        candidate_height = page_bbox[3] - page_bbox[1]
        if 72.0 <= candidate_width <= 2_000.0 and 72.0 <= candidate_height <= 2_000.0:
            width = candidate_width
            height = candidate_height
    section = output.sections[0]
    section.page_width = Pt(width)
    section.page_height = Pt(height)
    section.orientation = WD_ORIENT.LANDSCAPE if width > height else WD_ORIENT.PORTRAIT
    section.top_margin = Pt(_DEFAULT_MARGIN_POINTS)
    section.right_margin = Pt(_DEFAULT_MARGIN_POINTS)
    section.bottom_margin = Pt(_DEFAULT_MARGIN_POINTS)
    section.left_margin = Pt(_DEFAULT_MARGIN_POINTS)


def _publish_atomic(target: Path, content: bytes) -> None:
    if not target.is_absolute() or target.name in {"", ".", ".."} or "\x00" in target.name:
        raise ExportError("Target path must be an absolute DOCX path.")
    if target.suffix.lower() != ".docx":
        raise ExportError("Target path must use the .docx extension.")
    parent_fd = _open_absolute_directory(
        target.parent,
        "Target parent directory is unsafe or missing.",
    )
    temp_name = f".{target.name}.{uuid4().hex}.uploading"
    temp_fd = -1
    target_fd = -1
    linked = False
    inode: tuple[int, int] | None = None
    try:
        temp_fd = os.open(
            temp_name,
            os.O_WRONLY | os.O_CREAT | os.O_EXCL | os.O_NOFOLLOW,
            0o600,
            dir_fd=parent_fd,
        )
        _write_all(temp_fd, content)
        os.fsync(temp_fd)
        temp_stat = os.fstat(temp_fd)
        inode = temp_stat.st_dev, temp_stat.st_ino
        try:
            os.link(
                temp_name,
                target.name,
                src_dir_fd=parent_fd,
                dst_dir_fd=parent_fd,
                follow_symlinks=False,
            )
            linked = True
        except FileExistsError as error:
            raise ExportError("Target already exists and was not overwritten.") from error
        target_fd = os.open(
            target.name,
            os.O_RDONLY | os.O_NOFOLLOW,
            dir_fd=parent_fd,
        )
        target_stat = os.fstat(target_fd)
        if (target_stat.st_dev, target_stat.st_ino) != inode:
            raise ExportError("Published target does not match the reconstructed document.")
        os.unlink(temp_name, dir_fd=parent_fd)
        temp_name = ""
        os.fsync(parent_fd)
    except ExportError:
        if linked and inode is not None:
            _unlink_if_inode(parent_fd, target.name, inode)
        raise
    except OSError as error:
        if linked and inode is not None:
            _unlink_if_inode(parent_fd, target.name, inode)
        raise ExportError("Reconstructed DOCX could not be published safely.") from error
    finally:
        if target_fd >= 0:
            os.close(target_fd)
        if temp_fd >= 0:
            os.close(temp_fd)
        if temp_name:
            try:
                os.unlink(temp_name, dir_fd=parent_fd)
            except FileNotFoundError:
                pass
        os.close(parent_fd)


def _open_absolute_directory(path: Path, error_message: str) -> int:
    if not path.is_absolute() or ".." in path.parts:
        raise ExportError(error_message)
    current_fd = -1
    try:
        current_fd = os.open("/", os.O_RDONLY | os.O_DIRECTORY)
        for part in path.parts[1:]:
            next_fd = os.open(
                part,
                os.O_RDONLY | os.O_DIRECTORY | os.O_NOFOLLOW,
                dir_fd=current_fd,
            )
            os.close(current_fd)
            current_fd = next_fd
        return current_fd
    except OSError as error:
        if current_fd >= 0:
            os.close(current_fd)
        raise ExportError(error_message) from error


def _write_all(file_fd: int, content: bytes) -> None:
    view = memoryview(content)
    written = 0
    while written < len(view):
        count = os.write(file_fd, view[written:])
        if count <= 0:
            raise OSError("short write")
        written += count


def _unlink_if_inode(
    parent_fd: int,
    name: str,
    expected_inode: tuple[int, int],
) -> None:
    try:
        current = os.stat(name, dir_fd=parent_fd, follow_symlinks=False)
        if (current.st_dev, current.st_ino) == expected_inode:
            os.unlink(name, dir_fd=parent_fd)
    except FileNotFoundError:
        return
