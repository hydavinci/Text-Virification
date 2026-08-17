from __future__ import annotations

from collections import defaultdict
from contextlib import suppress
from dataclasses import dataclass
from pathlib import Path
from shutil import copyfile
from uuid import UUID

from docx import Document as WordDocument
from docx.document import Document as WordProcessingDocument
from docx.table import _Cell
from docx.text.run import Run

from text_verification.domain.documents import DocumentModel, TextBlock

from .replacements import ExportWarning, Replacement, ReplacementPlan


@dataclass(frozen=True, slots=True)
class ExportResult:
    path: Path
    warnings: list[ExportWarning]


class ExportError(Exception):
    def __init__(self, code: str, public_message: str) -> None:
        super().__init__(public_message)
        self.code = code
        self.public_message = public_message


@dataclass(frozen=True, slots=True)
class _RunTarget:
    paragraph_index: int | None = None
    table_index: int | None = None
    row_index: int | None = None
    column_index: int | None = None
    cell_paragraph_index: int | None = None
    run_index: int = 0


@dataclass(frozen=True, slots=True)
class _RunEdit:
    target: _RunTarget
    start: int
    end: int
    original: str
    value: str
    issue_id: UUID
    block_id: str


class DocxExporter:
    def export(
        self,
        source: Path,
        document: DocumentModel,
        plan: ReplacementPlan,
        target: Path,
    ) -> ExportResult:
        warnings = [*plan.warnings]
        edits_by_target: dict[_RunTarget, list[_RunEdit]] = defaultdict(list)
        block_by_id = {block.block_id: block for block in document.blocks}

        for replacement in plan.applicable:
            block = block_by_id.get(replacement.block_id)
            if block is None:
                raise ExportError("docx_export_failed", "无法导出 DOCX 文件。")

            resolved = self._resolve_edit(block, replacement)
            if isinstance(resolved, ExportWarning):
                warnings.append(resolved)
                continue
            edits_by_target[resolved.target].append(resolved)

        try:
            if edits_by_target:
                rendered = WordDocument(str(source))
                edited = False
                for target_ref, edits in edits_by_target.items():
                    run = self._resolve_run(rendered, target_ref)
                    updated_text = self._apply_run_edits(run.text, edits)
                    if updated_text == run.text:
                        continue
                    run.text = updated_text
                    edited = True

                if edited:
                    self._write_verified_document(rendered, target)
                else:
                    self._copy_verified_document(source, target)
            else:
                self._copy_verified_document(source, target)
        except ExportError:
            raise
        except Exception as error:
            raise ExportError("docx_export_failed", "无法导出 DOCX 文件。") from error

        return ExportResult(path=target, warnings=warnings)

    def _resolve_edit(
        self,
        block: TextBlock,
        replacement: Replacement,
    ) -> _RunEdit | ExportWarning:
        locator = block.source_locator
        if "runs" in locator and "paragraph_index" in locator:
            run_locator = self._find_run_locator(locator["runs"], replacement)
            if run_locator is None:
                return _unsafe_run_boundary_warning(replacement)
            return _RunEdit(
                target=_RunTarget(
                    paragraph_index=int(locator["paragraph_index"]),
                    run_index=int(run_locator["run_index"]),
                ),
                start=int(replacement.start) - int(run_locator["start"]),
                end=int(replacement.end) - int(run_locator["start"]),
                original=replacement.original,
                value=replacement.value,
                issue_id=replacement.issue_id,
                block_id=replacement.block_id,
            )

        if "paragraphs" in locator:
            for paragraph_locator in locator["paragraphs"]:
                run_locator = self._find_run_locator(paragraph_locator["runs"], replacement)
                if run_locator is None:
                    continue
                return _RunEdit(
                    target=_RunTarget(
                        table_index=int(locator["table_index"]),
                        row_index=int(locator["row_index"]),
                        column_index=int(locator["column_index"]),
                        cell_paragraph_index=int(paragraph_locator["paragraph_index"]),
                        run_index=int(run_locator["run_index"]),
                    ),
                    start=int(replacement.start) - int(run_locator["start"]),
                    end=int(replacement.end) - int(run_locator["start"]),
                    original=replacement.original,
                    value=replacement.value,
                    issue_id=replacement.issue_id,
                    block_id=replacement.block_id,
                )
            return _unsafe_run_boundary_warning(replacement)

        raise ExportError("docx_export_failed", "无法导出 DOCX 文件。")

    def _find_run_locator(
        self,
        run_locators: list[dict[str, int]],
        replacement: Replacement,
    ) -> dict[str, int] | None:
        for run_locator in run_locators:
            if (
                int(run_locator["start"]) <= replacement.start
                and replacement.end <= int(run_locator["end"])
            ):
                return run_locator
        return None

    def _resolve_run(
        self,
        document: WordProcessingDocument,
        target: _RunTarget,
    ) -> Run:
        if target.paragraph_index is not None:
            paragraph = document.paragraphs[target.paragraph_index]
            return paragraph.runs[target.run_index]

        if (
            target.table_index is None
            or target.row_index is None
            or target.column_index is None
            or target.cell_paragraph_index is None
        ):
            raise ExportError("docx_export_failed", "无法导出 DOCX 文件。")

        cell = self._resolve_cell(
            document,
            table_index=target.table_index,
            row_index=target.row_index,
            column_index=target.column_index,
        )
        paragraph = cell.paragraphs[target.cell_paragraph_index]
        return paragraph.runs[target.run_index]

    def _resolve_cell(
        self,
        document: WordProcessingDocument,
        *,
        table_index: int,
        row_index: int,
        column_index: int,
    ) -> _Cell:
        table = document.tables[table_index]
        return table.rows[row_index].cells[column_index]

    def _apply_run_edits(
        self,
        text: str,
        edits: list[_RunEdit],
    ) -> str:
        rendered = text
        for edit in sorted(
            edits,
            key=lambda item: (item.start, item.end, str(item.issue_id)),
            reverse=True,
        ):
            if rendered[edit.start : edit.end] != edit.original:
                raise ExportError("docx_export_failed", "无法导出 DOCX 文件。")
            rendered = rendered[: edit.start] + edit.value + rendered[edit.end :]
        return rendered

    def _write_verified_document(
        self,
        document: WordProcessingDocument,
        target: Path,
    ) -> None:
        temp_target = self._prepare_temp_target(target)

        try:
            document.save(str(temp_target))
        except Exception as error:
            with suppress(FileNotFoundError):
                temp_target.unlink()
            raise ExportError("docx_export_failed", "无法导出 DOCX 文件。") from error
        self._replace_verified_temp_target(temp_target, target)

    def _copy_verified_document(
        self,
        source: Path,
        target: Path,
    ) -> None:
        temp_target = self._prepare_temp_target(target)

        try:
            copyfile(source, temp_target)
        except Exception as error:
            with suppress(FileNotFoundError):
                temp_target.unlink()
            raise ExportError("docx_export_failed", "无法导出 DOCX 文件。") from error
        self._replace_verified_temp_target(temp_target, target)

    def _prepare_temp_target(self, target: Path) -> Path:
        target.parent.mkdir(parents=True, exist_ok=True)
        temp_target = target.with_name(f"{target.name}.tmp")
        with suppress(FileNotFoundError):
            temp_target.unlink()
        return temp_target

    def _replace_verified_temp_target(
        self,
        temp_target: Path,
        target: Path,
    ) -> None:
        try:
            WordDocument(str(temp_target))
            temp_target.replace(target)
        except Exception as error:
            with suppress(FileNotFoundError):
                temp_target.unlink()
            raise ExportError("docx_export_failed", "无法导出 DOCX 文件。") from error


def _unsafe_run_boundary_warning(replacement: Replacement) -> ExportWarning:
    return ExportWarning(
        code="unsafe_docx_run_boundary",
        message="修改范围跨越多个 DOCX 文本运行，无法在保留格式的前提下自动应用。",
        issue_id=replacement.issue_id,
        block_id=replacement.block_id,
    )
