"""
啄木鸟·中英文字智能检查 - Flask 后端
提供文件上传、文本分析 API、原地格式保留导出。
"""

import os
import io
import re
import html
import json
import uuid
import tempfile
import unicodedata
from datetime import datetime
from flask import Flask, request, jsonify, render_template, send_file
from analyzer import TextAnalyzer
from file_parser import parse_file, get_supported_formats, strip_html, _convert_doc_to_docx
from llm_review import review_issues, is_llm_review_enabled

app = Flask(__name__)
app.config['MAX_CONTENT_LENGTH'] = 25 * 1024 * 1024  # 25MB (支持大PDF)

analyzer = TextAnalyzer()

# 原始文件保存目录（用于导出时原地替换）
UPLOAD_DIR = os.path.join(tempfile.gettempdir(), 'translation_pre_checker')
os.makedirs(UPLOAD_DIR, exist_ok=True)


# ---------------------------------------------------------------------------
# Word 风格字数统计
# Word 的「字数」口径：
#   字数 = 中文字符 + 非中文单词（英文单词、数字串）
#   中文字符 = East Asian Width 为 W/F 的字符，或 EAW=A 的标点符号
#             （如中文引号 “” ‘’、破折号 —、省略号 …、中点 · 等）。
#   英文单词 = 连续英文字母（可包含 ' 或 -）。
#   数字串   = 连续数字（可包含 . 或 , 作为小数/千分位）。
# 该实现与 Microsoft Word 中文版的「字数」字段保持一致。
# ---------------------------------------------------------------------------
_EN_WORD_RE = re.compile(r"[a-zA-Z]+(?:['\-][a-zA-Z]+)*")
_NUM_TOKEN_RE = re.compile(r'\d+(?:[.,]\d+)*')


def _is_word_cjk_char(ch: str) -> bool:
    """判定字符是否被 Word 计入「中文字符」。

    Word 对中日韩文字及全角符号统一计为「中文字符」；
    对 East Asian Width = A（Ambiguous）的字符，仅标点符号（category 以 P 开头）
    也会被计入，例如中文引号 U+201C/U+201D、破折号 U+2014、省略号 U+2026。
    """
    eaw = unicodedata.east_asian_width(ch)
    if eaw in ('W', 'F'):
        return True
    if eaw == 'A' and unicodedata.category(ch).startswith('P'):
        return True
    return False


def count_word_words(text: str) -> int:
    """模拟 Microsoft Word 的「字数」统计。"""
    cjk_chars = sum(1 for ch in text if _is_word_cjk_char(ch))
    en_words = len(_EN_WORD_RE.findall(text))
    num_words = len(_NUM_TOKEN_RE.findall(text))
    return cjk_chars + en_words + num_words


def count_chars_no_space(text: str) -> int:
    """模拟 Microsoft Word 的「字符数（不计空格）」。"""
    return len(re.sub(r'\s', '', text))


def _export_filename(base_name: str, ext: str) -> str:
    """生成导出文件名：原文件名_修改版_月日时分秒.扩展名"""
    ts = datetime.now().strftime('%m%d%H%M%S')
    return f'{base_name}_修改版_{ts}{ext}'


@app.route('/')
def index():
    return render_template('index.html', formats=get_supported_formats())


@app.route('/api/analyze', methods=['POST'])
def analyze():
    """分析文本或文件"""

    # 方式1：直接提交文本
    if 'text' in request.form and request.form['text'].strip():
        text = request.form['text']
        filename = '直接输入文本'
        file_id = None
        file_ext = None
        page_map = []
    # 方式2：上传文件
    elif 'file' in request.files:
        file = request.files['file']
        if file.filename == '':
            return jsonify({'error': '未选择文件'}), 400

        suffix = os.path.splitext(file.filename)[1]
        file_id = str(uuid.uuid4())
        saved_path = os.path.join(UPLOAD_DIR, f'{file_id}{suffix}')
        file.save(saved_path)

        try:
            text, file_format, page_map = parse_file(saved_path)
            filename = file.filename
            file_ext = suffix
        except Exception as e:
            if os.path.exists(saved_path):
                os.unlink(saved_path)
            return jsonify({'error': f'文件解析失败: {str(e)}'}), 400

        if not text or not text.strip():
            if os.path.exists(saved_path):
                os.unlink(saved_path)
            return jsonify({'error': '文件内容为空或无法提取文本'}), 400
    else:
        return jsonify({'error': '请上传文件或输入文本'}), 400

    # 场景参数
    scenario = request.form.get('scenario', 'general')

    # 合规/安全层（PII）扫描开关；默认开启（前端未传或传非 'false' 即视为开启）
    enable_security = request.form.get('enable_security', 'true') != 'false'

    # 敏感内容（涉政/民族宗教/领土规范表述）检查开关；默认开启
    enable_sensitive = request.form.get('enable_sensitive', 'true') != 'false'

    # 广告法极限词（营销材料）检查开关；默认关闭，需用户显式开启
    enable_ad_extreme = request.form.get('enable_ad_extreme', 'false') != 'false'

    # 自定义术语表（JSON 字符串：[{"original":"AI","standard":"人工智能"}, ...]）
    custom_glossary = None
    glossary_json = request.form.get('custom_glossary', '')
    if glossary_json:
        try:
            custom_glossary = json.loads(glossary_json)
            if not isinstance(custom_glossary, list):
                custom_glossary = None
        except (json.JSONDecodeError, TypeError):
            custom_glossary = None

    # 禁用词库（JSON 字符串：["词汇1", "词汇2", ...]）
    banned_words = None
    banned_json = request.form.get('banned_words', '')
    if banned_json:
        try:
            banned_words = json.loads(banned_json)
            if not isinstance(banned_words, list):
                banned_words = None
        except (json.JSONDecodeError, TypeError):
            banned_words = None

    # 执行分析
    issues = analyzer.analyze(text, scenario=scenario,
                              custom_glossary=custom_glossary,
                              banned_words=banned_words,
                              enable_security=enable_security,
                              enable_sensitive=enable_sensitive,
                              enable_ad_extreme=enable_ad_extreme)

    # 云端语义复核（基于上下文理解降低误报）：未配置 key / 调用失败均自动降级
    review_stats = None
    if is_llm_review_enabled():
        try:
            issues, review_stats = review_issues(text, issues)
        except Exception as e:
            review_stats = {
                'enabled': True, 'failed': True, 'candidates': 0,
                'removed': 0, 'downgraded': 0, 'kept': 0,
                'reason': f'复核异常已回退纯规则结果: {str(e)[:80]}',
            }

    summary = analyzer.get_summary(issues)
    if review_stats:
        summary['llm_review'] = review_stats

    # 语言检测和智能计数（采用 Word「字数」口径）
    cjk_chars = sum(1 for ch in text if _is_word_cjk_char(ch))
    en_words = len(_EN_WORD_RE.findall(text))
    word_count = count_word_words(text)
    if cjk_chars >= en_words:
        language = 'zh'
        primary_count = word_count
        primary_label = '总字数'
    else:
        language = 'en'
        primary_count = word_count
        primary_label = '总单词数'

    # 文本统计
    stats = {
        'char_count': len(text),
        'char_count_no_space': count_chars_no_space(text),
        'line_count': text.count('\n') + 1,
        'paragraph_count': len([p for p in text.split('\n') if p.strip()]),
        'language': language,
        'primary_count': primary_count,
        'primary_label': primary_label,
    }

    issue_dicts = []
    for issue in issues:
        d = issue.to_dict()
        issue_dicts.append(d)

    return jsonify({
        'success': True,
        'filename': filename,
        'text': text,
        'stats': stats,
        'issues': issue_dicts,
        'summary': summary,
        'file_id': file_id,
        'file_ext': file_ext,
        'scenario': scenario,
    })


@app.route('/api/export', methods=['POST'])
def export_report():
    """导出分析报告为 HTML 文件"""
    data = request.get_json()
    if not data:
        return jsonify({'error': '无数据'}), 400

    html = generate_report_html(data)

    report_path = os.path.join(tempfile.gettempdir(), 'analysis_report.html')
    with open(report_path, 'w', encoding='utf-8') as f:
        f.write(html)

    return send_file(
        report_path,
        as_attachment=True,
        download_name='原文检查报告.html',
        mimetype='text/html'
    )


@app.route('/api/export-original', methods=['POST'])
def export_original():
    """在原始文件上做查找替换，保留原格式导出"""
    data = request.get_json()
    if not data:
        return jsonify({'error': '无数据'}), 400

    file_id = data.get('file_id')
    replacements = data.get('replacements', [])
    filename = data.get('filename', '修改后文本')
    track_changes = data.get('track_changes', False)

    # 将 dict 列表转为 tuple 列表，并清理可能混入的 HTML 代码
    replacements = [
        (strip_html(r.get('original', '')), strip_html(r.get('suggestion', '')))
        for r in replacements
    ]

    if not file_id:
        return jsonify({'error': '无文件信息，请重新上传文件后导出'}), 400

    # 查找保存的原始文件
    saved_path = None
    file_ext = None
    for f in os.listdir(UPLOAD_DIR):
        if f.startswith(file_id):
            saved_path = os.path.join(UPLOAD_DIR, f)
            file_ext = os.path.splitext(f)[1]
            break

    if not saved_path or not os.path.exists(saved_path):
        return jsonify({'error': '原始文件未找到，请重新上传文件'}), 400

    base_name = filename.rsplit('.', 1)[0] if '.' in filename else filename

    try:
        if file_ext in ('.docx', '.doc'):
            # 旧版 .doc 需先借助 textutil/LibreOffice 转换为 docx 再处理
            working_path = saved_path
            if file_ext == '.doc':
                working_path = _convert_doc_to_docx(saved_path)
            try:
                if track_changes:
                    return _export_docx_track_changes(working_path, replacements, base_name)
                else:
                    return _export_docx_inplace(working_path, replacements, base_name)
            finally:
                if file_ext == '.doc' and working_path != saved_path:
                    try:
                        os.unlink(working_path)
                    except OSError:
                        pass
        elif file_ext in ('.txt', '.md', '.csv'):
            return _export_text_inplace(saved_path, replacements, base_name, file_ext, track_changes)
        elif file_ext == '.pdf':
            if track_changes:
                return _export_pdf_track_changes(saved_path, replacements, base_name)
            else:
                return _export_pdf_inplace(saved_path, replacements, base_name)
        elif file_ext == '.rtf':
            return _export_rtf_inplace(saved_path, replacements, base_name, track_changes)
        else:
            return jsonify({'error': f'不支持的导出格式: {file_ext}'}), 400
    except Exception as e:
        return jsonify({'error': f'导出失败: {str(e)}'}), 500


# ============================================================
# 格式原地替换导出
# ============================================================

def _export_docx_inplace(file_path, replacements, base_name):
    """在原始 Word 文档上做查找替换，保留所有格式，并清理混入的 HTML 代码"""
    from docx import Document
    doc = Document(file_path)

    def process_paragraphs(paragraphs):
        for paragraph in paragraphs:
            # 第一层：run 级别替换（保留 run 内格式），同时清理 HTML
            for run in paragraph.runs:
                run.text = strip_html(run.text)
                for orig, sugg in replacements:
                    if orig and orig in run.text:
                        run.text = run.text.replace(orig, sugg)

            # 第二层：段落级别替换（处理跨 run 的文本），同时清理 HTML
            full_text = ''.join(run.text for run in paragraph.runs)
            full_text = strip_html(full_text)
            modified = full_text
            for orig, sugg in replacements:
                if orig and orig in modified:
                    modified = modified.replace(orig, sugg)

            if modified != full_text and paragraph.runs:
                paragraph.runs[0].text = modified
                for run in paragraph.runs[1:]:
                    run.text = ''

    # 正文段落
    process_paragraphs(doc.paragraphs)

    # 表格
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                process_paragraphs(cell.paragraphs)

    # 页眉页脚
    for section in doc.sections:
        for hf in [section.header, section.first_page_header, section.even_page_header,
                    section.footer, section.first_page_footer, section.even_page_footer]:
            if hf:
                process_paragraphs(hf.paragraphs)

    output_path = os.path.join(UPLOAD_DIR, f'output_{uuid.uuid4()}.docx')
    doc.save(output_path)

    return send_file(
        output_path,
        as_attachment=True,
        download_name=_export_filename(base_name, '.docx')
    )


def _export_docx_track_changes(file_path, replacements, base_name):
    """在原始 Word 文档上添加修订痕迹（Track Changes），保留所有格式"""
    from docx import Document
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    import copy
    from datetime import datetime

    doc = Document(file_path)
    author = "啄木鸟·中英文字智能检查"
    date_str = datetime.now().strftime("%Y-%m-%dT%H:%M:%SZ")
    rev_counter = [1000]
    XML_SPACE = '{http://www.w3.org/XML/1998/namespace}space'

    def next_rev_id():
        rev_counter[0] += 1
        return str(rev_counter[0])

    def make_run(text, rPr_template):
        text = strip_html(text)
        r = OxmlElement('w:r')
        if rPr_template is not None:
            r.append(copy.deepcopy(rPr_template))
        t = OxmlElement('w:t')
        t.set(XML_SPACE, 'preserve')
        t.text = text
        r.append(t)
        return r

    def make_del(text, rPr_template):
        text = strip_html(text)
        del_elem = OxmlElement('w:del')
        del_elem.set(qn('w:id'), next_rev_id())
        del_elem.set(qn('w:author'), author)
        del_elem.set(qn('w:date'), date_str)
        r = OxmlElement('w:r')
        if rPr_template is not None:
            r.append(copy.deepcopy(rPr_template))
        dt = OxmlElement('w:delText')
        dt.set(XML_SPACE, 'preserve')
        dt.text = text
        r.append(dt)
        del_elem.append(r)
        return del_elem

    def make_ins(text, rPr_template):
        text = strip_html(text)
        ins_elem = OxmlElement('w:ins')
        ins_elem.set(qn('w:id'), next_rev_id())
        ins_elem.set(qn('w:author'), author)
        ins_elem.set(qn('w:date'), date_str)
        r = OxmlElement('w:r')
        if rPr_template is not None:
            r.append(copy.deepcopy(rPr_template))
        t = OxmlElement('w:t')
        t.set(XML_SPACE, 'preserve')
        t.text = text
        r.append(t)
        ins_elem.append(r)
        return ins_elem

    def process_paragraph(paragraph):
        # 第一遍：run 级别匹配（保留 run 内格式）
        runs = list(paragraph.runs)
        for run in runs:
            text = strip_html(run.text)
            if not text:
                continue

            matches = []
            for orig, sugg in replacements:
                if not orig or orig == sugg:
                    continue
                start = 0
                while True:
                    idx = text.find(orig, start)
                    if idx == -1:
                        break
                    matches.append((idx, idx + len(orig), orig, sugg))
                    start = idx + len(orig)

            if not matches:
                continue

            # 排序、去重重叠匹配
            matches.sort(key=lambda x: (x[0], -(x[1] - x[0])))
            filtered = []
            last_end = 0
            for start, end, orig, sugg in matches:
                if start >= last_end:
                    filtered.append((start, end, orig, sugg))
                    last_end = end

            # 构建分段
            segments = []
            pos = 0
            for start, end, orig, sugg in filtered:
                if start > pos:
                    segments.append(('keep', text[pos:start]))
                segments.append(('replace', orig, sugg))
                pos = end
            if pos < len(text):
                segments.append(('keep', text[pos:]))

            rPr = run._element.find(qn('w:rPr'))
            run_elem = run._element
            insert_after = run_elem

            for seg in segments:
                if seg[0] == 'keep':
                    new_r = make_run(seg[1], rPr)
                    insert_after.addnext(new_r)
                    insert_after = new_r
                else:
                    orig_text, sugg_text = seg[1], seg[2]
                    del_elem = make_del(orig_text, rPr)
                    insert_after.addnext(del_elem)
                    insert_after = del_elem
                    if sugg_text:
                        ins_elem = make_ins(sugg_text, rPr)
                        insert_after.addnext(ins_elem)
                        insert_after = ins_elem

            run_elem.getparent().remove(run_elem)

        # 第二遍：段落级别匹配（处理跨 run 的文本）
        full_text = strip_html(''.join(run.text for run in paragraph.runs))
        remaining = []
        for orig, sugg in replacements:
            if not orig or orig == sugg:
                continue
            if orig in full_text:
                remaining.append((orig, sugg))

        if remaining and paragraph.runs:
            for orig, sugg in remaining:
                if orig not in full_text:
                    continue
                # 在段落级别做替换，保留第一个 run 的格式
                rPr = paragraph.runs[0]._element.find(qn('w:rPr'))
                modified = full_text.replace(orig, '\x00DEL\x00' + orig + '\x00INS\x00' + (sugg or '') + '\x00END\x00')
                full_text = modified

            # 重建段落
            parts = full_text.split('\x00')
            new_elements = []
            i = 0
            while i < len(parts):
                if parts[i] == 'DEL':
                    del_elem = make_del(parts[i + 1], rPr)
                    new_elements.append(del_elem)
                    i += 2
                elif parts[i] == 'INS':
                    if parts[i + 1]:
                        ins_elem = make_ins(parts[i + 1], rPr)
                        new_elements.append(ins_elem)
                    i += 2
                elif parts[i] == 'END':
                    i += 1
                else:
                    if parts[i]:
                        new_elements.append(make_run(parts[i], rPr))
                    i += 1

            # 清除现有 runs，插入新元素
            p_elem = paragraph._element
            for run in list(paragraph.runs):
                p_elem.remove(run._element)
            for elem in new_elements:
                p_elem.append(elem)

    def process_paragraphs(paragraphs):
        for p in paragraphs:
            process_paragraph(p)

    process_paragraphs(doc.paragraphs)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                process_paragraphs(cell.paragraphs)

    for section in doc.sections:
        for hf in [section.header, section.first_page_header, section.even_page_header,
                    section.footer, section.first_page_footer, section.even_page_footer]:
            if hf:
                process_paragraphs(hf.paragraphs)

    output_path = os.path.join(UPLOAD_DIR, f'output_{uuid.uuid4()}.docx')
    doc.save(output_path)

    return send_file(
        output_path,
        as_attachment=True,
        download_name=_export_filename(base_name, '.docx')
    )


def _export_text_inplace(file_path, replacements, base_name, ext, track_changes=False):
    """纯文本文件原地替换"""
    # 尝试多种编码读取
    content = None
    used_encoding = 'utf-8'
    for encoding in ['utf-8', 'gbk', 'gb2312', 'big5', 'utf-16', 'latin-1']:
        try:
            with open(file_path, 'r', encoding=encoding) as f:
                content = f.read()
            used_encoding = encoding
            break
        except (UnicodeDecodeError, UnicodeError):
            continue

    if content is None:
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            content = f.read()

    if track_changes:
        for orig, sugg in replacements:
            if orig and orig != sugg:
                if sugg:
                    content = content.replace(orig, f'【删除:{orig}→修改为:{sugg}】')
                else:
                    content = content.replace(orig, f'【删除:{orig}】')
    else:
        for orig, sugg in replacements:
            if orig:
                content = content.replace(orig, sugg)

    buffer = io.BytesIO(content.encode(used_encoding))
    return send_file(
        buffer,
        as_attachment=True,
        download_name=_export_filename(base_name, ext)
    )


def _export_pdf_inplace(file_path, replacements, base_name):
    """在原始 PDF 上做文本替换"""
    try:
        import fitz  # PyMuPDF
    except ImportError:
        return jsonify({'error': 'PDF 格式需要 PyMuPDF 库支持，请联系管理员安装'}), 500

    doc = fitz.open(file_path)

    def has_cjk(text):
        for ch in text:
            if '\u4e00' <= ch <= '\u9fff' or '\u3400' <= ch <= '\u4dbf':
                return True
        return False

    for page in doc:
        # 收集所有需要替换的位置
        redactions = []
        for orig, sugg in replacements:
            if not orig:
                continue
            instances = page.search_for(orig)
            for inst in instances:
                # 获取原始文字的字号信息
                font_size = 11
                blocks = page.get_text('dict').get('blocks', [])
                for block in blocks:
                    if 'lines' not in block:
                        continue
                    for line in block['lines']:
                        for span in line['spans']:
                            span_rect = fitz.Rect(span['bbox'])
                            if span_rect.intersects(inst):
                                font_size = span['size']
                                break
                redactions.append((inst, sugg, font_size))

        # 先添加所有 redaction 注解
        for inst, sugg, font_size in redactions:
            page.add_redact_annot(inst)

        # 应用 redaction（删除原文）
        page.apply_redactions()

        # 在原位置插入替换文本
        for inst, sugg, font_size in redactions:
            if not sugg:
                continue
            font_name = 'china-s' if has_cjk(sugg) else 'helv'
            try:
                page.insert_text(
                    fitz.Point(inst.x0, inst.y1 - 2),
                    sugg,
                    fontsize=font_size,
                    fontname=font_name,
                )
            except Exception:
                try:
                    page.insert_text(
                        fitz.Point(inst.x0, inst.y1 - 2),
                        sugg,
                        fontsize=font_size,
                        fontname='helv',
                    )
                except Exception:
                    pass

    output_path = os.path.join(UPLOAD_DIR, f'output_{uuid.uuid4()}.pdf')
    doc.save(output_path)
    doc.close()

    return send_file(
        output_path,
        as_attachment=True,
        download_name=_export_filename(base_name, '.pdf')
    )


def _export_pdf_track_changes(file_path, replacements, base_name):
    """在原始 PDF 上添加高亮注释标记修改位置"""
    try:
        import fitz  # PyMuPDF
    except ImportError:
        return jsonify({'error': 'PDF 格式需要 PyMuPDF 库支持，请联系管理员安装'}), 500

    doc = fitz.open(file_path)

    for page in doc:
        for orig, sugg in replacements:
            if not orig or orig == sugg:
                continue
            instances = page.search_for(orig)
            for inst in instances:
                # 添加高亮注释（黄色高亮）
                annot = page.add_highlight_annot(inst)
                # 设置注释信息（包含修改建议）
                info_content = f"原文: {orig}"
                if sugg:
                    info_content += f"\n建议修改为: {sugg}"
                else:
                    info_content += "\n建议: 删除"
                annot.set_info(title="啄木鸟·中英文字智能检查 - 修改建议", content=info_content)
                annot.update()

                # 如果有修改建议，在原文旁边插入红色标注文本
                if sugg:
                    try:
                        # 获取原文的字体大小
                        font_size = 11
                        blocks = page.get_text('dict').get('blocks', [])
                        for block in blocks:
                            if 'lines' not in block:
                                continue
                            for line in block['lines']:
                                for span in line['spans']:
                                    span_rect = fitz.Rect(span['bbox'])
                                    if span_rect.intersects(inst):
                                        font_size = span['size']
                                        break

                        # 在原文下方插入建议文本（红色）
                        annot_point = fitz.Point(inst.x0, inst.y1 + font_size + 2)
                        page.insert_text(
                            annot_point,
                            f"→{sugg}",
                            fontsize=font_size * 0.8,
                            fontname='helv',
                            color=(0.8, 0.0, 0.0),
                        )
                    except Exception:
                        pass

    output_path = os.path.join(UPLOAD_DIR, f'output_{uuid.uuid4()}.pdf')
    doc.save(output_path)
    doc.close()

    return send_file(
        output_path,
        as_attachment=True,
        download_name=_export_filename(base_name, '.pdf')
    )


def _export_rtf_inplace(file_path, replacements, base_name, track_changes=False):
    """在原始 RTF 文件上做文本替换"""

    def text_to_rtf_unicode(text):
        """将文本转换为 RTF Unicode 转义格式"""
        result = []
        for ch in text:
            code = ord(ch)
            if code > 127:
                result.append(f'\\u{code}?')
            elif ch in ('\\', '{', '}'):
                result.append('\\' + ch)
            else:
                result.append(ch)
        return ''.join(result)

    # 读取原始 RTF
    content = None
    for encoding in ['utf-8', 'latin-1']:
        try:
            with open(file_path, 'r', encoding=encoding) as f:
                content = f.read()
            break
        except (UnicodeDecodeError, UnicodeError):
            continue

    if content is None:
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            content = f.read()

    if track_changes:
        # 修订模式：用删除线标记原文 + 下划线标记修改后文本
        for orig, sugg in replacements:
            if not orig or orig == sugg:
                continue
            orig_rtf = text_to_rtf_unicode(orig)
            sugg_rtf = text_to_rtf_unicode(sugg) if sugg else ''
            # 尝试直接替换
            if orig in content:
                if sugg:
                    replacement = '{\\strike ' + orig + '}{\\uld\\cf1 ' + sugg + '}'
                else:
                    replacement = '{\\strike ' + orig + '}'
                content = content.replace(orig, replacement)
            elif orig_rtf in content:
                if sugg:
                    replacement = '{\\strike ' + orig_rtf + '}{\\uld\\cf1 ' + sugg_rtf + '}'
                else:
                    replacement = '{\\strike ' + orig_rtf + '}'
                content = content.replace(orig_rtf, replacement)
    else:
        for orig, sugg in replacements:
            if not orig:
                continue
            if orig in content:
                content = content.replace(orig, sugg)
            else:
                orig_rtf = text_to_rtf_unicode(orig)
                sugg_rtf = text_to_rtf_unicode(sugg)
                if orig_rtf in content:
                    content = content.replace(orig_rtf, sugg_rtf)

    buffer = io.BytesIO(content.encode('utf-8'))
    return send_file(
        buffer,
        as_attachment=True,
        download_name=_export_filename(base_name, '.rtf')
    )


def generate_report_html(data: dict) -> str:
    """生成 HTML 格式的分析报告"""
    filename = data.get('filename', '未知')
    stats = data.get('stats', {})
    summary = data.get('summary', {})
    issues = data.get('issues', [])

    type_colors = {
        '错别字': '#e74c3c',
        '异形词': '#fb7185',
        '全半角混用': '#ec4899',
        '漏字/缺字': '#e91e63',
        '成语误用': '#d946ef',
        '语病/表达': '#6366f1',
        '语法': '#9b59b6',
        '逻辑': '#e91e63',
        '标点符号': '#e67e22',
        '多余空格': '#3498db',
        '数字/格式': '#0ea5e9',
        '重复词语': '#f39c12',
        '文风/格式': '#1abc9c',
        '口语化': '#65a30d',
    }
    sev_colors = {
        '错误': '#e74c3c',
        '警告': '#f39c12',
        '建议': '#3498db',
    }

    issues_html = ''
    for i, issue in enumerate(issues, 1):
        type_color = type_colors.get(issue.get('type', ''), '#999')
        sev_color = sev_colors.get(issue.get('severity', ''), '#999')
        layer = issue.get('layer', '')
        layer_names = {
            'character': '字符层',
            'vocabulary': '词汇层',
            'sentence': '句子层',
            'format': '标点/格式层',
            'discourse': '语篇/语体层',
        }
        layer_cn = layer_names.get(layer, layer)
        issues_html += f'''
        <tr>
            <td>{i}</td>
            <td><span class="badge" style="background:#64748b">{html.escape(str(layer_cn))}</span></td>
            <td><span class="badge" style="background:{type_color}">{html.escape(str(issue.get('type', '')))}</span></td>
            <td><span class="badge" style="background:{sev_color}">{html.escape(str(issue.get('severity', '')))}</span></td>
            <td class="original">{html.escape(str(issue.get('original', '')))}</td>
            <td class="suggestion">{html.escape(str(issue.get('suggestion', '')))}</td>
            <td class="desc">{html.escape(str(issue.get('description', '')))}</td>
            <td class="context">{html.escape(str(issue.get('context', '')))}</td>
        </tr>'''

    by_type_html = ''
    for t, count in summary.get('by_type', {}).items():
        color = type_colors.get(t, '#999')
        by_type_html += f'<span class="stat-chip" style="border-color:{color};color:{color}">{t}: {count}</span>'

    by_sev_html = ''
    for s, count in summary.get('by_severity', {}).items():
        color = sev_colors.get(s, '#999')
        by_sev_html += f'<span class="stat-chip" style="border-color:{color};color:{color}">{s}: {count}</span>'

    return f'''<!DOCTYPE html>
<html lang="zh-CN">
<head>
<meta charset="UTF-8">
<title>原文检查报告 - {filename}</title>
<style>
    * {{ margin: 0; padding: 0; box-sizing: border-box; }}
    body {{ font-family: -apple-system, "PingFang SC", "Microsoft YaHei", sans-serif; background: #f5f6fa; color: #333; padding: 30px; }}
    .container {{ max-width: 1100px; margin: 0 auto; }}
    h1 {{ font-size: 24px; margin-bottom: 5px; }}
    .subtitle {{ color: #888; font-size: 14px; margin-bottom: 20px; }}
    .stats-bar {{ background: #fff; border-radius: 10px; padding: 16px 24px; margin-bottom: 20px; display: flex; gap: 30px; flex-wrap: wrap; box-shadow: 0 1px 3px rgba(0,0,0,0.08); }}
    .stat-item {{ font-size: 14px; }}
    .stat-item strong {{ font-size: 20px; color: #2c3e50; display: block; }}
    .chips {{ margin-bottom: 20px; display: flex; gap: 8px; flex-wrap: wrap; }}
    .stat-chip {{ padding: 4px 12px; border: 1px solid; border-radius: 20px; font-size: 13px; font-weight: 600; }}
    table {{ width: 100%; border-collapse: collapse; background: #fff; border-radius: 10px; overflow: hidden; box-shadow: 0 1px 3px rgba(0,0,0,0.08); }}
    th {{ background: #2c3e50; color: #fff; padding: 12px 14px; text-align: left; font-size: 13px; white-space: nowrap; }}
    td {{ padding: 10px 14px; border-bottom: 1px solid #eee; font-size: 13px; vertical-align: top; }}
    tr:hover {{ background: #f8f9fa; }}
    .badge {{ color: #fff; padding: 2px 8px; border-radius: 4px; font-size: 12px; white-space: nowrap; }}
    .original {{ color: #e74c3c; font-weight: 600; }}
    .suggestion {{ color: #27ae60; font-weight: 600; }}
    .desc {{ color: #555; max-width: 250px; }}
    .context {{ color: #999; font-size: 12px; max-width: 200px; }}
    .footer {{ text-align: center; margin-top: 30px; color: #aaa; font-size: 12px; }}
</style>
</head>
<body>
<div class="container">
    <h1>原文检查报告</h1>
    <p class="subtitle">文件名: {filename}</p>
    <div class="stats-bar">
        <div class="stat-item"><strong>{stats.get('char_count', 0)}</strong>总字符数</div>
        <div class="stat-item"><strong>{stats.get('char_count_no_space', 0)}</strong>不含空格字符数</div>
        <div class="stat-item"><strong>{stats.get('line_count', 0)}</strong>行数</div>
        <div class="stat-item"><strong>{summary.get('total', 0)}</strong>发现问题数</div>
    </div>
    <div class="chips">{by_type_html}</div>
    <div class="chips">{by_sev_html}</div>
    <table>
        <thead>
            <tr><th>#</th><th>层级</th><th>类型</th><th>级别</th><th>原文</th><th>建议修改</th><th>说明</th><th>上下文</th></tr>
        </thead>
        <tbody>
            {issues_html if issues_html else '<tr><td colspan="9" style="text-align:center;padding:40px;color:#999;">未检测到明显问题</td></tr>'}
        </tbody>
    </table>
    <div class="footer">由啄木鸟·中英文字智能检查生成</div>
</div>
</body>
</html>'''


if __name__ == '__main__':
    app.run(host='0.0.0.0', port=5088, debug=False)
