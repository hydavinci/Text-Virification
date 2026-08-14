"""
文件解析模块
支持 txt、docx、pdf、rtf 等格式的文件解析。
解析时追踪文本在原文中的位置信息（页码/段落号），方便用户定位。
"""

import os
import re
import shutil
import subprocess
import tempfile
from html.parser import HTMLParser
from typing import Tuple, Optional, List


class _HTMLStripper(HTMLParser):
    """用于剥离 HTML/XML 标签，只保留可见文本。"""

    def __init__(self):
        super().__init__()
        self.text = []

    def handle_data(self, data):
        self.text.append(data)

    def get_data(self):
        return ''.join(self.text)


def _strip_html(text: str) -> str:
    """
    安全地移除文本中的 HTML/XML 标签，只保留可见文本。

    如果标准库解析失败，则退回到正则表达式兜底。
    同时修复标签被错误截断后留下的零散属性文本（如 class="..."）。
    """
    if not text or '<' not in text or '>' not in text:
        return text
    try:
        stripper = _HTMLStripper()
        stripper.feed(text)
        result = stripper.get_data()
    except Exception:
        result = re.sub(r'<[^>]+>', '', text)

    # 防御性清理：去除标签被错误截断后残留的零散 HTML 属性片段
    # 例如 "class=\"issue-highlight\""、"data-issue-id=..." 等
    result = re.sub(r'\b[a-zA-Z][\w\-]*\s*=\s*"[^"]*"', '', result)
    result = re.sub(r'\b[a-zA-Z][\w\-]*\s*=\s*\'[^\']*\'', '', result)
    result = re.sub(r'[<>]', '', result)
    return result


# page_map 中每项为 (start_offset, end_offset, label)
PageMap = List[Tuple[int, int, str]]


def parse_file(file_path: str, file_extension: Optional[str] = None) -> Tuple[str, str, PageMap]:
    """
    解析文件，提取纯文本内容，并记录各段文本在原文中的位置信息。

    Args:
        file_path: 文件路径
        file_extension: 文件扩展名（如未提供则从路径推断）

    Returns:
        (提取的文本内容, 文件格式名称, page_map)
        page_map 为 (起始偏移, 结束偏移, 位置标签) 列表，如 [(0, 120, '第1页'), ...]
        对于不支持位置追踪的格式，page_map 为空列表。
    """
    if file_extension is None:
        file_extension = os.path.splitext(file_path)[1].lower().lstrip('.')

    if file_extension == 'txt':
        return _parse_txt(file_path), 'txt', []
    elif file_extension == 'docx':
        text, page_map = _parse_docx(file_path)
        return text, 'docx', page_map
    elif file_extension == 'doc':
        # 旧版 Word 二进制格式（.doc）：先转为 .docx 再复用 docx 解析逻辑
        text, page_map = _parse_doc(file_path)
        return text, 'doc', page_map
    elif file_extension == 'pdf':
        text, page_map = _parse_pdf(file_path)
        return text, 'pdf', page_map
    elif file_extension == 'rtf':
        return _parse_rtf(file_path), 'rtf', []
    elif file_extension == 'md':
        return _parse_txt(file_path), 'markdown', []
    elif file_extension == 'csv':
        return _parse_txt(file_path), 'csv', []
    else:
        # 默认尝试以文本方式读取
        try:
            return _parse_txt(file_path), file_extension, []
        except Exception:
            raise ValueError(f'不支持的文件格式: .{file_extension}，目前支持 txt、docx、doc、pdf、rtf、md、csv')


def _parse_txt(file_path: str) -> str:
    """解析纯文本文件，自动检测编码，并清理可能混入的 HTML 代码"""
    encodings = ['utf-8', 'gbk', 'gb2312', 'big5', 'utf-16', 'latin-1']
    content = None
    for encoding in encodings:
        try:
            with open(file_path, 'r', encoding=encoding) as f:
                content = f.read()
            break
        except (UnicodeDecodeError, UnicodeError):
            continue
    if content is None:
        # 如果所有编码都失败，用 utf-8 忽略错误
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            content = f.read()
    return _strip_html(content)


def _parse_docx(file_path: str) -> Tuple[str, PageMap]:
    """解析 Word (.docx) 文件，追踪段落号，并清理混入的 HTML 代码"""
    from docx import Document
    doc = Document(file_path)
    paragraphs = []
    page_map: PageMap = []
    offset = 0
    para_num = 0

    for para in doc.paragraphs:
        text = _strip_html(para.text).strip()
        if text:
            para_num += 1
            page_map.append((offset, offset + len(text), f'第{para_num}段'))
            paragraphs.append(text)
            offset += len(text) + 1  # +1 for \n

    # 也提取表格中的文本
    for table in doc.tables:
        for row in table.rows:
            row_text = _strip_html(' | '.join(cell.text.strip() for cell in row.cells)).strip()
            if row_text:
                para_num += 1
                page_map.append((offset, offset + len(row_text), f'第{para_num}段'))
                paragraphs.append(row_text)
                offset += len(row_text) + 1

    return '\n'.join(paragraphs), page_map


def _convert_doc_to_docx(file_path: str) -> str:
    """将旧版 Word (.doc) 二进制格式转换为 .docx，便于解析与导出。

    优先使用 macOS 内置的 textutil；若不可用则尝试 LibreOffice (soffice)。
    返回转换后的 .docx 临时文件路径（调用方负责清理）。
    """
    fd, out_path = tempfile.mkstemp(suffix='.docx')
    os.close(fd)

    # 1) macOS 内置 textutil（最可靠、零依赖）
    if shutil.which('textutil'):
        try:
            subprocess.run(
                ['textutil', '-convert', 'docx', file_path, '-output', out_path],
                check=True, capture_output=True, timeout=120,
            )
            if os.path.exists(out_path) and os.path.getsize(out_path) > 0:
                return out_path
        except (subprocess.CalledProcessError, subprocess.TimeoutExpired, OSError):
            pass

    # 2) LibreOffice 回退（Linux / 未安装 textutil 的环境）
    soffice = shutil.which('soffice') or shutil.which('libreoffice')
    if soffice:
        work_dir = os.path.dirname(out_path)
        try:
            subprocess.run(
                [soffice, '--headless', '--convert-to', 'docx', '--outdir', work_dir, file_path],
                check=True, capture_output=True, timeout=180,
            )
            candidate = os.path.join(
                work_dir,
                os.path.splitext(os.path.basename(file_path))[0] + '.docx'
            )
            if os.path.exists(candidate) and os.path.getsize(candidate) > 0:
                if candidate != out_path:
                    os.replace(candidate, out_path)
                return out_path
        except (subprocess.CalledProcessError, subprocess.TimeoutExpired, OSError):
            pass

    # 转换失败，清理临时文件
    try:
        if os.path.exists(out_path):
            os.unlink(out_path)
    except OSError:
        pass
    raise ValueError('无法处理 .doc 文件：本机未安装 textutil 或 LibreOffice，请先安装其中之一')


def _parse_doc(file_path: str) -> Tuple[str, PageMap]:
    """解析旧版 Word (.doc) 文件：借助 textutil/LibreOffice 转为 docx 后复用 docx 解析。"""
    docx_path = _convert_doc_to_docx(file_path)
    try:
        return _parse_docx(docx_path)
    finally:
        try:
            if docx_path and os.path.exists(docx_path):
                os.unlink(docx_path)
        except OSError:
            pass


def _parse_pdf(file_path: str) -> Tuple[str, PageMap]:
    """解析 PDF 文件，追踪页码，并清理混入的 HTML 代码"""
    import pdfplumber
    texts = []
    page_map: PageMap = []
    offset = 0
    with pdfplumber.open(file_path) as pdf:
        for i, page in enumerate(pdf.pages, 1):
            text = _strip_html(page.extract_text() or '').strip()
            if text:
                page_map.append((offset, offset + len(text), f'第{i}页'))
                texts.append(text)
                offset += len(text) + 1  # +1 for \n join
    return '\n'.join(texts), page_map


def _parse_rtf(file_path: str) -> str:
    """简单解析 RTF 文件"""
    try:
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            content = f.read()
    except Exception:
        with open(file_path, 'r', encoding='latin-1') as f:
            content = f.read()

    # 简单的 RTF 标签清理
    # 移除 RTF 控制字
    content = re.sub(r'\\\'[0-9a-fA-F]{2}', '', content)
    # 移除 RTF 控制词
    content = re.sub(r'\\[a-zA-Z]+-?\d*\s?', '', content)
    # 移除特殊字符
    content = re.sub(r'[{}]', '', content)
    # 处理转义字符
    content = content.replace('\\\\', '\\').replace('\\{', '{').replace('\\}', '}')
    # 清理多余空行
    content = re.sub(r'\n{3,}', '\n\n', content)
    # 清理可能混入的 HTML 代码
    return _strip_html(content)


# 对外导出 HTML 清理函数，供导出模块复用
strip_html = _strip_html


def get_supported_formats() -> list:
    """返回支持的文件格式列表"""
    return [
        {'ext': 'txt', 'name': '纯文本文件', 'accept': '.txt'},
        {'ext': 'docx', 'name': 'Word 文档', 'accept': '.docx'},
        {'ext': 'doc', 'name': '旧版 Word 文档', 'accept': '.doc'},
        {'ext': 'pdf', 'name': 'PDF 文档', 'accept': '.pdf'},
        {'ext': 'rtf', 'name': 'RTF 文档', 'accept': '.rtf'},
        {'ext': 'md', 'name': 'Markdown 文件', 'accept': '.md'},
        {'ext': 'csv', 'name': 'CSV 文件', 'accept': '.csv'},
    ]
