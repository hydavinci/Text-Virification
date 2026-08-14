"""
Word 文档生成模块
将 DocElement 列表转换为可编辑 .docx 文件
最大程度保留原文档样式（字号、标题层级、表格、图片）
"""

import io
from typing import List
from docx import Document
from docx.shared import Pt, Inches, Cm, Emu, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from ocr_engine import DocElement


def build_word(elements: List[DocElement], output_path: str = None,
               original_name: str = '') -> bytes:
    """将 DocElement 列表生成 .docx 文件

    Args:
        elements: 文档元素列表
        output_path: 输出路径（如果为 None，返回 bytes）
        original_name: 原始文件名（用于元数据）

    Returns:
        bytes (如果 output_path is None) or None (文件已保存)
    """
    doc = Document()

    # 设置默认字体
    style = doc.styles['Normal']
    font = style.font
    font.name = '宋体'
    font.size = Pt(11)
    # 设置东亚字体
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # 设置页面边距
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.17)
        section.right_margin = Cm(3.17)

    # 文档元数据
    if original_name:
        doc.core_properties.title = original_name

    for elem in elements:
        if elem.etype == 'heading':
            _add_heading(doc, elem)
        elif elem.etype == 'paragraph':
            _add_paragraph(doc, elem)
        elif elem.etype == 'table':
            _add_table(doc, elem)
        elif elem.etype == 'image':
            _add_image(doc, elem)

    # 保存
    if output_path:
        doc.save(output_path)
        return None
    else:
        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()


def _add_heading(doc: Document, elem: DocElement):
    """添加标题"""
    level = min(elem.level, 6) if elem.level > 0 else 1
    heading = doc.add_heading(level=level)

    run = heading.add_run(elem.text)
    run.font.size = Pt(elem.font_size if elem.font_size > 0 else 14)

    # 设置东亚字体
    _set_east_asia_font(run, _get_font_name(elem.text))


def _add_paragraph(doc: Document, elem: DocElement):
    """添加段落"""
    # 处理多行文本
    lines = elem.text.split('\n')
    if not lines:
        return

    para = doc.add_paragraph()
    para_format = para.paragraph_format
    para_format.space_before = Pt(0)
    para_format.space_after = Pt(6)
    para_format.line_spacing = 1.15

    for i, line in enumerate(lines):
        if i > 0:
            run = para.add_run()
            run.add_break()
        run = para.add_run(line)
        font_size = elem.font_size if elem.font_size > 0 else 11
        run.font.size = Pt(font_size)
        _set_east_asia_font(run, _get_font_name(line))


def _add_table(doc: Document, elem: DocElement):
    """添加表格"""
    if not elem.rows:
        return

    num_cols = max(len(row) for row in elem.rows) if elem.rows else 0
    if num_cols == 0:
        return

    # 创建表格
    table = doc.add_table(rows=len(elem.rows), cols=num_cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 填充数据
    for i, row_data in enumerate(elem.rows):
        for j, cell_text in enumerate(row_data):
            if j < num_cols:
                cell = table.cell(i, j)
                cell.text = cell_text or ''
                # 设置单元格字体
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.size = Pt(10)
                        _set_east_asia_font(run, _get_font_name(cell_text or ''))

    # 添加空段落防止表格粘连
    doc.add_paragraph()


def _add_image(doc: Document, elem: DocElement):
    """添加图片"""
    if not elem.image_data:
        return

    try:
        buf = io.BytesIO(elem.image_data)
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run()

        # 根据图片尺寸调整宽度
        from PIL import Image
        img = Image.open(io.BytesIO(elem.image_data))
        img_w, img_h = img.size

        # 页面可用宽度（A4 - 左右边距）
        max_width_inch = 5.5  # 约 14cm
        width_inch = min(img_w / 96.0, max_width_inch)

        run.add_picture(buf, width=Inches(width_inch))
    except Exception:
        # 图片插入失败时添加占位文本
        para = doc.add_paragraph('[图片]')
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _get_font_name(text: str) -> str:
    """根据文本内容选择合适的字体"""
    has_cjk = any('\u4e00' <= ch <= '\u9fff' for ch in text)
    has_kana = any('\u3040' <= ch <= '\u30ff' for ch in text)

    if has_kana:
        return 'MS Mincho'  # 日文用明朝体
    elif has_cjk:
        return '宋体'  # 中文用宋体
    else:
        return 'Times New Roman'


def _set_east_asia_font(run, font_name: str):
    """设置 run 的东亚字体"""
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), font_name)
    # 同时设置 ascii 和 hAnsi
    if font_name not in ('宋体',):
        rFonts.set(qn('w:ascii'), font_name)
        rFonts.set(qn('w:hAnsi'), font_name)
